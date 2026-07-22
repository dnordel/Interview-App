from datetime import date
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from data_store import InterviewHistoryStore
from hiring_pipeline import (
    HiringOfferNotificationAdapter,
    HiringPipelineStore,
    HiringStage,
    HiringWorkflowService,
    calculate_offer_approval_dates,
    normalize_candidate_phone,
)
from onboarding_service import OnboardingAccess, OnboardingService
from onboarding_store import OnboardingStore


def test_offer_template_snapshot_survives_missing_dropbox_source(tmp_path: Path) -> None:
    store = HiringPipelineStore(tmp_path / "interview_history.sqlite3")
    source = tmp_path / "Dropbox" / "offer-template.docx"
    source.parent.mkdir()
    Document().save(source)
    expected = source.read_bytes()

    template_id = store.snapshot_offer_template(
        school="Palmdale",
        template_kind="full_time",
        source_path=source,
    )
    source.unlink()

    materialized = store.materialize_offer_template(
        school="Palmdale",
        template_kind="full_time",
        destination_dir=tmp_path / "template-cache",
    )

    assert template_id
    assert materialized.parent == (tmp_path / "template-cache").resolve()
    assert materialized.suffix == ".docx"
    assert materialized.read_bytes() == expected


def test_offer_version_uses_its_exact_template_snapshot_after_active_template_changes(
    tmp_path: Path,
) -> None:
    store = HiringPipelineStore(tmp_path / "interview_history.sqlite3")
    first = tmp_path / "first.docx"
    first_doc = Document()
    first_doc.add_paragraph("first template")
    first_doc.save(first)
    second = tmp_path / "second.docx"
    second_doc = Document()
    second_doc.add_paragraph("second template")
    second_doc.save(second)

    store.snapshot_offer_template(
        school="Palmdale",
        template_kind="full_time",
        source_path=first,
        version_id="offer-v1",
    )
    store.snapshot_offer_template(
        school="Palmdale",
        template_kind="full_time",
        source_path=second,
        version_id="offer-v2",
    )

    exact = store.materialize_offer_template(
        school="Palmdale",
        template_kind="full_time",
        destination_dir=tmp_path / "exact-cache",
        version_id="offer-v1",
    )
    active = store.materialize_offer_template(
        school="Palmdale",
        template_kind="full_time",
        destination_dir=tmp_path / "active-cache",
    )

    assert exact.read_bytes() == first.read_bytes()
    assert active.read_bytes() == second.read_bytes()


def test_offer_template_version_binding_cannot_cross_school_scope(tmp_path: Path) -> None:
    store = HiringPipelineStore(tmp_path / "interview_history.sqlite3")
    source = tmp_path / "template.docx"
    Document().save(source)
    store.snapshot_offer_template(
        school="Palmdale",
        template_kind="full_time",
        source_path=source,
        version_id="offer-v1",
    )

    with pytest.raises(ValueError, match="No stored full-time offer template"):
        store.materialize_offer_template(
            school="Hawthorne",
            template_kind="full_time",
            destination_dir=tmp_path / "cache",
            version_id="offer-v1",
        )


def test_offer_template_snapshot_rejects_non_word_content(tmp_path: Path) -> None:
    store = HiringPipelineStore(tmp_path / "interview_history.sqlite3")
    source = tmp_path / "fake.docx"
    source.write_bytes(b"not a Word package")

    with pytest.raises(ValueError, match="not a valid Word document"):
        store.snapshot_offer_template(
            school="Palmdale",
            template_kind="full_time",
            source_path=source,
        )


def test_offer_approval_artifact_stage_promotes_previewed_bytes_unchanged(tmp_path: Path) -> None:
    from hiring_pipeline import OfferApprovalArtifactStage

    final_docx = tmp_path / "offers" / "approved-offer.docx"
    stage = OfferApprovalArtifactStage(final_docx)
    stage.staged_docx_path.write_bytes(b"approved docx bytes")
    stage.staged_pdf_path.write_bytes(b"approved pdf bytes")
    before = stage.hashes()

    promoted = stage.promote()

    assert promoted.docx_path == final_docx.resolve()
    assert promoted.pdf_path == final_docx.with_suffix(".pdf").resolve()
    assert promoted.docx_sha256 == before[0]
    assert promoted.pdf_sha256 == before[1]
    assert promoted.docx_path.read_bytes() == b"approved docx bytes"
    assert promoted.pdf_path.read_bytes() == b"approved pdf bytes"
    assert not stage.staging_dir.exists()


def test_completed_qualifying_interview_enters_director_review(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-65",
            "candidate_name": "  Maya   Patel ",
            "candidate_email": "maya@example.com",
            "school": "Palmdale",
            "position": "Preschool",
            "interview_date": "2026-07-14",
            "score": 65.0,
            "outcome": "Borderline",
        }
    )

    service = HiringWorkflowService(HiringPipelineStore(history_path))
    application = service.record_initial_interview(
        history_id="hist-65",
        legal_name="  Maya   Patel ",
        email="maya@example.com",
        phone="",
        school="Palmdale",
        position="Preschool",
        score=65.0,
        outcome="Borderline",
    )

    assert application.stage is HiringStage.DIRECTOR_REVIEW
    assert application.history_id == "hist-65"
    assert application.candidate_id
    assert application.application_id
    assert service.store.get_candidate(application.candidate_id).legal_name == "Maya Patel"
    linked_history = history.load()[0]
    assert linked_history["candidate_id"] == application.candidate_id
    assert linked_history["application_id"] == application.application_id


def test_backfill_groups_exact_name_school_and_keeps_application_cycles(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    history = InterviewHistoryStore(history_path)
    for history_id, position, score, outcome in (
        ("hist-1", "Preschool", 64.99, "No Hire"),
        ("hist-2", "Infant/Toddler", 65.0, "Borderline"),
        ("hist-3", "Preschool", 80.0, "Hire"),
    ):
        history.append(
            {
                "history_id": history_id,
                "candidate_name": "Maya Patel",
                "school": "Palmdale",
                "position": position,
                "interview_date": "2026-07-14",
                "score": score,
                "outcome": outcome,
            }
        )

    service = HiringWorkflowService(HiringPipelineStore(history_path))
    service.backfill_history()
    service.backfill_history()

    applications = service.store.list_applications()
    assert len(applications) == 3
    assert len({item.candidate_id for item in applications}) == 1
    assert [(item.position, item.cycle_number, item.stage) for item in applications] == [
        ("Preschool", 1, HiringStage.CLOSED),
        ("Infant/Toddler", 1, HiringStage.DIRECTOR_REVIEW),
        ("Preschool", 2, HiringStage.DIRECTOR_REVIEW),
    ]


def test_offer_approval_dates_use_three_calendar_days_and_first_monday() -> None:
    dates = calculate_offer_approval_dates(date(2026, 7, 14))

    assert dates.offer_date == date(2026, 7, 14)
    assert dates.reply_by_date == date(2026, 7, 17)
    assert dates.start_date == date(2026, 8, 3)


def test_normalize_candidate_phone_accepts_common_us_format() -> None:
    assert normalize_candidate_phone("310.555.0199") == "(310) 555-0199"


def test_calculated_external_offer_for_existing_candidate_ignores_prior_application_stage(tmp_path: Path) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "existing-candidate-offer.sqlite3"))
    prior = service.start_application(
        legal_name="Maya Patel",
        email="maya@example.com",
        phone="(310) 555-0199",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )

    offer = service.create_calculated_external_offer_for_candidate(
        candidate_id=prior.candidate_id,
        school="Hawthorne",
        position_id="teacher",
        qualification={
            "has_degree": False,
            "degree_type": "",
            "degree_in_ece": False,
            "ece_units_completed": "24",
            "infant_toddler_class_completed": True,
            "total_units_completed": "40",
            "years_experience": 3,
        },
        terms={"weekly_hours": "40", "template_path": "offer.docx", "output_dir": "offers"},
        actor="Admin",
    )

    created = service.store.get_application(offer.application_id)
    assert created.candidate_id == prior.candidate_id
    assert created.school == "Hawthorne"
    assert created.stage is HiringStage.EXECUTIVE_APPROVAL
    assert service.store.get_application(prior.application_id).stage is HiringStage.INITIAL_INTERVIEW
    assert offer.status == "pending_approval"
    assert offer.terms["qualification_snapshot"]["infant_toddler_class_completed"] is True


def test_offer_drafts_are_versioned_and_selected_version_is_submitted(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-offer",
            "candidate_name": "Maya Patel",
            "school": "Palmdale",
            "position": "Preschool",
            "score": 80,
            "outcome": "Hire",
        }
    )
    service = HiringWorkflowService(HiringPipelineStore(history_path))
    application = service.record_initial_interview(
        history_id="hist-offer",
        legal_name="Maya Patel",
        email="maya@example.com",
        phone="",
        school="Palmdale",
        position="Preschool",
        score=80,
        outcome="Hire",
    )
    service.record_director_decision(application.application_id, decision="hire", actor="Director")

    first = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": 24.0, "hours_week": 40},
        actor="HR Manager",
    )
    second = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": 25.0, "hours_week": 40},
        actor="HR Manager",
    )
    submitted = service.submit_offer_for_approval(
        application.application_id,
        second.version_id,
        actor="HR Manager",
    )

    assert (first.version_number, second.version_number) == (1, 2)
    assert first.version_id != second.version_id
    assert submitted.status == "pending_approval"
    assert service.store.get_application(application.application_id).stage is HiringStage.EXECUTIVE_APPROVAL
    assert [version.status for version in service.store.list_offer_versions(application.application_id)] == [
        "draft",
        "pending_approval",
    ]


def test_offer_submission_generates_and_persists_docx_and_pdf_together(tmp_path: Path) -> None:
    calls: list[str] = []

    def prepare(_application, _candidate, version):
        calls.append(version.version_id)
        docx = tmp_path / "offer.docx"
        pdf = tmp_path / "offer.pdf"
        docx.write_bytes(b"docx")
        pdf.write_bytes(b"%PDF")
        return docx, pdf

    service = HiringWorkflowService(
        HiringPipelineStore(tmp_path / "history.sqlite3"),
        prepare_offer_artifacts=prepare,
    )
    application = service.start_external_offer_application(
        legal_name="Maya Patel",
        email="maya@example.com",
        phone="",
        school="Palmdale",
        position="Teacher",
        actor="Admin",
        honorific="Ms.",
    )
    draft = service.create_offer_draft(
        application.application_id,
        terms={"school": "Palmdale", "position": "Teacher"},
        actor="Admin",
    )

    submitted = service.submit_offer_for_approval(
        application.application_id,
        draft.version_id,
        actor="Admin",
    )

    assert calls == [draft.version_id]
    assert submitted.docx_path == str((tmp_path / "offer.docx").resolve())
    assert submitted.pdf_path == str((tmp_path / "offer.pdf").resolve())
    assert submitted.docx_sha256
    assert submitted.pdf_sha256


def test_director_offer_submission_is_idempotent_by_source_key(tmp_path: Path) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))
    application = service.start_application(
        legal_name="Maya Patel", email="maya@example.com", phone="(310) 555-0199",
        school="Palmdale", position="Teacher", actor="Admin", honorific="Ms.",
    )
    service.finalize_initial_interview(
        application.application_id, history_id="hist-director-offer", score=80, outcome="Hire", actor="Admin",
    )
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    terms = {"hourly_pay": "24.00", "weekly_hours": "40"}

    first = service.ensure_director_offer_submitted(
        application.application_id, source_key="director-interview:17:v1", terms=terms, actor="Director",
    )
    second = service.ensure_director_offer_submitted(
        application.application_id, source_key="director-interview:17:v1", terms=terms, actor="Director",
    )

    assert first.version_id == second.version_id
    assert first.status == "pending_approval"
    assert len(service.store.list_offer_versions(application.application_id)) == 1


def test_external_offer_creates_offer_ready_application_and_submits(tmp_path: Path) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))

    submitted = service.create_external_offer(
        legal_name="External Candidate",
        email="external@example.com",
        phone="555-0100",
        honorific="Ms.",
        school="Palmdale",
        position="Preschool",
        terms={
            "candidate_name": "External Candidate",
            "candidate_email": "external@example.com",
            "school": "Palmdale",
            "position": "Preschool",
            "hourly_pay": "24.00",
            "weekly_hours": "40",
            "template_path": "offer.docx",
            "output_dir": "offers",
        },
        actor="Admin",
    )

    application = service.store.get_application(submitted.application_id)
    candidate = service.store.get_candidate(application.candidate_id)
    events = [event.event_type for event in service.store.list_events(application.application_id)]

    assert candidate.legal_name == "External Candidate"
    assert application.history_id.startswith("external_offer:")
    assert application.stage is HiringStage.EXECUTIVE_APPROVAL
    assert submitted.status == "pending_approval"
    assert submitted.version_number == 1
    assert events == [
        "external_offer_application_created",
        "offer_draft_created",
        "offer_submitted_for_approval",
    ]


def test_external_offer_calculates_pay_from_qualification_snapshot(tmp_path: Path) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))

    created = service.create_calculated_external_offer(
        legal_name="Jordan Lee",
        email="jordan@example.org",
        phone="",
        school="Palmdale",
        position_id="teacher",
        qualification={
            "has_degree": False,
            "degree_type": "",
            "degree_in_ece": False,
            "ece_units_completed": "24",
            "total_units_completed": "40",
            "years_experience": 3,
        },
        terms={"weekly_hours": "40"},
        actor="Admin",
    )

    assert created.terms["position"] == "Teacher"
    assert created.terms["hourly_pay"] == "19.00"
    assert created.terms["pay_calculation"]["career_lattice_level"] == 5


def test_correcting_pending_offer_qualification_recalculates_numbered_version(
    tmp_path: Path,
) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))
    original = service.create_calculated_external_offer(
        legal_name="Allison Example",
        email="allison@example.org",
        phone="",
        school="Palmdale",
        position_id="teacher",
        qualification={
            "has_degree": True,
            "degree_type": "AA",
            "degree_in_ece": False,
            "ece_units_completed": "18",
            "total_units_completed": None,
            "years_experience": 4,
        },
        terms={"weekly_hours": "40"},
        actor="Admin",
    )

    corrected = service.revise_pending_offer_qualification(
        original.application_id,
        original.version_id,
        qualification={
            **original.terms["qualification_snapshot"],
            "degree_in_ece": True,
        },
        actor="Executive",
    )

    versions = service.store.list_offer_versions(original.application_id)
    assert [(version.version_number, version.status) for version in versions] == [
        (1, "superseded"),
        (2, "pending_approval"),
    ]
    assert corrected.terms["qualification_snapshot"]["degree_in_ece"] is True
    assert corrected.terms["pay_calculation"]["career_lattice_level"] == 6
    assert corrected.terms["pay_calculation"]["base_hourly_rate"] == "19.00"
    assert corrected.terms["pay_calculation"]["experience_adjustment"] == "1.50"
    assert corrected.terms["hourly_pay"] == "20.50"


def test_approved_offer_advances_only_after_pdf_delivery(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-approve",
            "candidate_name": "Maya Patel",
            "candidate_email": "maya@example.com",
            "school": "Palmdale",
            "position": "Preschool",
            "score": 80,
            "outcome": "Hire",
        }
    )
    sent: list[tuple[str, str]] = []

    def send_offer(candidate, version, pdf_path, idempotency_key):
        sent.append((candidate.email, Path(pdf_path).suffix))
        assert idempotency_key == f"offer-version:{version.version_id}"
        return [SimpleNamespace(status="sent", error="")]

    service = HiringWorkflowService(HiringPipelineStore(history_path), send_offer=send_offer)
    application = service.record_initial_interview(
        history_id="hist-approve",
        legal_name="Maya Patel",
        email="maya@example.com",
        phone="",
        school="Palmdale",
        position="Preschool",
        score=80,
        outcome="Hire",
    )
    service.record_director_decision(application.application_id, decision="hire", actor="Director")
    draft = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": 24.0, "hours_week": 40},
        actor="HR Manager",
    )
    service.submit_offer_for_approval(application.application_id, draft.version_id, actor="HR Manager")
    docx_path = tmp_path / "offer-v1.docx"
    pdf_path = tmp_path / "offer-v1.pdf"
    docx_path.write_bytes(b"docx")
    pdf_path.write_bytes(b"pdf")

    approved = service.approve_offer(
        application.application_id,
        draft.version_id,
        approver_name="Deidre Nordel",
        approver_role="Executive Director",
        approval_date=date(2026, 7, 14),
        docx_path=docx_path,
        pdf_path=pdf_path,
        rendered_email="Subject: Offer\n\nAttached is your approved offer.",
    )

    assert approved.status == "sent"
    assert approved.approval_date == "2026-07-14"
    assert approved.document_reply_by_date == "2026-07-17"
    assert approved.operational_reply_by_date == "2026-07-17"
    assert approved.start_date == "2026-08-03"
    assert len(approved.docx_sha256) == 64
    assert len(approved.pdf_sha256) == 64
    assert approved.rendered_email.startswith("Subject: Offer")
    assert approved.approved_at
    assert approved.sent_at
    assert sent == [("maya@example.com", ".pdf")]
    assert service.store.get_application(application.application_id).stage is HiringStage.OFFER_SENT


def test_offer_can_be_approved_without_candidate_email_and_waits_for_send(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    service = HiringWorkflowService(
        HiringPipelineStore(history_path),
        send_offer=lambda *_args: pytest.fail("Offer must not send without candidate email."),
    )
    application = service.start_application(
        legal_name="Maya Patel",
        email="",
        phone="",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )
    service.finalize_initial_interview(
        application.application_id,
        history_id="hist-approve-no-email",
        score=80,
        outcome="Hire",
        actor="Admin",
    )
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    draft = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "24.00", "weekly_hours": "40"},
        actor="Admin",
    )
    service.submit_offer_for_approval(application.application_id, draft.version_id, actor="Admin")
    docx_path = tmp_path / "offer.docx"
    pdf_path = tmp_path / "offer.pdf"
    docx_path.write_bytes(b"docx")
    pdf_path.write_bytes(b"pdf")

    approved = service.approve_offer(
        application.application_id,
        draft.version_id,
        approver_name="Executive",
        approver_role="Executive Director",
        approval_date=date(2026, 7, 14),
        docx_path=docx_path,
        pdf_path=pdf_path,
    )

    assert approved.status == "approved"
    assert approved.send_status == "pending"
    assert service.store.get_application(application.application_id).stage is HiringStage.EXECUTIVE_APPROVAL


def test_approved_unsent_offer_saves_candidate_email_then_sends(tmp_path: Path) -> None:
    sent: list[tuple[str, str]] = []

    def send_offer(candidate, _version, pdf_path, _idempotency_key):
        sent.append((candidate.email, Path(pdf_path).name))
        return [SimpleNamespace(status="sent", error="")]

    service = HiringWorkflowService(
        HiringPipelineStore(tmp_path / "history.sqlite3"),
        send_offer=send_offer,
    )
    application = service.start_application(
        legal_name="Maya Patel",
        email="",
        phone="555-0100",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )
    service.finalize_initial_interview(
        application.application_id,
        history_id="hist-send-after-approval",
        score=80,
        outcome="Hire",
        actor="Admin",
    )
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    draft = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "24.00", "weekly_hours": "40"},
        actor="Admin",
    )
    service.submit_offer_for_approval(application.application_id, draft.version_id, actor="Admin")
    docx_path = tmp_path / "offer.docx"
    pdf_path = tmp_path / "offer.pdf"
    docx_path.write_bytes(b"docx")
    pdf_path.write_bytes(b"pdf")
    service.approve_offer(
        application.application_id,
        draft.version_id,
        approver_name="Executive",
        approver_role="Executive Director",
        approval_date=date(2026, 7, 14),
        docx_path=docx_path,
        pdf_path=pdf_path,
    )

    delivered = service.send_approved_offer(
        application.application_id,
        draft.version_id,
        candidate_email="maya@example.com",
        actor="Admin",
    )

    assert delivered.status == "sent"
    assert service.store.get_candidate(application.candidate_id).email == "maya@example.com"
    assert sent == [("maya@example.com", "offer.pdf")]
    assert service.store.get_application(application.application_id).stage is HiringStage.OFFER_SENT


def test_failed_offer_send_preserves_approval_and_retry_is_idempotent(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-retry",
            "candidate_name": "Ana Ruiz",
            "candidate_email": "ana@example.com",
            "school": "Hawthorne",
            "position": "Preschool",
            "score": 75,
            "outcome": "Hire",
        }
    )
    attempts = 0

    def send_offer(candidate, version, pdf_path, idempotency_key):
        nonlocal attempts
        attempts += 1
        status = "failed" if attempts == 1 else "sent"
        return [SimpleNamespace(status=status, error=f"SMTP failed for {candidate.email}")]

    service = HiringWorkflowService(HiringPipelineStore(history_path), send_offer=send_offer)
    application = service.backfill_history()[0]
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    draft = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "24.00", "weekly_hours": "40"},
        actor="HR",
    )
    service.submit_offer_for_approval(application.application_id, draft.version_id, actor="HR")
    docx_path = tmp_path / "offer.docx"
    pdf_path = tmp_path / "offer.pdf"
    docx_path.write_bytes(b"docx")
    pdf_path.write_bytes(b"pdf")

    failed = service.approve_offer(
        application.application_id,
        draft.version_id,
        approver_name="Executive",
        approver_role="Executive",
        approval_date=date(2026, 7, 14),
        docx_path=docx_path,
        pdf_path=pdf_path,
    )

    assert failed.status == "approved"
    assert failed.send_status == "failed"
    assert "ana@example.com" not in failed.send_error
    blocked = service.store.get_application(application.application_id)
    assert blocked.stage is HiringStage.EXECUTIVE_APPROVAL
    assert blocked.attention_code == "approved_send_failed"

    sent = service.retry_offer_send(application.application_id, draft.version_id, actor="Admin")
    duplicate = service.retry_offer_send(application.application_id, draft.version_id, actor="Admin")

    assert sent.status == "sent"
    assert duplicate.status == "sent"
    assert attempts == 2


def test_admin_compensation_revision_only_changes_pay_and_hours(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-revision",
            "candidate_name": "Jo Lee",
            "candidate_email": "jo@example.com",
            "school": "Palmdale",
            "position": "Preschool",
            "score": 90,
            "outcome": "Hire",
        }
    )
    sent_attachments: list[tuple[str, Path]] = []

    def send_offer(_candidate, version, pdf_path, _idempotency_key):
        sent_attachments.append((version.version_id, Path(pdf_path)))
        return [SimpleNamespace(status="sent", error="")]

    service = HiringWorkflowService(HiringPipelineStore(history_path), send_offer=send_offer)
    application = service.backfill_history()[0]
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    original = service.create_offer_draft(
        application.application_id,
        terms={
            "hourly_pay": "22.00",
            "weekly_hours": "40",
            "position": "Preschool Teacher",
            "school": "Palmdale",
        },
        actor="HR",
    )
    service.store.set_offer_status(original.version_id, "sent")
    service.store.update_application_stage(application.application_id, HiringStage.OFFER_SENT)

    with pytest.raises(ValueError, match="Admin"):
        service.create_compensation_revision(
            application.application_id,
            hourly_pay="25.00",
            weekly_hours="35",
            actor="HR User",
            actor_role="HR",
        )

    revision = service.create_compensation_revision(
        application.application_id,
        hourly_pay="25.00",
        weekly_hours="35",
        actor="Admin User",
        actor_role="Admin",
    )

    assert revision.version_number == 2
    assert revision.terms == {
        "hourly_pay": "25.00",
        "weekly_hours": "35",
        "position": "Preschool Teacher",
        "school": "Palmdale",
    }
    assert revision.status == "pending_approval"

    edited_docx = tmp_path / "offer-v2-edited.docx"
    edited_pdf = tmp_path / "offer-v2-edited.pdf"
    edited_docx.write_bytes(b"internal editable record")
    edited_pdf.write_bytes(b"%PDF-1.4\nfinal edited offer")
    delivered = service.approve_compensation_revision(
        application.application_id,
        revision.version_id,
        admin_name="Admin User",
        approval_date=date(2026, 7, 14),
        docx_path=edited_docx,
        pdf_path=edited_pdf,
    )

    assert delivered.status == "sent"
    assert sent_attachments == [(revision.version_id, edited_pdf.resolve())]
    assert sent_attachments[0][1].suffix.casefold() == ".pdf"
    assert sent_attachments[0][1] != edited_docx.resolve()


def test_pending_offer_pay_edit_creates_new_previewed_version(tmp_path: Path) -> None:
    generated: list[tuple[int, str]] = []

    def prepare(_application, _candidate, version):
        docx_path = tmp_path / f"offer-v{version.version_number}.docx"
        pdf_path = tmp_path / f"offer-v{version.version_number}.pdf"
        docx_path.write_bytes(f"pay={version.terms['hourly_pay']}".encode())
        pdf_path.write_bytes(f"pay={version.terms['hourly_pay']}".encode())
        generated.append((version.version_number, version.terms["hourly_pay"]))
        return docx_path, pdf_path

    service = HiringWorkflowService(
        HiringPipelineStore(tmp_path / "history.sqlite3"),
        prepare_offer_artifacts=prepare,
    )
    application = service.start_application(
        legal_name="Maya Patel",
        email="",
        phone="",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )
    service.finalize_initial_interview(
        application.application_id,
        history_id="hist-edit-pending-pay",
        score=80,
        outcome="Hire",
        actor="Admin",
    )
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    original = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "22.00", "weekly_hours": "40", "position": "Preschool"},
        actor="Admin",
    )
    service.submit_offer_for_approval(application.application_id, original.version_id, actor="Admin")

    with pytest.raises(ValueError, match="increment"):
        service.revise_pending_offer_pay(
            application.application_id,
            original.version_id,
            hourly_pay="24.10",
            actor="Executive",
        )

    revised = service.revise_pending_offer_pay(
        application.application_id,
        original.version_id,
        hourly_pay="24.50",
        actor="Executive",
    )

    assert service.store.get_offer_version(original.version_id).status == "superseded"
    assert revised.status == "pending_approval"
    assert revised.version_number == 2
    assert revised.terms == {
        "hourly_pay": "24.50",
        "weekly_hours": "40",
        "position": "Preschool",
    }
    assert Path(revised.pdf_path).read_bytes() == b"pay=24.50"
    assert generated == [(1, "22.00"), (2, "24.50")]


def test_delete_offer_removes_only_offer_and_generated_files(tmp_path: Path) -> None:
    output_dir = tmp_path / "offers"
    output_dir.mkdir()
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))
    application = service.start_external_offer_application(
        legal_name="Maya Patel",
        email="",
        phone="",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )
    draft = service.create_offer_draft(
        application.application_id,
        terms={
            "hourly_pay": "24.00",
            "weekly_hours": "40",
            "output_dir": str(output_dir),
        },
        actor="Admin",
    )
    service.submit_offer_for_approval(application.application_id, draft.version_id, actor="Admin")
    docx_path = output_dir / "offer.docx"
    pdf_path = output_dir / "offer.pdf"
    unrelated_path = output_dir / "keep.docx"
    docx_path.write_bytes(b"docx")
    pdf_path.write_bytes(b"pdf")
    unrelated_path.write_bytes(b"keep")
    service.store.record_offer_artifacts(
        draft.version_id,
        docx_path=docx_path,
        pdf_path=pdf_path,
    )

    remaining_application = service.delete_offer(
        application.application_id,
        draft.version_id,
        actor="Admin",
    )

    assert service.store.list_offer_versions(application.application_id) == []
    assert docx_path.exists() is False
    assert pdf_path.exists() is False
    assert unrelated_path.read_bytes() == b"keep"
    assert remaining_application.archived_at == ""
    assert remaining_application.stage is HiringStage.OFFER_DRAFT
    assert service.store.get_candidate(application.candidate_id).legal_name == "Maya Patel"


def test_delete_offer_rejects_artifact_outside_configured_output_directory(tmp_path: Path) -> None:
    output_dir = tmp_path / "offers"
    output_dir.mkdir()
    outside_docx = tmp_path / "outside.docx"
    outside_pdf = tmp_path / "outside.pdf"
    outside_docx.write_bytes(b"docx")
    outside_pdf.write_bytes(b"pdf")
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))
    application = service.start_external_offer_application(
        legal_name="Maya Patel",
        email="",
        phone="",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )
    draft = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "24.00", "weekly_hours": "40", "output_dir": str(output_dir)},
        actor="Admin",
    )
    service.submit_offer_for_approval(application.application_id, draft.version_id, actor="Admin")
    service.store.record_offer_artifacts(
        draft.version_id,
        docx_path=outside_docx,
        pdf_path=outside_pdf,
    )

    with pytest.raises(ValueError, match="outside"):
        service.delete_offer(application.application_id, draft.version_id, actor="Admin")

    assert outside_docx.exists()
    assert outside_pdf.exists()
    assert service.store.get_offer_version(draft.version_id).status == "pending_approval"


def test_deadline_extension_does_not_change_document_and_only_latest_sent_can_be_accepted(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-accept",
            "candidate_name": "Nia King",
            "school": "Hawthorne",
            "position": "Teacher",
            "score": 82,
            "outcome": "Hire",
        }
    )
    service = HiringWorkflowService(HiringPipelineStore(history_path))
    application = service.backfill_history()[0]
    service.record_director_decision(application.application_id, decision="Hire", actor="Director")
    first = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "22", "weekly_hours": "40"},
        actor="HR",
    )
    service.store.set_offer_status(first.version_id, "sent")
    service.store.update_application_stage(application.application_id, HiringStage.OFFER_SENT)
    second = service.create_offer_draft(
        application.application_id,
        terms={"hourly_pay": "24", "weekly_hours": "40"},
        actor="Admin",
    )
    service.store.set_offer_status(second.version_id, "sent")

    with pytest.raises(ValueError, match="latest"):
        service.accept_offer(application.application_id, first.version_id, actor="Admin")

    service.extend_offer_deadline(
        application.application_id,
        second.version_id,
        reply_by_date=date(2026, 7, 25),
        actor="Admin",
    )
    extended = service.store.get_offer_version(second.version_id)
    assert extended.operational_reply_by_date == "2026-07-25"
    assert extended.document_reply_by_date == ""

    accepted = service.accept_offer(application.application_id, second.version_id, actor="Admin")
    assert accepted.stage is HiringStage.ACCEPTED


def test_archive_hides_application_and_linked_profile_deletion_fails_closed(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-archive",
            "candidate_name": "Kim Park",
            "school": "Palmdale",
            "position": "Teacher",
            "score": 50,
            "outcome": "No Hire",
        }
    )
    service = HiringWorkflowService(HiringPipelineStore(history_path))
    application = service.backfill_history()[0]

    archived = service.archive_application(application.application_id, actor="Admin")

    assert archived.archived_at
    assert service.store.list_applications() == []
    assert service.store.list_applications(include_archived=True) == [archived]
    with pytest.raises(ValueError, match="linked"):
        service.store.delete_candidate_profile(application.candidate_id)


def test_offer_acceptance_is_preserved_when_notification_blocks_and_retries(tmp_path: Path) -> None:
    store = HiringPipelineStore(tmp_path / "history.sqlite3")
    attempts = iter(["blocked", "sent"])
    service = HiringWorkflowService(
        store,
        notify_offer_accepted=lambda candidate, version, key: [
            SimpleNamespace(status=next(attempts), error="attachment unavailable")
        ],
    )
    application = service.start_application(
        legal_name="Ari Lane", honorific="Ms.", email="ari@example.org", phone="",
        school="Palmdale", position="Teacher", actor="Interviewer",
    )
    store.update_application_stage(application.application_id, HiringStage.OFFER_DRAFT)
    offer = service.create_offer_draft(
        application.application_id,
        terms={"school": "Palmdale", "position": "Teacher"},
        actor="HR",
    )
    store.set_offer_status(offer.version_id, "sent")
    store.update_application_stage(application.application_id, HiringStage.OFFER_SENT)

    accepted = service.accept_offer(application.application_id, offer.version_id, actor="Admin")

    assert accepted.stage is HiringStage.ACCEPTED
    assert store.get_application(application.application_id).attention_code == "accepted_notification_pending"
    retried = service.retry_accepted_notification(application.application_id, offer.version_id)
    assert retried is True
    assert store.get_application(application.application_id).attention_code == ""
    assert service.retry_pending_accepted_notifications() == 0


def test_offer_acceptance_calls_idempotent_onboarding_handoff_with_director_snapshot(tmp_path: Path) -> None:
    store = HiringPipelineStore(tmp_path / "history.sqlite3")
    handoffs = []
    service = HiringWorkflowService(
        store,
        onboarding_accept_offer=lambda application, candidate, version, director_id, director_name: handoffs.append(
            (application, candidate, version, director_id, director_name)
        ),
    )
    application = service.start_application(
        legal_name="Ari Lane",
        email="ari@example.org",
        phone="6615550123",
        school="Palmdale",
        position="Teacher",
        actor="Interviewer",
    )
    store.update_application_stage(application.application_id, HiringStage.DIRECTOR_REVIEW)
    service.record_director_decision(
        application.application_id,
        decision="hire",
        actor="Director Jones",
        actor_id="staffing-interview-42",
    )
    offer = service.create_offer_draft(
        application.application_id,
        terms={"school": "Palmdale", "position": "Teacher"},
        actor="HR",
    )
    store.set_offer_status(offer.version_id, "sent")
    store.update_application_stage(application.application_id, HiringStage.OFFER_SENT)

    service.accept_offer(application.application_id, offer.version_id, actor="Admin")

    assert len(handoffs) == 1
    assert handoffs[0][3:] == ("staffing-interview-42", "Director Jones")


def test_notification_retry_does_not_clear_failed_onboarding_handoff_attention(tmp_path: Path) -> None:
    store = HiringPipelineStore(tmp_path / "history.sqlite3")
    service = HiringWorkflowService(
        store,
        onboarding_accept_offer=lambda *_args: (_ for _ in ()).throw(RuntimeError("handoff failed")),
        notify_offer_accepted=lambda *_args: [SimpleNamespace(status="sent", error="")],
    )
    application = service.start_application(
        legal_name="Ari Lane", email="ari@example.org", phone="6615550123",
        school="Palmdale", position="Teacher", actor="Interviewer",
    )
    store.update_application_stage(application.application_id, HiringStage.DIRECTOR_REVIEW)
    service.record_director_decision(application.application_id, decision="hire", actor="Director")
    offer = service.create_offer_draft(
        application.application_id,
        terms={"school": "Palmdale", "position": "Teacher"},
        actor="HR",
    )
    store.set_offer_status(offer.version_id, "sent")
    store.update_application_stage(application.application_id, HiringStage.OFFER_SENT)

    accepted = service.accept_offer(application.application_id, offer.version_id, actor="Admin")

    assert accepted.attention_code == "onboarding_handoff_pending"


def test_failed_accepted_offer_handoff_retries_into_real_onboarding_without_duplicates(tmp_path: Path) -> None:
    onboarding = OnboardingService(
        OnboardingStore(tmp_path / "onboarding.sqlite3"),
        OnboardingAccess("admin", "admin-1"),
    )
    template = onboarding.create_task_template_draft(
        template_key="orientation", school="Palmdale", title="Orientation",
        owner_role="Director", due_offset_days=0,
    )
    onboarding.publish_task_template(template.id)
    attempts = {"count": 0}

    def handoff(application, candidate, version, director_id, director_name):
        attempts["count"] += 1
        if attempts["count"] == 1:
            raise OSError("temporary handoff failure")
        return onboarding.accept_offer(
            application_id=application.application_id,
            legal_name=candidate.legal_name, school=application.school,
            role=application.position, acceptance_date="2026-07-20",
            start_date=str(version.start_date or version.terms["start_date"]),
            email=candidate.email, phone=candidate.phone,
            hiring_director_id=director_id, hiring_director_name=director_name,
        )

    store = HiringPipelineStore(tmp_path / "hiring.sqlite3")
    service = HiringWorkflowService(store, onboarding_accept_offer=handoff)
    application = service.start_application(
        legal_name="Ari Lane", email="ari@example.org", phone="6615550123",
        school="Palmdale", position="Teacher", actor="Interviewer",
    )
    store.update_application_stage(application.application_id, HiringStage.DIRECTOR_REVIEW)
    service.record_director_decision(
        application.application_id, decision="hire", actor="Director Jones",
        actor_id="staffing-director-42",
    )
    offer = service.create_offer_draft(
        application.application_id,
        terms={"school": "Palmdale", "position": "Teacher", "start_date": "2026-08-03"},
        actor="HR",
    )
    store.set_offer_status(offer.version_id, "sent")
    store.update_application_stage(application.application_id, HiringStage.OFFER_SENT)

    assert service.accept_offer(application.application_id, offer.version_id, actor="Admin").attention_code == "onboarding_handoff_pending"
    assert service.retry_onboarding_handoff(application.application_id, actor="Admin") is True
    assert store.get_application(application.application_id).attention_code == ""
    assert len(onboarding.list_employees()) == 1
    assert len(onboarding.list_tasks()) == 1
    assert onboarding.list_employees()[0].hiring_director_id == "staffing-director-42"


def test_start_then_finalize_keeps_one_application_cycle(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    service = HiringWorkflowService(HiringPipelineStore(history_path))

    started = service.start_application(
        legal_name="Tara Moss",
        email="tara@example.com",
        phone="555-0100",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )
    completed = service.finalize_initial_interview(
        started.application_id,
        history_id="hist-tara",
        score=65.0,
        outcome="Borderline",
        actor="Admin",
    )

    assert started.stage is HiringStage.INITIAL_INTERVIEW
    assert completed.application_id == started.application_id
    assert completed.history_id == "hist-tara"
    assert completed.stage is HiringStage.DIRECTOR_REVIEW
    assert len(service.store.list_applications()) == 1


def test_start_application_persists_candidate_honorific(tmp_path: Path) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))

    application = service.start_application(
        legal_name="Tara Moss",
        honorific="Mr.",
        email="",
        phone="",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )

    assert service.store.get_candidate(application.candidate_id).honorific == "Mr."


def test_candidate_profiles_can_be_searched_and_contact_fields_updated(tmp_path: Path) -> None:
    service = HiringWorkflowService(HiringPipelineStore(tmp_path / "history.sqlite3"))
    application = service.start_application(
        legal_name="Maria Hernandez",
        email="old@example.com",
        phone="555-0100",
        school="Palmdale",
        position="Preschool",
        actor="Admin",
    )

    updated = service.update_candidate_profile(
        application.candidate_id,
        legal_name="Maria Hernandez",
        preferred_name="Mari",
        email="mari@example.com",
        phone="555-0199",
    )

    assert updated.preferred_name == "Mari"
    assert updated.email == "mari@example.com"
    assert service.search_candidate_profiles("mari") == [updated]
    assert service.search_candidate_profiles("555-0199") == [updated]


def test_migration_report_infers_strongest_stage_and_flags_conflicts(tmp_path: Path) -> None:
    history_path = tmp_path / "history.sqlite3"
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-conflict",
            "candidate_name": "Alex Kim",
            "school": "Palmdale",
            "position": "Teacher",
            "score": 40,
            "outcome": "No Hire",
            "offer_status": "accepted",
        }
    )
    history.append({"history_id": "bad-row", "candidate_name": "Missing School"})
    service = HiringWorkflowService(HiringPipelineStore(history_path))

    report = service.reconcile_history()
    migrated = service.store.application_for_history("hist-conflict")

    assert report.source_rows == 2
    assert report.application_count == 1
    assert report.profile_count == 1
    assert report.skipped_rows == 1
    assert report.conflict_count == 1
    assert report.stage_counts == {"accepted": 1}
    assert migrated is not None
    assert migrated.stage is HiringStage.ACCEPTED
    assert migrated.attention_code == "migration_conflict"


def test_offer_notification_adapter_sends_candidate_pdf_only() -> None:
    calls = []

    class Notifications:
        def emit_event(self, event_type, payload, idempotency_key):
            calls.append((event_type, payload, idempotency_key))
            return [SimpleNamespace(status="sent", error="")]

    adapter = HiringOfferNotificationAdapter(Notifications())
    candidate = SimpleNamespace(legal_name="Maya Patel", email="maya@example.com", honorific="Ms.")
    version = SimpleNamespace(
        version_id="v1",
        approval_date="2026-07-14",
        document_reply_by_date="2026-07-17",
        start_date="2026-08-03",
        terms={"hourly_pay": "24.00", "school": "Palmdale"},
    )

    result = adapter(candidate, version, Path("approved.pdf"), "offer-version:v1")

    assert result[0].status == "sent"
    event_type, payload, key = calls[0]
    assert event_type == "offer.approved"
    assert key == "offer-version:v1"
    assert payload["candidate_email"] == "maya@example.com"
    assert payload["honorific"] == "Ms."
    assert payload["school"] == "Palmdale"
    assert payload["offer_pdf_path"] == "approved.pdf"
    assert payload["attachments"] == ["approved.pdf"]


def test_offer_notification_adapter_rejects_word_document_attachment() -> None:
    class Notifications:
        def emit_event(self, *_args, **_kwargs):
            raise AssertionError("DOCX must never reach notification delivery")

    adapter = HiringOfferNotificationAdapter(Notifications())
    candidate = SimpleNamespace(legal_name="Maya Patel", email="maya@example.com", honorific="Ms.")
    version = SimpleNamespace(
        version_id="v2",
        approval_date="2026-07-14",
        document_reply_by_date="2026-07-17",
        start_date="2026-08-03",
        terms={"school": "Palmdale"},
    )

    with pytest.raises(ValueError, match="PDF"):
        adapter(candidate, version, Path("edited-offer.docx"), "offer-version:v2")
