from datetime import date
from pathlib import Path

import pytest
from docx import Document

from scoring_reporting import (
    OfferInput,
    OfferLetterService,
    OfferTemplateError,
    build_school_offer_filename,
    derive_offer_schedule,
    next_available_offer_path,
    parse_requested_hourly_pay,
)


@pytest.mark.parametrize("suffix", [".docx", ".docm", ".DOCM"])
def test_validate_template_path_accepts_allowed_word_templates(tmp_path: Path, suffix: str):
    template_path = tmp_path / f"template{suffix}"
    template_path.write_bytes(b"placeholder")

    OfferLetterService.validate_template_path(template_path)


def test_validate_template_path_rejects_unsupported_template_type(tmp_path: Path):
    template_path = tmp_path / "template.txt"
    template_path.write_text("not a word file", encoding="utf-8")

    with pytest.raises(OfferTemplateError, match=r"\.docx or \.docm"):
        OfferLetterService.validate_template_path(template_path)


def test_offer_replacements_include_candidate_title() -> None:
    data = OfferInput(
        first_name="Tatiana",
        last_name="Zuluaga",
        city="Palmdale",
        position="Teacher",
        start_date=date(2026, 7, 27),
        start_time_12h="09:30 AM",
        end_time_12h="06:30 PM",
        hourly_pay=24.0,
        hours=40,
        created_on=date(2026, 7, 13),
        title="Ms.",
    )

    replacements = OfferLetterService.build_replacements(data)

    assert replacements["[Title]"] == "Ms."


def test_offer_replacements_preserve_fractional_weekly_hours_and_pto() -> None:
    data = OfferInput(
        first_name="Tatiana",
        last_name="Zuluaga",
        city="Palmdale",
        position="Teacher",
        start_date=date(2026, 7, 27),
        start_time_12h="08:15 AM",
        end_time_12h="05:00 PM",
        hourly_pay=24.0,
        hours=38.75,
        created_on=date(2026, 7, 13),
        title="Ms.",
    )

    replacements = OfferLetterService.build_replacements(data)

    assert replacements["[Hours]"] == "38.75"
    assert replacements["[PTO]"] == "77.5"
    assert replacements["[PTO2]"] == "155"


def test_derive_offer_schedule_subtracts_lunch_for_long_shift() -> None:
    schedule = derive_offer_schedule("8:00 AM", "5:00 PM")

    assert schedule.gross_daily_hours == 9
    assert schedule.net_daily_hours == 8
    assert schedule.weekly_hours == 40
    assert schedule.employment_type == "full_time"


def test_parse_requested_hourly_pay_accepts_one_clear_amount() -> None:
    assert parse_requested_hourly_pay("I am looking for $24/hour, negotiable") == 24


@pytest.mark.parametrize(
    ("school", "expected"),
    [
        ("Palmdale", "Launch Pad Learning PMD Offer of Employment to Tatiana Zuluaga.docx"),
        ("North Long Beach", "Launch Pad Learning NLB Offer of Employment to Tatiana Zuluaga.docx"),
        ("Hawthorne", "Preschool Partners, LLC Offer of Employment to Tatiana Zuluaga.docx"),
    ],
)
def test_school_offer_filename_uses_school_legal_name(school: str, expected: str) -> None:
    assert build_school_offer_filename(school, "Tatiana Zuluaga") == expected


def test_next_available_offer_path_preserves_existing_revisions(tmp_path: Path) -> None:
    filename = "Launch Pad Learning PMD Offer of Employment to Tatiana Zuluaga.docx"
    (tmp_path / filename).touch()
    (tmp_path / "Launch Pad Learning PMD Offer of Employment to Tatiana Zuluaga (2).docx").touch()

    assert next_available_offer_path(tmp_path, filename).name == (
        "Launch Pad Learning PMD Offer of Employment to Tatiana Zuluaga (3).docx"
    )


def test_approved_offer_replaces_hardcoded_date_deadline_and_computed_start(tmp_path: Path) -> None:
    template_path = tmp_path / "offer-template.docx"
    template = Document()
    template.add_paragraph("Launch Pad Learning")
    template.add_paragraph("July 13, 2026")
    template.add_paragraph("Dear [Title] [Last Name]:")
    template.add_paragraph("Start date: [StartDate]")
    template.add_paragraph("Reply by 5 p.m., [OfferDeadline].")
    template.save(template_path)
    output_path = tmp_path / "approved.docx"
    data = OfferInput(
        first_name="Tatiana",
        last_name="Zuluaga",
        city="Palmdale",
        position="Teacher",
        start_date=date(2026, 7, 27),
        start_time_12h="09:30 AM",
        end_time_12h="06:30 PM",
        hourly_pay=24.0,
        hours=40,
        created_on=date(2026, 7, 13),
        title="Ms.",
    )

    OfferLetterService.render_approved_offer(
        template_path,
        output_path,
        data,
        approval_date=date(2026, 7, 14),
    )

    text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
    assert "July 14, 2026" in text
    assert "Start date: 08/03/2026" in text
    assert "Reply by 5 p.m., 07/17/2026." in text
    assert "July 13, 2026" not in text
