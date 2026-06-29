import tempfile
import unittest
from pathlib import Path

from docx.shared import Inches, Pt
from docx_compat import Document
from scoring_reporting import DocxExporter, ReportingValidationError, _executive_summary_sections_from_structured


def _cell_xml_contains(cell, text: str) -> bool:
    return text in cell._tc.xml


def _doc_text(doc) -> str:
    paragraph_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    return f"{paragraph_text}\n{table_text}"


class TestDocxExporterValidation(unittest.TestCase):
    def test_structured_executive_sections_flatten_nested_new_prompt_schema(self):
        sections = _executive_summary_sections_from_structured(
            {
                "recommendation": {"rating": "Recommend", "rationale": "Evidence supports fit."},
                "key_strengths": [{"strength": "Warm routines", "evidence": "Uses greeting rituals."}],
                "key_concerns_or_risks": [{"concern": "Safety detail", "evidence_or_gap": "Needs more supervision examples."}],
            }
        )

        self.assertEqual(sections["recommendation"], ["Recommend: Evidence supports fit."])
        self.assertEqual(sections["strengths"], ["Warm routines: Uses greeting rituals."])
        self.assertEqual(sections["concerns"], ["Safety detail: Needs more supervision examples."])

    def _rubric(self):
        return {
            "traits": [],
            "tracks": {
                "general": {
                    "label": "General",
                    "max_weighted_total": 10,
                }
            },
            "absolute_disqualifiers": [],
        }

    def _scoring(self):
        return {
            "rows": [],
            "weighted_total": 0,
            "max_weighted_total": 10,
            "percent_of_max": 0.0,
            "critical_eq_1": False,
            "disqualifier_present": False,
            "locked_rule": None,
            "outcome": "No Hire",
        }

    def _trait_scoring(self):
        return {
            "rows": [
                {
                    "trait_id": "trait_1",
                    "trait_name": "Empathy",
                    "priority": "Critical",
                    "weight": 3,
                    "raw_score": 5,
                    "weighted_score": 15,
                    "system_checkbox_score": 15,
                    "net_signal_score": 6,
                    "suggested_raw_score": 4,
                    "final_raw_score": 5,
                    "interviewer_adjusted": True,
                    "adjustment_reason": "Full answer was stronger than signal evidence.",
                    "deepseek_calculated_score": 9,
                    "deepseek_raw_score": 3,
                    "model_trait_score": {
                        "raw_score": 3,
                        "evidence_quote": "Calm, concrete example.",
                        "rationale": "Some descriptor match.",
                        "risks_or_gaps": "Limited detail.",
                    },
                    "primary_question": "How do you support a child in distress?",
                    "no_example_after_followups": False,
                    "question_notes": "",
                    "trait_notes": "",
                    "verbatim_notes": "Calm, concrete example.",
                    "absolute_disqualifier": False,
                    "selected_signal_ids": ["S_CHILD_CENTERED", "S_COREGULATION"],
                    "model_signal_suggestions": [
                        {
                            "signal_id": "S_CHILD_CENTERED",
                            "confidence": 0.8,
                            "evidence_quote": "Calm, concrete example.",
                            "rationale": "Matched calm routine.",
                        },
                        {
                            "signal_id": "S_MODEL_ONLY",
                            "confidence": 0.4,
                            "evidence_quote": "Possible observation.",
                            "rationale": "Possible observation.",
                        },
                    ],
                    "model_signal_override": {
                        "accepted_signal_ids": ["S_CHILD_CENTERED"],
                        "rejected_signal_ids": ["S_MODEL_ONLY"],
                        "manual_only_signal_ids": ["S_COREGULATION"],
                    },
                }
            ],
            "weighted_total": 15,
            "max_weighted_total": 15,
            "percent_of_max": 100.0,
            "percent_of_max_label": "100%",
            "critical_eq_1": False,
            "disqualifier_present": False,
            "locked_rule": None,
            "outcome": "Hire",
        }

    def test_export_falls_back_for_stale_track_key(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "stale-track",
            }
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            path = exporter.export(self._rubric(), payload, self._scoring())
            self.assertTrue(path.exists())

    def test_export_raises_when_candidate_name_missing(self):
        payload = {
            "candidate": {
                "interview_date": "2026-02-20",
                "track": "general",
            }
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            with self.assertRaisesRegex(ReportingValidationError, "required candidate field: 'name'"):
                exporter.export(self._rubric(), payload, self._scoring())

    def test_extract_full_candidate_transcript_prefers_flow_transcript(self):
        payload = {
            "flow_transcript": [
                {"candidate_transcript": "Whisper segment one."},
                {"candidate_transcript": "Whisper segment two."},
            ],
            "audio_recording": [
                {"candidate_transcript": "Stale recording transcript."},
            ],
        }

        transcript = DocxExporter._extract_full_candidate_transcript(payload)

        self.assertEqual(transcript, "Whisper segment one.\n\nWhisper segment two.")

    def test_export_uses_candidate_transcript_for_custom_flow_entry(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "flow_transcript": [
                {
                    "type": "custom",
                    "title": "Custom Question",
                    "question": "How do you partner with families?",
                    "candidate_transcript": "I send weekly updates and hold check-ins.",
                    "answer": "Interviewer note that should not replace transcript.",
                }
            ],
            "custom_answers": [],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)
            doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Full Candidate Answer (auto-transcribed)", doc_text)
        self.assertIn("I send weekly updates and hold check-ins.", table_text)
        self.assertNotIn("Interviewer note that should not replace transcript.", doc_text + table_text)

    def test_export_omits_custom_questions_table_from_notes_document(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "custom_answers": [
                {"question_text": "", "answer": ""},
                {"question_text": "", "answer": "full-time"},
                {"question_text": "", "answer": "$20"},
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)

        doc_text = _doc_text(doc)
        self.assertNotIn("Custom Questions (Non-scored)", doc_text)
        self.assertNotIn("Custom question 1", doc_text)
        self.assertNotIn("full-time", doc_text)
        self.assertNotIn("$20", doc_text)

    def test_export_collapses_duplicate_transcription_attempts(self):
        duplicate_transcript = (
            "[Q1 Attempt 1]\n"
            "I use visual routines and songs.\n\n"
            "[Q1 Attempt 2]\n"
            "I use visual routines and songs."
        )
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "flow_transcript": [
                {
                    "type": "custom",
                    "title": "Custom Question",
                    "question": "How do you transition children?",
                    "candidate_transcript": duplicate_transcript,
                }
            ],
            "custom_answers": [],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertEqual(table_text.count("I use visual routines and songs."), 1)
        self.assertNotIn("[Q1 Attempt 2]", table_text)

    def test_export_marks_missing_trait_transcript_without_empty_answer_page(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "trait",
                    "id": "trait_1",
                    "title": "Empathy",
                    "question": "How do you help a child?",
                    "candidate_transcript": "",
                    "raw_score": 5,
                }
            ],
            "custom_answers": [],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._trait_scoring())
            doc = Document(out_path)
            doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Full Candidate Answer (auto-transcribed): Not captured", doc_text)
        self.assertNotIn("(No candidate transcript captured)", table_text)

    def test_export_places_executive_summary_before_report_sections(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "executive_summary": "Candidate uses calm routines and family communication.",
            "interview_highlights": ["Uses visual timers.", "Communicates consistently with families."],
            "answer_summaries": [
                {
                    "flow_index": 1,
                    "summary": "Gave a concrete routine example.",
                    "evidence_quotes": ["songs and visual timers"],
                    "rubric_alignment": "Uses predictable transition supports.",
                    "risks_or_gaps": "",
                }
            ],
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "custom",
                    "title": "Custom Question",
                    "question": "How do you support transitions?",
                    "candidate_transcript": "I use songs and visual timers.",
                }
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            doc_text = "\n".join(paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertLess(paragraphs.index("Candidate Snapshot"), paragraphs.index("Executive Summary"))
        self.assertLess(paragraphs.index("Candidate Snapshot"), paragraphs.index("AI-generated Evidence Summary"))
        self.assertLess(paragraphs.index("Scorecard Snapshot"), paragraphs.index("Consolidated Answer Summaries"))
        self.assertIn("Candidate uses calm routines and family communication.", doc_text)
        self.assertIn("Uses visual timers.", doc_text)
        self.assertIn("Communicates consistently with families.", doc_text)
        self.assertIn("Answer summary", table_text)
        self.assertIn("Gave a concrete routine example.", table_text)
        summary_cards = [
            table
            for table in doc.tables
            if table.rows[0].cells[0].text.startswith("Question: How do you support transitions?")
        ]
        summary_card_text = "\n".join(cell.text for row in summary_cards[0].rows for cell in row.cells)
        self.assertNotIn("Evidence", summary_card_text)
        self.assertNotIn("songs and visual timers", summary_card_text)
        self.assertNotIn("Uses predictable transition supports.", table_text)

    def test_export_renders_generated_executive_summary_as_structured_decision_brief(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "executive_summary": "\n".join(
                [
                    "**Recommendation**: Recommend with reservations.",
                    "**Overall Fit**: Calm and practical, with several classroom examples.",
                    "**Role-Specific Match**: Matches toddler routines and family communication.",
                    "**Score Pattern**: High empathy evidence, lower safety specificity.",
                    "**Key Strengths**:",
                    "- Uses visual routines during transitions.",
                    "- Communicates with families early.",
                    "**Key Concerns or Risks**:",
                    "- Needs more safety-specific detail.",
                    "- Limited examples of accountability.",
                    "**Suggested Follow-Up Questions**:",
                    "1. How would you handle a child resisting diapering?",
                    "2. What would you document after a safety incident?",
                    "3. How do you incorporate parent feedback into classroom practice?",
                    "4. How do you manage stress while maintaining student care?",
                    "**Final Hiring Notes**: Verify safety judgment before final offer.",
                ]
            ),
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)

        full_text = _doc_text(doc)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        recommendation_cells = [
            table.rows[0].cells[0]
            for table in doc.tables
            if table.rows[0].cells[0].text == "Recommendation: Recommend with reservations."
        ]
        at_a_glance_table = next(table for table in doc.tables if table.rows[0].cells[0].text == "Overall Fit")
        strengths_table = next(table for table in doc.tables if table.rows[0].cells[0].text == "Key Strengths")
        final_note_cells = [
            table.rows[0].cells[0]
            for table in doc.tables
            if table.rows[0].cells[0].text.startswith("Final Hiring Notes:")
        ]

        self.assertIn("Executive Summary", paragraphs)
        self.assertEqual(recommendation_cells[0].text, "Recommendation: Recommend with reservations.")
        self.assertTrue(_cell_xml_contains(recommendation_cells[0], 'w:fill="3A3100"'))
        self.assertTrue(_cell_xml_contains(recommendation_cells[0], 'w:color w:val="FFFFFF"'))
        self.assertEqual([row.cells[0].text for row in at_a_glance_table.rows], ["Overall Fit", "Role-Specific Match", "Score Pattern"])
        self.assertTrue(_cell_xml_contains(at_a_glance_table.rows[0].cells[0], 'w:fill="263940"'))
        self.assertTrue(_cell_xml_contains(at_a_glance_table.rows[0].cells[0], 'w:color w:val="B7D4FF"'))
        self.assertEqual(strengths_table.rows[0].cells[0].text, "Key Strengths")
        self.assertEqual(strengths_table.rows[0].cells[1].text, "Key Concerns or Risks")
        self.assertTrue(_cell_xml_contains(strengths_table.rows[0].cells[0], 'w:fill="203B23"'))
        self.assertTrue(_cell_xml_contains(strengths_table.rows[0].cells[1], 'w:fill="3A3100"'))
        self.assertTrue(final_note_cells)
        self.assertTrue(_cell_xml_contains(final_note_cells[0], 'w:fill="263940"'))
        self.assertTrue(_cell_xml_contains(final_note_cells[0], 'w:color w:val="FFFFFF"'))
        self.assertIn("Uses visual routines during transitions.", full_text)
        self.assertIn("How would you handle a child resisting diapering?", full_text)
        self.assertIn("How do you manage stress while maintaining student care?", full_text)
        self.assertNotIn("**Recommendation**", full_text)
        self.assertNotIn("- Uses visual routines", full_text)
        self.assertNotIn("1. How would you handle", full_text)

    def test_export_cleans_freeform_executive_summary_fallback(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "executive_summary": "**Unexpected Notes**\n- Candidate stayed calm.\nPlain follow-up needed.",
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)

        full_text = _doc_text(doc)
        self.assertIn("Additional Notes", full_text)
        self.assertIn("Unexpected Notes", full_text)
        self.assertIn("Candidate stayed calm.", full_text)
        self.assertIn("Plain follow-up needed.", full_text)
        self.assertNotIn("**Unexpected Notes**", full_text)
        self.assertNotIn("- Candidate stayed calm.", full_text)

    def test_export_formats_structured_executive_summary_sections(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "executive_summary_sections": {
                "recommendation": "Recommend with reservations.",
                "overall_fit": "Calm and practical.",
                "role_specific_match": "Matches toddler routines.",
                "score_pattern": "High empathy, lower specificity.",
                "key_strengths": ["Uses visual routines.", "Communicates early.", "Stays calm."],
                "key_concerns_or_risks": ["Needs safety detail.", "Verify reliability.", "Probe coachability."],
                "suggested_follow_up_questions": ["Q1?", "Q2?", "Q3?", "Q4?"],
                "final_hiring_notes": "Verify safety judgment.",
            },
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)

        full_text = _doc_text(doc)
        recommendation_cells = [
            table.rows[0].cells[0]
            for table in doc.tables
            if table.rows[0].cells[0].text == "Recommendation: Recommend with reservations."
        ]
        self.assertTrue(recommendation_cells)
        self.assertTrue(_cell_xml_contains(recommendation_cells[0], 'w:fill="3A3100"'))
        self.assertIn("Role-Specific Match", full_text)
        self.assertIn("Q4?", full_text)
        self.assertIn("Final Hiring Notes: Verify safety judgment.", full_text)

    def test_export_uses_interview_notes_template_layout(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "school": "Palmdale",
                "track": "general",
            },
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)

        section = doc.sections[0]
        self.assertEqual(section.top_margin, Inches(0.55))
        self.assertEqual(section.right_margin, Inches(0.65))
        self.assertEqual(section.bottom_margin, Inches(0.55))
        self.assertEqual(section.left_margin, Inches(0.65))
        self.assertEqual(doc.paragraphs[0].text, "Candidate Interview Decision Brief")
        self.assertEqual(doc.paragraphs[0].alignment, 1)
        self.assertEqual(doc.paragraphs[0].runs[0].font.name, "Arial")
        self.assertEqual(doc.paragraphs[0].runs[0].font.size, Pt(17))
        self.assertEqual(str(doc.paragraphs[0].runs[0].font.color.rgb), "1F4E79")
        self.assertIn("Ada | Palmdale | General | Interview Date: 2026-02-20", doc.paragraphs[1].text)
        self.assertEqual(doc.paragraphs[1].alignment, 1)
        top_summary_cells = [
            table.rows[0].cells[0]
            for table in doc.tables
            if table.rows[0].cells[0].text.startswith("Interviewer score:")
        ]
        self.assertTrue(top_summary_cells)
        self.assertIn("Interviewer score: 0 / 10", top_summary_cells[0].text)
        self.assertIn("AI advisory score: N/A (incomplete)", top_summary_cells[0].text)
        self.assertIn("AI-trait-based score: not generated", top_summary_cells[0].text)
        self.assertTrue(_cell_xml_contains(top_summary_cells[0], 'w:fill="FFF7D6"'))
        snapshot_tables = [
            table
            for table in doc.tables
            if table.rows[0].cells[0].text == "Has degree"
        ]
        self.assertTrue(snapshot_tables)
        self.assertEqual(snapshot_tables[0].rows[0].cells[1].text, "Not provided")
        snapshot_text = "\n".join(cell.text for row in snapshot_tables[0].rows for cell in row.cells)
        self.assertNotIn("Candidate Name", snapshot_text)
        self.assertNotIn("School/Location", snapshot_text)
        self.assertNotIn("Palmdale", snapshot_text)
        self.assertNotIn("Track", snapshot_text)
        self.assertNotIn("Interview Date", snapshot_text)
        self.assertNotIn("2026-02-20", snapshot_text)
        header_cells = [
            table.rows[0].cells
            for table in doc.tables
            if table.rows[0].cells[0].text == "Trait"
        ][0]
        self.assertEqual(header_cells[3].text, "Raw\nScore")
        self.assertEqual(header_cells[4].text, "Weighted\nScore")
        self.assertEqual(header_cells[5].text, "AI Raw\nScore")
        self.assertEqual(header_cells[6].text, "AI Weighted\nScore")
        self.assertTrue(_cell_xml_contains(header_cells[0], 'w:fill="EAF3F8"'))
        self.assertTrue(_cell_xml_contains(header_cells[0], "<w:tcMar>"))

    def test_export_snapshot_conditionally_shows_infant_toddler_and_total_units_rows(self):
        rubric = self._rubric()
        rubric["tracks"]["infant_toddler"] = {"label": "Infant/Toddler", "max_weighted_total": 10}
        rubric["tracks"]["behavior_support_specialist"] = {"label": "Behavior Support Specialist", "max_weighted_total": 10}
        base_candidate = {
            "name": "Ada",
            "interview_date": "2026-02-20",
            "school": "Palmdale",
            "qualification": {
                "has_degree": True,
                "degree_type": "BA",
                "degree_in_ece": True,
                "ece_units_completed": 12,
                "infant_toddler_class_completed": True,
                "total_units_completed": 24,
                "years_experience": 3,
            },
        }

        with tempfile.TemporaryDirectory() as general_td, tempfile.TemporaryDirectory() as infant_td, tempfile.TemporaryDirectory() as bss_td:
            general_path = DocxExporter(Path(general_td)).export(
                rubric,
                {"candidate": {**base_candidate, "track": "general"}},
                self._scoring(),
            )
            infant_path = DocxExporter(Path(infant_td)).export(
                rubric,
                {
                    "candidate": {
                        **base_candidate,
                        "track": "infant_toddler",
                        "qualification": {**base_candidate["qualification"], "has_degree": False},
                    }
                },
                self._scoring(),
            )
            bss_path = DocxExporter(Path(bss_td)).export(
                rubric,
                {"candidate": {**base_candidate, "track": "behavior_support_specialist"}},
                self._scoring(),
            )
            general_doc = Document(general_path)
            infant_doc = Document(infant_path)
            bss_doc = Document(bss_path)

        general_text = _doc_text(general_doc)
        infant_text = _doc_text(infant_doc)
        bss_text = _doc_text(bss_doc)
        self.assertNotIn("Infant/toddler class completed", general_text)
        self.assertNotIn("Total units completed (if no degree)", general_text)
        self.assertIn("Infant/toddler class completed", infant_text)
        self.assertIn("Total units completed (if no degree)", infant_text)
        self.assertIn("Infant/toddler class completed", bss_text)

    def test_export_uses_director_ready_section_order_and_appendices(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "school": "Palmdale",
                "track": "general",
            },
            "executive_summary": "Candidate uses calm routines.",
            "answer_summaries": [
                {
                    "flow_index": 1,
                    "summary": "Uses gentle redirection.",
                    "evidence_quotes": ["gentle voice"],
                    "rubric_alignment": "Gentle classroom guidance.",
                    "risks_or_gaps": "Needs more safety detail.",
                }
            ],
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "trait",
                    "id": "trait_1",
                    "title": "Empathy",
                    "question": "How do you redirect?",
                    "candidate_transcript": "I use a gentle voice.",
                    "raw_score": 5,
                }
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._trait_scoring())
            doc = Document(out_path)

        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        full_text = _doc_text(doc)

        self.assertEqual(paragraphs[0], "Candidate Interview Decision Brief")
        self.assertTrue(doc.tables[0].rows[0].cells[0].text.startswith("Interviewer score:"))
        self.assertLess(paragraphs.index("Candidate Snapshot"), paragraphs.index("Executive Summary"))
        self.assertLess(paragraphs.index("Candidate Snapshot"), paragraphs.index("Scorecard Snapshot"))
        self.assertLess(paragraphs.index("Scorecard Snapshot"), paragraphs.index("Consolidated Answer Summaries"))
        self.assertLess(paragraphs.index("Consolidated Answer Summaries"), paragraphs.index("Director Decision Brief"))
        self.assertLess(paragraphs.index("Director Decision Brief"), paragraphs.index("Critical Safety Review"))
        self.assertLess(paragraphs.index("Critical Safety Review"), paragraphs.index("Interview Transcript Appendix"))
        self.assertLess(paragraphs.index("Interview Transcript Appendix"), paragraphs.index("AI Advisory Appendix"))
        self.assertEqual(full_text.count("Uses gentle redirection."), 1)
        self.assertGreater(full_text.index("I use a gentle voice."), full_text.index("Interview Transcript Appendix"))
        self.assertGreater(full_text.index("AI-trait-based signal observations:"), full_text.index("AI Advisory Appendix"))
        self.assertNotIn("Hiring Manager Evidence Notes", paragraphs)
        self.assertNotIn("Global disqualifiers reviewed:", full_text)
        self.assertNotIn("Answer summary evidence", full_text)
        scorecard = [table for table in doc.tables if table.rows[0].cells[0].text == "Trait"][0]
        self.assertTrue(_cell_xml_contains(scorecard.rows[1].cells[0], 'w:fill="FFF7D6"'))
        self.assertTrue(_cell_xml_contains(scorecard.rows[1].cells[5], 'w:fill="E6F4F1"'))
        self.assertNotIn("DeepSeek", full_text)

    def test_export_renders_ai_risk_flag_evidence_in_critical_safety_review(self):
        scoring = self._trait_scoring()
        scoring["rows"][0]["raw_score"] = 2
        scoring["rows"][0]["verbatim_notes"] = ""
        scoring["rows"][0]["question_notes"] = ""
        scoring["rows"][0]["trait_notes"] = ""
        scoring["rows"][0]["model_trait_score"] = {
            "raw_score": 2,
            "evidence_quote": "I yell first",
            "rationale": "Matches low safety descriptor.",
            "risks_or_gaps": "Unsafe first response.",
            "risk_flag_evidence": "Candidate says they would yell first.",
        }
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "trait",
                    "id": "trait_1",
                    "title": "Empathy",
                    "question": "How do you respond?",
                    "candidate_transcript": "I yell first.",
                }
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)

        safety_tables = [
            table
            for table in doc.tables
            if table.rows[0].cells[0].text == "Critical Trait"
        ]
        self.assertTrue(safety_tables)
        safety_text = "\n".join(cell.text for row in safety_tables[0].rows for cell in row.cells)
        self.assertIn("Yes", safety_text)
        self.assertIn("Candidate says they would yell first.", safety_text)
        self.assertIn("I yell first", safety_text)
        self.assertNotIn("None recorded", safety_text)

    def test_export_renders_answer_summaries_as_readable_cards(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "answer_summaries": [
                {
                    "flow_index": 1,
                    "summary": "Candidate explained a detailed classroom routine without forcing narrow column wrapping.",
                    "evidence_quotes": ["uses visual timer", "sings cleanup song"],
                    "rubric_alignment": "Predictable routine support.",
                    "risks_or_gaps": "Needs more safety detail.",
                }
            ],
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "trait",
                    "id": "trait_1",
                    "title": "Empathy",
                    "question": "How do you support transitions?",
                    "candidate_transcript": "I use visual timers.",
                }
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._trait_scoring())
            doc = Document(out_path)

        full_text = _doc_text(doc)
        summary_cards = [
            table
            for table in doc.tables
            if table.rows[0].cells[0].text.startswith("Question: How do you support transitions?")
        ]

        self.assertTrue(summary_cards)
        self.assertEqual(len(summary_cards[0].columns), 2)
        summary_card_text = "\n".join(cell.text for row in summary_cards[0].rows for cell in row.cells)
        self.assertNotIn("Question text", summary_card_text)
        self.assertIn("Ratings", full_text)
        self.assertIn("Interviewer: 5/5 | AI-Advisor: 3/5 | AI-trait-based: 4/5", summary_card_text)
        self.assertNotIn("Interviewer rating", full_text)
        self.assertNotIn("AI-advisory rating", full_text)
        self.assertNotIn("AI-trait-based rating", full_text)
        self.assertIn("Answer summary", full_text)
        self.assertNotIn("Evidence", summary_card_text)
        self.assertNotIn("uses visual timer; sings cleanup song", summary_card_text)
        self.assertIn("Predictable routine support.", summary_card_text)
        self.assertIn("Needs more safety detail.", summary_card_text)
        transcript_index = [paragraph.text for paragraph in doc.paragraphs].index("Interview Transcript Appendix")
        later_paragraphs = "\n".join(paragraph.text for paragraph in doc.paragraphs[transcript_index:])
        self.assertNotIn("Answer summary evidence", later_paragraphs)
        self.assertNotIn("Needs more safety detail.", later_paragraphs)

    def test_export_omits_ratings_for_non_scored_answer_summary_and_intro_sections(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "answer_summaries": [
                {
                    "flow_index": 1,
                    "summary": "Intro summary should not render.",
                },
                {
                    "flow_index": 2,
                    "summary": "Candidate described family partnership.",
                    "evidence_quotes": ["weekly check-ins"],
                    "rubric_alignment": "Specific observed competency: Family communication",
                    "risks_or_gaps": "",
                },
            ],
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "intro",
                    "title": "Intro Script",
                    "question": "Introductory script",
                    "candidate_transcript": "Intro text should not render.",
                },
                {
                    "flow_index": 2,
                    "type": "custom",
                    "title": "Family Partnership",
                    "question": "",
                    "prompt": "How do you partner with families?",
                    "candidate_transcript": "I use weekly check-ins.",
                },
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._scoring())
            doc = Document(out_path)

        full_text = _doc_text(doc)
        summary_cards = [
            table
            for table in doc.tables
            if table.rows[0].cells[0].text.startswith("Question: How do you partner with families?")
        ]

        self.assertTrue(summary_cards)
        custom_card_text = "\n".join(cell.text for row in summary_cards[0].rows for cell in row.cells)
        self.assertNotIn("Ratings", custom_card_text)
        self.assertNotIn("Interviewer rating", custom_card_text)
        self.assertNotIn("AI-advisory rating", custom_card_text)
        self.assertNotIn("AI-trait-based rating", custom_card_text)
        self.assertIn("Answer summary", custom_card_text)
        self.assertNotIn("Evidence", custom_card_text)
        self.assertNotIn("weekly check-ins", custom_card_text)
        self.assertNotIn("Rubric alignment", custom_card_text)
        self.assertNotIn("Specific observed competency: Family communication", custom_card_text)
        self.assertNotIn("Intro Script", full_text)
        self.assertNotIn("Introductory script", full_text)
        self.assertNotIn("Intro text should not render.", full_text)
        self.assertNotIn("Intro summary should not render.", full_text)

    def test_export_collapses_missing_ai_advisory_scoring(self):
        scoring = self._trait_scoring()
        for row in scoring["rows"]:
            row["net_signal_score"] = None
            row["suggested_raw_score"] = None
            row["deepseek_calculated_score"] = None
            row["deepseek_raw_score"] = None
            row["model_trait_score"] = {}
            row["model_signal_suggestions"] = []

        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "model_suggestion_status": "processing",
            "model_scoring_status": "processing",
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)

        full_text = _doc_text(doc)

        self.assertIn("AI scoring not generated", full_text)
        self.assertIn("suggestions: processing", full_text)
        self.assertNotIn("AI net signal score\nN/A", full_text)

    def test_answer_summary_card_explains_missing_signal_advisory_rating(self):
        scoring = self._trait_scoring()
        scoring["rows"][0]["suggested_raw_score"] = None
        scoring["rows"][0]["net_signal_score"] = None
        scoring["rows"][0]["model_signal_suggestions"] = []
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "answer_summaries": [
                {
                    "flow_index": 1,
                    "summary": "Candidate described support.",
                }
            ],
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "trait",
                    "id": "trait_1",
                    "question": "How do you support transitions?",
                    "candidate_transcript": "I use visual timers.",
                }
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)

        table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Ratings\nInterviewer: 5/5 | AI-Advisor: 3/5 | AI-trait-based: No supported signals", table_text)
        self.assertNotIn("AI-advisory rating\nNo supported signals", table_text)
        self.assertNotIn("AI-trait-based rating\n3/5", table_text)
        self.assertNotIn("AI-advisory rating\nN/A", table_text)

    def test_export_reports_complete_deepseek_total_out_of_max(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._trait_scoring())
            doc = Document(out_path)
            doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Interviewer score: 15 / 15", table_text)
        self.assertIn("AI advisory score: 9 / 15 (60.0%).", table_text)
        self.assertIn("AI-trait-based score: 9 / 15 (60.0%).", table_text)
        self.assertIn("AI advisory Total: 9 / 15", table_text)
        self.assertIn("Weighted Total: 15 / 15", table_text)
        self.assertNotIn("AI advisory score: N/A (incomplete)", doc_text + table_text)
        self.assertNotIn("DeepSeek", doc_text + table_text)

    def test_score_callout_reports_interviewer_advisory_and_trait_based_scores_without_recommendation(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
        }
        scoring = self._trait_scoring()
        scoring["rows"][0]["deepseek_calculated_score"] = 12
        scoring["rows"][0]["model_trait_score"]["raw_score"] = 3

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        first_table_text = doc.tables[0].rows[0].cells[0].text
        self.assertNotIn("Recommendation:", first_table_text)
        self.assertIn("Interviewer score: 15 / 15 (100.0%).", table_text)
        self.assertIn("AI advisory score: 12 / 15 (80.0%).", table_text)
        self.assertIn("AI-trait-based score: 9 / 15 (60.0%).", table_text)

    def test_export_reports_incomplete_deepseek_total_when_any_included_trait_missing_score(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
        }
        scoring = self._trait_scoring()
        scoring["rows"][0]["deepseek_calculated_score"] = None

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Interviewer score: 15 / 15", table_text)
        self.assertIn("AI advisory score: N/A (incomplete).", table_text)
        self.assertIn("AI-trait-based score: 9 / 15 (60.0%).", table_text)
        self.assertIn("AI advisory Total: N/A (incomplete)", table_text)

    def test_export_explains_incomplete_human_and_missing_ai_recommendation_status(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "model_suggestion_status": "processing",
            "model_scoring_status": "processing",
        }
        scoring = self._trait_scoring()
        scoring["outcome"] = "Incomplete"
        scoring["locked_rule"] = "One or more applicable traits are missing final raw scores"
        scoring["rows"].append(
            {
                "trait_name": "Team Orientation",
                "priority": "High",
                "weight": 2,
                "raw_score": None,
                "weighted_score": 0,
                "system_checkbox_score": 0,
                "deepseek_calculated_score": None,
                "deepseek_raw_score": None,
                "net_signal_score": None,
                "suggested_raw_score": None,
                "model_trait_score": {},
                "model_signal_suggestions": [],
                "primary_question": "How do you handle conflict?",
                "no_example_after_followups": False,
                "question_notes": "",
                "trait_notes": "",
                "verbatim_notes": "",
                "absolute_disqualifier": False,
                "selected_signal_ids": [],
            }
        )
        for row in scoring["rows"]:
            row["deepseek_calculated_score"] = None

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Skipped scored questions: 1", table_text)
        self.assertNotIn("missing final raw score", table_text)
        self.assertIn("AI advisory score: not generated (suggestions: processing; scoring: processing).", table_text)
        self.assertNotIn("AI advisory score: N/A (incomplete)", table_text)

    def test_export_treats_legacy_unrated_rows_as_skipped_and_omits_notes(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "flow_transcript": [
                {
                    "flow_index": 2,
                    "type": "trait",
                    "id": "trait_b",
                    "title": "Team Orientation",
                    "question": "How do you handle conflict?",
                    "candidate_transcript": "Should not render.",
                }
            ],
        }
        scoring = self._trait_scoring()
        scoring["rows"].append(
            {
                "trait_name": "Team Orientation",
                "trait_id": "trait_b",
                "priority": "High",
                "weight": 2,
                "raw_score": None,
                "weighted_score": 0,
                "system_checkbox_score": 0,
                "deepseek_calculated_score": None,
                "deepseek_raw_score": None,
                "net_signal_score": None,
                "suggested_raw_score": None,
                "model_trait_score": {},
                "model_signal_suggestions": [],
                "primary_question": "How do you handle conflict?",
                "no_example_after_followups": False,
                "question_notes": "Should not render.",
                "trait_notes": "",
                "verbatim_notes": "",
                "absolute_disqualifier": False,
                "selected_signal_ids": [],
            }
        )

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, scoring)
            doc = Document(out_path)
            full_text = _doc_text(doc)

        self.assertIn("skipped scored questions: 1", full_text)
        self.assertNotIn("Team Orientation", full_text)
        self.assertNotIn("Should not render.", full_text)

    def test_export_reports_legacy_selected_signals_and_deepseek_trait_suggestions(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._trait_scoring())
            doc = Document(out_path)
            doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Compatibility selected signal IDs:", doc_text)
        self.assertIn("AI-trait-based signal observations:", doc_text)
        self.assertIn("Raw\nScore", table_text)
        self.assertIn("Weighted\nScore", table_text)
        self.assertIn("AI Raw\nScore", table_text)
        self.assertIn("AI Weighted\nScore", table_text)
        self.assertNotIn("Final interviewer raw score", table_text)
        self.assertNotIn("Human weighted score", table_text)
        self.assertIn("AI-trait-based net signal score", table_text)
        self.assertIn("AI-trait-based raw score", table_text)
        self.assertNotIn("Interviewer adjusted from AI-trait-based score", table_text)
        self.assertNotIn("Adjustment reason", table_text)
        self.assertIn("AI-Advisor raw score", table_text)
        self.assertIn("AI-trait-based weighted score", table_text)
        self.assertIn("AI-Advisor score evidence", table_text)
        self.assertIn("S_CHILD_CENTERED, S_COREGULATION", table_text)
        self.assertIn("AI signal-scored observations", table_text)
        self.assertIn("AI-Advisor observations", table_text)
        self.assertIn("Compatibility selected-only observations", table_text)
        self.assertIn("Used by AI-trait-based scoring", table_text)
        self.assertIn("AI suggestion", table_text)
        self.assertNotIn("DeepSeek", doc_text + table_text)

    def test_export_renders_generated_deepseek_summary_and_trait_scores_together(self):
        payload = {
            "candidate": {
                "name": "Ada",
                "interview_date": "2026-02-20",
                "track": "general",
            },
            "executive_summary": "Candidate uses calm routines.",
            "interview_highlights": ["Uses a gentle voice."],
            "answer_summaries": [
                {
                    "flow_index": 1,
                    "summary": "Uses gentle redirection.",
                    "evidence_quotes": ["gentle voice"],
                    "rubric_alignment": "Gentle classroom guidance.",
                    "risks_or_gaps": "Needs more safety detail.",
                }
            ],
            "flow_transcript": [
                {
                    "flow_index": 1,
                    "type": "trait",
                    "id": "trait_1",
                    "title": "Empathy",
                    "question": "How do you redirect?",
                    "candidate_transcript": "I use a gentle voice.",
                    "raw_score": 5,
                }
            ],
        }

        with tempfile.TemporaryDirectory() as td:
            exporter = DocxExporter(Path(td))
            out_path = exporter.export(self._rubric(), payload, self._trait_scoring())
            doc = Document(out_path)
            doc_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("Executive Summary", doc_text)
        self.assertIn("Candidate uses calm routines.", doc_text)
        self.assertIn("AI-generated Evidence Summary", doc_text)
        self.assertIn("Uses a gentle voice.", doc_text)
        self.assertIn("Answer summary", table_text)
        self.assertIn("Uses gentle redirection.", table_text)
        self.assertIn("AI-trait-based raw score", table_text)
        self.assertIn("AI-trait-based weighted score", table_text)
        self.assertIn("9", table_text)
        self.assertIn("AI-Advisor raw score", table_text)
        self.assertIn("Calm, concrete example.", table_text)
        self.assertNotIn("DeepSeek", doc_text + table_text)


if __name__ == "__main__":
    unittest.main()
