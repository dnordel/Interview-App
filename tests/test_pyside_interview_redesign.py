import json
import os
import queue
import sqlite3
import sys
import threading
import time
import wave
from datetime import date, timedelta
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest
import pyside_interview_app
from candidate_report import CandidateReportRepository
from data_store import InterviewHistoryStore, SchoolOfferSettingsStore
from docx import Document
from interview_runtime import map_indeed_transcript_to_questions, parse_indeed_transcript_text
from notification_models import NotificationRecipient, NotificationRule, NotificationSendResult
from notification_store import NotificationStore
from staffing_dashboard_v2 import (
    StaffingDashboardV2Page,
    _display_date,
    _validation_issues_from_rows,
    configure_v2_scroll_areas,
)
from staffing_models import StaffingMetricRow
from staffing_service import StaffingChangeConflict
from source_update_monitor import source_digest

from onboarding_operations import Employee, EmployeeTask
from pyside_interview_app import (
    PySideInterviewSession,
    build_interview_redesign_model,
    build_pyside_candidate_board,
    build_pyside_onboarding_board,
    latest_pyside_draft_path,
    standard_window_control_flags,
)
from scoring_reporting import CandidateQualification
from visual_test_support import (
    TYPOGRAPHY_STRESS_TEXT,
    VisualTestDatabaseRegistry,
    assert_no_large_unpainted_region,
    assert_vertical_text_fits,
    assert_widget_text_glyphs_supported,
    configure_visual_test_app,
)


def _pyside_window_on_page(model, page_name: str):
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    target = "Staffing v2" if page_name in {"Interviews", "Candidates", "Offers"} else page_name
    window.sidebar.setCurrentRow(nav_items.index(target))
    if page_name == "Interviews":
        window.staffing_v2_dashboard.show_external_page("interviews")
        window.hiring_v2_router.show_interview()
    elif page_name in {"Candidates", "Offers"}:
        window.staffing_v2_dashboard.show_external_page(page_name.casefold())
    window.QtWidgets.QApplication.processEvents()
    return window


@pytest.mark.pyside_gui
def test_source_update_reload_banner_dashboard_scenario(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    source_root = tmp_path / "src"
    source_root.mkdir()
    source_file = source_root / "feature.py"
    source_file.write_text("VERSION = 1\n", encoding="utf-8")
    version_file = tmp_path / "source_version.txt"
    version_file.write_text(
        f"updated_at=2026-07-16T12:00:00Z\nsource_sha256={source_digest(source_root)}\n",
        encoding="utf-8",
    )
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text('{"schools":[]}', encoding="utf-8")
    monkeypatch.setattr(pyside_interview_app, "SOURCE_VERSION_PATH", version_file, raising=False)
    monkeypatch.setattr(pyside_interview_app, "SOURCE_UPDATE_ROOT", source_root, raising=False)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", tmp_path / "notifications.sqlite3")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)

    window.window.show()
    app.processEvents()
    window._start_source_update_monitoring()
    assert window.source_update_banner.isHidden()
    source_file.write_text("VERSION = 2\n", encoding="utf-8")
    version_file.write_text(
        f"updated_at=2026-07-16T12:05:00Z\nsource_sha256={source_digest(source_root)}\n",
        encoding="utf-8",
    )
    window._poll_source_updates()
    app.processEvents()

    assert window.source_update_banner.isVisible()
    assert window.source_update_restart_button.text() == "Restart App"
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
def test_staffing_conflict_prompt_identifies_user_and_remote_changes(monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    captured: dict[str, Any] = {}

    def question(*args: Any) -> Any:
        captured["title"] = args[1]
        captured["message"] = args[2]
        return qt_widgets.QMessageBox.StandardButton.Yes

    monkeypatch.setattr(qt_widgets.QMessageBox, "question", question)
    owner = SimpleNamespace(QtWidgets=qt_widgets, window=None)
    conflict = StaffingChangeConflict(
        event_id="event-1",
        source_replica="director:palmdale:pmd",
        school="Palmdale",
        operation="update_assignment_details",
        base_snapshot={},
        local_snapshot={},
        remote_payload={"assignment_id": 12, "classroom": "Tranquility", "notes": "Director note"},
    )

    accepted = pyside_interview_app.PySideInterviewWindow._resolve_staffing_change_conflict(owner, conflict)

    assert accepted is True
    assert captured["title"] == "Staffing Change Conflict"
    assert "pmd" in captured["message"]
    assert "Classroom: Tranquility" in captured["message"]
    assert "Notes: Director note" in captured["message"]
    assert "accept these changes" in captured["message"].lower()


@pytest.mark.pyside_gui
def test_pyside_new_interview_setup_scenario(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text('{"schools":[]}', encoding="utf-8")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", tmp_path / "notifications.sqlite3")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)

    window.staffing_v2_dashboard.external_nav_buttons["interviews"].click()
    app.processEvents()

    assert window.hiring_v2_router.current_route == "interview"
    assert window.interview_tabs.currentIndex() == 0
    assert window.interview_tabs.currentWidget().findChild(qt_widgets.QLabel, "Title").text() == "New Interview"
    window.window.close()
    app.processEvents()


def test_pyside_new_interview_setup_matches_mockup_controls(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "Palmdale"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    setup = window.interview_tabs.widget(0)

    assert setup.objectName() == "HiringV2NewInterviewSetup"
    assert setup.findChild(qt_widgets.QLineEdit, "HiringV2SetupCandidateName").text() == ""
    assert setup.findChild(qt_widgets.QComboBox, "HiringV2SetupSchool").currentText() == "Hawthorne"
    assert setup.findChild(qt_widgets.QComboBox, "HiringV2SetupTrack").currentText() == next(
        iter(model.track_labels.values())
    )
    interview_type = setup.findChild(qt_widgets.QComboBox, "HiringV2SetupInterviewType")
    assert [interview_type.itemText(index) for index in range(interview_type.count())] == ["First Interview"]
    progress = setup.findChild(qt_widgets.QFrame, "HiringV2SetupProgress")
    assert [label.text() for label in progress.findChildren(qt_widgets.QLabel, "HiringV2SetupProgressLabel")] == [
        "Setup",
        "Introduction",
        "Questions",
        "Review",
    ]
    buttons = {button.text() for button in setup.findChildren(qt_widgets.QPushButton)}
    assert buttons == {"Test Audio", "Cancel", "Begin Interview"}
    assert setup.findChild(qt_widgets.QComboBox, "HiringV2SetupAudioSource") is not None
    assert setup.findChild(qt_widgets.QLabel, "HiringV2SetupMicrophoneStatus") is not None
    assert setup.findChild(qt_widgets.QLabel, "HiringV2SetupSystemAudioStatus") is not None
    assert setup.findChild(qt_widgets.QLabel, "HiringV2SetupTranscriptStatus") is not None
    window.window.close()
    app.processEvents()


def test_pyside_new_interview_setup_begins_first_interview_without_contact_fields(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    from hiring_pipeline import HiringPipelineStore

    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    capture_started: list[bool] = []
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: capture_started.append(True))

    window.home_begin_button.click()
    app.processEvents()
    assert window.session is None
    assert HiringPipelineStore(history_path).list_applications() == []

    window.home_candidate_input.setText("Sofia Ramirez")

    window.home_begin_button.click()
    app.processEvents()

    applications = HiringPipelineStore(history_path).list_applications()
    candidate = HiringPipelineStore(history_path).get_candidate(applications[0].candidate_id)
    assert candidate.legal_name == "Sofia Ramirez"
    assert candidate.email == ""
    assert candidate.phone == ""
    assert window.session is not None
    assert window.session.candidate_name == "Sofia Ramirez"
    assert window.interview_tabs.currentIndex() == 2
    assert capture_started == [True]
    window.window.close()
    app.processEvents()


def test_pyside_new_interview_cancel_discards_changes_and_returns_dashboard(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    no = window.QtWidgets.QMessageBox.StandardButton.No
    yes = window.QtWidgets.QMessageBox.StandardButton.Yes
    answers = iter((no, yes))
    monkeypatch.setattr(window.QtWidgets.QMessageBox, "question", lambda *_args, **_kwargs: next(answers))
    window.home_candidate_input.setText("Discard Me")

    window.interview_tabs.widget(0).findChild(qt_widgets.QPushButton, "HiringV2SetupCancel").click()
    app.processEvents()

    assert window.home_candidate_input.text() == "Discard Me"
    assert window.staffing_v2_dashboard.current_page_id == "interviews"

    window.interview_tabs.widget(0).findChild(qt_widgets.QPushButton, "HiringV2SetupCancel").click()
    app.processEvents()

    assert window.home_candidate_input.text() == ""
    assert window.staffing_v2_dashboard.current_page_id == "staffing_dashboard"
    window.window.close()
    app.processEvents()


def test_pyside_exit_live_interview_saves_draft_and_returns_fresh_setup(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    draft_path = tmp_path / "drafts" / "application.json"
    session = PySideInterviewSession(model=model, draft_path=draft_path, application_id="application")
    session.start(candidate_name="Sofia Ramirez", school="Hawthorne", track_key=next(iter(model.flows)))
    window.session = session
    window.session_index = session.current_index
    window._render_live_question_page()
    window.interview_tabs.setCurrentIndex(2)

    window.interview_tabs.widget(2).findChild(qt_widgets.QPushButton, "LiveInterviewExit").click()
    app.processEvents()

    assert draft_path.exists()
    assert window.session is None
    assert window.hiring_v2_router.current_route == "interview"
    assert window.interview_tabs.currentIndex() == 0
    assert window.home_candidate_input.text() == ""
    assert window.staffing_v2_dashboard.current_page_id == "interviews"
    window.window.close()
    app.processEvents()


def test_pyside_manual_audio_preflight_updates_setup_without_starting_interview(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    from PySide6 import QtTest
    from interview_audio_preflight import AudioPreflightResult

    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    tested_sources: list[str] = []

    def ready_result(source: str):
        tested_sources.append(source)
        return AudioPreflightResult(True, True, True, "")

    monkeypatch.setattr(window, "_run_manual_audio_preflight", ready_result)
    window.home_audio_source_combo.clear()
    window.home_audio_source_combo.addItem("CABLE Output", "CABLE Output")

    window.home_test_audio_button.click()
    assert window.home_test_audio_button.isEnabled() is False
    for _ in range(20):
        app.processEvents()
        QtTest.QTest.qWait(10)
        if window.home_test_audio_button.isEnabled():
            break

    assert tested_sources == ["CABLE Output"]
    assert window.home_microphone_status.text() == "Microphone connected"
    assert window.home_system_audio_status.text() == "System audio connected"
    assert window.home_transcript_status.text() == "Live transcription ready"
    assert window.session is None
    window.window.close()
    app.processEvents()


def test_pyside_manual_audio_preflight_locks_only_duplicate_test_action(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    from PySide6 import QtTest
    from interview_audio_preflight import AudioPreflightResult

    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    release = threading.Event()
    calls: list[str] = []

    def delayed_result(source: str) -> AudioPreflightResult:
        calls.append(source)
        release.wait(1.0)
        return AudioPreflightResult(True, True, True, "")

    monkeypatch.setattr(window, "_run_manual_audio_preflight", delayed_result)
    window.home_audio_source_combo.clear()
    window.home_audio_source_combo.addItem("CABLE Output", "CABLE Output")

    window.home_test_audio_button.click()
    app.processEvents()

    assert not window.home_test_audio_button.isEnabled()
    assert window.home_begin_button.isEnabled()
    assert window.interview_tabs.widget(0).findChild(
        qt_widgets.QPushButton, "HiringV2SetupCancel"
    ).isEnabled()
    window.home_test_audio_button.click()
    assert calls == ["CABLE Output"]

    release.set()
    for _attempt in range(100):
        app.processEvents()
        QtTest.QTest.qWait(10)
        if window.home_test_audio_button.isEnabled():
            break
    assert window.home_test_audio_button.isEnabled()
    assert window.home_begin_button.isEnabled()
    assert calls == ["CABLE Output"]
    window.window.close()
    app.processEvents()


def test_pyside_manual_audio_preflight_runs_five_seconds_and_removes_artifacts(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    delays: list[float] = []
    artifact_dirs: list[Path] = []
    import interview_audio_recorder

    def write_signal(path: Path) -> None:
        with wave.open(str(path), "wb") as wav_file:
            wav_file.setnchannels(1)
            wav_file.setsampwidth(2)
            wav_file.setframerate(16000)
            wav_file.writeframes((100).to_bytes(2, "little", signed=True) * 16000)

    def fake_start_recording(**kwargs: Any) -> Any:
        output_dir = Path(kwargs["output_dir"])
        artifact_dirs.append(output_dir)
        mic = output_dir / "microphone.wav"
        system = output_dir / "system.wav"
        write_signal(mic)
        write_signal(system)
        return SimpleNamespace(
            mic_wav=mic,
            sys_wav=system,
            sys_label="CANDIDATE",
            stop=lambda: None,
            transcribe_new_segments=lambda **_kwargs: [
                SimpleNamespace(speaker="CANDIDATE", text="private candidate speech")
            ],
        )

    monkeypatch.setattr(interview_audio_recorder, "start_recording", fake_start_recording)
    monkeypatch.setattr(pyside_interview_app.time, "sleep", delays.append)
    monkeypatch.setattr(
        pyside_interview_app,
        "resolve_runtime",
        lambda _settings: SimpleNamespace(model="tiny", device="cpu", compute_type="int8", backend="test"),
    )

    result = window._run_manual_audio_preflight("CABLE Output")

    assert result.ready
    assert delays == [5]
    assert len(artifact_dirs) == 1
    assert not artifact_dirs[0].exists()
    window.window.close()
    app.processEvents()


def test_pyside_window_close_stops_active_manual_audio_preflight(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    stopped = threading.Event()
    window._manual_audio_preflight_session = SimpleNamespace(stop=stopped.set)

    assert window._request_window_close() is True
    assert stopped.wait(0.5)
    window._manual_audio_preflight_session = None
    window.window.close()
    app.processEvents()


def test_pyside_window_close_stops_all_interview_capture_work(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    recording_stopped = threading.Event()
    preflight_stopped = threading.Event()
    window.recording_session = SimpleNamespace(stop=recording_stopped.set)
    window._manual_audio_preflight_session = SimpleNamespace(stop=preflight_stopped.set)
    window._manual_audio_preflight_queue = queue.Queue()
    window._live_transcript_queue = queue.Queue()
    window._manual_audio_preflight_timer = window.QtCore.QTimer(window.window)
    window._manual_audio_preflight_timer.start(1000)
    window._live_transcript_timer = window.QtCore.QTimer(window.window)
    window._live_transcript_timer.start(1000)

    window.window.close()
    app.processEvents()

    assert recording_stopped.wait(0.5)
    assert preflight_stopped.wait(0.5)
    assert window.recording_session is None
    assert window._manual_audio_preflight_session is None
    assert window._manual_audio_preflight_queue is None
    assert window._live_transcript_queue is None
    assert window._manual_audio_preflight_timer is None
    assert window._live_transcript_timer is None


def test_pyside_setup_audio_probe_populates_detected_source_and_readiness(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")

    window._apply_setup_audio_probe(
        available_devices=["Default Microphone", "CABLE Output (VB-Audio Virtual Cable)"],
        microphone_device="Default Microphone",
        system_device="CABLE Output (VB-Audio Virtual Cable)",
        transcription_ready=True,
    )

    assert window.home_audio_source_combo.currentText() == "CABLE Output (VB-Audio Virtual Cable)"
    assert window.home_audio_source_combo.count() == 2
    assert window.home_microphone_status.text() == "Microphone connected"
    assert window.home_system_audio_status.text() == "System audio connected"
    assert window.home_transcript_status.text() == "Live transcription ready"
    assert window.session is None
    window.window.close()
    app.processEvents()


def test_pyside_begin_rejects_audio_source_that_is_no_longer_available(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    from hiring_pipeline import HiringPipelineStore

    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    monkeypatch.setattr(pyside_interview_app.sys, "platform", "win32")
    monkeypatch.setattr(
        pyside_interview_app,
        "list_windows_dshow_audio_devices",
        lambda: ["Different Available Device"],
    )
    window.home_audio_source_combo.clear()
    window.home_audio_source_combo.addItem("Disconnected CABLE Output", "Disconnected CABLE Output")
    window.home_candidate_input.setText("Sofia Ramirez")

    window.home_begin_button.click()
    app.processEvents()

    assert window.session is None
    assert HiringPipelineStore(history_path).list_applications() == []
    assert "no longer available" in window.home_setup_validation.text().lower()
    assert window.interview_tabs.currentIndex() == 0
    window.window.close()
    app.processEvents()


def test_pyside_begin_cancels_running_manual_audio_test_before_capture(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    monkeypatch.setattr(pyside_interview_app.sys, "platform", "linux")
    capture_started: list[bool] = []
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: capture_started.append(True))
    preflight_stopped = threading.Event()
    cancel_event = threading.Event()
    window._manual_audio_preflight_session = SimpleNamespace(stop=preflight_stopped.set)
    window._manual_audio_preflight_cancel_event = cancel_event
    window._manual_audio_preflight_queue = queue.Queue()
    window._manual_audio_preflight_timer = window.QtCore.QTimer(window.window)
    window._manual_audio_preflight_timer.start(1000)
    window.home_candidate_input.setText("Sofia Ramirez")

    window.home_begin_button.click()
    app.processEvents()

    assert preflight_stopped.wait(0.5)
    assert cancel_event.is_set()
    assert window._manual_audio_preflight_session is None
    assert window._manual_audio_preflight_queue is None
    assert window._manual_audio_preflight_timer is None
    assert capture_started == [True]
    assert window.session is not None
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.visual_inspection
def test_pyside_new_interview_setup_responsive_visual_scenario(
    tmp_path: Path,
    visual_test_databases: VisualTestDatabaseRegistry,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    from PySide6 import QtGui

    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    configure_visual_test_app(app)
    original_font = app.font()
    history_path = visual_test_databases.database("new_interview_setup.sqlite3")
    visual_test_databases.expect_seeded(history_path, table="interview_history")
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "setup-visual-fixture",
            "candidate_name": "Sofia Ramirez",
            "school": "Hawthorne",
            "position": "Infant/Toddler Teacher",
            "interview_date": "2026-07-15",
            "outcome": "Incomplete",
            "score": "0.0%",
        }
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    setup = window.interview_tabs.widget(0)
    window.window.resize(1680, 945)
    window.window.show()
    app.processEvents()
    realistic = window.window.grab()
    assert_no_large_unpainted_region(realistic)
    assert_widget_text_glyphs_supported(setup)
    assert realistic.save(str(tmp_path / "new-interview-realistic.png"))
    window.home_candidate_input.setText(TYPOGRAPHY_STRESS_TEXT)
    window.home_microphone_status.setText(TYPOGRAPHY_STRESS_TEXT)
    window.home_begin_button.setText(f"Begin {TYPOGRAPHY_STRESS_TEXT}")
    try:
        for width, height, scale in (
            (1680, 945, 1.0),
            (1366, 768, 1.0),
            (1366, 768, 1.25),
            (1366, 768, 1.5),
        ):
            font = QtGui.QFont(original_font)
            font.setPointSizeF(max(8.0, original_font.pointSizeF() * scale))
            app.setFont(font)
            window.window.resize(width, height)
            window.window.show()
            app.processEvents()
            rendered = window.window.grab()
            assert_no_large_unpainted_region(rendered)
            assert_widget_text_glyphs_supported(setup)
            assert rendered.save(str(tmp_path / f"new-interview-{width}x{height}-{scale}.png"))
            assert setup.horizontalScrollBar().maximum() == 0
            card_width = setup.findChild(qt_widgets.QFrame, "HiringV2CandidateSetupCard").width()
            assert 820 <= card_width <= 940
            assert_vertical_text_fits(window.home_candidate_input, TYPOGRAPHY_STRESS_TEXT)
            assert_vertical_text_fits(window.home_microphone_status, TYPOGRAPHY_STRESS_TEXT)
            assert_vertical_text_fits(window.home_begin_button, f"Begin {TYPOGRAPHY_STRESS_TEXT}")
    finally:
        app.setFont(original_font)
        window.window.close()
        app.processEvents()


def test_pyside_candidates_owns_roster_pipeline_history_and_legacy_interview_actions(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Candidates")
    candidates = window.staffing_v2_dashboard.external_pages["candidates"]
    tabs = candidates.findChild(qt_widgets.QTabWidget, "HiringV2CandidatesTabs")

    assert [tabs.tabText(index) for index in range(tabs.count())] == ["Candidate Roster", "Pipeline / History"]
    pipeline_tab = tabs.widget(1)
    assert pipeline_tab.findChild(qt_widgets.QTableWidget, "HiringV2ApplicationList") is not None
    assert {button.text() for button in pipeline_tab.findChildren(qt_widgets.QPushButton)} >= {
        "Import Indeed Transcript",
        "Continue Saved Draft",
        "Delete Saved Draft",
    }
    interviews = window.staffing_v2_dashboard.external_pages["interviews"]
    assert interviews.findChild(qt_widgets.QTableWidget, "HiringV2ApplicationList") is None
    assert interviews.findChild(qt_widgets.QPushButton, "ImportIndeedTranscriptButton") is None
    window.window.close()
    app.processEvents()


def test_pyside_candidates_resume_button_routes_saved_draft_into_live_interview(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_test = pytest.importorskip("PySide6.QtTest")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Hawthorne"],
    )
    draft_path = history_path.parent / "pyside_drafts" / "sofia.json"
    session = PySideInterviewSession(model=model, draft_path=draft_path)
    session.start(candidate_name="Sofia Ramirez", school="Hawthorne", track_key=next(iter(model.flows)))
    window = _pyside_window_on_page(model, "Candidates")
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    tabs = window.hiring_v2_candidates_tabs
    window.window.show()
    app.processEvents()

    qt_test.QTest.mouseClick(
        tabs.tabBar(),
        qt_core.Qt.MouseButton.LeftButton,
        pos=tabs.tabBar().tabRect(0).center(),
    )
    app.processEvents()
    assert tabs.currentIndex() == 0

    qt_test.QTest.mouseClick(
        tabs.tabBar(),
        qt_core.Qt.MouseButton.LeftButton,
        pos=tabs.tabBar().tabRect(1).center(),
    )
    app.processEvents()
    assert tabs.currentIndex() == 1
    assert window.candidate_continue_draft_button.isEnabled()

    window.candidate_continue_draft_button.click()
    app.processEvents()

    assert window.session is not None
    assert window.session.candidate_name == "Sofia Ramirez"
    assert window.staffing_v2_dashboard.current_page_id == "interviews"
    assert window.interview_tabs.currentIndex() == 2
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
def test_pyside_staffing_v2_hiring_live_focus_rail_scenario(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text('{"schools":[]}', encoding="utf-8")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", tmp_path / "notifications.sqlite3")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    dashboard = window.staffing_v2_dashboard
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]

    assert "Hiring v2" not in nav_items
    assert set(dashboard.external_pages) == {"interviews", "candidates", "offers"}
    assert dashboard.page_stack.currentWidget() is dashboard.dashboard_view
    dashboard.show_external_page("interviews")
    window._set_hiring_focus_mode(True)
    app.processEvents()
    assert dashboard.staffing_sidebar.width() == 64
    assert dashboard.external_nav_buttons["interviews"].isEnabled()
    assert not dashboard.dashboard_nav_button.isEnabled()
    window._set_hiring_focus_mode(False)
    assert dashboard.staffing_sidebar.width() == 252
    assert dashboard.dashboard_nav_button.isEnabled()
    window.window.close()
    app.processEvents()


def test_pyside_session_imports_indeed_transcript_for_rating_review(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Miriam", school="Palmdale", track_key="preschool")

    result = session.import_indeed_transcript_text(
        """
Speaker 1: So, Miriam, why are you applying specifically to our company, Launchpad?

Speaker 0: You are closer to home and the ratios look safe and intentional.

Speaker 1: Tell me about a time a child was having a hard moment emotionally.

Speaker 0: I got low, helped them breathe, and supported words for what happened.
""",
    )

    assert result.interviewer_speaker == "Speaker 1"
    assert result.candidate_speaker == "Speaker 0"
    assert result.mapped_count == 2
    assert session.answers["Why-LPL"]["notes"].startswith("You are closer")
    assert session.answers["trait_1"]["notes"].startswith("I got low")
    assert session.answers["trait_1"]["score"] == ""
    assert session.flow_candidate_transcripts[2].startswith("I got low")
    assert session.active_question().question_id == "trait_1"

def test_pyside_session_writes_interview_session_snapshots(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Miriam", school="Palmdale", track_key="preschool")
    session.flow_candidate_transcripts[0] = "Candidate wants to work close to home."

    session.save_answer_and_advance(notes="Strong intro answer", score="")

    session_files = list((tmp_path / "interview_sessions").glob("*.json"))
    assert len(session_files) == 1
    payload = json.loads(session_files[0].read_text(encoding="utf-8"))
    assert payload["interview"]["candidate_name"] == "Miriam"
    assert payload["interview"]["interview_date"] == session.interview_date
    assert payload["questions"]["0"]["item_id"] == "intro_script"
    assert payload["questions"]["0"]["candidate_transcript"] == "Candidate wants to work close to home."
    assert "Company Statement" in payload["questions"]["0"]["notes"]["question_text"]

def test_pyside_session_imports_indeed_transcript_when_speaker_numbers_flip(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Chantelle", school="Palmdale", track_key="preschool")

    result = session.import_indeed_transcript_text(
        """
Speaker 0: Why are you applying specifically to our company, Launchpad?

Speaker 1: The school feels aligned with my experience and commute.

Speaker 0: Tell me about a time a child was having a hard moment emotionally.

Speaker 1: I noticed frustration, paused, and helped the child use words.
""",
    )

    assert result.interviewer_speaker == "Speaker 0"
    assert result.candidate_speaker == "Speaker 1"
    assert session.answers["trait_1"]["notes"].startswith("I noticed frustration")
    assert session.answers["trait_1"]["score"] == ""


def test_indeed_import_keeps_follow_up_with_active_scored_answer() -> None:
    turns = parse_indeed_transcript_text(
        """
Speaker 0: Tell me about a time a child was having a really hard moment emotionally. What did you notice and how did you respond?
Speaker 1: I stayed calm and gave the child space.
Speaker 0: What did you do at that time after the child calmed down?
Speaker 1: I checked in and helped the child return to play.
Speaker 0: Tell me about a time you felt overwhelmed or stressed at work with children present. What did you do to regulate yourself?
Speaker 1: I took a breath and asked another teacher for support.
"""
    )
    result = map_indeed_transcript_to_questions(
        turns,
        [
            {
                "flow_index": 3,
                "question_id": "trait_1",
                "prompt": "Tell me about a time a child was having a really hard moment emotionally. "
                "What did you notice about what they were feeling, and how did you respond?",
            },
            {
                "flow_index": 4,
                "question_id": "trait_2",
                "prompt": "Tell me about a time you felt overwhelmed or stressed at work with children present. "
                "What did you do to regulate yourself and manage your stress?",
            },
            {
                "flow_index": 12,
                "question_id": "FT-or-PT",
                "prompt": "Are you looking for full-time or part-time?",
            },
        ],
    )
    by_id = {match.question_id: match.candidate_transcript for match in result.matches}

    assert by_id["trait_1"] == (
        "I stayed calm and gave the child space. "
        "I checked in and helped the child return to play."
    )
    assert by_id["trait_2"] == "I took a breath and asked another teacher for support."
    assert result.unmatched_question_ids == ["FT-or-PT"]


def test_indeed_import_uses_questions_supplied_by_new_track() -> None:
    result = map_indeed_transcript_to_questions(
        parse_indeed_transcript_text(
            "Speaker 3: Describe how you safely operate a pottery kiln.\n"
            "Speaker 7: I inspect ventilation, verify temperature controls, and use protective equipment.\n"
        ),
        [
            {
                "flow_index": 1,
                "question_id": "ceramics_safety",
                "prompt": "How would you safely operate and monitor a pottery kiln?",
            }
        ],
    )

    assert result.mapped_count == 1
    assert result.matches[0].question_id == "ceramics_safety"
    assert result.matches[0].candidate_transcript.startswith("I inspect ventilation")


def _widget_text(widget) -> str:
    from PySide6 import QtWidgets

    labels = widget.findChildren(QtWidgets.QLabel)
    return " ".join(label.text() for label in labels)


def _table_text(table) -> str:
    if table is None:
        return ""
    values: list[str] = []
    for row in range(table.rowCount()):
        for column in range(table.columnCount()):
            item = table.item(row, column)
            if item is not None:
                values.append(item.text())
    return " ".join(values)


def _icon_has_primary_blue(icon, size: int = 18) -> bool:
    image = icon.pixmap(size, size).toImage()
    for y in range(image.height()):
        for x in range(image.width()):
            color = image.pixelColor(x, y)
            if color.alpha() > 0 and color.red() == 37 and color.green() == 99 and color.blue() == 235:
                return True
    return False

def test_staffing_v2_validation_helpers_hide_seed_dates_and_build_worklist() -> None:
    rows = [
        StaffingMetricRow(
            assignment_id=1,
            school="Hawthorne",
            classroom="Harmony 1",
            position_name="Teacher 1",
            position_type="Teacher",
            status="need_now",
            start_date="1970-01-01T00:00:00Z",
        ),
        StaffingMetricRow(
            assignment_id=2,
            school="Hawthorne",
            classroom="Harmony 1",
            position_name="Teacher 2",
            position_type="Teacher",
            status="filled",
            person_name="Stephanie",
            permit_status="no_permit_or_application",
        ),
    ]

    issues = _validation_issues_from_rows(rows)

    assert _display_date("1970-01-01T00:00:00Z") == "-"
    assert [issue["issue"] for issue in issues[:1]] == ["Unfilled Need Now position"]
    assert {issue["issue"] for issue in issues[1:]} == {
        "Placeholder start date",
        "Permit status needs review",
    }
    assert {issue["type"] for issue in issues} == {"Coverage", "Lifecycle", "Compliance"}
    assert issues[0]["severity"] == "Critical"

def test_staffing_v2_scroll_helper_relays_wheel_events_from_child_widgets() -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    root = qt_widgets.QWidget()
    root_layout = qt_widgets.QVBoxLayout(root)
    scroll = qt_widgets.QScrollArea()
    scroll.setWidgetResizable(True)
    content = qt_widgets.QWidget()
    content_layout = qt_widgets.QVBoxLayout(content)
    first_child = qt_widgets.QLabel("Top row")
    content_layout.addWidget(first_child)
    for index in range(40):
        content_layout.addWidget(qt_widgets.QLabel(f"Scrollable row {index}"))
    scroll.setWidget(content)
    scroll.resize(220, 120)
    root_layout.addWidget(scroll)
    root.show()
    app.processEvents()

    configure_v2_scroll_areas(qt_widgets, root, qt_core)
    wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(10, 10),
        qt_core.QPointF(10, 10),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(first_child, wheel)
    app.processEvents()

    assert scroll.verticalScrollBar().value() > 0
    assert wheel.isAccepted()
    root.close()
    app.processEvents()

def test_staffing_v2_scroll_helper_relays_from_nested_item_view_without_range() -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    root = qt_widgets.QWidget()
    root_layout = qt_widgets.QVBoxLayout(root)
    scroll = qt_widgets.QScrollArea()
    scroll.setWidgetResizable(True)
    content = qt_widgets.QWidget()
    content_layout = qt_widgets.QVBoxLayout(content)
    table = qt_widgets.QTableWidget(1, 1)
    table.setItem(0, 0, qt_widgets.QTableWidgetItem("Visible table row"))
    table.setFixedHeight(90)
    content_layout.addWidget(table)
    for index in range(40):
        content_layout.addWidget(qt_widgets.QLabel(f"Scrollable row {index}"))
    scroll.setWidget(content)
    scroll.resize(260, 140)
    root_layout.addWidget(scroll)
    root.show()
    app.processEvents()

    configure_v2_scroll_areas(qt_widgets, root, qt_core)
    wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(10, 10),
        qt_core.QPointF(10, 10),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(table.viewport(), wheel)
    app.processEvents()

    assert table.verticalScrollBar().maximum() == 0
    assert scroll.verticalScrollBar().value() > 0
    assert wheel.isAccepted()
    root.close()
    app.processEvents()


def test_staffing_v2_application_wheel_router_does_not_retain_root() -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    root = qt_widgets.QWidget()

    configure_v2_scroll_areas(qt_widgets, root, qt_core)

    routers = app._staffing_v2_application_wheel_routers
    router = routers[id(root)]
    assert not hasattr(router, "root")
    assert router._root_ref() is root

def test_redesign_model_prioritizes_guided_interview_workflow(tmp_path: Path) -> None:
    rubric_path = tmp_path / "rubric.json"
    overrides_path = tmp_path / "question_overrides.json"
    history_path = tmp_path / "interview_history.sqlite3"

    rubric_path.write_text(
        json.dumps(
            {
                "metadata": {"version": "test"},
                "scoring": {},
                "tracks": {"preschool": {"label": "Preschool", "max_weighted_total": 5}},
                "absolute_disqualifiers": ["Unsafe handling"],
                "traits": [
                    {
                        "id": "trait_1",
                        "name": "Empathy",
                        "priority": "Critical",
                        "weight": 1,
                        "applicable_tracks": ["preschool"],
                        "primary_question": "Tell me about a hard child moment.",
                        "descriptors": {
                            "1": "Serious concern",
                            "2": "Weak",
                            "3": "Mixed / acceptable",
                            "4": "Strong",
                            "5": "Excellent",
                        },
                        "sample_answers": {},
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    overrides_path.write_text(
        json.dumps(
            {
                "track_trait_order": {},
                "trait_question_overrides": {},
                "custom_questions": {
                    "preschool": [{"id": "Why-LPL", "text": "Why Launch Pad Learning?", "order": 1}]
                },
                "track_question_flow": {
                    "preschool": [
                        {"type": "custom", "id": "Why-LPL"},
                        {"type": "trait", "id": "trait_1"},
                    ]
                },
            }
        ),
        encoding="utf-8",
    )
    InterviewHistoryStore(history_path).append(
        {
            "candidate_name": "Latoya Nugent",
            "school": "Palmdale",
            "track": "Preschool",
            "score": 60.95,
            "status": "Finalized",
            "next_action": "Generate Offer",
        }
    )

    model = build_interview_redesign_model(
        rubric_path=rubric_path,
        overrides_path=overrides_path,
        history_path=history_path,
        school_options=["Palmdale"],
    )

    assert model.app_title == "Interview Assistant"
    assert model.navigation == [
        "Staffing",
        "Staffing v2",
        "Onboarding",
    ]
    assert model.home.primary_action == "Start a New Interview"
    assert model.home.admin_visible_on_home is False
    assert model.home.recent_interviews[0].next_action == "Generate Offer"
    assert history_path.exists()
    assert not history_path.with_suffix(".json").exists()
    assert model.setup_steps == ["Candidate", "Interview Plan", "Ready"]

    preschool_flow = model.flows["preschool"]
    assert [item.kind for item in preschool_flow.items] == ["custom", "trait"]
    assert preschool_flow.items[0].prompt == "Why Launch Pad Learning?"
    scored = preschool_flow.items[1]
    assert scored.score_cards[0].label == "1"
    assert scored.score_cards[0].description == "Serious concern"
    assert "Needs follow-up" in scored.quick_actions
    assert "Disqualifier observed" in scored.quick_actions

def test_pyside_model_accepts_revision_probe_and_gateway_fields(tmp_path: Path) -> None:
    rubric_path = tmp_path / "rubric.json"
    overrides_path = tmp_path / "question_overrides.json"
    history_path = tmp_path / "interview_history.sqlite3"

    rubric_path.write_text(
        json.dumps(
            {
                "metadata": {"version": "test"},
                "scoring": {},
                "tracks": {
                    "behavior_support_specialist": {
                        "label": "Behavior Support Specialist",
                        "max_weighted_total": 5,
                        "gateway_requirements": ["Meets schedule requirements"],
                    }
                },
                "absolute_disqualifiers": ["Unsafe handling"],
                "interviewer_guidance": {},
                "traits": [
                    {
                        "id": "bss_trait_1",
                        "name": "Behavior Support",
                        "priority": "Critical",
                        "weight": 1,
                        "applicable_tracks": ["behavior_support_specialist"],
                        "primary_question": "Tell me about behavior support.",
                        "follow_up_probes": ["What did you try first?", "What changed after that?"],
                        "descriptors": {
                            "1": "Serious concern. Unsafe or blaming.",
                            "2": "Weak. Thin or adult-centered.",
                            "3": "Mixed. Safe but basic.",
                            "4": "Strong. Practical and child-centered.",
                            "5": "Excellent. Reflective, specific, and preventive.",
                        },
                        "sample_answers": {
                            "1": "Low",
                            "2": "Weak",
                            "3": "Mixed",
                            "4": "Strong",
                            "5": "Excellent",
                        },
                    }
                ],
                "final_record_fields": ["Gateway Requirements Status"],
                "canonical_text": "canonical",
            }
        ),
        encoding="utf-8",
    )
    overrides_path.write_text(
        json.dumps(
            {
                "track_trait_order": {},
                "trait_question_overrides": {},
                "custom_questions": {},
                "track_question_flow": {
                    "behavior_support_specialist": [{"type": "trait", "id": "bss_trait_1"}]
                },
            }
        ),
        encoding="utf-8",
    )
    InterviewHistoryStore(history_path).load()

    model = build_interview_redesign_model(
        rubric_path=rubric_path,
        overrides_path=overrides_path,
        history_path=history_path,
        school_options=["Palmdale"],
    )

    scored = model.flows["behavior_support_specialist"].items[0]
    assert scored.followups == ["What did you try first?", "What changed after that?"]
    assert scored.score_cards[0].description == "Serious concern. Unsafe or blaming."

def test_pyside_model_loads_legacy_history_into_history_rows(tmp_path: Path) -> None:
    canonical_dir = tmp_path / "user_artifacts"
    canonical_dir.mkdir()
    legacy_path = tmp_path / "interview_history.json"
    legacy_path.write_text(
        json.dumps(
            [
                {
                    "history_id": "legacy-1",
                    "candidate_name": "Legacy Candidate",
                    "school": "Palmdale",
                    "position": "Preschool Teacher",
                    "offer_status": "not_generated",
                }
            ]
        ),
        encoding="utf-8",
    )

    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=canonical_dir / "interview_history.json",
        school_options=["Palmdale"],
    )

    assert [row.candidate for row in model.home.history_rows] == ["Legacy Candidate"]
    assert model.home.recent_interviews[0].candidate == "Legacy Candidate"

def test_pyside_history_rows_expose_school_position_and_offer_action(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-1",
            "candidate_name": "Latoya Nugent",
            "school": "Palmdale",
            "position": "Preschool Teacher",
            "percent_of_max": 88.5,
            "outcome": "Hire",
            "offer_status": "not_generated",
            "interview_notes_path": str(tmp_path / "notes.docx"),
        }
    )

    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )

    row = model.home.history_rows[0]

    assert row.row_key == "hist-1"
    assert row.school == "Palmdale"
    assert row.position == "Preschool Teacher"
    assert row.offer_status == "not_generated"
    assert row.offer_action == "Generate Offer"
    assert row.interview_date == ""
    assert row.notes_path.endswith("notes.docx")


def test_pyside_history_rows_show_newest_interviews_first(tmp_path: Path) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    store = InterviewHistoryStore(history_path)
    store.append(
        {
            "history_id": "hist-old",
            "candidate_name": "Old Candidate",
            "interview_date": "2026-07-07",
        }
    )
    store.append(
        {
            "history_id": "hist-today",
            "candidate_name": "Today Candidate",
            "interview_date": "2026-07-08",
        }
    )

    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )

    assert [row.candidate for row in model.home.history_rows] == ["Today Candidate", "Old Candidate"]
    assert model.home.recent_interviews[0].candidate == "Today Candidate"

def test_pyside_session_autosaves_and_resumes_guided_interview(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    draft_path = tmp_path / "drafts" / "latoya-preschool.json"

    session = PySideInterviewSession(model=model, draft_path=draft_path)
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    first = session.active_question()

    session.save_answer_and_advance(notes="Wants values-aligned school.", score="")

    assert draft_path.exists()
    assert first is not None
    assert session.current_index == 1

    resumed = PySideInterviewSession.load(model=model, draft_path=draft_path)

    assert resumed.candidate_name == "Latoya Nugent"
    assert resumed.school == "Palmdale"
    assert resumed.track_key == "preschool"
    assert resumed.current_index == 1
    assert resumed.answers[first.question_id]["notes"] == "Wants values-aligned school."
    assert resumed.active_question().kind == "custom"

def test_pyside_session_back_and_skip_preserve_answers_and_move_through_flow(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")

    intro = session.active_question()
    session.save_answer_and_advance(notes="Intro complete.")
    custom = session.active_question()
    session.save_answer_and_advance(notes="Values aligned.")
    scored = session.active_question()

    session.go_back()

    assert session.active_question() == custom
    assert session.answers[custom.question_id]["notes"] == "Values aligned."

    session.skip_active_question(notes="No extra custom notes.")
    assert session.active_question() == scored

    session.skip_active_question(notes="Skipped rating after discussion.")
    assert session.active_question() is None
    assert session.answers[scored.question_id]["score"] == ""
    assert session.answers[scored.question_id]["skipped"] is True
    assert intro is not None
    assert session.answers[intro.question_id]["notes"] == "Intro complete."

def test_pyside_live_footer_blocks_unrated_scored_next_and_keeps_skip_enabled(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro complete.")
    session.save_answer_and_advance(notes="Values aligned.")
    window.session = session
    window.session_track_key = session.track_key
    window.session_index = session.current_index
    window._render_live_question_page()

    footer_buttons = {
        button.property("pyside_live_footer_action"): button
        for button in window.interview_tabs.widget(2).findChildren(qt_widgets.QPushButton)
        if button.property("pyside_live_footer_action")
    }
    capture_bar = window.interview_tabs.widget(2).findChild(qt_widgets.QFrame, "LiveInterviewHeader")
    question_rail = window.interview_tabs.widget(2).findChild(qt_widgets.QFrame, "LiveInterviewStageRail")

    assert footer_buttons["back"].isEnabled()
    assert footer_buttons["skip"].isEnabled()
    assert not footer_buttons["finalize"].isEnabled()
    assert window.interview_tabs.tabBar().isHidden()
    assert capture_bar is not None
    assert capture_bar.property("captureState") in {"inactive", "warning", "recording"}
    stage_states = [row.property("stageState") for row in question_rail.findChildren(qt_widgets.QFrame)]
    assert "complete" in stage_states
    assert "active" in stage_states

    window.score_group.buttons()[0].setChecked(True)
    app.processEvents()

    assert footer_buttons["finalize"].isEnabled()
    window.window.close()
    app.processEvents()

def test_pyside_home_import_indeed_transcript_opens_rating_flow(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    transcript_path = tmp_path / "indeed.txt"
    transcript_path.write_text(
        """
Speaker 1: Tell me about a time a child was having a hard moment emotionally.

Speaker 0: I noticed the child was upset and helped them name the feeling.
""",
        encoding="utf-8",
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(
        window,
        "_collect_indeed_transcript_import_request",
        lambda: {
            "candidate_name": "Miriam",
            "interview_date": "2026-07-10",
            "school": "Palmdale",
            "track_key": "preschool",
            "transcript_path": transcript_path,
        },
    )
    window._start_pyside_interview_recording = lambda: (_ for _ in ()).throw(AssertionError("recording should not start"))

    window.staffing_v2_dashboard.external_pages["candidates"].findChild(
        qt_widgets.QPushButton, "ImportIndeedTranscriptButton"
    ).click()
    app.processEvents()

    assert window.session is not None
    assert window.session.active_question().question_id == "trait_1"
    assert "helped them name" in window.live_notes.toPlainText()
    assert not window.live_next_button.isEnabled()
    assert window.interview_tabs.currentIndex() == 2
    window.window.close()
    app.processEvents()

def test_pyside_home_import_indeed_transcript_creates_history_row_from_metadata(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale", "Hawthorne"],
    )
    transcript_path = tmp_path / "indeed.txt"
    transcript_path.write_text(
        """
Speaker 1: Tell me about a time a child was having a hard moment emotionally.

Speaker 0: I noticed the child was upset and helped them name the feeling.
""",
        encoding="utf-8",
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(
        window,
        "_collect_indeed_transcript_import_request",
        lambda: {
            "candidate_name": "Miriam Rivera",
            "interview_date": "2026-07-10",
            "school": "Hawthorne",
            "track_key": "preschool",
            "transcript_path": transcript_path,
        },
    )
    window._start_pyside_interview_recording = lambda: (_ for _ in ()).throw(AssertionError("recording should not start"))

    window._import_indeed_transcript_from_home()

    rows = InterviewHistoryStore(history_path).load()
    assert len(rows) == 1
    row = rows[0]
    assert row["candidate_name"] == "Miriam Rivera"
    assert row["interview_date"] == "2026-07-10"
    assert row["school"] == "Hawthorne"
    assert row["track_key"] == "preschool"
    assert row["outcome"] == "Incomplete"
    assert row["imported_indeed_transcript"]["source_path"] == str(transcript_path)
    assert row["answers"]["trait_1"]["notes"].startswith("I noticed")
    assert any(
        str(item.get("candidate_transcript") or "").startswith("I noticed")
        for item in row["flow_transcript"]
    )
    assert row["flow_recordings"][0]["source"] == "indeed_transcript_import"
    assert window.session is not None
    assert window.session.candidate_name == "Miriam Rivera"
    assert window.session.interview_date == "2026-07-10"
    assert window.session.active_question().question_id == "trait_1"
    assert "helped them name" in window.live_notes.toPlainText()
    assert window.interview_tabs.currentIndex() == 2
    assert window._review_history_id == row["history_id"]
    window.window.close()
    app.processEvents()


def test_pyside_home_import_reuses_matching_scored_history_row(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    store = InterviewHistoryStore(history_path)
    history_entry = (
        {
            "history_id": "hist-michelle",
            "candidate_name": "Michelle Oropeza",
            "school": "Palmdale",
            "position": "Preschool",
            "track_key": "preschool",
            "interview_date": "2026-07-14",
            "outcome": "Hire",
            "review_scores": {"trait_1": "4"},
            "answers": {"trait_1": {"kind": "trait", "score": "4", "notes": "Prior notes"}},
        }
    )
    store.append_with_candidate_report(
        history_entry,
        {
            "schema_version": 1,
            "history_id": "hist-michelle",
            "candidate": {"candidate_name": "Michelle Oropeza", "school": "Palmdale", "track": "preschool"},
            "questions": [
                {
                    "question_id": "trait_1",
                    "type": "trait",
                    "title": "Empathy",
                    "prompt": "Tell me about a hard child moment.",
                    "transcript": "",
                    "original_transcript": "",
                    "rating": 4,
                }
            ],
            "scoring": {"percent_of_max": 80.0, "outcome": "Hire", "rows": []},
            "summaries": {},
            "report_path": "",
        },
        actor="admin-user",
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    transcript_path = tmp_path / "michelle.txt"
    transcript_path.write_text(
        "Speaker 0: Tell me about a hard child moment.\n"
        "Speaker 1: I stayed calm, named the feeling, and helped the child recover.\n",
        encoding="utf-8",
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(
        window,
        "_collect_indeed_transcript_import_request",
        lambda: {
            "candidate_name": " Michelle Oropeza ",
            "interview_date": "2026-07-14",
            "school": "Palmdale",
            "track_key": "preschool",
            "transcript_path": transcript_path,
        },
    )
    monkeypatch.setattr(window, "_regenerate_history_import_artifacts", lambda _row, _session: {})

    window._import_indeed_transcript_from_home()

    rows = store.load()
    assert len(rows) == 1
    assert rows[0]["history_id"] == "hist-michelle"
    assert rows[0]["answers"]["trait_1"]["score"] == "4"
    assert rows[0]["answers"]["trait_1"]["notes"].startswith("I stayed calm")
    assert window._review_history_id == "hist-michelle"
    director_report = CandidateReportRepository(history_path).load_visible_version(
        "hist-michelle", role="director", school_scope="Palmdale"
    )
    assert director_report.snapshot["questions"][0]["transcript"].startswith("I stayed calm")
    window.window.close()
    app.processEvents()



def test_pyside_history_import_indeed_transcript_opens_review_for_existing_candidate(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    history_path = tmp_path / "interview_history.sqlite3"
    store = InterviewHistoryStore(history_path)
    store.append(
        {
            "history_id": "hist-grace",
            "candidate_name": "Grace Morales",
            "school": "Palmdale",
            "position": "Preschool Teacher",
            "interview_date": "2026-07-10",
            "outcome": "Incomplete",
            "score": "80.0%",
            "review_scores": {"trait_1": "4"},
            "candidate": {
                "qualification": {
                    "has_degree": True,
                    "degree_type": "AA",
                    "degree_in_ece": True,
                    "ece_units_completed": 35,
                    "infant_toddler_class_completed": True,
                    "total_units_completed": None,
                    "years_experience": 30,
                }
            },
            "answers": {
                "FT-or-PT": {
                    "kind": "custom",
                    "title": "Non-scored question",
                    "prompt": "Are you looking for full-time or part-time?",
                    "notes": "part-time for now, full-time eventually.",
                    "score": "",
                    "quick_actions": [],
                },
                "trait_1": {
                    "kind": "trait",
                    "title": "Empathy",
                    "prompt": "Tell me about a time a child was having a hard moment emotionally.",
                    "notes": "Prior notes",
                    "score": "4",
                    "quick_actions": [],
                }
            },
        }
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    transcript_path = tmp_path / "grace-indeed.txt"
    transcript_path.write_text(
        """
Speaker 0: Tell me about a time a child was having a hard moment emotionally.

Speaker 1: I got low, named the feeling, and helped the child find words.

Speaker 0: Are you looking for full-time or part-time?

Speaker 1: I gave a long transcript answer that should not replace the saved logistics field.
""",
        encoding="utf-8",
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(
        window.QtWidgets.QFileDialog,
        "getOpenFileName",
        lambda *_args, **_kwargs: (str(transcript_path), "Text files (*.txt)"),
    )
    row = window.model.home.history_rows[0]

    window._import_indeed_transcript_for_history_row(row)

    assert window.session is not None
    assert window.session.candidate_name == "Grace Morales"
    assert window.session.school == "Palmdale"
    assert window.session.answers["trait_1"]["score"] == "4"
    assert window.interview_tabs.currentIndex() == 3
    review_page = window.interview_tabs.widget(3)
    assert "I got low, named the feeling" in _widget_text(review_page)
    trait_card = next(
        card
        for card in review_page.findChildren(qt_widgets.QFrame, "CompletedTranscriptCard")
        if card.property("questionId") == "trait_1"
    )
    assert "Rating 4 / 5" in _widget_text(trait_card)
    assert trait_card.findChild(qt_widgets.QPushButton, "CompletedTranscriptDetail") is not None

    stored = InterviewHistoryStore(history_path).load()[0]
    assert stored["review_scores"]["trait_1"] == "4"
    assert stored["scoring"]["rows"][0]["raw_score"] == 4
    assert stored["score"] == "80.0%"
    assert stored["candidate"]["qualification"]["degree_type"] == "AA"
    assert stored["answers"]["FT-or-PT"]["notes"] == "part-time for now, full-time eventually."
    assert stored["imported_indeed_transcript"]["mapped_count"] == 1
    assert stored["answers"]["trait_1"]["notes"].startswith("I got low")
    assert any(
        str(item.get("candidate_transcript") or "").startswith("I got low")
        for item in stored["flow_transcript"]
    )
    assert stored["flow_recordings"][0]["candidate_transcript"].startswith("I got low")
    assert Path(stored["report_path"]).resolve().is_relative_to(tmp_path.resolve())
    window.window.close()
    app.processEvents()

def test_pyside_history_import_regenerates_basic_notes(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    history_path = tmp_path / "interview_history.sqlite3"
    notes_dir = tmp_path / "notes"
    notes_dir.mkdir()
    stale_notes = notes_dir / "2026-07-10 - Palmdale - Grace Morales - Basic Interview Notes.docx"
    stale_notes.write_text("stale", encoding="utf-8")
    store = InterviewHistoryStore(history_path)
    store.append(
        {
            "history_id": "hist-grace",
            "candidate_name": "Grace Morales",
            "school": "Palmdale",
            "position": "Preschool Teacher",
            "interview_date": "2026-07-10",
            "interview_notes_path": str(stale_notes),
            "saved_report_path": str(stale_notes),
        }
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    transcript_path = tmp_path / "grace-indeed.txt"
    transcript_path.write_text(
        """
Speaker 0: Tell me about a time a child was having a hard moment emotionally.

Speaker 1: I got low, named the feeling, and helped the child find words.
""",
        encoding="utf-8",
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(
        window.QtWidgets.QFileDialog,
        "getOpenFileName",
        lambda *_args, **_kwargs: (str(transcript_path), "Text files (*.txt)"),
    )

    window._import_indeed_transcript_for_history_row(window.model.home.history_rows[0])

    row = InterviewHistoryStore(history_path).load()[0]
    notes_path = Path(row["interview_notes_path"])
    assert notes_path.exists()
    assert notes_path == stale_notes
    assert "Structured Behavioral Interview Notes" in _docx_text(notes_path)
    assert "I got low, named the feeling" in _docx_text(notes_path)
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
def test_candidate_screen_regenerate_notes_restores_structured_report_content_scenario(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing-staffing-seed.json")
    history_path = tmp_path / "interview_history.sqlite3"
    notes_path = tmp_path / "notes" / "candidate-notes.docx"
    store = InterviewHistoryStore(history_path)
    store.append_with_candidate_report(
        {
            "history_id": "hist-regenerate",
            "candidate_name": "Jordan Lee",
            "school": "Palmdale",
            "position": "Preschool",
            "track": "preschool",
            "interview_date": "2026-07-14",
            "interview_notes_path": str(notes_path),
            "saved_report_path": str(notes_path),
        },
        {
            "schema_version": 1,
            "history_id": "hist-regenerate",
            "candidate": {
                "candidate_name": "Jordan Lee",
                "school": "Palmdale",
                "track": "preschool",
                "interview_date": "2026-07-14",
                "qualification": {
                    "has_degree": True,
                    "degree_type": "AA",
                    "degree_in_ece": True,
                    "ece_units_completed": 24,
                    "years_experience": 4,
                },
            },
            "questions": [
                {
                    "flow_index": 1,
                    "question_id": "Why-LPL",
                    "type": "custom",
                    "title": "Why Launch Pad Learning?",
                    "prompt": "Why Launch Pad Learning?",
                    "transcript": "I value the school mission.",
                    "interviewer_notes": "Mission aligned.",
                },
                {
                    "flow_index": 2,
                    "question_id": "trait_1",
                    "type": "trait",
                    "title": "Empathy",
                    "prompt": "Tell me about a hard child moment.",
                    "transcript": "I named the feeling and stayed close.",
                    "interviewer_notes": "Strong specific example.",
                    "rating": 4,
                    "weight": 1,
                    "weighted_score": 4,
                    "priority": "Critical",
                    "skipped": False,
                    "absolute_disqualifier": False,
                    "no_example_after_followups": False,
                },
            ],
            "scoring": {
                "weighted_total": 4,
                "max_weighted_total": 5,
                "percent_of_max": 80.0,
                "outcome": "Hire",
                "rows": [
                    {
                        "trait_id": "trait_1",
                        "raw_score": 4,
                        "weight": 1,
                        "weighted_score": 4,
                    }
                ],
            },
            "summaries": {},
            "report_path": str(notes_path),
        },
        actor="admin-user",
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    warnings: list[str] = []
    monkeypatch.setattr(
        window.QtWidgets.QMessageBox,
        "warning",
        lambda _parent, _title, message: warnings.append(str(message)),
    )

    window._regenerate_history_notes(window.model.home.history_rows[0])

    assert warnings == []
    regenerated_path = Path(store.load()[0]["report_path"])
    rendered = _docx_text(regenerated_path)
    assert "AA" in rendered
    assert "I value the school mission." in rendered
    assert "I named the feeling and stayed close." in rendered
    assert "4 / 5" in rendered
    window.window.close()
    app.processEvents()

def test_pyside_show_schedules_recording_interface_preload_once(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    scheduled = []

    def _single_shot(_delay, callback):
        scheduled.append(callback)

    monkeypatch.setattr(window.QtCore.QTimer, "singleShot", _single_shot)
    monkeypatch.setattr(window.window, "showMaximized", lambda: None)

    window.show()
    window.show()

    assert window._preload_recording_interface_async in scheduled
    assert scheduled.count(window._preload_recording_interface_async) == 1
    window.window.close()
    app.processEvents()

def test_pyside_home_delete_saved_draft_requires_confirmation(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    draft_path = tmp_path / "drafts" / "latoya.json"
    draft_path.parent.mkdir()
    draft_path.write_text('{"schema":"pyside_interview_draft.v1"}', encoding="utf-8")
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(pyside_interview_app, "latest_pyside_draft_path", lambda _drafts_dir=None: draft_path)
    window._refresh_home_draft_panel()
    no = window.QtWidgets.QMessageBox.StandardButton.No
    yes = window.QtWidgets.QMessageBox.StandardButton.Yes

    monkeypatch.setattr(window.QtWidgets.QMessageBox, "question", lambda *_args, **_kwargs: no)
    window.candidate_delete_draft_button.click()
    app.processEvents()
    assert draft_path.exists()

    monkeypatch.setattr(window.QtWidgets.QMessageBox, "question", lambda *_args, **_kwargs: yes)
    window.candidate_delete_draft_button.click()
    app.processEvents()
    assert not draft_path.exists()
    assert not window.home_continue_button.isEnabled()
    assert not window.home_delete_draft_button.isEnabled()
    window.window.close()
    app.processEvents()

def test_pyside_draft_actions_live_under_candidates_not_new_interview(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)

    start_buttons = {
        button.text() for button in window.interview_tabs.widget(0).findChildren(qt_widgets.QPushButton)
    }
    candidate_tabs = window.staffing_v2_dashboard.external_pages["candidates"].findChild(
        qt_widgets.QTabWidget,
        "HiringV2CandidatesTabs",
    )
    draft_buttons = {button.text() for button in candidate_tabs.widget(1).findChildren(qt_widgets.QPushButton)}

    assert start_buttons == {"Test Audio", "Cancel", "Begin Interview"}
    assert {"Continue Saved Draft", "Delete Saved Draft"} <= draft_buttons
    assert window.window.findChild(qt_widgets.QTableWidget, "PySideHistoryGrid") is None
    window.window.close()
    app.processEvents()

def test_pyside_back_reentry_overwrites_existing_flow_timestamp(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    window.session = session

    window._mark_flow_timestamp_at(1, 12.0)
    window._mark_flow_timestamp_at(1, 22.0, overwrite=True)

    assert session.flow_time_marks == [{"flow_index": 1, "t": 22.0}]
    window.window.close()
    app.processEvents()

def test_pyside_back_reentry_overwrites_existing_flow_end_timestamp(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    window.session = session
    window.recording_started_monotonic = 100.0
    session.flow_time_marks = [{"flow_index": 1, "t": 2.0, "end_t": 5.0}]
    window._overwrite_next_live_boundary_timestamp = True
    monkeypatch.setattr(pyside_interview_app.time, "monotonic", lambda: 112.0)

    assert window._close_flow_timestamp(1) == 12.0
    assert session.flow_time_marks == [{"flow_index": 1, "t": 2.0, "end_t": 12.0}]
    window.window.close()
    app.processEvents()

def test_pyside_back_reentry_next_overwrites_following_question_start(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    monkeypatch.setattr(pyside_interview_app.time, "monotonic", lambda: 112.0)
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.current_index = 1
    session.flow_time_marks = [
        {"flow_index": 1, "t": 2.0, "end_t": 5.0},
        {"flow_index": 2, "t": 5.0},
    ]
    window.session = session
    window.recording_started_monotonic = 100.0
    window._overwrite_next_live_boundary_timestamp = True

    window._save_and_next()

    assert session.flow_time_marks == [
        {"flow_index": 1, "t": 2.0, "end_t": 12.0},
        {"flow_index": 2, "t": 12.0},
    ]
    window.window.close()
    app.processEvents()

def test_pyside_session_enforces_intro_custom_scored_final_custom_workflow(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_workflow_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")

    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")

    seen: list[tuple[str, str, str]] = []
    while (active := session.active_question()) is not None:
        seen.append((active.kind, active.question_id, active.progress_label))
        session.save_answer_and_advance(notes=f"notes for {active.question_id}", score="5" if active.kind == "trait" else "")

    assert session.candidate_name == "Latoya Nugent"
    assert session.school == "Palmdale"
    assert session.track_key == "preschool"
    assert seen == [
        ("intro", "intro_script", "Question 1 of 8"),
        ("qualification", "Why-ECE", "Question 2 of 8"),
        ("custom", "Why-LPL", "Question 3 of 8"),
        ("trait", "trait_1", "Question 4 of 8"),
        ("custom", "FT-or-PT", "Question 5 of 8"),
        ("custom", "Not-Avail", "Question 6 of 8"),
        ("custom", "Pay", "Question 7 of 8"),
        ("custom", "Start", "Question 8 of 8"),
    ]
    assert "Palmdale is open weekdays" in session.answers["intro_script"]["prompt"]


def test_pyside_session_persists_candidate_honorific_in_draft(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_workflow_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    draft_path = tmp_path / "draft.json"
    session = PySideInterviewSession(model=model, draft_path=draft_path)

    session.start(candidate_name="Jordan Lee", honorific="Mr.", school="Palmdale", track_key="preschool")

    resumed = PySideInterviewSession.load(model=model, draft_path=draft_path)
    assert resumed.honorific == "Mr."

def test_pyside_recording_starts_on_begin_for_intro_audio_check(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    starts: list[str] = []
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: starts.append("start"))

    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()

    assert starts == ["start"]
    assert window.session.active_question().kind == "intro"

    window._save_and_next()

    assert starts == ["start"]
    assert window.session.active_question().kind == "custom"
    window.window.close()
    app.processEvents()

def test_pyside_recording_start_failure_shows_audio_device_warning(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(pyside_interview_app.sys, "platform", "win32")
    monkeypatch.setattr(pyside_interview_app, "resolve_default_windows_system_device", lambda: "Wrong Output Device")
    monkeypatch.setattr(
        pyside_interview_app,
        "resolve_default_windows_microphone_device",
        lambda: "Microphone (3- Realtek(R) Audio)",
    )

    import interview_audio_recorder

    recording_kwargs = {}

    def _raise_bad_device(**kwargs: object) -> object:
        recording_kwargs.update(kwargs)
        raise RuntimeError("Recording process exited immediately. Check configured audio device names.")

    monkeypatch.setattr(interview_audio_recorder, "start_recording", _raise_bad_device)

    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()
    app.processEvents()

    warning = window.window.findChild(qt_widgets.QLabel, "PySideRecordingWarning")
    assert warning is not None
    assert "Wrong Output Device" in warning.text()
    assert "Check Windows/meeting output" in warning.text()
    assert window.recording_session is None
    assert window.recording_started_monotonic is None
    assert window.recording_system_device == "Wrong Output Device"
    assert recording_kwargs["win_mic_device"] == "Microphone (3- Realtek(R) Audio)"
    window.window.close()
    app.processEvents()

def test_pyside_begin_uses_selected_setup_system_audio_source(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(pyside_interview_app.sys, "platform", "win32")
    monkeypatch.setattr(pyside_interview_app, "list_windows_dshow_audio_devices", lambda: ["Selected Cable"])
    monkeypatch.setattr(pyside_interview_app, "resolve_default_windows_system_device", lambda: "Default Cable")
    monkeypatch.setattr(
        pyside_interview_app,
        "resolve_default_windows_microphone_device",
        lambda: "Default Microphone",
    )
    import interview_audio_recorder

    recording_kwargs: dict[str, object] = {}

    def capture_then_fail(**kwargs: object) -> object:
        recording_kwargs.update(kwargs)
        raise RuntimeError("test stop")

    monkeypatch.setattr(interview_audio_recorder, "start_recording", capture_then_fail)
    window.home_audio_source_combo.clear()
    window.home_audio_source_combo.addItem("Selected Cable", "Selected Cable")
    window.home_candidate_input.setText("Latoya Nugent")

    window._begin_selected_interview()

    assert recording_kwargs["win_sys_device"] == "Selected Cable"
    assert window.recording_system_device == "Selected Cable"
    window.window.close()
    app.processEvents()

def test_pyside_intro_audio_check_warns_on_blank_transcription(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()

    window._apply_pyside_intro_audio_check_result("")
    app.processEvents()

    warning = window.window.findChild(qt_widgets.QLabel, "PySideRecordingWarning")
    assert warning is not None
    assert "No speech was transcribed from the first 15 seconds" in warning.text()
    assert "Record the interview in Zoom" in warning.text()
    window.window.close()
    app.processEvents()

def test_pyside_intro_audio_preflight_warns_without_blocking_active_interview(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    from interview_audio_preflight import AudioPreflightResult

    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()
    result = AudioPreflightResult(
        microphone_ready=True,
        system_audio_ready=False,
        transcription_ready=False,
        warning=(
            "Audio check did not verify candidate/system audio, candidate transcription. "
            "Check audio settings and record the interview in Zoom as a backup."
        ),
    )

    window._apply_pyside_intro_audio_preflight_result(result)
    app.processEvents()

    warning = window.interview_tabs.widget(2).findChild(qt_widgets.QLabel, "PySideRecordingWarning")
    assert warning is not None
    assert "candidate/system audio" in warning.text()
    assert window.session is not None
    assert window.session.active_question() is not None
    window.window.close()
    app.processEvents()

def test_pyside_intro_audio_check_ignores_interviewer_only_segments(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()

    session = SimpleNamespace(
        transcribe_new_segments=lambda **_kwargs: [SimpleNamespace(speaker="INTERVIEWER", text="Did you see our website?")]
    )

    assert window._transcribe_pyside_intro_audio_sample(session) == ""
    window.window.close()
    app.processEvents()

def test_pyside_system_audio_route_check_warns_when_capture_is_silent(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()
    window.session.save_answer_and_advance(notes="Intro read.")
    window.session_index = window.session.current_index
    window._render_live_question_page()

    sys_wav = tmp_path / "silent_sys.wav"
    with wave.open(str(sys_wav), "wb") as wav_file:
        wav_file.setnchannels(1)
        wav_file.setsampwidth(2)
        wav_file.setframerate(16000)
        wav_file.writeframes(b"\x00\x00" * 16000)

    window.recording_session = SimpleNamespace(sys_wav=sys_wav)
    window.recording_system_device = "VB-Audio Virtual Cable (CABLE Input)"
    window._check_pyside_system_audio_capture()
    app.processEvents()

    warning = window.window.findChild(qt_widgets.QLabel, "PySideRecordingWarning")
    assert warning is not None
    assert "No meeting/system audio detected yet" in warning.text()
    assert "switch Zoom/Windows output to VB-CABLE" in warning.text()
    window.window.close()
    app.processEvents()

def test_pyside_qualification_screen_keeps_education_and_experience_with_why_ece(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_workflow_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    draft_path = tmp_path / "draft.json"
    session = PySideInterviewSession(model=model, draft_path=draft_path)
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro read.")

    qualification_screen = session.active_question()
    assert qualification_screen is not None
    assert qualification_screen.kind == "qualification"
    assert qualification_screen.question_id == "Why-ECE"
    assert "why early childhood education" in qualification_screen.prompt.lower()

    session.save_answer_and_advance(
        notes="Grew up helping younger siblings.",
        qualification={
            "has_degree": True,
            "degree_type": "BA",
            "degree_in_ece": False,
            "ece_units_completed": 18,
            "infant_toddler_class_completed": True,
            "total_units_completed": None,
            "years_experience": 4,
        },
    )

    resumed = PySideInterviewSession.load(model=model, draft_path=draft_path)
    answer = resumed.answers["Why-ECE"]

    assert answer["notes"] == "Grew up helping younger siblings."
    assert answer["qualification"]["degree_type"] == "BA"
    assert answer["qualification"]["ece_units_completed"] == 18
    assert answer["qualification"]["years_experience"] == 4

def test_pyside_finalize_payload_uses_qualification_fields_and_keeps_why_ece_transcript(
    tmp_path: Path,
    monkeypatch,
) -> None:
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_workflow_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Lisvelia Pazos-Hilario", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro read.")
    session.save_answer_and_advance(
        notes="",
        qualification={
            "has_degree": True,
            "degree_type": "AA",
            "degree_in_ece": True,
            "ece_units_completed": None,
            "infant_toddler_class_completed": False,
            "total_units_completed": None,
            "years_experience": 7,
        },
    )
    session.qualification = {}
    session.flow_candidate_transcripts[1] = "Why-ECE candidate answer from transcript."
    session.save_answer_and_advance(notes="Why LPL manual note.")
    session.save_answer_and_advance(notes="Scored example.", score="5")
    session.save_answer_and_advance(notes="full-time")
    session.flow_candidate_transcripts[4] = "Noisy ASR for FT-or-PT should not win."
    session.save_answer_and_advance(notes="no")
    session.flow_candidate_transcripts[5] = "Noisy ASR for Not-Avail should not win."
    session.save_answer_and_advance(notes="Asked for 35/hour.")
    session.flow_candidate_transcripts[6] = "Noisy ASR for Pay should not win."
    session.save_answer_and_advance(notes="asap")
    session.flow_candidate_transcripts[7] = "Noisy ASR for Start should not win."
    adapter = pyside_interview_app._PySideFinalizeAdapter(
        session,
        base_dir=tmp_path,
        history_path=tmp_path / "history.sqlite3",
    )
    scoring = pyside_interview_app.ScoringEngine.evaluate(
        adapter._rubric_with_question_overrides(),
        adapter.state.track,
        adapter.state.trait_inputs,
    )

    context = pyside_interview_app.build_finalize_context(
        adapter,
        scoring,
        [],
        session._transcript_metadata(),
    )

    assert context.payload["candidate"]["qualification"] == {
        "has_degree": True,
        "degree_type": "AA",
        "degree_in_ece": True,
        "ece_units_completed": None,
        "infant_toddler_class_completed": False,
        "total_units_completed": None,
        "years_experience": 7,
    }
    by_id = {item["id"]: item for item in context.payload["flow_transcript"]}
    assert by_id["Why-ECE"]["candidate_transcript"] == "Why-ECE candidate answer from transcript."
    by_id = {item["id"]: item for item in context.payload["flow_transcript"]}
    assert by_id["FT-or-PT"]["candidate_transcript"] == "full-time"
    assert by_id["Not-Avail"]["candidate_transcript"] == "no"
    assert by_id["Pay"]["candidate_transcript"] == "Asked for 35/hour."
    assert by_id["Start"]["candidate_transcript"] == "asap"

def test_latest_pyside_draft_path_returns_newest_json(tmp_path: Path) -> None:
    old = tmp_path / "old.json"
    new = tmp_path / "new.json"
    old.write_text("{}", encoding="utf-8")
    new.write_text("{}", encoding="utf-8")

    assert latest_pyside_draft_path(tmp_path) == new

def test_pyside_session_review_summary_uses_scoring_rules(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")

    summary = session.review_summary()

    assert summary.percent_of_max == 100.0
    assert summary.outcome == "Hire"
    assert summary.next_action == "Generate Offer"
    assert summary.missing_scores == []
    assert summary.strongest_evidence == ["Warm child-centered example."]

def test_pyside_session_review_score_update_recalculates_and_persists(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="")
    session.flow_candidate_transcripts[2] = "I helped a child breathe and name their feelings."

    session.update_review_score("trait_1", 5)

    assert session.answers["trait_1"]["score"] == "5"
    assert session.review_summary().outcome == "Hire"
    draft = json.loads((tmp_path / "draft.json").read_text(encoding="utf-8"))
    assert draft["answers"]["trait_1"]["score"] == "5"
    session_file = next((tmp_path / "interview_sessions").glob("*.json"))
    payload = json.loads(session_file.read_text(encoding="utf-8"))
    assert payload["questions"]["2"]["notes"]["raw_score"] == 5
    assert payload["questions"]["2"]["candidate_transcript"] == "I helped a child breathe and name their feelings."

def test_pyside_offer_review_defaults_are_prefilled_from_completed_session(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")

    defaults = session.offer_review_defaults()

    assert defaults["candidate"] == "Latoya Nugent"
    assert defaults["school"] == "Palmdale"
    assert defaults["position"] == "Preschool"
    assert defaults["determination"] == "Hire"
    assert defaults["next_action"] == "Generate Offer"


def test_pyside_back_then_next_preserves_answer_start_and_moves_boundary(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.current_index = 2
    session.flow_time_marks = [
        {"flow_index": 1, "t": 2.0, "end_t": 5.0},
        {"flow_index": 2, "t": 5.0},
    ]
    window.session = session
    window.recording_started_monotonic = 100.0
    monkeypatch.setattr(pyside_interview_app.time, "monotonic", lambda: 112.0)

    window._go_back_live_question()
    window._save_and_next()

    assert session.flow_time_marks == [
        {"flow_index": 1, "t": 2.0, "end_t": 12.0},
        {"flow_index": 2, "t": 12.0},
    ]
    window.window.close()
    app.processEvents()

def test_pyside_review_screen_hides_finalize_and_offer_actions(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")
    window = _pyside_window_on_page(model, "Offers")
    window.session = session
    window._render_review_page()

    review_button_text = [
        child.text()
        for child in window.interview_tabs.widget(3).findChildren(qt_widgets.QPushButton)
    ]
    assert "Finalize Interview" not in review_button_text
    assert "Generate Offer" not in review_button_text

    assert "Offers" not in [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.window.close()
    app.processEvents()

def test_pyside_review_screen_shows_interviewer_closeout_without_slow_outputs(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro complete.", score="")
    session.skip_active_question(notes="No extra qualification notes.")
    session.save_answer_and_advance(notes="Candidate transcript should stay hidden.", score="")
    session.flow_candidate_transcripts[2] = "Transcript text should not render."
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.session = session

    window._render_review_page()

    review_page = window.interview_tabs.widget(3)
    visible_text = _widget_text(review_page)
    assert "Interview Complete" in visible_text
    assert "Latoya Nugent" in visible_text
    assert "Palmdale" in visible_text
    assert "Finalizing interview" in visible_text
    assert "Score Summary" in visible_text
    assert "Captured Transcripts" in visible_text
    assert "Transcript text should not render." in visible_text

    table = review_page.findChild(qt_widgets.QTableWidget, "CompletedInterviewTraitTable")
    assert table is not None
    assert table.columnCount() == 4
    assert table.item(0, 0).text() == "Empathy"
    assert table.item(0, 1).text() == "Missing"
    window.window.close()
    app.processEvents()

def test_pyside_completed_review_exposes_transcript_detail_for_rating_change(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_BASE_DIR", tmp_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro complete.", score="")
    session.save_answer_and_advance(notes="No extra qualification notes.", score="")
    session.save_answer_and_advance(notes="Review this against transcript.", score="")
    session.flow_candidate_transcripts[2] = "Candidate described calming a child and helping them use words."
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.session = session

    window._render_review_page()
    review_page = window.interview_tabs.widget(3)
    card = next(
        card
        for card in review_page.findChildren(qt_widgets.QFrame, "CompletedTranscriptCard")
        if card.property("questionId") == "trait_1"
    )
    assert "Candidate described calming a child" in _widget_text(card)
    assert card.findChild(qt_widgets.QPushButton, "CompletedTranscriptDetail") is not None
    window.window.close()
    app.processEvents()

def test_pyside_session_generates_offer_document_from_template(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("[Title] [First Name] [Last Name] | [City] | [Position] | [HourlyPay] | [Hours]")
    doc.save(template_path)

    output_path = session.generate_offer_document(
        template_path=template_path,
        output_dir=tmp_path / "offers",
        start_date=date(2026, 6, 23),
        start_time_12h="08:00 AM",
        end_time_12h="05:00 PM",
        hourly_pay=22.5,
        hours=40,
        created_on=date(2026, 6, 20),
        title="Ms.",
    )

    assert output_path.exists()
    assert output_path.name == "Launch Pad Learning PMD Offer of Employment to Latoya Nugent.docx"
    rendered = _docx_text(output_path)
    assert "Ms. Latoya Nugent | Palmdale | Preschool | 22.50 | 40" in rendered


def test_pyside_session_offer_generation_preserves_prior_offer_revision(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("[Title] [First Name] [Last Name]")
    document.save(template_path)
    kwargs = {
        "template_path": template_path,
        "output_dir": tmp_path / "offers",
        "start_date": date(2026, 6, 23),
        "start_time_12h": "08:00 AM",
        "end_time_12h": "05:00 PM",
        "hourly_pay": 22.5,
        "hours": 40,
        "created_on": date(2026, 6, 20),
        "title": "Ms.",
    }

    first = session.generate_offer_document(**kwargs)
    second = session.generate_offer_document(**kwargs)

    assert first.name == "Launch Pad Learning PMD Offer of Employment to Latoya Nugent.docx"
    assert second.name == "Launch Pad Learning PMD Offer of Employment to Latoya Nugent (2).docx"

def test_pyside_window_runs_due_notification_schedule_on_startup(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    runs = []

    class FakeNotificationStore:
        def ensure_default_rules(self):
            return None

        def get_or_create_rollout_date(self, today):
            return today

    class FakeNotifications:
        store = FakeNotificationStore()
        email_settings = SimpleNamespace(smtp_host="smtp.example.org", sender_email="sender@example.org")

        def activate_ready_system_rules(self):
            return 0

        def run_due_notifications(self):
            runs.append("ran")
            return []

    monkeypatch.setattr(
        pyside_interview_app,
        "notification_service_from_email_account_settings",
        lambda **_kwargs: FakeNotifications(),
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window._schedule_startup_notifications()
    app.processEvents()

    assert runs == ["ran"]
    window.window.close()
    app.processEvents()

def test_pyside_session_generates_interview_notes_document(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")

    history_path = tmp_path / "interview_history.sqlite3"

    output_path = session.generate_interview_notes_document(output_dir=tmp_path / "notes", history_path=history_path)

    assert output_path.exists()
    rendered = _docx_text(output_path)
    assert "Latoya Nugent" in rendered
    assert "Warm child-centered example." in rendered
    assert "Final Outcome" in rendered
    assert "Hire" in rendered
    assert history_path.exists()
    assert not (tmp_path / "notes" / "interview_history.sqlite3").exists()
    assert not (tmp_path / "notes" / "interview_history.json").exists()

def test_pyside_finalize_writes_basic_notes(tmp_path: Path) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")
    history_path = tmp_path / "interview_history.sqlite3"
    result = session.finalize_interview(base_dir=tmp_path, history_path=history_path)

    report_path = Path(result["out_path"])
    integration_path = Path(result["integration_path"])
    rows = InterviewHistoryStore(history_path).load()

    assert report_path.exists()
    assert integration_path.exists()
    assert "pyside_notes" not in str(report_path)
    rendered = _docx_text(report_path)
    assert "Structured Behavioral Interview Notes" in rendered
    assert "1. Candidate Snapshot" in rendered
    assert "2. Candidate Education and Experience Summary" in rendered
    assert "3. Score Summary" in rendered
    assert "4. Candidate Answers" in rendered
    assert "Warm child-centered example." in rendered
    assert "Weighted Total" in rendered
    assert "Final Outcome" in rendered
    assert "Consolidated Answer Summaries" not in rendered
    assert rows[0]["candidate_name"] == "Latoya Nugent"
    assert Path(rows[0]["interview_notes_path"]) == report_path

def test_pyside_finalize_writes_interview_notes_to_school_dropbox_folder(tmp_path: Path, monkeypatch) -> None:
    dropbox_root = tmp_path / "Dropbox"
    base_dir = dropbox_root / "App" / "interviews"
    settings_path = tmp_path / "school_offer_settings.json"
    settings_path.write_text(
        json.dumps(
            {
                "Palmdale": {
                    "interview_notes_dir": r"\Dropbox\LPL PMD Office Shared\Staff\Candidates",
                }
            }
        ),
        encoding="utf-8",
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Mission aligned.", score="")
    session.save_answer_and_advance(notes="Values aligned.", score="")
    session.save_answer_and_advance(notes="Warm child-centered example.", score="5")
    monkeypatch.setattr(pyside_interview_app, "SCHOOL_OFFER_SETTINGS_PATH", settings_path)

    result = session.finalize_interview(base_dir=base_dir, history_path=tmp_path / "interview_history.sqlite3")

    report_path = Path(result["out_path"])
    assert report_path.parent == dropbox_root / "LPL PMD Office Shared" / "Staff" / "Candidates"
    assert report_path.exists()

def test_pyside_finalize_excludes_intro_from_basic_notes_and_preserves_candidate_answers(
    tmp_path: Path, monkeypatch
) -> None:
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro notes.", score="")
    session.save_answer_and_advance(notes="Custom notes.", score="")
    session.save_answer_and_advance(notes="Manual fallback notes.", score="5")
    session.flow_recordings = {
        index: {
            "flow_index": index,
            "base_name": "pyside-recording",
            "transcript_jsonl": str(tmp_path / "transcript.jsonl"),
            "candidate_transcript": transcript,
        }
        for index, transcript in {
            0: "Candidate heard the intro.",
            1: "Candidate gave custom answer.",
            2: "Candidate described a warm child-centered example.",
        }.items()
    }
    session.flow_candidate_transcripts = {
        0: "Candidate heard the intro.",
        1: "Candidate gave custom answer.",
        2: "Candidate described a warm child-centered example.",
    }
    result = session.finalize_interview(base_dir=tmp_path, history_path=tmp_path / "interview_history.sqlite3")

    assert result["transcript_complete"] is True
    rendered = _docx_text(Path(result["out_path"]))
    assert "Candidate heard the intro." not in rendered
    assert "Candidate gave custom answer." in rendered
    assert "Candidate described a warm child-centered example." in rendered
    rows = InterviewHistoryStore(tmp_path / "interview_history.sqlite3").load()
    assert rows[0]["flow_recordings"] == list(session.flow_recordings.values())

def test_pyside_last_question_footer_finalizes_and_shows_complete_home(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    finalized: list[bool] = []

    def _fake_generate() -> None:
        finalized.append(True)
        window.review_status_label.setText("Interview finalized: fake.docx")

    monkeypatch.setattr(window, "_generate_interview_notes_from_session", _fake_generate)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()

    while window.session is not None and window.session.active_question() is not None:
        question = window.session.active_question()
        buttons = {
            button.property("pyside_live_footer_action"): button
            for button in window.window.findChildren(qt_widgets.QPushButton)
            if button.property("pyside_live_footer_action")
        }
        assert "exit" in buttons
        if question == window.session._workflow_items()[-1]:
            assert "finalize" in buttons
            if question.score_cards:
                window.score_group.buttons()[0].setChecked(True)
            buttons["finalize"].click()
            break
        assert "next" in buttons
        notes = window.interview_tabs.widget(2).findChild(qt_widgets.QTextEdit, "LiveInterviewerNotes")
        if notes is not None:
            notes.setPlainText(f"notes for {question.question_id}")
        if question.score_cards:
            window.score_group.buttons()[0].setChecked(True)
        buttons["next"].click()

    app.processEvents()

    assert finalized == [True]
    assert window.interview_tabs.currentIndex() == 3
    assert window.interview_tabs.currentWidget().findChild(qt_widgets.QLabel, "CompletedInterviewTitle").text() == "Interview Complete"
    assert window.interview_tabs.currentWidget().findChild(
        qt_widgets.QLabel, "CompletedInterviewStatus"
    ).text() == "Finalizing interview"
    finish = window.interview_tabs.currentWidget().findChild(
        qt_widgets.QPushButton, "CompletedInterviewFinish"
    )
    assert finish is not None
    assert not finish.isEnabled()
    window.window.close()
    app.processEvents()

def test_pyside_next_marks_new_question_at_click_boundary(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    ticks = iter([100.0, 105.0, 109.0])
    monkeypatch.setattr(pyside_interview_app.time, "monotonic", lambda: next(ticks))

    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()
    window.recording_started_monotonic = 100.0
    window._render_live_question_page()
    window._save_and_next()

    marks = window.session.flow_time_marks
    assert marks[0]["flow_index"] == 0
    assert marks[0]["end_t"] == 5.0
    assert marks[1]["flow_index"] == 1
    assert marks[1]["t"] == 5.0
    window.window.close()
    app.processEvents()

def test_pyside_review_table_shows_generated_transcripts(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    session.save_answer_and_advance(notes="Intro notes.")
    session.save_answer_and_advance(notes="Manual custom notes.")
    session.save_answer_and_advance(notes="Manual scored notes.", score="5")
    session.flow_candidate_transcripts[1] = "Generated custom transcript."
    session.flow_candidate_transcripts[2] = "Generated scored transcript."
    window.session = session

    window._render_review_page()

    review_page = window.interview_tabs.widget(3)
    transcript_text = " ".join(
        label.text()
        for card in review_page.findChildren(qt_widgets.QFrame, "CompletedTranscriptCard")
        for label in card.findChildren(qt_widgets.QLabel)
    )
    assert "Generated custom transcript." in transcript_text
    assert "Generated scored transcript." in transcript_text
    window.window.close()
    app.processEvents()

def test_pyside_keeps_recording_active_after_last_scored_question_navigation(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_workflow_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = _pyside_window_on_page(model, "Interviews")
    session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    session.start(candidate_name="Latoya Nugent", school="Palmdale", track_key="preschool")
    window.session = session
    window.recording_session = object()

    workflow = session._workflow_items()
    first_custom_index = next(index for index, item in enumerate(workflow) if item.kind == "custom")
    trait_index = next(index for index, item in enumerate(workflow) if item.kind == "trait")

    assert window._should_stop_recording_after_question(first_custom_index, workflow[first_custom_index]) is False
    assert window._should_stop_recording_after_question(trait_index, workflow[trait_index]) is False
    window.window.close()
    app.processEvents()

def test_pyside_last_scored_question_routes_to_review_before_transcription(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()

    while window.session is not None and window.session.active_question() is not None:
        question = window.session.active_question()
        notes = window.interview_tabs.widget(2).findChild(qt_widgets.QTextEdit, "LiveInterviewerNotes")
        if notes is not None:
            notes.setPlainText(f"notes for {question.question_id}")
        if question.score_cards:
            window.score_group.buttons()[0].setChecked(True)
        if question == window.session._workflow_items()[-1]:
            break
        window._save_and_next()

    stopped: list[bool] = []

    class BlockingRecordingSession:
        def stop_and_transcribe(self, **_kwargs):
            stopped.append(True)
            raise AssertionError("recording stop must run from finalize worker")

    window.recording_session = BlockingRecordingSession()
    monkeypatch.setattr(window, "_generate_interview_notes_from_session", lambda: None)
    window._save_and_next(finalize=True)

    assert stopped == []
    assert window.interview_tabs.currentIndex() == 3
    assert window.interview_tabs.currentWidget().findChild(qt_widgets.QLabel, "CompletedInterviewTitle").text() == "Interview Complete"
    assert window.recording_session is not None
    window.window.close()
    app.processEvents()

def test_pyside_finalize_returns_while_recording_transcription_finishes(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    monkeypatch.setattr(window, "_start_pyside_interview_recording", lambda: None)
    window.home_candidate_input.setText("Latoya Nugent")
    window._begin_selected_interview()
    while window.session.active_question() is not None:
        window.session.save_answer_and_advance(
            notes="Evidence.",
            score="5" if window.session.current_index == 2 else "",
        )
    window._render_review_page()
    started = threading.Event()
    release = threading.Event()

    class SlowRecordingSession:
        def stop_and_transcribe(self, **_kwargs):
            started.set()
            release.wait(timeout=2)
            raise RuntimeError("synthetic transcription delay")

    window.recording_session = SlowRecordingSession()
    finalized: list[bool] = []
    scheduled_closes: list[bool] = []
    monkeypatch.setattr(window.session, "finalize_interview", lambda **_kwargs: finalized.append(True) or {"out_path": "done.docx"})
    monkeypatch.setattr(window, "_schedule_close_pyside_finalize_progress", lambda: scheduled_closes.append(True))

    window._generate_interview_notes_from_session()

    assert started.wait(timeout=1)
    assert finalized == []
    assert "Finalizing interview" in window.review_status_label.text()
    release.set()
    for _ in range(50):
        app.processEvents()
        if "Interview finalized:" in window.review_status_label.text():
            break

    assert finalized == [True]
    assert "Interview finalized: done.docx" in window.review_status_label.text()
    assert scheduled_closes == [True]
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
def test_finalize_success_immediately_refreshes_candidate_history_scenario(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing-staffing-seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    assert window.hiring_v2_page.candidates_table.rowCount() == 0
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-alexandra",
            "candidate_name": "Newly Finalized Candidate",
            "school": "Palmdale",
            "position": "Teacher",
            "interview_date": "2026-07-16",
            "score": "88%",
            "status": "Hire",
        }
    )
    pyside_interview_app.HiringWorkflowService(pyside_interview_app.HiringPipelineStore(history_path)).record_initial_interview(
        history_id="hist-alexandra",
        legal_name="Newly Finalized Candidate",
        email="",
        phone="",
        school="Palmdale",
        position="Teacher",
        score=88.0,
        outcome="Hire",
    )
    messages: queue.Queue[dict[str, Any]] = queue.Queue()
    messages.put(
        {
            "ok": True,
            "artifact_update": True,
            "result": {"history_id": "hist-alexandra", "out_path": "done.docx", "scoring": {"outcome": "Hire"}},
        }
    )
    fake_timer = SimpleNamespace(stop=lambda: None, deleteLater=lambda: None)

    window._poll_pyside_finalize_worker(messages, fake_timer)
    app.processEvents()

    candidates = [
        window.hiring_v2_page.candidates_table.item(row, 0).text()
        for row in range(window.hiring_v2_page.candidates_table.rowCount())
    ]
    assert candidates == ["Newly Finalized Candidate"]
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
def test_finalize_success_immediately_refreshes_director_interview_candidates_scenario(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing-staffing-seed.json")
    monkeypatch.setattr(
        pyside_interview_app,
        "STAFFING_REFERRAL_QUEUE_PATH",
        tmp_path / "staffing_referrals.pending.jsonl",
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    table = window.staffing_v2_dashboard.director_interview_pending_table
    assert table.rowCount() == 0
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "hist-alexandra",
            "candidate_name": "Newly Finalized Candidate",
            "school": "Palmdale",
            "position": "Teacher",
            "interview_date": "2026-07-16",
            "score": "88%",
            "status": "Hire",
        }
    )
    window.session = SimpleNamespace(
        candidate_name="Newly Finalized Candidate",
        school="Palmdale",
        position="Teacher",
        interview_date="2026-07-16",
    )
    messages: queue.Queue[dict[str, Any]] = queue.Queue()
    messages.put(
        {
            "ok": True,
            "artifact_update": False,
            "result": {
                "history_id": "hist-alexandra",
                "out_path": "done.docx",
                "scoring": {"outcome": "Hire", "percent_of_max": 88.0},
            },
        }
    )
    fake_timer = SimpleNamespace(stop=lambda: None, deleteLater=lambda: None)
    monkeypatch.setattr(window, "_prompt_candidate_contact_handoff", lambda _result: None)
    monkeypatch.setattr(window, "_emit_pyside_rating_notification", lambda _result: None)
    monkeypatch.setattr(window, "_render_review_page", lambda: None)

    window._poll_pyside_finalize_worker(messages, fake_timer)
    app.processEvents()

    candidates = [table.item(row, 0).text() for row in range(table.rowCount())]
    assert candidates == ["Newly Finalized Candidate"]
    window.window.close()
    app.processEvents()


def test_pyside_finalize_uses_inline_status_without_progress_window(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window._show_pyside_finalize_progress("Stopping recording")
    assert window.pyside_finalize_progress_dialog is None
    assert window.pyside_finalize_progress_label is None
    window._report_pyside_finalize_progress("Building interview notes")
    assert window._pyside_finalize_progress_step == "Building interview notes"
    window.window.close()
    app.processEvents()


def test_candidate_contact_handoff_popup_threshold_excludes_exact_sixty_five() -> None:
    assert pyside_interview_app.should_prompt_candidate_contact_handoff(65.01) is True
    assert pyside_interview_app.should_prompt_candidate_contact_handoff("88%") is True
    assert pyside_interview_app.should_prompt_candidate_contact_handoff(65) is False


def test_director_hire_sync_creates_one_pending_offer_with_derived_terms(tmp_path: Path, monkeypatch) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-auto-offer",
            "candidate_name": "Jordan Lee",
            "candidate_email": "jordan@example.org",
            "school": "Palmdale",
            "position": "Teacher",
            "score": 88,
            "outcome": "Hire",
            "candidate": {
                "qualification": {
                    "has_degree": True,
                    "degree_type": "BA",
                    "degree_in_ece": True,
                    "years_experience": 6,
                }
            },
            "answers": {"Pay": {"notes": "$24.50 per hour"}},
        }
    )
    hiring = pyside_interview_app.HiringWorkflowService(
        pyside_interview_app.HiringPipelineStore(history_path)
    )
    application = hiring.record_initial_interview(
        history_id="hist-auto-offer",
        legal_name="Jordan Lee",
        honorific="Mr.",
        email="jordan@example.org",
        phone="",
        school="Palmdale",
        position="Teacher",
        score=88,
        outcome="Hire",
    )
    staffing_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    staffing_store = pyside_interview_app.StaffingStore(staffing_path)
    staffing_store.initialize()
    staffing = pyside_interview_app.StaffingService(staffing_store)
    referral = staffing.upsert_director_candidate_referral(
        history_id="hist-auto-offer",
        candidate_name="Jordan Lee",
        candidate_email="jordan@example.org",
        school="Palmdale",
        interviewer_outcome="hire",
    )
    staffing.record_director_interview(
        referral.id,
        director_name="Avery Director",
        completed_date="2026-07-16",
        rating=9,
        decision="hire",
        decision_notes="Hire.",
        proposed_shift_start="8:00 AM",
        proposed_shift_end="5:00 PM",
        proposed_classroom="Chef",
    )
    offer_settings = SchoolOfferSettingsStore(tmp_path / "offer-settings.json")
    offer_settings.save(
        {
            "Palmdale": {
                "full_time_template": str(tmp_path / "full-time.docx"),
                "part_time_template": str(tmp_path / "part-time.docx"),
                "offer_output_dir": str(tmp_path / "offers"),
            }
        }
    )
    window = object.__new__(pyside_interview_app.PySideInterviewWindow)
    window.model = SimpleNamespace(history_path=history_path)
    window.school_offer_store = offer_settings
    notifications = []

    class FakeNotifications:
        def emit_event(self, event_type, payload, idempotency_key):
            notifications.append((event_type, payload, idempotency_key))
            return []

    window.notification_service = FakeNotifications()

    window._sync_hiring_v2_director_decisions(hiring)
    window._sync_hiring_v2_director_decisions(hiring)

    offers = hiring.store.list_offer_versions(application.application_id)
    assert len(offers) == 1
    assert offers[0].status == "pending_approval"
    assert offers[0].terms["weekly_hours"] == "40"
    assert offers[0].terms["employment_type"] == "full_time"
    assert offers[0].terms["proposed_classroom"] == "Chef"
    assert offers[0].terms["hourly_pay"] == "24.5"
    assert offers[0].terms["honorific"] == "Mr."
    hire_payload = notifications[0][1]
    assert hire_payload["interview_score"] == "88"
    assert hire_payload["degree_display"] == "BA"
    assert hire_payload["degree_in_ece_display"] == "\nDegree in ECE: Yes"
    assert hire_payload["experience"] == "6"


def test_director_hire_sync_migrates_existing_staffing_db_before_contact_read(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    history_path = tmp_path / "interview_history.sqlite3"
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-legacy-staffing",
            "candidate_name": "Legacy Candidate",
            "candidate_email": "legacy@example.org",
            "school": "Palmdale",
            "position": "Teacher",
            "score": 86,
            "outcome": "Hire",
            "answers": {"Pay": {"notes": "$23.00 per hour"}},
        }
    )
    hiring = pyside_interview_app.HiringWorkflowService(
        pyside_interview_app.HiringPipelineStore(history_path)
    )
    application = hiring.record_initial_interview(
        history_id="hist-legacy-staffing",
        legal_name="Legacy Candidate",
        honorific="Ms.",
        email="legacy@example.org",
        phone="",
        school="Palmdale",
        position="Teacher",
        score=86,
        outcome="Hire",
    )
    staffing_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    staffing_store = pyside_interview_app.StaffingStore(staffing_path)
    staffing_store.initialize()
    staffing = pyside_interview_app.StaffingService(staffing_store)
    referral = staffing.upsert_director_candidate_referral(
        history_id="hist-legacy-staffing",
        candidate_name="Legacy Candidate",
        candidate_email="legacy@example.org",
        school="Palmdale",
        interviewer_outcome="hire",
    )
    staffing.record_director_interview(
        referral.id,
        director_name="Avery Director",
        completed_date="2026-07-16",
        rating=9,
        decision="hire",
        decision_notes="Hire.",
        proposed_shift_start="8:00 AM",
        proposed_shift_end="5:00 PM",
        proposed_classroom="Chef",
    )
    with sqlite3.connect(staffing_path) as conn:
        conn.execute("ALTER TABLE director_candidate_referrals DROP COLUMN candidate_phone")
    offer_settings = SchoolOfferSettingsStore(tmp_path / "offer-settings.json")
    offer_settings.save(
        {
            "Palmdale": {
                "full_time_template": str(tmp_path / "full-time.docx"),
                "part_time_template": str(tmp_path / "part-time.docx"),
                "offer_output_dir": str(tmp_path / "offers"),
            }
        }
    )
    window = object.__new__(pyside_interview_app.PySideInterviewWindow)
    window.model = SimpleNamespace(history_path=history_path)
    window.school_offer_store = offer_settings
    window.notification_service = SimpleNamespace(emit_event=lambda *_args, **_kwargs: [])

    window._sync_hiring_v2_director_decisions(hiring)

    offers = hiring.store.list_offer_versions(application.application_id)
    assert len(offers) == 1
    with sqlite3.connect(staffing_path) as conn:
        columns = {row[1] for row in conn.execute("PRAGMA table_info(director_candidate_referrals)").fetchall()}
    assert "candidate_phone" in columns


def test_pyside_review_table_updates_transcript_when_finalize_worker_reports_transcripts(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    window.session.start(candidate_name="Live Transcript", school="Palmdale", track_key="preschool")
    for item in window.session._workflow_items():
        window.session.answers[item.question_id] = {
            "kind": item.kind,
            "title": item.title,
            "prompt": item.prompt,
            "notes": "Saved note",
            "score": "3" if item.kind == "trait" else "",
            "quick_actions": [],
        }
    window.session.current_index = len(window.session._workflow_items())
    window._render_review_page()
    review_page = window.interview_tabs.widget(3)
    trait_card = next(
        card
        for card in review_page.findChildren(qt_widgets.QFrame, "CompletedTranscriptCard")
        if card.property("category") == "Scored"
    )
    question_id = str(trait_card.property("questionId"))
    trait_row = next(
        index for index, item in enumerate(window.session._workflow_items()) if item.question_id == question_id
    )
    assert "No candidate transcript captured" in _widget_text(trait_card)

    messages: queue.Queue[dict[str, Any]] = queue.Queue()
    window.session.flow_candidate_transcripts[trait_row] = "Candidate described a calm redirect."
    messages.put({"ok": True, "event": "transcripts_updated"})
    fake_timer = SimpleNamespace(stop=lambda: None, deleteLater=lambda: None)

    window._poll_pyside_finalize_worker(messages, fake_timer)
    app.processEvents()

    refreshed_page = window.interview_tabs.widget(3)
    refreshed_card = next(
        card
        for card in refreshed_page.findChildren(qt_widgets.QFrame, "CompletedTranscriptCard")
        if card.property("questionId") == question_id
    )
    assert "Candidate described a calm redirect." in _widget_text(refreshed_card)
    window.window.close()
    app.processEvents()


def _write_two_trait_rubric(tmp_path: Path) -> Path:
    rubric_path = _write_test_rubric(tmp_path)
    rubric = json.loads(rubric_path.read_text(encoding="utf-8"))
    rubric["tracks"]["preschool"]["max_weighted_total"] = 10
    rubric["traits"][0]["priority"] = "Standard"
    rubric["traits"].append(
        {
            "id": "trait_2",
            "name": "Reliability",
            "priority": "Standard",
            "weight": 1,
            "applicable_tracks": ["preschool"],
            "primary_question": "Tell me about showing up reliably.",
            "descriptors": {
                "1": "Serious concern",
                "2": "Weak",
                "3": "Mixed / acceptable",
                "4": "Strong",
                "5": "Excellent",
            },
            "sample_answers": {},
        }
    )
    rubric_path.write_text(json.dumps(rubric), encoding="utf-8")
    return rubric_path

def test_pyside_completed_review_has_no_inline_score_update_or_referral_side_effect(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    referral_queue_path = tmp_path / "staffing_referrals.pending.jsonl"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_REFERRAL_QUEUE_PATH", referral_queue_path)
    model = build_interview_redesign_model(
        rubric_path=_write_two_trait_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Palmdale"],
    )
    history = InterviewHistoryStore(history_path)
    history.append(
        {
            "history_id": "hist-review",
            "candidate_name": "Review Candidate",
            "school": "Palmdale",
            "track": "Preschool",
            "interview_date": date.today().isoformat(),
            "outcome": "No Hire",
            "score": "60.0%",
        }
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.session = PySideInterviewSession(model=model, draft_path=tmp_path / "draft.json")
    window.session.start(candidate_name="Review Candidate", school="Palmdale", track_key="preschool")
    for item in window.session._workflow_items():
        window.session.answers[item.question_id] = {
            "kind": item.kind,
            "title": item.title,
            "prompt": item.prompt,
            "notes": "Saved note",
            "score": "3" if item.kind == "trait" else "",
            "quick_actions": [],
        }
    window._review_history_id = "hist-review"
    window._render_review_page()

    table = window.window.findChild(qt_widgets.QTableWidget, "CompletedInterviewTraitTable")
    assert table is not None
    assert window.window.findChild(qt_widgets.QPushButton, "PySideReviewApplyScoresButton") is None
    assert pyside_interview_app._pop_staffing_referral_queue_for_school(
        "Palmdale", queue_path=referral_queue_path
    ) == []
    window.window.close()
    app.processEvents()


def test_pyside_review_source_exposes_finalize_button_not_placeholder_notes() -> None:
    source = Path("src/pyside_interview_app.py").read_text(encoding="utf-8")

    assert 'self._primary_button("Finalize Interview")' not in source
    assert 'output_dir=DEFAULT_BASE_DIR / "pyside_notes"' not in source


def test_pyside_onboarding_board_surfaces_next_required_task() -> None:
    employee = Employee(
        id="emp-1",
        name="Taylor Green",
        acceptance_date="2026-06-01",
        start_date="2026-06-10",
        school="Palmdale",
        tasks=[
            EmployeeTask(id="past", template_id="", title="Past task", due_date="2026-06-09"),
            EmployeeTask(id="today", template_id="", title="Today task", due_date="2026-06-10"),
            EmployeeTask(id="done", template_id="", title="Done task", due_date="2026-06-08", completed=True),
        ],
    )

    board = build_pyside_onboarding_board(
        employees=[employee],
        scheduler_settings={"critical_window_days": 2},
        today=date(2026, 6, 10),
    )

    assert board.overdue == 1
    assert board.due_today == 1
    assert board.next_task == "Taylor Green: Past task"
    assert board.rows[0]["employee"] == "Taylor Green"
    assert board.rows[0]["next_task"] == "Past task"

def test_pyside_candidate_board_groups_history_by_candidate() -> None:
    rows = [
        {
            "candidate_name": "Latoya Nugent",
            "school": "Palmdale",
            "track": "Preschool",
            "outcome": "Hire",
            "percent_of_max": 85.0,
            "next_action": "Generate Offer",
        },
        {
            "candidate_name": "Dalia Gaspar",
            "school": "Hawthorne",
            "track": "Preschool",
            "outcome": "No Hire",
            "percent_of_max": 60.0,
            "next_action": "Return Home",
        },
    ]

    board = build_pyside_candidate_board(rows)

    assert board.total_candidates == 2
    assert board.rows[0]["candidate"] == "Latoya Nugent"
    assert board.rows[0]["status"] == "Hire"
    assert board.rows[0]["next_action"] == "Generate Offer"

def test_pyside_candidate_board_recomputes_stale_history_status_from_score() -> None:
    rows = [
        {
            "candidate_name": "Tatiana",
            "school": "Palmdale",
            "track": "Preschool",
            "determination": "No Hire",
            "interview_score": 70.0,
            "offer_status": "not_generated",
        },
    ]

    board = build_pyside_candidate_board(rows)

    assert board.rows[0]["score"] == "70.0"
    assert board.rows[0]["status"] == "Borderline"
    assert board.history_rows[0].status == "Borderline"

def test_pyside_pages_do_not_force_content_wider_than_viewport(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    rubric_path = _write_test_rubric(tmp_path)
    overrides_path = _write_test_overrides(tmp_path)
    settings_path = tmp_path / "school_offer_settings.json"
    settings_path.write_text(json.dumps({}), encoding="utf-8")
    notification_rules_path = tmp_path / "notification_rules.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "DEFAULT_RUBRIC_PATH", rubric_path)
    monkeypatch.setattr(pyside_interview_app, "QUESTIONS_OVERRIDE_PATH", overrides_path)
    monkeypatch.setattr(pyside_interview_app, "SCHOOL_OFFER_SETTINGS_PATH", settings_path)
    model = build_interview_redesign_model(
        rubric_path=rubric_path,
        overrides_path=overrides_path,
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.window.show()
    window.window.resize(640, 480)
    app.processEvents()

    page_scrolls = window.window.findChildren(qt_widgets.QScrollArea, "PySidePageScrollArea")

    assert page_scrolls
    for scroll in page_scrolls:
        assert scroll.horizontalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAsNeeded
        assert scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAsNeeded
        assert scroll.widgetResizable() is True
        assert scroll.minimumWidth() == 0
        assert scroll.widget().minimumWidth() <= scroll.viewport().width()
        assert scroll.widget().width() <= scroll.viewport().width() + 2
        assert scroll.widget().sizePolicy().horizontalPolicy() == qt_widgets.QSizePolicy.Policy.Expanding
    window.window.close()
    app.processEvents()

@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_initial_window_fits_available_screen_after_display_scaling(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    available = app.primaryScreen().availableGeometry()

    assert window.window.width() <= max(640, available.width() - 40)
    assert window.window.height() <= max(480, available.height() - 40)
    assert window.window.maximumWidth() == 16777215
    assert window.window.maximumHeight() == 16777215
    assert window.window.isMaximized() is False
    legacy_label = "Switch to " + "Tk UI"
    assert all(button.text() != legacy_label for button in window.window.findChildren(qt_widgets.QPushButton))
    window.window.close()
    app.processEvents()


def test_pyside_window_uses_native_title_minimize_and_maximize_controls() -> None:
    class FakeWindowType:
        Window = 1
        WindowTitleHint = 2
        WindowSystemMenuHint = 4
        WindowMinimizeButtonHint = 8
        WindowMaximizeButtonHint = 16
        WindowCloseButtonHint = 32
        FramelessWindowHint = 64

    class FakeQt:
        WindowType = FakeWindowType

    class FakeQtCore:
        Qt = FakeQt

    flags = standard_window_control_flags(FakeQtCore)

    assert flags & FakeWindowType.WindowTitleHint
    assert flags & FakeWindowType.WindowSystemMenuHint
    assert flags & FakeWindowType.WindowMinimizeButtonHint
    assert flags & FakeWindowType.WindowMaximizeButtonHint
    assert flags & FakeWindowType.WindowCloseButtonHint
    assert not flags & FakeWindowType.FramelessWindowHint

def test_pyside_window_show_opens_main_window_maximized() -> None:
    window = pyside_interview_app.PySideInterviewWindow.__new__(pyside_interview_app.PySideInterviewWindow)
    calls: list[str] = []

    class FakeWindow:
        def showMaximized(self) -> None:
            calls.append("showMaximized")

    window.window = FakeWindow()
    window._fit_window_to_available_screen = lambda *, fill_available=False: calls.append(f"fit:{fill_available}")
    window._schedule_startup_notifications = lambda: calls.append("schedule_notifications")
    window._schedule_recording_interface_preload = lambda: calls.append("schedule_recording_preload")

    window.show()

    assert calls == ["fit:True", "showMaximized", "schedule_notifications", "schedule_recording_preload"]

def test_pyside_window_show_primes_window_to_available_screen_before_maximize() -> None:
    window_class = getattr(pyside_interview_app, "PySide" + "InterviewWindow")
    window = window_class.__new__(window_class)
    calls: list[tuple[str, int, int, int, int] | tuple[str]] = []

    class FakeRect:
        def __init__(self, x: int, y: int, width: int, height: int) -> None:
            self._x = x
            self._y = y
            self._width = width
            self._height = height

        def x(self) -> int:
            return self._x

        def y(self) -> int:
            return self._y

        def width(self) -> int:
            return self._width

        def height(self) -> int:
            return self._height

        def right(self) -> int:
            return self._x + self._width - 1

        def bottom(self) -> int:
            return self._y + self._height - 1

    class FakeScreen:
        def availableGeometry(self) -> FakeRect:
            return FakeRect(0, 0, 1920, 1040)

    class FakeWindow:
        def isMaximized(self) -> bool:
            return False

        def isFullScreen(self) -> bool:
            return False

        def screen(self) -> FakeScreen:
            return FakeScreen()

        def width(self) -> int:
            return 1180

        def height(self) -> int:
            return 760

        def resize(self, *_args) -> None:
            raise AssertionError("startup maximize should set the full available geometry")

        def geometry(self) -> FakeRect:
            return FakeRect(120, 80, 1180, 760)

        def move(self, *_args) -> None:
            raise AssertionError("startup maximize should set the full available geometry")

        def setGeometry(self, rect: FakeRect) -> None:
            calls.append(("setGeometry", rect.x(), rect.y(), rect.width(), rect.height()))

        def showMaximized(self) -> None:
            calls.append(("showMaximized",))

    window.window = FakeWindow()
    window._schedule_startup_notifications = lambda: None
    window._schedule_recording_interface_preload = lambda: None

    window.show()

    assert calls == [("setGeometry", 0, 0, 1920, 1040), ("showMaximized",)]

def test_pyside_window_schedules_startup_notifications_once_after_show() -> None:
    window = pyside_interview_app.PySideInterviewWindow.__new__(pyside_interview_app.PySideInterviewWindow)
    calls: list[str] = []

    class FakeSignal:
        def connect(self, _callback) -> None:
            return None

    class FakeTimer:
        def __init__(self, _parent=None) -> None:
            self.timeout = FakeSignal()

        @staticmethod
        def singleShot(delay_ms: int, callback) -> None:
            calls.append(f"timer:{delay_ms}")
            callback()

        def setInterval(self, _interval_ms: int) -> None:
            return None

        def start(self) -> None:
            return None

    class FakeQtCore:
        QTimer = FakeTimer

    window.QtCore = FakeQtCore
    window.window = object()
    window._startup_notifications_scheduled = False
    window._run_due_notifications_safely = lambda: calls.append("notifications")

    window._schedule_startup_notifications()
    window._schedule_startup_notifications()

    assert calls == ["timer:0", "notifications"]

def test_pyside_window_can_defer_secondary_pages_until_navigation(tmp_path: Path, monkeypatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    built: list[str] = []

    def page(name: str):
        def build(self):
            built.append(name)
            widget = qt_widgets.QWidget()
            widget.setObjectName(f"{name}Page")
            return widget

        return build

    monkeypatch.setattr(pyside_interview_app.PySideInterviewWindow, "_staffing_page", page("Staffing"))
    monkeypatch.setattr(pyside_interview_app.PySideInterviewWindow, "_staffing_v2_page", page("Staffing v2"))
    monkeypatch.setattr(pyside_interview_app.PySideInterviewWindow, "_onboarding_page", page("Onboarding"))
    monkeypatch.setattr(pyside_interview_app.PySideInterviewWindow, "_run_due_notifications_safely", lambda self: None)
    model = pyside_interview_app.InterviewRedesignModel(
        app_title="Test",
        navigation=["Staffing", "Staffing v2", "Onboarding"],
        setup_steps=[],
        school_options=[],
        track_labels={},
        readiness_checks=[],
        home=pyside_interview_app.HomeModel(
            primary_action="Start",
            continue_action="Continue",
            admin_visible_on_home=False,
            recent_interviews=[],
            history_rows=[],
        ),
        flows={},
        rubric={},
        history_path=tmp_path / "history.json",
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)

    assert built == ["Staffing v2"]
    assert window.stack.count() == 3
    assert window.stack.widget(1).objectName() == "Staffing v2Page"

    window._select_main_nav_row(2)
    window._select_main_nav_row(2)

    assert built == ["Staffing v2", "Onboarding"]
    assert window.stack.currentWidget().objectName() == "OnboardingPage"
    window.window.close()
    app.processEvents()

def test_pyside_window_fit_keeps_maximized_state_intact() -> None:
    window = pyside_interview_app.PySideInterviewWindow.__new__(pyside_interview_app.PySideInterviewWindow)

    class FakeWindow:
        def __init__(self) -> None:
            self.maximum_size_calls = 0
            self.resize_calls = 0
            self.move_calls = 0

        def isMaximized(self) -> bool:
            return True

        def isFullScreen(self) -> bool:
            return False

        def screen(self):
            raise AssertionError("maximized windows should not be clamped")

        def setMaximumSize(self, *_args) -> None:
            self.maximum_size_calls += 1

        def resize(self, *_args) -> None:
            self.resize_calls += 1

        def move(self, *_args) -> None:
            self.move_calls += 1

    fake_window = FakeWindow()
    window.window = fake_window

    window._fit_window_to_available_screen()

    assert fake_window.maximum_size_calls == 0
    assert fake_window.resize_calls == 0
    assert fake_window.move_calls == 0

def test_pyside_window_fit_does_not_cap_normal_window_maximum_size() -> None:
    window = pyside_interview_app.PySideInterviewWindow.__new__(pyside_interview_app.PySideInterviewWindow)

    class FakeRect:
        def __init__(self, width: int, height: int) -> None:
            self._width = width
            self._height = height

        def width(self) -> int:
            return self._width

        def height(self) -> int:
            return self._height

        def x(self) -> int:
            return 0

        def y(self) -> int:
            return 0

        def right(self) -> int:
            return self._width - 1

        def bottom(self) -> int:
            return self._height - 1

    class FakeScreen:
        def availableGeometry(self) -> FakeRect:
            return FakeRect(1200, 800)

    class FakeWindow:
        def __init__(self) -> None:
            self.maximum_size_calls = 0
            self.resize_calls = 0
            self.move_calls = 0
            self._geometry = FakeRect(1180, 760)

        def isMaximized(self) -> bool:
            return False

        def isFullScreen(self) -> bool:
            return False

        def screen(self) -> FakeScreen:
            return FakeScreen()

        def setMaximumSize(self, *_args) -> None:
            self.maximum_size_calls += 1

        def width(self) -> int:
            return 1180

        def height(self) -> int:
            return 760

        def resize(self, *_args) -> None:
            self.resize_calls += 1

        def geometry(self) -> FakeRect:
            return self._geometry

        def move(self, *_args) -> None:
            self.move_calls += 1

    fake_window = FakeWindow()
    window.window = fake_window

    window._fit_window_to_available_screen()

    assert fake_window.maximum_size_calls == 0
    assert fake_window.resize_calls == 0
    assert fake_window.move_calls == 1

def test_pyside_rating_notification_emits_for_hire_and_borderline_only(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    notifications = []

    class FakeNotifications:
        directory = SimpleNamespace(director_names={"palmdale": "Edith"})

        def emit_event(self, event_type, payload, idempotency_key):
            notifications.append((event_type, payload, idempotency_key))
            return []

    window.notification_service = FakeNotifications()
    window.session = SimpleNamespace(
        candidate_name="Jane Doe",
        school="Palmdale",
        position="Teacher",
        interview_date="2026-07-02",
        qualification=CandidateQualification(
            has_degree=True,
            degree_type="BA",
            degree_in_ece=False,
            ece_units_completed=18,
            years_experience=4,
        ),
    )

    window._emit_pyside_rating_notification({"scoring": {"outcome": "Hire", "percent_of_max_label": "85%"}, "history_id": "hist-1"})
    window._emit_pyside_rating_notification({"scoring": {"outcome": "Borderline", "percent_of_max": 70}, "history_id": "hist-2"})
    window._emit_pyside_rating_notification({"scoring": {"outcome": "No Hire", "percent_of_max": 50}, "history_id": "hist-3"})

    assert [event[0] for event in notifications] == ["interview.rating.qualified", "interview.rating.qualified"]
    assert notifications[0][1]["candidate_name"] == "Jane Doe"
    assert notifications[0][1]["director_name"] == "Edith"
    assert notifications[0][1]["score"] == "85%"
    assert notifications[0][1]["degree_type"] == "BA"
    assert notifications[0][1]["ece_units_completed"] == "18"
    assert notifications[0][1]["years_experience"] == "4"
    assert notifications[1][2] == "hist-2:interview.rating.qualified"
    window.window.close()
    app.processEvents()

def test_pyside_live_question_wraps_scores_inside_vertical_scroll_area(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    model = build_interview_redesign_model(
        rubric_path=_write_long_descriptor_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "missing-history.json",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)

    window.session_index = 1
    window._render_live_question_page()

    scroll_areas = window.interview_tabs.widget(2).findChildren(qt_widgets.QScrollArea)
    score_labels = window.interview_tabs.widget(2).findChildren(qt_widgets.QLabel, "ScoreOptionText")

    assert scroll_areas
    assert scroll_areas[0].widgetResizable()
    assert scroll_areas[0].horizontalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOff
    assert score_labels
    assert all(label.wordWrap() for label in score_labels)
    assert any("consent-oriented strategies" in label.text() for label in score_labels)
    window.window.close()
    app.processEvents()

def test_pyside_contract_documents_supported_desktop_surface() -> None:
    contract_text = Path("contracts/pyside_interview_app.contract.yaml").read_text(encoding="utf-8")
    required_categories = [
        "pyside_desktop_surface",
        "interview finalize/history",
        "onboarding",
        "Staffing v2 Settings",
        "Staffing v2 notification routing",
        "runtime_wrapper",
        "setup_and_run",
    ]

    for category in required_categories:
        assert category in contract_text


def _write_test_rubric(tmp_path: Path) -> Path:
    rubric_path = tmp_path / "rubric.json"
    rubric_path.write_text(
        json.dumps(
            {
                "metadata": {"version": "test"},
                "scoring": {},
                "tracks": {"preschool": {"label": "Preschool", "max_weighted_total": 5}},
                "absolute_disqualifiers": ["Unsafe handling"],
                "traits": [
                    {
                        "id": "trait_1",
                        "name": "Empathy",
                        "priority": "Critical",
                        "weight": 1,
                        "applicable_tracks": ["preschool"],
                        "primary_question": "Tell me about a hard child moment.",
                        "descriptors": {"1": "Concern", "2": "Weak", "3": "Mixed", "4": "Strong", "5": "Excellent"},
                        "sample_answers": {},
                    }
                ],
            }
        ),
        encoding="utf-8",
    )
    return rubric_path


def _write_long_descriptor_rubric(tmp_path: Path) -> Path:
    rubric_path = _write_test_rubric(tmp_path)
    payload = json.loads(rubric_path.read_text(encoding="utf-8"))
    payload["traits"][0]["descriptors"]["4"] = (
        "Acknowledges the child's discomfort and follows safety rules, uses respectful language "
        "but may not fully describe consent-oriented strategies, focuses on keeping the child "
        "safe while trying to respond calmly and professionally."
    )
    rubric_path.write_text(json.dumps(payload), encoding="utf-8")
    return rubric_path


def _docx_text(path: Path) -> str:
    doc = Document(path)
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)

def test_pyside_staffing_dashboard_imports_seed_and_shows_metrics(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "dont_need_now"},
                                    {"position_name": "Teacher 2", "position_type": "Teacher", "status": "need_now"},
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    labels = [label.text() for label in window.stack.currentWidget().findChildren(qt_widgets.QLabel)]
    table = window.window.findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")

    assert any("Open positions: 1" in text for text in labels)
    assert table is not None
    assert table.rowCount() == 2
    assert table.horizontalHeaderItem(2).text() == "Person"
    assert table.horizontalHeaderItem(5).text() == "Permit Status"
    assert table.horizontalHeaderItem(6).text() == "Details"
    assert table.horizontalHeaderItem(7).text() == "Action"
    assert table.item(1, 1).text() == "Tranquility"
    assert table.item(1, 2).text() == "OPEN POSITION"
    assert table.item(1, 3).text() == "need_now"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_dashboard_renders_parallel_main_dashboard_without_mutating_db(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"},
                                    {
                                        "position_name": "Teacher 2",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Imgard", "permit_status": "permit_in_process"},
                                    },
                                ],
                            },
                            {
                                "name": "Tranquility",
                                "program": "Preschool",
                                "licensed_capacity": 18,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "coming", "person": {"name": "James"}},
                                ],
                            },
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    before_count = len(store.list_assignments())
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    assert "Staffing" in nav_items
    assert "Staffing v2" in nav_items
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()

    assert window.sidebar_panel.isHidden()
    assert page.findChild(qt_widgets.QFrame, "StaffingV2Shell") is not None
    assert "QScrollBar:vertical" in page.styleSheet()
    assert "QScrollBar::handle:vertical" in page.styleSheet()
    staffing_sidebar = page.findChild(qt_widgets.QFrame, "StaffingV2Sidebar")
    assert staffing_sidebar is not None
    assert not staffing_sidebar.isHidden()
    sidebar_text = _widget_text(staffing_sidebar)
    assert "Admin Studio" not in sidebar_text
    assert "Launch Pad Learning" in sidebar_text
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2DashboardNavButton").text() == "Staffing Dashboard"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2DashboardNavButton").property("staffingV2ActiveNav") is True
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2HomeNavButton").text() == "Dashboard"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleNavButton").text() == "People"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryNavButton").text() == "Assignment History"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2AnalyticsNavButton").text() == "Analytics"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").text() == "Notifications"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2IntegrationsNavButton").text() == "Integrations"
    for object_name in (
        "StaffingV2HomeNavButton",
        "StaffingV2DashboardNavButton",
        "StaffingV2ClassroomsNavButton",
        "StaffingV2PeopleNavButton",
        "StaffingV2HistoryNavButton",
        "StaffingV2AnalyticsNavButton",
        "StaffingV2NotificationsNavButton",
        "StaffingV2ValidationNavButton",
        "StaffingV2IntegrationsNavButton",
        "StaffingV2SettingsNavButton",
    ):
        assert not page.findChild(qt_widgets.QPushButton, object_name).icon().isNull()
    for object_name in (
        "StaffingV2HomeNavButton",
        "StaffingV2AnalyticsNavButton",
        "StaffingV2IntegrationsNavButton",
    ):
        assert not page.findChild(qt_widgets.QPushButton, object_name).isEnabled()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2SettingsNavButton").isEnabled()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").isEnabled()
    assert page.findChild(qt_widgets.QFrame, "StaffingV2TopTabBar") is None
    header_top_row = page.findChild(qt_widgets.QFrame, "StaffingV2DashboardHeaderTopRow")
    summary_action_row = page.findChild(qt_widgets.QFrame, "StaffingV2DashboardSummaryActionRow")
    assert header_top_row is not None
    assert summary_action_row is not None
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2SchoolFilter").parent() is header_top_row
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ProgramFilter").parent() is header_top_row
    assert page.findChild(qt_widgets.QLineEdit, "StaffingV2Search").parent() is header_top_row
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").parent() is header_top_row
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ExportButton").parent() is summary_action_row
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ViewHistoryButton").parent() is summary_action_row
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ExportButton").text() == "Export"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ViewHistoryButton").text() == "View History"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").text() == "Add Position"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").minimumHeight() >= 40
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2ExportButton").icon().isNull()
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2ViewHistoryButton").icon().isNull()
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").icon().isNull()
    export_button = page.findChild(qt_widgets.QPushButton, "StaffingV2ExportButton")
    assert export_button.isEnabled()
    export_button.click()
    app.processEvents()
    export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2DashboardExportDialog")
    assert export_dialog is not None
    export_text = _widget_text(export_dialog)
    assert "Export Staffing Dashboard" in export_text
    assert "Schools 1" in export_text
    assert "School filter Hawthorne" in export_text
    assert "Program filter All Programs" in export_text
    assert "Permit issues 1" in export_text
    assert "Open positions 1" in export_text
    assert "Harmony 1" in export_text
    export_dialog.accept()
    app.processEvents()
    assert _icon_has_primary_blue(page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").icon())
    assert page.findChild(qt_widgets.QFrame, "StaffingV2Sidebar").minimumWidth() >= 240
    assert page.objectName() == "PySideStaffingV2Page"
    assert not window.staffing_v2_dashboard.settings_nav_button.isHidden()
    assert page.findChild(qt_widgets.QWidget, "StaffingV2SettingsPage") is None
    assert page.findChild(qt_widgets.QLabel, "StaffingV2PageTitle").text() == "Staffing Dashboard"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2SchoolFilter").currentText() == "Hawthorne"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ProgramFilter").currentText() == "All Programs"
    assert page.findChild(qt_widgets.QLineEdit, "StaffingV2Search").placeholderText() == "Search classrooms"
    assert page.findChild(qt_widgets.QLineEdit, "StaffingV2Search").actions()
    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2MetricCard"))
    assert "Schools: 1" in metric_text
    assert "Open positions: 1" in metric_text
    assert "Avg fill time:" in metric_text
    assert "Open > 7 days:" in metric_text
    assert "Validation: 3 issues" in metric_text
    assert "20639" not in metric_text
    summary_chips = page.findChildren(qt_widgets.QFrame, "StaffingV2MetricCard")
    assert all(card.parent() is summary_action_row for card in summary_chips)
    chip_variants = {chip.accessibleName(): chip.property("staffingV2SummaryVariant") for chip in summary_chips}
    assert chip_variants["Schools: 1"] == "info"
    assert chip_variants["Open positions: 1"] == "info"
    assert chip_variants["Open > 7 days: 0"] == "danger"
    assert chip_variants["Validation: 3 issues"] == "danger"
    for card in summary_chips:
        assert card.maximumHeight() <= 50
        icon = card.findChild(qt_widgets.QLabel, "StaffingV2SummaryIcon")
        assert icon is not None
        assert icon.pixmap() is not None
        assert not icon.pixmap().isNull()
    classroom_list = page.findChild(qt_widgets.QListWidget, "StaffingV2ClassroomList")
    detail_scroll = page.findChild(qt_widgets.QScrollArea, "StaffingV2DashboardDetailScroll")
    assert classroom_list.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    assert classroom_list.horizontalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOff
    assert detail_scroll is not None
    assert detail_scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    assert detail_scroll.horizontalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOff
    assert detail_scroll.widget().objectName() == "StaffingV2DashboardDetailContent"
    list_filter = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomListFilterButton")
    assert list_filter is not None
    assert list_filter.text() == ""
    assert list_filter.isEnabled()
    assert not list_filter.icon().isNull()
    assert list_filter.toolTip() == "Classroom filters"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomListFooter").text() == (
        "Showing 1-2 of 2 classrooms"
    )
    assert classroom_list.count() == 2
    assert "Harmony 1" in classroom_list.item(0).text()
    assert "Need 1" in classroom_list.item(0).text()
    assert classroom_list.item(0).sizeHint().height() >= 60
    first_row_widget = classroom_list.itemWidget(classroom_list.item(0))
    assert first_row_widget is not None
    assert first_row_widget.objectName() == "StaffingV2ClassroomListItem"
    assert first_row_widget.property("staffingV2StatusFill") == "need_now"
    assert first_row_widget.grab().toImage().pixelColor(4, first_row_widget.height() // 2).name().upper() == "#FEE2E2"
    assert first_row_widget.testAttribute(qt_core.Qt.WidgetAttribute.WA_TransparentForMouseEvents)
    assert first_row_widget.findChild(qt_widgets.QFrame, "StaffingV2ClassroomStatusDot") is not None
    assert first_row_widget.findChild(qt_widgets.QFrame, "StaffingV2ClassroomStatusDot").property("staffingV2Status") == "need_now"
    assert first_row_widget.findChild(qt_widgets.QLabel, "StaffingV2ClassroomItemTitle").text() == "Harmony 1"
    assert first_row_widget.findChild(qt_widgets.QLabel, "StaffingV2ClassroomItemTitle").testAttribute(
        qt_core.Qt.WidgetAttribute.WA_TransparentForMouseEvents
    )
    count_label = first_row_widget.findChild(qt_widgets.QLabel, "StaffingV2ClassroomItemCounts")
    assert count_label.text() == "Need 1 · Replace 0\nComing 0 · Filled 1\nDon't Need 0"
    assert count_label.wordWrap()
    assert first_row_widget.minimumHeight() >= 120
    assert count_label.minimumHeight() >= count_label.fontMetrics().lineSpacing() * 3
    assert classroom_list.item(0).sizeHint().height() >= 136
    assert first_row_widget.findChild(qt_widgets.QLabel, "StaffingV2ClassroomItemChevron").text() == ">"
    list_filter.click()
    app.processEvents()
    filter_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2DashboardClassroomFilterDrawer")
    assert filter_dialog is not None
    assert "Classroom Filters" in _widget_text(filter_dialog)
    status_filter = filter_dialog.findChild(qt_widgets.QComboBox, "StaffingV2DashboardClassroomStatusFilter")
    status_filter.setCurrentText("Coming")
    filter_dialog.findChild(qt_widgets.QPushButton, "StaffingV2FilterApplyButton").click()
    app.processEvents()
    assert classroom_list.count() == 1
    assert "Tranquility" in classroom_list.item(0).text()
    export_button.click()
    app.processEvents()
    filtered_export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2DashboardExportDialog")
    assert filtered_export_dialog is not None
    filtered_export_text = _widget_text(filtered_export_dialog)
    assert "Classroom status filter Coming" in filtered_export_text
    assert "Tranquility" in filtered_export_text
    assert "Harmony 1" not in filtered_export_text
    filtered_export_dialog.accept()
    app.processEvents()
    filter_dialog = None
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomListFilterButton").click()
    app.processEvents()
    filter_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2DashboardClassroomFilterDrawer")
    filter_dialog.findChild(qt_widgets.QPushButton, "StaffingV2FilterResetButton").click()
    filter_dialog.findChild(qt_widgets.QPushButton, "StaffingV2FilterApplyButton").click()
    app.processEvents()
    assert classroom_list.count() == 2
    classroom_list.setCurrentRow(0)
    app.processEvents()
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomTitle").text() == "Harmony 1"
    overview_cards = page.findChildren(qt_widgets.QFrame, "StaffingV2OverviewCard")
    assert len(overview_cards) == 6
    for overview_card in overview_cards:
        icon = overview_card.findChild(qt_widgets.QLabel, "StaffingV2CardIcon")
        assert icon is not None
        assert icon.pixmap() is not None
        assert not icon.pixmap().isNull()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert table.maximumHeight() <= 230
    assert table.columnCount() == 8
    assert [table.horizontalHeaderItem(column).text() for column in range(table.columnCount())] == [
        "",
        "Position",
        "Person",
        "Status",
        "Start Date",
        "Days Open",
        "Permit Status",
        "Next Action",
    ]
    assert table.verticalHeader().isHidden()
    assert table.horizontalHeader().sectionResizeMode(0) == qt_widgets.QHeaderView.ResizeMode.Fixed
    assert [table.columnWidth(column) for column in (3, 6, 7)] == [170, 205, 210]
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    table_widget_text = {
        table.cellWidget(row, column).text().strip()
        if hasattr(table.cellWidget(row, column), "text")
        else _widget_text(table.cellWidget(row, column)).strip()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.cellWidget(row, column) is not None
    }
    table_text |= table_widget_text
    assert {"Teacher 1", "OPEN POSITION", "Need Now", "Teacher 2", "Imgard", "Filled"} <= table_text
    assert "20639" not in table_text
    assert table.editTriggers() == qt_widgets.QAbstractItemView.EditTrigger.NoEditTriggers
    need_now_row = _staffing_row_for_position(table, "Teacher 1")
    filled_row = _staffing_row_for_position(table, "Teacher 2")
    assert table.item(need_now_row, 0).text() == "1"
    assert table.item(filled_row, 0).text() == "2"
    assert table.cellWidget(need_now_row, 3).objectName() == "StaffingV2NeedNowChip"
    assert table.cellWidget(filled_row, 3).objectName() == "StaffingV2FilledChip"
    assert table.cellWidget(need_now_row, 3).property("staffingV2StatusFill") == "need_now"
    assert table.cellWidget(filled_row, 3).property("staffingV2StatusFill") == "filled"
    assert table.cellWidget(need_now_row, 3).grab().toImage().pixelColor(4, 12).name().upper() == "#FEE2E2"
    assert table.cellWidget(filled_row, 3).grab().toImage().pixelColor(4, 12).name().upper() == "#DCFCE7"
    assert table.cellWidget(need_now_row, 6).objectName() == "StaffingV2NeutralChip"
    assert table.cellWidget(filled_row, 6).objectName() == "StaffingV2ComingChip"
    for chip in (
        table.cellWidget(need_now_row, 3),
        table.cellWidget(filled_row, 3),
        table.cellWidget(need_now_row, 6),
        table.cellWidget(filled_row, 6),
    ):
        assert chip.findChild(qt_widgets.QLabel, "StaffingV2ChipIcon") is not None
        assert chip.findChild(qt_widgets.QLabel, "StaffingV2ChipText") is not None
    assert table.item(need_now_row, 3) is None
    assert table.item(filled_row, 6) is None
    need_now_action = table.cellWidget(need_now_row, table.columnCount() - 1)
    filled_action = table.cellWidget(filled_row, table.columnCount() - 1)
    assert need_now_action.text() == "Mark Coming"
    assert need_now_action.minimumWidth() >= 198
    assert need_now_action.menu() is not None
    assert [action.text() for action in need_now_action.menu().actions()] == [
        "Mark Coming",
        "Mark Not Needed",
        "Delete Position",
        "View Details",
    ]
    assert filled_action.text() == "Manage Filled"
    assert filled_action.minimumWidth() >= 198
    assert filled_action.menu() is not None
    assert [action.text() for action in filled_action.menu().actions()] == ["Manage Filled", "Replace", "Update Permit", "View Details"]
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").text() == "Add Position"
    drop_zone = page.findChild(qt_widgets.QFrame, "StaffingV2AddPositionDropZone")
    assert drop_zone is not None
    drop_zone_button = drop_zone.findChild(qt_widgets.QPushButton, "StaffingV2DropZoneAddButton")
    assert drop_zone_button.text() == "Add Position"
    assert not drop_zone_button.icon().isNull()
    priority_chip = page.findChild(qt_widgets.QFrame, "StaffingV2PriorityChip")
    assert priority_chip is not None
    assert priority_chip.findChild(qt_widgets.QLabel, "StaffingV2PriorityChipIcon") is not None
    assert priority_chip.findChild(qt_widgets.QLabel, "StaffingV2PriorityChipText").text() == "Need Now"
    assert "Need Now" in _widget_text(page.findChild(qt_widgets.QFrame, "StaffingV2StatusKey"))
    assert len(store.list_assignments()) == before_count
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_dashboard_scrollbars_have_scrollable_range(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    classrooms = [
        {
            "name": f"Classroom {index:02d}",
            "program": "Preschool",
            "licensed_capacity": 18,
            "slots": [
                {
                    "position_name": "Teacher 1",
                    "position_type": "Teacher",
                    "status": "filled",
                    "person": {"name": f"Teacher {index}"},
                }
            ],
        }
        for index in range(1, 18)
    ]
    seed_path.write_text(json.dumps({"schools": [{"name": "Hawthorne", "classrooms": classrooms}]}), encoding="utf-8")
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.window.resize(1280, 700)
    window.window.show()
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    classroom_list = page.findChild(qt_widgets.QListWidget, "StaffingV2ClassroomList")
    detail_scroll = page.findChild(qt_widgets.QScrollArea, "StaffingV2DashboardDetailScroll")

    assert classroom_list.count() == 17
    assert classroom_list.verticalScrollBar().maximum() > 0
    classroom_list.verticalScrollBar().setValue(classroom_list.verticalScrollBar().maximum())
    assert classroom_list.verticalScrollBar().value() == classroom_list.verticalScrollBar().maximum()
    assert detail_scroll.verticalScrollBar().maximum() > 0
    detail_scroll.verticalScrollBar().setValue(detail_scroll.verticalScrollBar().maximum())
    assert detail_scroll.verticalScrollBar().value() == detail_scroll.verticalScrollBar().maximum()
    scroll_areas = page.findChildren(qt_widgets.QAbstractScrollArea)
    assert scroll_areas
    for scroll_area in scroll_areas:
        assert scroll_area.sizeAdjustPolicy() == qt_widgets.QAbstractScrollArea.SizeAdjustPolicy.AdjustIgnored
        assert scroll_area.verticalScrollBar().singleStep() == 24
        if isinstance(scroll_area, qt_widgets.QAbstractItemView):
            assert scroll_area.verticalScrollMode() == qt_widgets.QAbstractItemView.ScrollMode.ScrollPerPixel
            assert scroll_area.horizontalScrollMode() == qt_widgets.QAbstractItemView.ScrollMode.ScrollPerPixel

    classroom_list.verticalScrollBar().setValue(0)
    list_center = classroom_list.viewport().rect().center()
    classroom_wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(list_center),
        qt_core.QPointF(classroom_list.viewport().mapToGlobal(list_center)),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(classroom_list.viewport(), classroom_wheel)
    assert classroom_list.verticalScrollBar().value() > 0

    classroom_list.verticalScrollBar().setValue(0)
    routed_item_wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(page.mapFromGlobal(classroom_list.viewport().mapToGlobal(list_center))),
        qt_core.QPointF(classroom_list.viewport().mapToGlobal(list_center)),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(page, routed_item_wheel)
    assert classroom_list.verticalScrollBar().value() > 0

    classroom_list.verticalScrollBar().setValue(0)
    window_item_wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(window.window.mapFromGlobal(classroom_list.viewport().mapToGlobal(list_center))),
        qt_core.QPointF(classroom_list.viewport().mapToGlobal(list_center)),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(window.window, window_item_wheel)
    assert classroom_list.verticalScrollBar().value() > 0

    detail_scroll.verticalScrollBar().setValue(0)
    detail_target = detail_scroll.widget().findChildren(qt_widgets.QLabel)[-1]
    detail_center = detail_target.rect().center()
    detail_wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(detail_center),
        qt_core.QPointF(detail_target.mapToGlobal(detail_center)),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(detail_target, detail_wheel)
    assert detail_scroll.verticalScrollBar().value() > 0

    detail_scroll.verticalScrollBar().setValue(0)
    visible_detail_center = detail_scroll.viewport().rect().center()
    window_detail_wheel = qt_gui.QWheelEvent(
        qt_core.QPointF(window.window.mapFromGlobal(detail_scroll.viewport().mapToGlobal(visible_detail_center))),
        qt_core.QPointF(detail_scroll.viewport().mapToGlobal(visible_detail_center)),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )
    app.sendEvent(window.window, window_detail_wheel)
    assert detail_scroll.verticalScrollBar().value() > 0

    assignment_id = store.list_assignments()[0].id
    window.staffing_v2_dashboard._show_position_drawer(assignment_id)
    for index in range(24):
        window.staffing_v2_dashboard.drawer_layout.addWidget(qt_widgets.QLabel(f"Wheel relay row {index}"))
    window.staffing_v2_dashboard.drawer_panel.show_overlay()
    app.processEvents()
    drawer_scroll = page.findChild(qt_widgets.QScrollArea, "StaffingV2PositionDrawerScroll")
    drawer_scroll.widget().setMinimumHeight(1400)
    window.staffing_v2_dashboard.drawer_panel.show_overlay()
    app.processEvents()
    assert drawer_scroll.verticalScrollBar().maximum() > 0
    wheel_target = drawer_scroll.widget().findChildren(qt_widgets.QLabel)[-1]
    drawer_scroll.verticalScrollBar().setValue(0)
    center = wheel_target.rect().center()
    wheel_event = qt_gui.QWheelEvent(
        qt_core.QPointF(center),
        qt_core.QPointF(wheel_target.mapToGlobal(center)),
        qt_core.QPoint(0, 0),
        qt_core.QPoint(0, -120),
        qt_core.Qt.MouseButton.NoButton,
        qt_core.Qt.KeyboardModifier.NoModifier,
        qt_core.Qt.ScrollPhase.ScrollUpdate,
        False,
    )

    app.sendEvent(wheel_target, wheel_event)

    assert drawer_scroll.verticalScrollBar().value() > 0

    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_full_app_selection_resyncs_shared_page_scrollbars(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": f"Classroom {index:02d}",
                                "positions": [
                                    {
                                        "position_name": "Teacher 1",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": f"Teacher {index}"},
                                    }
                                ],
                            }
                            for index in range(1, 14)
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    staffing_index = nav_items.index("Staffing v2")
    window.sidebar.setCurrentRow(staffing_index)
    app.processEvents()
    dashboard = window.staffing_v2_dashboard
    assert isinstance(dashboard, StaffingDashboardV2Page)
    assert dashboard.school_filter == ""
    page = window.stack.currentWidget()
    detail_scroll = page.findChild(qt_widgets.QScrollArea, "StaffingV2DashboardDetailScroll")
    assert detail_scroll.verticalScrollBar().maximum() > 0

    sync_calls: list[str] = []
    original_sync = dashboard._sync_staffing_v2_scroll_ranges

    def sync_spy() -> None:
        sync_calls.append("sync")
        original_sync()

    dashboard._sync_staffing_v2_scroll_ranges = sync_spy
    window.sidebar.setCurrentRow(nav_items.index("Staffing"))
    app.processEvents()
    window.sidebar.setCurrentRow(staffing_index)
    app.processEvents()

    assert sync_calls == ["sync"]
    assert isinstance(window.staffing_v2_dashboard, StaffingDashboardV2Page)
    assert window.stack.currentWidget() is page
    window.window.close()
    app.processEvents()

def _open_staffing_v2_notifications_test_window(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    rules: list[NotificationRule],
):
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text('{"schools":[]}', encoding="utf-8")
    notification_rules_path = tmp_path / "notification_rules.sqlite3"
    store = NotificationStore(notification_rules_path)
    saved_rules = [store.save_rule(rule) for rule in rules]
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", notification_rules_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").click()
    app.processEvents()
    return app, qt_widgets, window, page, store, saved_rules


def _open_notification_rule_card(app, qt_widgets, page, label: str) -> None:
    rule_list = page.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList")
    item = next(item for item in (rule_list.item(index) for index in range(rule_list.count())) if label in item.text())
    rule_list.itemClicked.emit(item)
    app.processEvents()


def _click_message_box_choice(
    qt_core,
    qt_widgets,
    choice: str,
    observed: dict[str, object],
    *,
    expected_title: str | None = None,
) -> None:
    attempts = {"count": 0}

    def click_when_visible() -> None:
        attempts["count"] += 1
        message_box = next(
            (
                widget
                for widget in qt_widgets.QApplication.topLevelWidgets()
                if isinstance(widget, qt_widgets.QMessageBox) and widget.isVisible()
                and (expected_title is None or widget.windowTitle() == expected_title)
            ),
            None,
        )
        if message_box is None:
            if attempts["count"] < 100:
                qt_core.QTimer.singleShot(10, click_when_visible)
            return
        observed["title"] = message_box.windowTitle()
        observed["text"] = message_box.text()
        observed["choices"] = {button.text().replace("&", "") for button in message_box.buttons()}
        button = next(
            button for button in message_box.buttons() if button.text().replace("&", "") == choice
        )
        button.click()

    qt_core.QTimer.singleShot(0, click_when_visible)


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
@pytest.mark.parametrize("choice", ["Save", "Discard", "Keep Editing"])
def test_pyside_staffing_v2_notifications_dirty_exit_choices_use_real_dialog(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, choice: str
) -> None:
    qt_core = pytest.importorskip("PySide6.QtCore")
    original_label = "Hiring manager: position needed now"
    edited_label = f"Edited via {choice}"
    app, qt_widgets, window, page, store, saved_rules = _open_staffing_v2_notifications_test_window(
        tmp_path,
        monkeypatch,
        [
            NotificationRule(
                event_type="staffing.assignment.need_now",
                label=original_label,
                subject_template="Position needed now: {position_name}",
                body_template="Please review {position_name}.",
                recipients=[NotificationRecipient(email="director@example.com", role_label="Director")],
                active=True,
            )
        ],
    )
    saved = saved_rules[0]
    _open_notification_rule_card(app, qt_widgets, page, original_label)
    editor = page.findChild(qt_widgets.QFrame, "StaffingV2NotificationEditor")
    label_editor = page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationRuleLabel")
    label_editor.setText(edited_label)
    observed: dict[str, object] = {}
    _click_message_box_choice(qt_core, qt_widgets, choice, observed)

    page.findChild(qt_widgets.QPushButton, "StaffingV2DashboardNavButton").click()
    app.processEvents()

    assert observed == {
        "title": "Unsaved Notification Changes",
        "text": "Save changes before leaving this notification rule?",
        "choices": {"Save", "Discard", "Keep Editing"},
    }
    persisted_label = store.get_rule(saved.id or 0).label
    if choice == "Keep Editing":
        assert page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").property(
            "staffingV2ActiveNav"
        ) is True
        assert not editor.isHidden()
        assert label_editor.text() == edited_label
        assert persisted_label == original_label
        label_editor.setText(original_label)
    else:
        assert page.findChild(qt_widgets.QPushButton, "StaffingV2DashboardNavButton").property(
            "staffingV2ActiveNav"
        ) is True
        assert editor.isHidden()
        assert persisted_label == (edited_label if choice == "Save" else original_label)

    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
@pytest.mark.parametrize("choice", ["No", "Yes"])
def test_pyside_staffing_v2_notifications_named_delete_confirmation_controls_deletion(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch, choice: str
) -> None:
    qt_core = pytest.importorskip("PySide6.QtCore")
    rule_label = "Leadership: offer accepted"
    app, qt_widgets, window, page, store, saved_rules = _open_staffing_v2_notifications_test_window(
        tmp_path,
        monkeypatch,
        [
            NotificationRule(
                event_type="offer.accepted",
                label=rule_label,
                subject_template="Offer accepted: {candidate_name}",
                body_template="{candidate_name} accepted the offer.",
                recipients=[NotificationRecipient(email="director@example.com", role_label="Director")],
                active=False,
            )
        ],
    )
    saved = saved_rules[0]
    _open_notification_rule_card(app, qt_widgets, page, rule_label)
    editor = page.findChild(qt_widgets.QFrame, "StaffingV2NotificationEditor")
    delete_button = page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationDelete")

    observed: dict[str, object] = {}
    _click_message_box_choice(
        qt_core, qt_widgets, choice, observed, expected_title="Delete Notification Rule"
    )
    delete_button.click()
    app.processEvents()

    assert observed["title"] == "Delete Notification Rule"
    assert rule_label in str(observed["text"])
    assert "cannot be undone" in str(observed["text"]).lower()
    assert {"Yes", "No"}.issubset(observed["choices"])
    if choice == "No":
        assert store.get_rule(saved.id or 0).label == rule_label
        assert not editor.isHidden()
    else:
        with pytest.raises(ValueError, match="not found"):
            store.get_rule(saved.id or 0)
        rule_list = page.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList")
        assert rule_label not in "\n".join(rule_list.item(index).text() for index in range(rule_list.count()))
        assert editor.isHidden()
        assert page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsStatus").text() == (
            "Notification rule deleted."
        )

    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_notifications_grid_collapses_and_restores_actual_columns(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    qt_test = pytest.importorskip("PySide6.QtTest")
    rules = [
        NotificationRule(
            event_type=f"custom.grid.{index}",
            label=f"Grid rule {index}",
            subject_template=f"Grid subject {index}",
            body_template=f"Grid body {index}",
            recipients=[NotificationRecipient(email="director@example.com", role_label="Director")],
            active=False,
        )
        for index in range(4)
    ]
    app, qt_widgets, window, page, _store, _saved_rules = _open_staffing_v2_notifications_test_window(
        tmp_path, monkeypatch, rules
    )
    window.window.show()
    app.processEvents()
    list_widget = page.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList")
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsViewToggle").setCurrentText("Grid")

    def first_two_rectangles(width: int):
        list_widget.setFixedWidth(width)
        app.processEvents()
        qt_test.QTest.qWait(30)
        list_widget.doItemsLayout()
        app.processEvents()
        return list_widget.visualItemRect(list_widget.item(0)), list_widget.visualItemRect(list_widget.item(1))

    wide_first, wide_second = first_two_rectangles(1000)
    assert wide_first.top() == wide_second.top()
    assert wide_first.left() != wide_second.left()

    narrow_first, narrow_second = first_two_rectangles(600)
    assert narrow_first.left() == narrow_second.left()
    assert narrow_second.top() > narrow_first.top()

    restored_first, restored_second = first_two_rectangles(1000)
    assert restored_first.top() == restored_second.top()
    assert restored_first.left() != restored_second.left()

    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_notifications_manual_payload_uses_attachment_picker_and_unsaved_draft(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    qt_core = pytest.importorskip("PySide6.QtCore")
    pdf_path = tmp_path / "Jordan Lee Offer.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n%%EOF\n")
    send_calls: list[dict[str, object]] = []

    class FakeNotificationService:
        def send_test_preview(
            self,
            rule: NotificationRule,
            payload: dict[str, str],
            recipient_email: str,
            idempotency_key: str,
        ) -> NotificationSendResult:
            send_calls.append(
                {
                    "rule": rule,
                    "payload": dict(payload),
                    "recipient": recipient_email,
                    "idempotency_key": idempotency_key,
                }
            )
            return NotificationSendResult(
                event_type=f"{rule.event_type}.test",
                rule_id=rule.id,
                status="sent",
                recipient_count=1,
            )

    monkeypatch.setattr(
        pyside_interview_app,
        "notification_service_from_email_account_settings",
        lambda **_kwargs: FakeNotificationService(),
    )
    app, qt_widgets, window, page, store, saved_rules = _open_staffing_v2_notifications_test_window(
        tmp_path,
        monkeypatch,
        [
            NotificationRule(
                event_type="offer.approved",
                label="Offer approved",
                subject_template="Offer approved: {candidate_name}",
                body_template="{candidate_name} for {position_name} at {school}.",
                recipients=[NotificationRecipient(email="director@example.com", role_label="Director")],
                active=False,
            )
        ],
    )
    saved = saved_rules[0]
    _open_notification_rule_card(app, qt_widgets, page, saved.label)
    draft_subject = "Draft offer for {candidate_name}"
    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationSubject").setText(draft_subject)
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationDeliveryToggle").click()
    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationTestRecipient").setText(
        "qa-recipient@example.org"
    )
    payload_selector = page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationTestPayload")
    assert payload_selector.currentText() == "Manual payload…"
    monkeypatch.setattr(
        qt_widgets.QFileDialog,
        "getOpenFileName",
        staticmethod(lambda *_args, **_kwargs: (str(pdf_path), "PDF Files (*.pdf)")),
    )
    observed_dialog: dict[str, object] = {}
    attempts = {"count": 0}

    def complete_manual_payload() -> None:
        attempts["count"] += 1
        dialog = next(
            (
                widget
                for widget in qt_widgets.QApplication.topLevelWidgets()
                if isinstance(widget, qt_widgets.QDialog)
                and widget.objectName() == "StaffingV2NotificationManualPayloadDialog"
                and widget.isVisible()
            ),
            None,
        )
        if dialog is None:
            if attempts["count"] < 100:
                qt_core.QTimer.singleShot(10, complete_manual_payload)
            return
        editors = {
            editor.objectName(): editor
            for editor in dialog.findChildren(qt_widgets.QLineEdit)
            if editor.objectName().startswith("StaffingV2NotificationManualPayload_")
        }
        observed_dialog["title"] = dialog.windowTitle()
        observed_dialog["fields"] = set(editors)
        required_fields = {
            "StaffingV2NotificationManualPayload_candidate_name",
            "StaffingV2NotificationManualPayload_position_name",
            "StaffingV2NotificationManualPayload_school",
            "StaffingV2NotificationManualPayload_offer_pdf_path",
        }
        if not required_fields.issubset(editors):
            dialog.reject()
            return
        editors["StaffingV2NotificationManualPayload_candidate_name"].setText("Jordan Lee")
        editors["StaffingV2NotificationManualPayload_position_name"].setText("Lead Teacher")
        editors["StaffingV2NotificationManualPayload_school"].setText("Hawthorne")
        browse = next(button for button in dialog.findChildren(qt_widgets.QPushButton) if button.text() == "Browse")
        browse.click()
        observed_dialog["attachment"] = editors[
            "StaffingV2NotificationManualPayload_offer_pdf_path"
        ].text()
        button_box = dialog.findChild(qt_widgets.QDialogButtonBox)
        button_box.button(qt_widgets.QDialogButtonBox.StandardButton.Ok).click()

    qt_core.QTimer.singleShot(0, complete_manual_payload)
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationSendTest").click()
    status = page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsStatus")
    deadline = time.monotonic() + 5
    while time.monotonic() < deadline and status.text() != "Test send sent.":
        app.processEvents()
        time.sleep(0.01)

    assert observed_dialog["title"] == "Manual Test Payload"
    assert {
        "StaffingV2NotificationManualPayload_candidate_name",
        "StaffingV2NotificationManualPayload_position_name",
        "StaffingV2NotificationManualPayload_school",
        "StaffingV2NotificationManualPayload_offer_pdf_path",
    }.issubset(observed_dialog["fields"])
    assert observed_dialog["attachment"] == str(pdf_path)
    assert len(send_calls) == 1
    call = send_calls[0]
    assert call["recipient"] == "qa-recipient@example.org"
    assert call["rule"].subject_template == draft_subject
    assert call["payload"] == {
        "candidate_name": "Jordan Lee",
        "position_name": "Lead Teacher",
        "school": "Hawthorne",
        "offer_pdf_path": str(pdf_path),
    }
    assert store.get_rule(saved.id or 0).subject_template == saved.subject_template
    assert status.text() == "Test send sent."

    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationSubject").setText(saved.subject_template)
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_notifications_manager_dashboard_scenario(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    notification_rules_path = tmp_path / "notification_rules.sqlite3"
    saved = NotificationStore(notification_rules_path).save_rule(
        NotificationRule(
            event_type="staffing.assignment.need_now",
            label="Hiring manager: position needed now",
            subject_template="Position needed now: {position_name}",
            body_template="Hi {hiring_manager_name},\n\nPlease review {position_name}.",
            recipients=[
                NotificationRecipient(
                    email="director@example.com",
                    name="Director",
                    role_label="Director",
                )
            ],
            active=False,
        )
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", notification_rules_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    page = window.stack.currentWidget()
    notifications_nav = page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton")

    assert notifications_nav.isEnabled()
    notifications_nav.click()
    app.processEvents()

    assert page.findChild(qt_widgets.QWidget, "StaffingV2NotificationsDashboard") is not None
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").property("staffingV2ActiveNav") is True
    assert page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsTitle").text() == "Notifications"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsRuleCount").text().endswith(" rules")
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEventFilter").currentText() == "All events"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEnabledFilter").currentText() == "All statuses"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsCreateButton").text() == "Create Rule"
    list_widget = page.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList")
    assert list_widget.count() >= 1
    editor = page.findChild(qt_widgets.QFrame, "StaffingV2NotificationEditor")
    assert editor.isHidden()
    list_widget.itemClicked.emit(list_widget.item(0))
    app.processEvents()
    assert not editor.isHidden()
    page.resize(1000, 700)
    window.staffing_v2_dashboard.notification_editor_overlay.reposition()
    app.processEvents()
    footer = page.findChild(qt_widgets.QWidget, "StaffingV2NotificationEditorFooter")
    drawer_scroll = page.findChild(qt_widgets.QScrollArea, "StaffingV2NotificationEditorScroll")
    assert editor.width() <= page.width()
    assert footer.geometry().bottom() <= editor.rect().bottom()
    assert drawer_scroll.geometry().bottom() < footer.geometry().top()
    list_text = "\n".join(list_widget.item(index).text() for index in range(list_widget.count()))
    assert "Hiring manager: position needed now" in list_text
    assert "staffing.assignment.need_now" in list_text

    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEventFilter").setCurrentText(saved.event_type)
    app.processEvents()
    assert page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationRuleLabel").text() == saved.label
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationRuleEvent").currentText() == saved.event_type
    assert page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationSubject").text() == saved.subject_template
    assert page.findChild(qt_widgets.QPlainTextEdit, "StaffingV2NotificationBody").toPlainText() == saved.body_template
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsRecipientsFilter").currentText() == "All recipients"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsTemplateFilter").currentText() == "All templates"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsSort").currentText() == "Event sort"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsViewToggle").currentText() == "List"
    assert "Director <director@example.com>" in _widget_text(
        page.findChild(qt_widgets.QFrame, "StaffingV2NotificationRecipientChips")
    )
    assert "{position_name}" in page.findChild(qt_widgets.QLabel, "StaffingV2NotificationVariablesPreview").text()
    notice_button = page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationVariable_notice_date")
    final_day_button = page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationVariable_final_day")
    assert notice_button is not None
    assert final_day_button is not None
    subject_editor = page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationSubject")
    subject_editor.setFocus()
    subject_editor.setCursorPosition(0)
    subject_editor.setCursorPosition(len(subject_editor.text()))
    notice_button.click()
    assert "{notice_date}" in subject_editor.text()
    body_editor = page.findChild(qt_widgets.QPlainTextEdit, "StaffingV2NotificationBody")
    assert "No issues found" in page.findChild(qt_widgets.QLabel, "StaffingV2NotificationValidation").text()
    audit_panel = page.findChild(qt_widgets.QFrame, "StaffingV2NotificationAuditPanel")
    assert audit_panel.isHidden()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationDeliveryToggle").click()
    app.processEvents()
    assert not audit_panel.isHidden()
    assert "Pending scheduled:" in page.findChild(qt_widgets.QLabel, "StaffingV2NotificationAuditSummary").text()
    subject_editor.setText(saved.subject_template)
    body_editor.setPlainText(saved.body_template)
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationCancel").click()
    app.processEvents()
    assert editor.isHidden()

    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_notifications_filters_chips_preview_and_test_send(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text('{"schools":[]}', encoding="utf-8")
    notification_rules_path = tmp_path / "notification_rules.sqlite3"
    store = NotificationStore(notification_rules_path)
    saved = store.save_rule(
        NotificationRule(
            event_type="staffing.assignment.need_now",
            label="Need now",
            subject_template="Need now: {position_name}",
            body_template="Please review {position_name} at {school}.",
            recipients=[NotificationRecipient(email="director@example.com", role_label="Director")],
            active=False,
        )
    )
    missing = store.save_rule(
        NotificationRule(
            event_type="custom.missing",
            label="Missing template",
            subject_template="",
            body_template="",
            recipients=[],
            active=False,
        )
    )
    store.set_rule_active(missing.id or 0, True)

    class FakeNotificationService:
        def send_test_preview(
            self,
            rule: NotificationRule,
            payload: dict[str, str],
            recipient_email: str,
            idempotency_key: str,
        ) -> NotificationSendResult:
            NotificationStore(notification_rules_path).record_send_attempt(
                event_type="staffing.assignment.need_now.test",
                rule_id=rule.id,
                idempotency_key=idempotency_key,
                recipient_count=1,
                status="sent",
            )
            return NotificationSendResult(
                event_type="staffing.assignment.need_now.test",
                rule_id=rule.id,
                status="sent",
                recipient_count=1,
            )

    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", notification_rules_path)
    monkeypatch.setattr(
        pyside_interview_app,
        "notification_service_from_email_account_settings",
        lambda **_kwargs: FakeNotificationService(),
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").click()
    app.processEvents()

    list_widget = page.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList")
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEnabledFilter").setCurrentText("Disabled")
    app.processEvents()
    assert "Need now" in "\n".join(list_widget.item(index).text() for index in range(list_widget.count()))
    assert "Missing template" not in "\n".join(list_widget.item(index).text() for index in range(list_widget.count()))
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsTemplateFilter").setCurrentText("Missing subject")
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEnabledFilter").setCurrentText("All statuses")
    app.processEvents()
    assert list_widget.count() == 1
    assert "Missing template" in list_widget.item(0).text()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsClearFilters").click()
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsSort").setCurrentText("Recipients sort")
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsViewToggle").setCurrentText("Grid")
    app.processEvents()
    assert list_widget.property("staffingV2NotificationViewMode") == "grid"

    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEventFilter").setCurrentText(saved.event_type)
    app.processEvents()
    recipients = page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationRecipients")
    recipients.setText("HR <hr@example.com>")
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationRecipientAdd").click()
    app.processEvents()
    chips_text = _widget_text(page.findChild(qt_widgets.QFrame, "StaffingV2NotificationRecipientChips"))
    assert "HR <hr@example.com>" in chips_text
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationRecipientRemove_hr_example_com").click()
    app.processEvents()
    assert "HR <hr@example.com>" not in _widget_text(page.findChild(qt_widgets.QFrame, "StaffingV2NotificationRecipientChips"))
    recipient_picker = page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationRecipientPicker")
    recipient_picker.setCurrentText("Hiring Manager")
    recipient_picker.setCurrentText("Director")
    app.processEvents()
    chips_text = _widget_text(page.findChild(qt_widgets.QFrame, "StaffingV2NotificationRecipientChips"))
    assert "Hiring Manager <recruiting@launchpadpreschool.com>" in chips_text
    assert "Director (school-based)" in chips_text

    subject = page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationSubject")
    subject.setCursorPosition(len(subject.text()))
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationSubjectVariable_position_name").click()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationBodyBold").click()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationVariable_school").click()
    assert "{position_name}" in subject.text()
    assert "{school}" in page.findChild(qt_widgets.QPlainTextEdit, "StaffingV2NotificationBody").toPlainText()

    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationPreview").click()
    app.processEvents()
    preview = page.findChild(qt_widgets.QDialog, "StaffingV2NotificationPreviewDialog")
    assert preview is not None
    assert "Teacher 1" in _widget_text(preview)
    preview.close()

    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationTestRecipient").setText("tester@example.org")
    window.staffing_v2_dashboard._manual_notification_test_payload = lambda: {
        "position_name": "Teacher 1",
        "school": "Hawthorne",
    }
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationSendTest").click()
    deadline = time.monotonic() + 2
    while time.monotonic() < deadline:
        app.processEvents()
        if page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsStatus").text() == "Test send sent.":
            break
        time.sleep(0.01)
    assert page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsStatus").text() == "Test send sent."
    assert "sent" in page.findChild(qt_widgets.QLabel, "StaffingV2NotificationAuditSummary").text()
    assert NotificationStore(notification_rules_path).get_rule(saved.id).active is False
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationSave").click()
    app.processEvents()
    reloaded = NotificationStore(notification_rules_path).get_rule(saved.id)
    assert [(recipient.recipient_type, recipient.role_key) for recipient in reloaded.recipients if recipient.recipient_type == "role"] == [
        ("role", "hiring_manager"),
        ("role", "director"),
    ]

    recipients.setText("not-an-email")
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationSave").click()
    app.processEvents()
    assert "Invalid recipient email" in page.findChild(qt_widgets.QLabel, "StaffingV2NotificationValidation").text()

    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_notifications_editor_saves_rule_changes(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text('{"schools":[]}', encoding="utf-8")
    notification_rules_path = tmp_path / "notification_rules.sqlite3"
    NotificationStore(notification_rules_path).save_rule(
        NotificationRule(
            event_type="staffing.assignment.need_now",
            label="Hiring manager: position needed now",
            subject_template="Position needed now: {position_name}",
            body_template="Please review {position_name}.",
            recipients=[NotificationRecipient(email="director@example.com", name="Director", role_label="Director")],
            active=True,
        )
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "NOTIFICATION_RULES_PATH", notification_rules_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationsNavButton").click()
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationsEventFilter").setCurrentText("staffing.assignment.need_now")
    app.processEvents()

    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationRuleLabel").setText("Director: permit updated")
    event = page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationRuleEvent")
    event.setCurrentText("staffing.permit.updated")
    page.findChild(qt_widgets.QCheckBox, "StaffingV2NotificationEnabled").setChecked(False)
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationTiming").setCurrentText("Reference date")
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationDateField").setCurrentText("start_date")
    page.findChild(qt_widgets.QComboBox, "StaffingV2NotificationOffsetDirection").setCurrentText("After")
    page.findChild(qt_widgets.QSpinBox, "StaffingV2NotificationOffsetDays").setValue(2)
    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationSubject").setText("Permit updated: {person_name}")
    page.findChild(qt_widgets.QPlainTextEdit, "StaffingV2NotificationBody").setPlainText("Permit: {permit_status}")
    page.findChild(qt_widgets.QLineEdit, "StaffingV2NotificationRecipients").setText(
        "Director <director@example.com>, HR <hr@example.com>"
    )
    page.findChild(qt_widgets.QPushButton, "StaffingV2NotificationSave").click()
    app.processEvents()

    rule = next(
        rule
        for rule in NotificationStore(notification_rules_path).list_rules("staffing.permit.updated")
        if rule.label == "Director: permit updated"
    )
    assert rule.label == "Director: permit updated"
    assert rule.active is False
    assert rule.trigger_timing == "date_offset"
    assert rule.date_field == "start_date"
    assert rule.offset_days == 2
    assert rule.subject_template == "Permit updated: {person_name}"
    assert [recipient.email for recipient in rule.recipients] == ["director@example.com", "hr@example.com"]
    assert page.findChild(qt_widgets.QLabel, "StaffingV2NotificationsStatus").text() == "Notification rule saved."

    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_classrooms_dashboard_uses_new_shell_and_db_rows(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"},
                                    {
                                        "position_name": "Teacher 2",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Maria Gonzalez", "permit_status": "teacher_permit_approved"},
                                    },
                                ],
                            },
                            {
                                "name": "Tranquility",
                                "program": "Infant",
                                "licensed_capacity": 18,
                                "slots": [
                                    {
                                        "position_name": "Teacher 1",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "James Mitchell", "permit_status": "permit_in_process"},
                                    },
                                    {
                                        "position_name": "Teacher 2",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Sofia Ramirez", "permit_status": "no_units_needed"},
                                    },
                                ],
                            },
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    before_assignments = store.list_assignments()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()

    classrooms_nav = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsNavButton")
    assert classrooms_nav.isEnabled()
    classrooms_nav.click()
    app.processEvents()

    assert window.sidebar_panel.isHidden()
    assert page.findChild(qt_widgets.QFrame, "StaffingV2Shell") is not None
    assert classrooms_nav.property("staffingV2ActiveNav") is True
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2DashboardNavButton").property("staffingV2ActiveNav") is False
    assert page.findChild(qt_widgets.QWidget, "StaffingV2ClassroomManagementDashboard") is not None
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsTitle").text() == "Classroom Management"
    assert "Manage classroom records" in page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsSubtitle").text()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsAddButton").text() == "Add Classroom"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsAddButton").isEnabled()
    classrooms_export = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsExportButton")
    assert classrooms_export.text() == "Export"
    assert not classrooms_export.icon().isNull()
    assert classrooms_export.isEnabled()
    classrooms_export.click()
    app.processEvents()
    classrooms_export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ClassroomsExportDialog")
    assert classrooms_export_dialog is not None
    classrooms_export_text = _widget_text(classrooms_export_dialog)
    assert "Export Classroom Management" in classrooms_export_text
    assert "Total classrooms 2" in classrooms_export_text
    assert "Open positions 1" in classrooms_export_text
    assert "Harmony 1" in classrooms_export_text
    assert "Tranquility" in classrooms_export_text
    classrooms_export_dialog.accept()
    app.processEvents()
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsAddButton").icon().isNull()
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsSchoolFilter").currentText() == "All Schools"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsProgramFilter").currentText() == "All Programs"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsStatusFilter").currentText() == "All Statuses"
    classrooms_search = page.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsSearch")
    assert classrooms_search.placeholderText() == "Search classrooms..."
    assert classrooms_search.actions()
    classrooms_filters = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsMoreFilters")
    assert classrooms_filters.text() == "Filters 0"
    assert classrooms_filters.property("staffingV2FilterActiveCount") == 0
    assert not classrooms_filters.icon().isNull()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsClear").text() == "Clear"
    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2ClassroomsMetricCard"))
    assert "Total Classrooms 2" in metric_text
    assert "Active 2" in metric_text
    assert "Avg Licensed Capacity 21.0" in metric_text
    assert "Total Positions 4" in metric_text
    assert "Open Positions 1" in metric_text
    assert len(page.findChildren(qt_widgets.QFrame, "StaffingV2ClassroomsMetricCard")) == 5
    for card in page.findChildren(qt_widgets.QFrame, "StaffingV2ClassroomsMetricCard"):
        assert card.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsMetricIcon") is not None
        assert card.findChild(qt_widgets.QLabel, "StaffingV2MetricValue") is not None
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ClassroomsTable")
    assert table.rowCount() == 2
    expected_headers = [
        "Classroom",
        "School",
        "Program",
        "Licensed Capacity",
        "Total Positions",
        "Filled",
        "Open",
        "Priority Status",
        "Active",
        "Actions",
    ]
    assert [table.horizontalHeaderItem(index).text() for index in range(table.columnCount())] == expected_headers
    assert table.columnWidth(0) >= 110
    assert table.columnWidth(3) >= 120
    assert table.columnWidth(7) >= 130
    assert table.currentRow() == 0
    assert table.horizontalScrollBar().value() == 0
    assert table.horizontalHeaderItem(0).text() == "Classroom"
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert {
        "Harmony 1",
        "Hawthorne",
        "Preschool",
        "24",
        "2",
        "1",
        "Need Now",
        "Yes",
        "Tranquility",
        "Infant",
        "18",
        "Filled",
    } <= table_text
    assert table.cellWidget(0, 7).objectName() in {"StaffingV2NeedNowChip", "StaffingV2ReplaceChip"}
    assert table.cellWidget(1, 7).text() == "Filled"
    assert table.cellWidget(0, 8).objectName() == "StaffingV2HealthyChip"
    for row in range(table.rowCount()):
        view_button = table.cellWidget(row, 9)
        assert isinstance(view_button, qt_widgets.QPushButton)
        assert view_button.text() == "View"
        assert not view_button.icon().isNull()
        assert view_button.isEnabled()
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsResultCount").text() == "Showing 1 to 2 of 2 classrooms"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsRowsPerPage").currentText() == "10 / page"
    detail = page.findChild(qt_widgets.QFrame, "StaffingV2ClassroomsDetailPanel")
    detail_text = _widget_text(detail)
    assert "Classroom Detail" in detail_text
    assert "Harmony 1" in detail_text
    assert "Staffing Summary" in detail_text
    assert "Current Positions" in detail_text
    assert "Teacher 1" in detail_text
    assert "OPEN POSITION" in detail_text
    detail_scroll = detail.findChild(qt_widgets.QScrollArea, "StaffingV2ClassroomsDetailScroll")
    assert detail_scroll is not None
    assert detail_scroll.widgetResizable()
    assert detail_scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    detail_viewport_width = detail_scroll.viewport().width()
    first_classroom_card = detail.findChild(qt_widgets.QFrame, "StaffingV2ClassroomsDetailCard")
    assert first_classroom_card.y() < 160
    assert first_classroom_card.width() <= detail_viewport_width
    assert detail_scroll.widget().findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsSaveButton") is None
    classroom_close = detail.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsDetailClose")
    assert classroom_close is not None
    assert classroom_close.text() == ""
    assert not classroom_close.icon().isNull()
    assert len(detail.findChildren(qt_widgets.QFrame, "StaffingV2ClassroomsDetailMetricCard")) == 4
    for detail_card in detail.findChildren(qt_widgets.QFrame, "StaffingV2ClassroomsDetailCard"):
        assert detail_card.sizePolicy().verticalPolicy() == qt_widgets.QSizePolicy.Policy.Maximum
    detail_footer = detail.findChild(qt_widgets.QWidget, "StaffingV2ClassroomsDetailFooter")
    assert detail_footer is not None
    assert detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsDetailClose") is None
    deactivate_button = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsDeactivateButton")
    assert deactivate_button.text() == "Deactivate Classroom"
    assert not deactivate_button.icon().isNull()
    assert detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsDeactivateButton") is deactivate_button
    save_button = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsSaveButton")
    assert save_button.text() == "Save Changes"
    assert not save_button.icon().isNull()
    assert save_button.parent() is not detail_scroll.widget()
    assert detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsSaveButton") is save_button
    validation_panel = page.findChild(qt_widgets.QFrame, "StaffingV2ClassroomsValidationPanel")
    assert "Classroom Validation & Health" in _widget_text(validation_panel)
    health_cards = validation_panel.findChildren(qt_widgets.QFrame, "StaffingV2ClassroomsHealthCard")
    assert len(health_cards) == 5
    assert {card.property("staffingV2HealthVariant") for card in health_cards} <= {"success", "warning", "danger"}
    assert "success" in {card.property("staffingV2HealthVariant") for card in health_cards}
    assert len(store.list_assignments()) == len(before_assignments)
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_add_classroom_dialog_creates_classroom_through_service(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 1",
        position_type="Teacher",
    )
    before_assignments = store.list_assignments()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsNavButton").click()
    app.processEvents()

    add_classroom = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsAddButton")
    assert add_classroom.isEnabled()
    add_classroom.click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2AddClassroomDialog")
    assert dialog is not None
    assert "Add Classroom" in _widget_text(dialog)
    assert dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddClassroomSchool").currentText() == "Hawthorne"
    dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddClassroomName").setText("Sunflower")
    dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddClassroomProgram").setCurrentText("Preschool")
    dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddClassroomCapacity").setText("18")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2AddClassroomSave").click()
    app.processEvents()

    classrooms = store.list_classrooms()
    created = [classroom for classroom in classrooms if classroom.name == "Sunflower"][0]
    assert created.school == "Hawthorne"
    assert created.program == "Preschool"
    assert created.licensed_capacity == 18
    assert len(store.list_assignments()) == len(before_assignments)
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ClassroomsTable")
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert {"Sunflower", "18", "0", "Don't Need"} <= table_text
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_classrooms_filter_side_panel_filters_rows_without_mutating_db(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [{"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}],
                            },
                            {
                                "name": "Tranquility",
                                "program": "Infant",
                                "licensed_capacity": 18,
                                "slots": [
                                    {
                                        "position_name": "Teacher 1",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "James Mitchell", "permit_status": "teacher_permit_approved"},
                                    }
                                ],
                            },
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    before_assignments = store.list_assignments()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsNavButton").click()
    app.processEvents()
    body_splitter = page.findChild(qt_widgets.QSplitter, "StaffingV2ClassroomsBodySplitter")
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ClassroomsTable")
    before_splitter_sizes = body_splitter.sizes()
    before_table_width = table.width()

    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsMoreFilters").click()
    app.processEvents()

    drawer = page.findChild(qt_widgets.QFrame, "StaffingV2ClassroomsFilterDrawer")
    assert drawer is not None
    assert not drawer.isHidden()
    assert drawer.width() >= 420
    assert drawer.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2ClassroomManagementDashboard")
    assert body_splitter.sizes() == before_splitter_sizes
    assert table.width() == before_table_width
    assert table.horizontalScrollBar().value() == 0
    drawer_text = _widget_text(drawer)
    assert "Filters" in drawer_text
    assert "School" in drawer_text
    assert "Program" in drawer_text
    assert "Status" in drawer_text
    assert "Open Positions" in drawer_text
    assert "Days Open" in drawer_text
    assert "Permit Status" in drawer_text
    assert "Assigned Staff" in drawer_text
    assert "Sort By" in drawer_text
    assert drawer.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsFilterSchool").currentText() == "All Schools"
    assert drawer.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsFilterProgram").currentText() == "All Programs"
    assert drawer.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsFilterPermit").currentText() == "All Permit Statuses"
    assert drawer.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsFilterAssignedStaff").currentText() == "All Staff"
    assert drawer.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsFilterSortBy").currentText() == "Default Order"
    assert drawer.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsFilterDaysFrom") is not None
    assert drawer.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsFilterDaysTo") is not None
    reset = drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterReset")
    assert reset.text() == "Reset"
    assert not reset.icon().isNull()
    assert drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterClose").text() == ""
    apply_button = drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterApply")
    assert apply_button.text() == "Apply Filters 0"
    assert apply_button.property("staffingV2FilterActiveCount") == 0
    assert not apply_button.icon().isNull()
    cancel_button = drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterCancel")
    assert not cancel_button.isHidden()
    assert not apply_button.isHidden()
    assert cancel_button.geometry().bottom() <= drawer.rect().bottom()
    assert apply_button.geometry().bottom() <= drawer.rect().bottom()
    need_now = drawer.findChild(qt_widgets.QCheckBox, "StaffingV2ClassroomsFilterNeedNow")
    filled = drawer.findChild(qt_widgets.QCheckBox, "StaffingV2ClassroomsFilterFilled")
    dont_need = drawer.findChild(qt_widgets.QCheckBox, "StaffingV2ClassroomsFilterDontNeed")
    assert need_now.isChecked()
    assert filled.isChecked()
    assert dont_need.isChecked()
    filled.setChecked(False)
    app.processEvents()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ClassroomsTable")
    assert table.rowCount() == 2
    drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterCancel").click()
    app.processEvents()
    assert drawer.isHidden()
    assert table.rowCount() == 2
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsMoreFilters").click()
    app.processEvents()
    filled = drawer.findChild(qt_widgets.QCheckBox, "StaffingV2ClassroomsFilterFilled")
    assert filled.isChecked()
    filled.setChecked(False)
    drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterApply").click()
    app.processEvents()

    assert table.rowCount() == 1
    assert table.item(0, 0).text() == "Harmony 1"
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsMoreFilters").click()
    app.processEvents()
    drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterReset").click()
    drawer.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsFilterApply").click()
    app.processEvents()
    assert table.rowCount() == 2
    assert len(store.list_assignments()) == len(before_assignments)
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_classrooms_paginates_and_saves_detail_without_assignment_mutation(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    for index in range(12):
        store.seed_assignment(
            school="Hawthorne",
            classroom=f"Room {index + 1:02d}",
            position_name="Teacher 1",
            position_type="Teacher",
            status="filled",
            person_name=f"Teacher {index + 1:02d}",
        )
    before_assignments = store.list_assignments()
    before_history = {assignment.id: store.active_history_count(assignment.id) for assignment in before_assignments}
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsNavButton").click()
    app.processEvents()

    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ClassroomsTable")
    assert table.horizontalHeaderItem(0).text() == "Classroom"
    assert table.rowCount() == 10
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsResultCount").text() == "Showing 1 to 10 of 12 classrooms"
    next_page = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsNextPage")
    previous_page = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsPreviousPage")
    current_page = page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsCurrentPage")
    rows_per_page = page.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsRowsPerPage")
    assert next_page.isEnabled()
    assert not previous_page.isEnabled()
    next_page.click()
    app.processEvents()
    assert current_page.text() == "2"
    assert table.rowCount() == 2
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsResultCount").text() == "Showing 11 to 12 of 12 classrooms"
    previous_page.click()
    app.processEvents()
    rows_per_page.setCurrentText("25 / page")
    app.processEvents()
    assert table.rowCount() == 12
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsResultCount").text() == "Showing 1 to 12 of 12 classrooms"

    name = page.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsDetailNameEdit")
    program = page.findChild(qt_widgets.QComboBox, "StaffingV2ClassroomsDetailProgramEdit")
    capacity = page.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsDetailCapacityEdit")
    display_order = page.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsDetailDisplayOrderEdit")
    assert name.text() == "Room 01"
    name.setText("Room 01A")
    program.setCurrentText("Pre-K")
    capacity.setText("22")
    display_order.setText("4")
    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsSaveButton").click()
    app.processEvents()

    updated = next(classroom for classroom in store.list_classrooms() if classroom.name == "Room 01A")
    assert updated.program == "Pre-K"
    assert updated.licensed_capacity == 22
    assert updated.display_order == 4
    assert len(store.list_assignments()) == len(before_assignments)
    assert {assignment.id: store.active_history_count(assignment.id) for assignment in before_assignments} == before_history
    assert "Room 01A" in _widget_text(page.findChild(qt_widgets.QFrame, "StaffingV2ClassroomsDetailPanel"))

    page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsDeactivateButton").click()
    app.processEvents()
    status = page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsDetailStatus")
    assert "active assignments" in status.text()
    assert any(classroom.id == updated.id for classroom in store.list_classrooms())

    page.findChild(qt_widgets.QLineEdit, "StaffingV2ClassroomsSearch").setText("Room 12")
    app.processEvents()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ClassroomsCurrentPage").text() == "1"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ClassroomsResultCount").text() == "Showing 1 to 1 of 1 classrooms"

    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_add_position_dialog_creates_need_now_position_through_service(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    before_count = len(store.list_assignments())
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()

    page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2AddPositionDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    dialog_text = _widget_text(dialog)
    assert "Add Position" in dialog_text
    assert "Create a new position for a classroom." in dialog_text
    assert "Status Definitions" in dialog_text
    assert "Need Now" in dialog_text
    assert "Don't Need Now" in dialog_text
    assert "Coming" in dialog_text
    assert "Filled" in dialog_text
    close_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionClose")
    assert close_button is not None
    assert close_button.text() == ""
    assert not close_button.icon().isNull()
    assert dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionSchool").currentText() == "Hawthorne"
    assert dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionClassroom").currentText() == "Harmony 1"
    dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionType").setCurrentText("Teacher")
    name = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddPositionName")
    name.setText("Teacher 2")
    dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionInitialStatus").setCurrentText("Need Now")
    notes = dialog.findChild(qt_widgets.QTextEdit, "StaffingV2AddPositionNotes")
    notes.setPlainText("Added from v2 dialog.")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionSubmit").click()
    app.processEvents()

    assignments = store.list_assignments()
    assert len(assignments) == before_count + 1
    created = next(row for row in assignments if row.position_name == "Teacher 2")
    assert created.status == "need_now"
    assert created.classroom == "Harmony 1"
    assert created.classroom_program == "Preschool"
    assert created.classroom_capacity == 24
    assert created.notes == "Added from v2 dialog."
    assert store.active_history_count(created.id) == 1
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert "Teacher 2" in table_text
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_add_position_submit_immediately_shows_created_position_when_filters_are_stale(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "filled"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    director_model = pyside_interview_app.build_director_staffing_model(full_model, school="Hawthorne")
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    page = window.stack.widget(0)
    page.findChild(qt_widgets.QLineEdit, "StaffingV2Search").setText("not-visible")
    app.processEvents()
    assert page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable").rowCount() == 0

    page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").click()
    app.processEvents()
    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2AddPositionDialog")
    assert dialog is not None
    position_type = dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionType")
    assert "Director" in [position_type.itemText(index) for index in range(position_type.count())]
    position_type.setCurrentText("Director")
    dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddPositionName").setText("Director")
    dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionInitialStatus").setCurrentText("Need Now")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionSubmit").click()
    app.processEvents()

    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    visible_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert page.findChild(qt_widgets.QLineEdit, "StaffingV2Search").text() == ""
    assert "Director" in visible_text
    assert page.findChild(qt_widgets.QLabel, "StaffingV2DrawerPositionName").text() == "Director"
    created = next(row for row in window.staffing_store.list_assignments() if row.position_name == "Director")
    assert created.position_type == "Director"
    page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkComing").click()
    app.processEvents()
    coming_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2MarkComingDialog")
    assert coming_dialog is not None
    role = coming_dialog.findChild(qt_widgets.QComboBox, "StaffingV2ComingRole")
    assert "Director" in [role.itemText(index) for index in range(role.count())]
    assert role.currentText() == "Director"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_delete_position_removes_accidental_director(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Office",
                                "program": "Support",
                                "slots": [
                                    {"position_name": "Office", "position_type": "Office", "status": "filled", "person_name": "Violet"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    director_model = pyside_interview_app.build_director_staffing_model(full_model, school="Hawthorne")
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    page = window.stack.widget(0)

    page.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionButton").click()
    app.processEvents()
    add_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2AddPositionDialog")
    add_dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionType").setCurrentText("Director")
    add_dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddPositionName").setText("Director")
    add_dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPositionInitialStatus").setCurrentText("Need Now")
    add_dialog.findChild(qt_widgets.QPushButton, "StaffingV2AddPositionSubmit").click()
    app.processEvents()

    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    director_button = _staffing_button_for_position(table, "Director")
    delete_action = next(action for action in director_button.menu().actions() if action.text() == "Delete Position")
    delete_action.trigger()
    app.processEvents()
    delete_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2DeletePositionDialog")
    assert delete_dialog is not None
    assert "Director" in _widget_text(delete_dialog)
    delete_dialog.findChild(qt_widgets.QPushButton, "StaffingV2DeletePositionConfirm").click()
    app.processEvents()

    assert [row.position_name for row in store.list_assignments()] == ["Office"]
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    refreshed_text = {
        refreshed_table.item(row, column).text()
        for row in range(refreshed_table.rowCount())
        for column in range(refreshed_table.columnCount())
        if refreshed_table.item(row, column) is not None
    }
    assert "Director" not in refreshed_text
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_position_detail_drawer_opens_from_position_row(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"},
                                    {
                                        "position_name": "Teacher 2",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "start_date": "2026-07-08",
                                        "person": {"name": "Imgard", "permit_status": "permit_in_process"},
                                    },
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    teacher_row = _staffing_row_for_position(table, "Teacher 1")

    table.cellClicked.emit(teacher_row, 0)
    app.processEvents()

    drawer = page.findChild(qt_widgets.QFrame, "StaffingV2PositionDrawer")
    assert drawer is not None
    assert not drawer.isHidden()
    assert page.findChild(qt_widgets.QLabel, "StaffingV2DrawerTitle").text() == "Position Detail"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2DrawerPositionName").text() == "Teacher 1"
    drawer_text = _widget_text(drawer)
    assert "Harmony 1" in drawer_text
    assert "Hawthorne" in drawer_text
    assert "Need Now" in drawer_text
    assert "OPEN POSITION" in drawer_text
    assert "Position Overview" in drawer_text
    assert "Available Next Actions" in drawer_text
    assert "Data Integrity / Validation" in drawer_text
    assert "Lifecycle History" in drawer_text
    assert "Related Person" in drawer_text
    assert "No person is currently assigned to this position." in drawer_text
    assign_person = drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerAssignPerson")
    assert assign_person is not None
    assert assign_person.text() == "Assign or Create Person"
    assert assign_person.isEnabled()
    assert not assign_person.icon().isNull()
    assert _icon_has_primary_blue(assign_person.icon())
    assert drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkComing").text() == "Mark Coming"
    assert drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkDontNeed").text() == "Mark Not Needed"
    history_button = drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerViewHistory")
    assert history_button.text() == "View Full History"
    assert history_button.isEnabled()
    assert not drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkComing").icon().isNull()
    assert not drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkDontNeed").icon().isNull()
    assert not history_button.icon().isNull()
    close = page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerClose")
    assert close is not None
    assert close.text() == ""
    assert not close.icon().isNull()
    footer_buttons = {
        button.objectName(): button.text()
        for button in drawer.findChildren(qt_widgets.QPushButton)
        if button.objectName().startswith("StaffingV2Drawer")
    }
    assert footer_buttons["StaffingV2DrawerCancel"] == "Cancel"
    assert not drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerCancel").icon().isNull()
    assert drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerSaveDraft") is None
    assert drawer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerSaveChanges") is None
    history_button.click()
    app.processEvents()
    assert page.findChild(qt_widgets.QWidget, "StaffingV2AssignmentHistoryDashboard") is not None
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryNavButton").property("staffingV2ActiveNav") is True
    assert len(store.list_assignments()) == 2
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_position_drawer_wires_open_and_edit_actions(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 1",
        position_type="Teacher",
        status="dont_need_now",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    table.cellClicked.emit(_staffing_row_for_position(table, "Teacher 1"), 0)
    app.processEvents()

    page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkNeedNow").click()
    app.processEvents()

    assert store.get_assignment(assignment_id).status == "need_now"
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(table, "Teacher 1").text() == "Mark Coming"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkComing") is not None

    page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerEditPosition").click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2EditPositionDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    name = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2EditPositionName")
    assert name.text() == "Teacher 1"
    name.setText("Teacher 1A")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2EditPositionSubmit").click()
    app.processEvents()

    assert store.get_assignment(assignment_id).position_name == "Teacher 1A"
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_row_for_position(refreshed_table, "Teacher 1A") == 0
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_mark_coming_dialog_saves_through_service(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "program": "Preschool",
                                "licensed_capacity": 24,
                                "slots": [
                                    {
                                        "position_name": "Aide 2",
                                        "position_type": "Aide",
                                        "status": "need_now",
                                        "current_opened_date": "2026-07-01",
                                    }
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    store.import_seed_file(seed_path)
    assignment_id = next(row.id for row in store.list_assignments() if row.position_name == "Aide 2")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")

    _staffing_button_for_position(table, "Aide 2").click()
    app.processEvents()

    assert window.window.findChild(qt_widgets.QDialog, "PySideStaffingMarkComingDialog") is None
    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2MarkComingDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    dialog_text = _widget_text(dialog)
    assert "Mark Coming" in dialog_text
    assert "Position Summary" in dialog_text
    assert "Candidate Selection" in dialog_text
    assert "Candidate Details" in dialog_text
    assert "Validation / Requirements" in dialog_text
    assert "What will happen on save" in dialog_text
    assert "This action does not close the open assignment history cycle." in dialog_text
    close_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingClose")
    assert close_button is not None
    assert close_button.text() == ""
    assert not close_button.icon().isNull()
    assert not dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingSelectPerson").icon().isNull()
    assert not dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingCreatePerson").icon().isNull()
    assert _icon_has_primary_blue(dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingCreatePerson").icon())
    assert not dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingSubmit").icon().isNull()
    people_search = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2ComingPeopleSearch")
    people_search.setText("Emily")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingSelectPerson").click()
    app.processEvents()
    selector = window.window.findChild(qt_widgets.QDialog, "StaffingV2SelectPersonDialog")
    assert selector is not None and not selector.isHidden()
    selector.findChild(qt_widgets.QPushButton, "StaffingV2SelectPersonCancel").click()
    app.processEvents()
    full_name = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2ComingFullName")
    full_name.setText("Temporary Name")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingCreatePerson").click()
    app.processEvents()
    assert full_name.text() == ""

    full_name.setText("Emily Carter")
    dialog.findChild(qt_widgets.QComboBox, "StaffingV2ComingRole").setCurrentText("Teacher")
    dialog.findChild(qt_widgets.QDateEdit, "StaffingV2ComingStartDate").setDate(qt_core.QDate(2026, 8, 1))
    dialog.findChild(qt_widgets.QComboBox, "StaffingV2ComingPermitStatus").setCurrentText("Permit in Process")
    dialog.findChild(qt_widgets.QSpinBox, "StaffingV2ComingUnits").setValue(12)
    assert dialog.findChild(qt_widgets.QCheckBox, "StaffingV2ComingActive").isChecked()
    dialog.findChild(qt_widgets.QTextEdit, "StaffingV2ComingNotes").setPlainText("Hiring notes stay UI-only for this slice.")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2ComingSubmit").click()
    app.processEvents()

    updated = store.get_assignment(assignment_id)
    assert updated.status == "coming"
    assert updated.position_type == "Teacher"
    assert updated.position_name == "Teacher 2"
    assert updated.person_name == "Emily Carter"
    assert updated.start_date == "2026-08-01"
    assert updated.permit_status == "permit_in_process"
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(refreshed_table, "Teacher 2").text() == "Mark Filled"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_mark_filled_dialog_accepts_actual_start_date(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_test = pytest.importorskip("PySide6.QtTest")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 1",
        position_type="Teacher",
    )
    service = pyside_interview_app.StaffingService(
        store,
        clock=lambda: "2026-07-01T09:00:00Z",
    )
    service.open_position(assignment_id)
    scheduled_start = date.today() + timedelta(days=1)
    actual_start = date.today()
    service.mark_coming(assignment_id, person_name="Emily Carter", start_date=scheduled_start.isoformat())
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")

    _staffing_button_for_position(table, "Teacher 1").click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2MarkFilledDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    dialog_text = _widget_text(dialog)
    assert "Mark Filled" in dialog_text
    assert "Position Summary" in dialog_text
    assert "Assigned Person" in dialog_text
    assert "Start Confirmation" in dialog_text
    assert "Validation / Requirements" in dialog_text
    assert "What will happen on save" in dialog_text
    assert "This action closes the current open assignment history cycle." in dialog_text
    assert window.window.findChild(qt_widgets.QDialog, "PySideStaffingMarkFilledDialog") is None
    close_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2FilledClose")
    assert close_button is not None
    assert close_button.text() == ""
    assert not close_button.icon().isNull()
    filled_date = dialog.findChild(qt_widgets.QDateEdit, "StaffingV2FilledDate")
    assert filled_date is not None
    assert filled_date.date() == qt_core.QDate.fromString(scheduled_start.isoformat(), "yyyy-MM-dd")
    assert filled_date.isEnabled()
    assert filled_date.calendarPopup()
    qt_test.QTest.mouseClick(filled_date, qt_core.Qt.MouseButton.LeftButton)
    app.processEvents()
    assert filled_date.calendarWidget().isVisible()
    filled_date.calendarWidget().parentWidget().hide()
    filled_date.setDate(qt_core.QDate.fromString(actual_start.isoformat(), "yyyy-MM-dd"))

    dialog.findChild(qt_widgets.QTextEdit, "StaffingV2FilledNotes").setPlainText("Started and verified.")
    assert dialog.findChild(qt_widgets.QCheckBox, "StaffingV2FilledStarted").isChecked()
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2FilledSubmit").click()
    app.processEvents()

    updated = store.get_assignment(assignment_id)
    assert updated.status == "filled"
    assert updated.person_name == "Emily Carter"
    assert updated.start_date == actual_start.isoformat()
    assert updated.current_filled_date == actual_start.isoformat()
    assert store.closed_days_to_fill() == [(actual_start - date(2026, 7, 1)).days]
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(refreshed_table, "Teacher 1").text() == "Manage Filled"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_manage_filled_dialog_selects_next_workflow(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_test = pytest.importorskip("PySide6.QtTest")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 2",
        position_type="Teacher",
        status="filled",
        person_name="Imgard",
        permit_status="permit_in_process",
    )
    before = store.get_assignment(assignment_id)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    button = _staffing_button_for_position(table, "Teacher 2")

    assert button.text() == "Manage Filled"
    assert button.isEnabled()
    button.click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ManageFilledDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    dialog_text = _widget_text(dialog)
    assert "Manage Filled Position" in dialog_text
    assert "Choose what you want to do with this filled position." in dialog_text
    assert "Imgard" in dialog_text
    assert "Teacher 2" in dialog_text
    assert "Filled" in dialog_text
    assert "Permit in Process" in dialog_text
    assert "Update Permit Status" in dialog_text
    assert "Replace Employee" in dialog_text
    assert "What happens next" in dialog_text
    assert "This step does not change anything until you continue" in dialog_text
    close_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2ManageFilledClose")
    assert close_button is not None
    assert close_button.text() == ""
    assert not close_button.icon().isNull()

    permit_option = dialog.findChild(qt_widgets.QRadioButton, "StaffingV2ManageFilledPermitOption")
    replace_option = dialog.findChild(qt_widgets.QRadioButton, "StaffingV2ManageFilledReplaceOption")
    continue_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2ManageFilledContinue")
    assert permit_option.isChecked()
    assert not replace_option.isChecked()
    replace_option.click()
    app.processEvents()
    assert replace_option.isChecked()
    assert continue_button.text() == "Continue"

    continue_button.click()
    app.processEvents()
    replace_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ReplaceEmployeeDialog")
    assert replace_dialog is not None
    assert not replace_dialog.isHidden()
    assert window.window.findChild(qt_widgets.QDialog, "PySideStaffingReplaceDialog") is None
    replace_text = _widget_text(replace_dialog)
    assert "Replace Employee" in replace_text
    assert "Position Summary" in replace_text
    assert "Replacement Details" in replace_text
    assert "Validation / Requirements" in replace_text
    assert "What will happen on save" in replace_text
    close = replace_dialog.findChild(qt_widgets.QPushButton, "StaffingV2ReplaceClose")
    assert close is not None
    assert close.text() == ""
    assert not close.icon().isNull()
    notice = replace_dialog.findChild(qt_widgets.QDateEdit, "StaffingV2ReplaceNotice")
    final_day = replace_dialog.findChild(qt_widgets.QDateEdit, "StaffingV2ReplaceFinalDay")
    reason = replace_dialog.findChild(qt_widgets.QComboBox, "StaffingV2ReplaceReason")
    assert notice is not None
    assert final_day is not None
    assert notice.calendarPopup()
    assert final_day.calendarPopup()
    assert notice.date() == qt_core.QDate.currentDate()
    assert final_day.date() == qt_core.QDate.currentDate()
    assert reason is not None
    assert [reason.itemText(index) for index in range(reason.count())] == [
        "Resignation",
        "Termination",
        "Leave of absence",
        "Transfer",
        "Other",
    ]
    submit = replace_dialog.findChild(qt_widgets.QPushButton, "StaffingV2ReplaceSubmit")
    qt_test.QTest.mouseClick(notice, qt_core.Qt.MouseButton.LeftButton, pos=notice.rect().center())
    app.processEvents()
    assert notice.calendarWidget().isVisible()
    after = store.get_assignment(assignment_id)
    assert after.status == before.status
    assert after.person_name == before.person_name
    assert after.permit_status == before.permit_status

    notice.setDate(qt_core.QDate(2026, 8, 1))
    final_day.setDate(qt_core.QDate(2026, 8, 15))
    reason.setCurrentText("Transfer")
    submit.click()
    app.processEvents()
    replaced = store.get_assignment(assignment_id)
    assert replaced.status == "replace"
    assert replaced.person_name == before.person_name
    assert replaced.notice_given == "2026-08-01"
    assert replaced.final_working_day == "2026-08-15"
    page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerMarkNeedNow").click()
    app.processEvents()
    need_now_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2MarkNeedNowDialog")
    assert need_now_dialog is not None
    need_now_submit = need_now_dialog.findChild(qt_widgets.QPushButton, "StaffingV2NeedNowSubmit")
    need_now_submit.click()
    need_now_submit.click()
    app.processEvents()
    reopened = store.get_assignment(assignment_id)
    assert reopened.status == "need_now"
    assert reopened.person_id is None
    assert reopened.person_name == ""
    assert reopened.start_date == ""
    assert need_now_dialog.isHidden()
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(refreshed_table, "Teacher 2").text() == "Mark Coming"
    drawer = page.findChild(qt_widgets.QFrame, "StaffingV2PositionDrawer")
    assert drawer is not None
    assert "Need Now" in _widget_text(drawer)
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_replace_today_reopens_need_now_and_clears_person(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 2",
        position_type="Teacher",
        status="filled",
        person_name="Imgard",
        permit_status="permit_in_process",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")

    _staffing_button_for_position(table, "Teacher 2").click()
    app.processEvents()
    manage_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ManageFilledDialog")
    manage_dialog.findChild(qt_widgets.QRadioButton, "StaffingV2ManageFilledReplaceOption").click()
    manage_dialog.findChild(qt_widgets.QPushButton, "StaffingV2ManageFilledContinue").click()
    app.processEvents()
    replace_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ReplaceEmployeeDialog")
    final_day = replace_dialog.findChild(qt_widgets.QDateEdit, "StaffingV2ReplaceFinalDay")
    assert final_day.date() == qt_core.QDate.currentDate()

    replace_dialog.findChild(qt_widgets.QPushButton, "StaffingV2ReplaceSubmit").click()
    app.processEvents()

    updated = store.get_assignment(assignment_id)
    assert updated.status == "need_now"
    assert updated.person_id is None
    assert updated.person_name == ""
    assert updated.start_date == ""
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(refreshed_table, "Teacher 2").text() == "Mark Coming"
    drawer = page.findChild(qt_widgets.QFrame, "StaffingV2PositionDrawer")
    assert drawer is not None
    drawer_text = _widget_text(drawer)
    assert "Need Now" in drawer_text
    assert "OPEN POSITION" in drawer_text
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_update_permit_dialog_saves_people_permit_details(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 2",
        position_type="Teacher",
        status="filled",
        person_name="Imgard",
        permit_status="permit_in_process",
    )
    person_id = store.get_assignment(assignment_id).person_id or 0
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    _staffing_button_for_position(table, "Teacher 2").click()
    app.processEvents()
    manage = window.window.findChild(qt_widgets.QDialog, "StaffingV2ManageFilledDialog")
    manage.findChild(qt_widgets.QPushButton, "StaffingV2ManageFilledContinue").click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2UpdatePermitDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    dialog_text = _widget_text(dialog)
    assert "Update Permit Status" in dialog_text
    assert "Update the employee permit level without reopening this position." in dialog_text
    assert "Position Summary" in dialog_text
    assert "Permit Update" in dialog_text
    assert "Validation / Requirements" in dialog_text
    assert "What will happen on save" in dialog_text
    assert "This action updates People only and does not reopen the staffing cycle." in dialog_text
    close_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2PermitClose")
    assert close_button is not None
    assert close_button.text() == ""
    assert not close_button.icon().isNull()
    assert dialog.findChild(qt_widgets.QLineEdit, "StaffingV2PermitEmployeeName").text() == "Imgard"
    assert not dialog.findChild(qt_widgets.QLineEdit, "StaffingV2PermitEmployeeName").isEnabled()
    assert dialog.findChild(qt_widgets.QLineEdit, "StaffingV2PermitRole").text() == "Teacher"
    assert not dialog.findChild(qt_widgets.QLineEdit, "StaffingV2PermitRole").isEnabled()

    dialog.findChild(qt_widgets.QComboBox, "StaffingV2PermitNewStatus").setCurrentText("Teacher Permit")
    dialog.findChild(qt_widgets.QDateEdit, "StaffingV2PermitEffectiveDate").setDate(qt_core.QDate(2026, 7, 6))
    dialog.findChild(qt_widgets.QSpinBox, "StaffingV2PermitUnits").setValue(24)
    dialog.findChild(qt_widgets.QTextEdit, "StaffingV2PermitNotes").setPlainText("Permit file received.")
    assert dialog.findChild(qt_widgets.QCheckBox, "StaffingV2PermitDocumentationReceived").isChecked()
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2PermitSubmit").click()
    app.processEvents()

    updated = store.get_assignment(assignment_id)
    with store.connect() as conn:
        person = store.person_context(conn, person_id)
    assert updated.status == "filled"
    assert updated.permit_status == "teacher_permit_approved"
    assert person.permit_effective_date == "2026-07-06"
    assert person.units == 24
    assert person.permit_documentation_received is True
    assert person.permit_notes == "Permit file received."
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_mark_need_now_dialog_clears_replacement(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 2",
        position_type="Teacher",
        status="filled",
        person_name="Maria Gonzalez",
        permit_status="teacher_permit_approved",
    )
    service = pyside_interview_app.StaffingService(store, clock=lambda: "2026-05-01T09:00:00Z")
    service.mark_replacing(assignment_id, notice_given="2026-05-01", final_working_day="2026-05-22")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    button = _staffing_button_for_position(table, "Teacher 2")

    assert button.text() == "Mark Need Now"
    button.click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2MarkNeedNowDialog")
    assert dialog is not None
    assert not dialog.isHidden()
    dialog_text = _widget_text(dialog)
    assert "Mark Position as Need Now" in dialog_text
    assert "This action will change the status from Replace to Need Now and reopen the position for hiring." in dialog_text
    assert "Position Summary" in dialog_text
    assert "Harmony 1 (Hawthorne)" in dialog_text
    assert "Maria Gonzalez" in dialog_text
    assert "May 1, 2026" in dialog_text
    assert "May 22, 2026" in dialog_text
    assert "What will happen" in dialog_text
    assert "Status will change from Replace to Need Now" in dialog_text
    assert "Teacher name and start date will be cleared" in dialog_text
    assert "Options" in dialog_text
    close_button = dialog.findChild(qt_widgets.QPushButton, "StaffingV2NeedNowClose")
    assert close_button is not None
    assert close_button.text() == ""
    assert not close_button.icon().isNull()
    checkbox = dialog.findChild(qt_widgets.QCheckBox, "StaffingV2NeedNowClearPerson")
    assert checkbox is not None
    assert checkbox.isChecked()
    assert "This will remove Maria Gonzalez as the assigned person for this position." in dialog_text

    dialog.findChild(qt_widgets.QPushButton, "StaffingV2NeedNowSubmit").click()
    app.processEvents()

    updated = store.get_assignment(assignment_id)
    assert updated.status == "need_now"
    assert updated.person_id is None
    assert updated.person_name == ""
    assert updated.start_date == ""
    assert store.active_history_count(assignment_id) == 1
    refreshed_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(refreshed_table, "Teacher 2").text() == "Mark Coming"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_people_dashboard_renders_employee_management_from_db(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    maria_assignment_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 1",
        position_type="Teacher",
        status="filled",
        person_name="Maria Gonzalez",
        permit_status="teacher_permit_approved",
    )
    sofia_assignment_id = store.seed_assignment(
        school="North Long Beach",
        classroom="Unity 1",
        position_name="Aide 1",
        position_type="Aide",
        status="filled",
        person_name="Sofia Ramirez",
        permit_status="permit_in_process",
    )
    with store.connect() as conn:
        maria_person_id = store.get_assignment(maria_assignment_id).person_id
        sofia_person_id = store.get_assignment(sofia_assignment_id).person_id
        conn.execute("UPDATE people SET units = 18 WHERE id = ?", (maria_person_id,))
        conn.execute("UPDATE people SET units = 6 WHERE id = ?", (sofia_person_id,))
    before_people = store.list_people()
    before_assignments = store.list_assignments()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "North Long Beach"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()

    page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleNavButton").click()
    app.processEvents()

    assert page.findChild(qt_widgets.QWidget, "StaffingV2PeopleDashboard") is not None
    assert page.findChild(qt_widgets.QLabel, "StaffingV2PeopleTitle").text() == "People / Employee Management"
    assert (
        page.findChild(qt_widgets.QLabel, "StaffingV2PeopleSubtitle").text()
        == "Manage employee records, permits, roles, and assignments."
    )
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleAddButton").text() == "Add Person"
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleAddButton").icon().isNull()
    people_search = page.findChild(qt_widgets.QLineEdit, "StaffingV2PeopleSearch")
    assert people_search.placeholderText() == "Search by name, role, or email..."
    assert people_search.actions()
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2PeopleActiveFilter").currentText() == "All"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2PeopleRoleFilter").currentText() == "All"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2PeoplePermitFilter").currentText() == "All"
    people_more_filters = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleMoreFilters")
    assert people_more_filters.text() == "More Filters"
    assert not people_more_filters.icon().isNull()
    assert people_more_filters.isEnabled()
    people_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PeopleTable")
    people_table_width = people_table.width()
    people_more_filters.click()
    app.processEvents()
    people_filter_drawer = page.findChild(qt_widgets.QFrame, "StaffingV2PeopleFilterDrawer")
    assert people_filter_drawer is not None
    assert not people_filter_drawer.isHidden()
    assert people_filter_drawer.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2PeopleDashboard")
    assert people_table.width() == people_table_width
    people_filter_scroll = people_filter_drawer.findChild(qt_widgets.QScrollArea, "StaffingV2PeopleFilterDrawerScroll")
    assert people_filter_scroll is not None
    assert people_filter_scroll.widget().findChild(qt_widgets.QPushButton, "StaffingV2PeopleFilterApply") is None
    people_filter_footer = people_filter_drawer.findChild(qt_widgets.QWidget, "StaffingV2PeopleFilterDrawerFooter")
    assert people_filter_footer is not None
    assert people_filter_footer.findChild(qt_widgets.QPushButton, "StaffingV2PeopleFilterApply") is not None
    people_filter_text = _widget_text(people_filter_drawer)
    assert "Filters" in people_filter_text
    assert "Active Status" in people_filter_text
    assert "Role" in people_filter_text
    assert "Permit Status" in people_filter_text
    people_drawer_role = people_filter_drawer.findChild(qt_widgets.QComboBox, "StaffingV2PeopleFilterRole")
    assert people_drawer_role.currentText() == "All"
    people_drawer_role.setCurrentText("Teacher")
    people_filter_drawer.findChild(qt_widgets.QPushButton, "StaffingV2PeopleFilterApply").click()
    app.processEvents()
    filtered_table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PeopleTable")
    assert filtered_table.rowCount() == 1
    assert filtered_table.item(0, 0).text() == "Maria Gonzalez"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2PeopleResultCount").text() == "Showing 1 to 1 of 1 people"
    page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleClear").click()
    app.processEvents()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleClear").text() == "Clear"
    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2PeopleMetricCard"))
    assert "Total People 2" in metric_text
    assert "Active 2" in metric_text
    assert "Teachers 1" in metric_text
    assert "Aides 1" in metric_text
    assert "Avg Units 12.0" in metric_text
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PeopleTable")
    assert table.rowCount() == 2
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert {
        "Maria Gonzalez",
        "Teacher",
        "Teacher Permit",
        "18",
        "Active",
        "Hawthorne\nHarmony 1 - Teacher 1",
        "Sofia Ramirez",
        "Aide",
        "Permit in Process",
    } <= table_text
    role_badges = {
        table.item(row, 0).text(): table.cellWidget(row, 1)
        for row in range(table.rowCount())
    }
    assert role_badges["Maria Gonzalez"].accessibleName() == "Role: Teacher"
    assert role_badges["Sofia Ramirez"].accessibleName() == "Role: Aide"
    assert role_badges["Maria Gonzalez"].property("staffingV2Role") == "teacher"
    assert role_badges["Sofia Ramirez"].property("staffingV2Role") == "aide"
    assert role_badges["Maria Gonzalez"].property("staffingV2Role") != role_badges["Sofia Ramirez"].property(
        "staffingV2Role"
    )
    for role_badge in role_badges.values():
        assert role_badge.objectName() == "StaffingV2RoleBadge"
        assert not role_badge.findChild(qt_widgets.QLabel, "StaffingV2RoleBadgeIcon").pixmap().isNull()
    for row in range(table.rowCount()):
        view_button = table.cellWidget(row, 6)
        assert isinstance(view_button, qt_widgets.QPushButton)
        assert view_button.text() == "View"
        assert not view_button.icon().isNull()
        assert view_button.isEnabled()
    sofia_row = next(row for row in range(table.rowCount()) if table.item(row, 0).text() == "Sofia Ramirez")
    sofia_view = table.cellWidget(sofia_row, 6)
    assert isinstance(sofia_view, qt_widgets.QPushButton)
    sofia_view.click()
    app.processEvents()
    assert page.findChild(qt_widgets.QLabel, "StaffingV2PeopleResultCount").text() == "Showing 1 to 2 of 2 people"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2PeopleRowsPerPage").currentText() == "10 / page"
    detail = page.findChild(qt_widgets.QFrame, "StaffingV2PeopleDetailPanel")
    assert detail.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2PeopleDashboard")
    people_detail_scroll = detail.findChild(qt_widgets.QScrollArea, "StaffingV2PeopleDetailPanelScroll")
    assert people_detail_scroll is not None
    assert people_detail_scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    people_tabs = page.findChild(qt_widgets.QFrame, "StaffingV2PeopleDetailTabs")
    assert people_tabs is not None
    assert people_tabs.y() < 160
    assert people_tabs.width() <= people_detail_scroll.viewport().width()
    assert people_detail_scroll.widget().findChild(qt_widgets.QPushButton, "StaffingV2PeopleEditButton") is None
    people_detail_close = detail.findChild(qt_widgets.QPushButton, "StaffingV2PeopleDetailClose")
    assert people_detail_close is not None
    assert people_detail_close.text() == ""
    assert not people_detail_close.icon().isNull()
    people_detail_footer = detail.findChild(qt_widgets.QWidget, "StaffingV2PeopleDetailPanelFooter")
    assert people_detail_footer is not None
    assert people_detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2PeopleDetailClose") is None
    detail_text = _widget_text(detail)
    assert "Sofia Ramirez" in detail_text
    assert "SR" in detail_text
    assert "Employee Information" in detail_text
    assert "Role Aide" in detail_text
    assert "Permit Status Permit in Process" in detail_text
    assert "Units 6" in detail_text
    assert "Current Assignment" in detail_text
    assert "North Long Beach" in detail_text
    assert "Unity 1 - Aide 1" in detail_text
    assert "Employment Status" in detail_text
    assert "Additional Information" in detail_text
    for detail_card in detail.findChildren(qt_widgets.QFrame, "StaffingV2PeopleDetailCard"):
        assert detail_card.sizePolicy().verticalPolicy() == qt_widgets.QSizePolicy.Policy.Maximum
    overview_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleOverviewTab")
    assignments_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleAssignmentsTab")
    history_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleHistoryTab")
    notes_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleNotesTab")
    documents_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleDocumentsTab")
    assert [tab.text() for tab in [overview_tab, assignments_tab, history_tab, notes_tab, documents_tab]] == [
        "Overview",
        "Assignments",
        "History",
        "Notes",
        "Documents",
    ]
    assert overview_tab.property("staffingV2ActivePeopleTab") is True
    for inactive_tab in [assignments_tab, history_tab, notes_tab, documents_tab]:
        assert inactive_tab.property("staffingV2ActivePeopleTab") is False
    history_tab.click()
    app.processEvents()
    assert history_tab.property("staffingV2ActivePeopleTab") is True
    assert overview_tab.property("staffingV2ActivePeopleTab") is False
    deactivate_button = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleDeactivateButton")
    assert deactivate_button.text() == "Deactivate Employee"
    assert not deactivate_button.icon().isNull()
    edit_button = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleEditButton")
    assert edit_button.text() == "Edit Person"
    assert not edit_button.icon().isNull()
    assert people_detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2PeopleEditButton") is edit_button
    assert len(store.list_people()) == len(before_people)
    assert len(store.list_assignments()) == len(before_assignments)
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_add_person_dialog_creates_person_through_service(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    before_people = store.list_people()
    before_assignments = store.list_assignments()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleNavButton").click()
    app.processEvents()

    add_person = page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleAddButton")
    assert add_person.isEnabled()
    add_person.click()
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2AddPersonDialog")
    assert dialog is not None
    assert "Add Person" in _widget_text(dialog)
    name = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddPersonName")
    role = dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPersonRole")
    permit = dialog.findChild(qt_widgets.QComboBox, "StaffingV2AddPersonPermit")
    units = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2AddPersonUnits")
    name.setText("Nina Patel")
    role.setCurrentText("Aide")
    permit.setCurrentText("Permit in Process")
    units.setText("5")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2AddPersonSave").click()
    app.processEvents()

    people = store.list_people()
    assert len(people) == len(before_people) + 1
    created = people[0]
    assert created.name == "Nina Patel"
    assert created.role == "Aide"
    assert created.permit_status == "permit_in_process"
    assert created.units == 5
    assert len(store.list_assignments()) == len(before_assignments)
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PeopleTable")
    assert table.rowCount() == 1
    assert table.item(0, 0).text() == "Nina Patel"
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_assignment_history_dashboard_renders_history_from_db(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    closed_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 1",
        position_type="Teacher",
    )
    open_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Quest",
        position_name="Teacher 2",
        position_type="Teacher",
        status="need_now",
    )
    with store.connect() as conn:
        person_id = store.ensure_person(conn, "Emily Carter", "Teacher", "permit_in_process", "2026-07-05T09:00:00Z")
        closed_classroom_id = conn.execute("SELECT classroom_id FROM assignments WHERE id = ?", (closed_id,)).fetchone()["classroom_id"]
        conn.execute(
            """
            UPDATE assignments
            SET person_id = ?, status = 'filled', current_opened_date = '2026-05-08',
                current_filled_date = '2026-05-20'
            WHERE id = ?
            """,
            (person_id, closed_id),
        )
        conn.execute(
            """
            INSERT INTO assignment_history (
                assignment_id, classroom_id, position_name, opened_date, filled_date,
                days_to_fill, closed_reason, created_at, updated_at
            ) VALUES (?, ?, 'Teacher 1', '2026-05-08', '2026-05-20', 12, 'filled',
                '2026-05-08T09:15:00Z', '2026-05-20T14:05:00Z')
            """,
            (closed_id, closed_classroom_id),
        )
        open_classroom_id = conn.execute("SELECT classroom_id FROM assignments WHERE id = ?", (open_id,)).fetchone()["classroom_id"]
        conn.execute(
            """
            INSERT INTO assignment_history (
                assignment_id, classroom_id, position_name, opened_date, created_at, updated_at
            ) VALUES (?, ?, 'Teacher 2', '2026-07-01', '2026-07-01T09:15:00Z', '2026-07-01T09:15:00Z')
            """,
            (open_id, open_classroom_id),
        )
    before_records = store.list_assignment_history()
    before_assignments = store.list_assignments()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()

    page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryNavButton").click()
    app.processEvents()

    assert page.findChild(qt_widgets.QWidget, "StaffingV2AssignmentHistoryDashboard") is not None
    assert page.findChild(qt_widgets.QLabel, "StaffingV2HistoryTitle").text() == "Assignment History"
    assert "Review open-to-fill staffing cycles" in page.findChild(qt_widgets.QLabel, "StaffingV2HistorySubtitle").text()
    history_export_button = page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryExportButton")
    assert history_export_button.text() == "Export"
    assert not history_export_button.icon().isNull()
    assert history_export_button.isEnabled()
    history_export_button.click()
    app.processEvents()
    history_export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2HistoryExportDialog")
    assert history_export_dialog is not None
    history_export_text = _widget_text(history_export_dialog)
    assert "Export Assignment History" in history_export_text
    assert "Total records 2" in history_export_text
    assert f"A-{closed_id:04d}" in history_export_text
    assert f"A-{open_id:04d}" in history_export_text
    history_export_dialog.accept()
    app.processEvents()
    history_validation_button = page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryValidationButton")
    assert history_validation_button.text() == "View Validation"
    assert history_validation_button.isEnabled()
    assert not history_validation_button.icon().isNull()
    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2HistoryMetricCard"))
    assert "Total Cycles 2" in metric_text
    assert "Open Cycles 1" in metric_text
    assert "Closed Cycles 1" in metric_text
    assert "Avg Days to Fill 12.0" in metric_text
    assert "Data Issues 1" in metric_text
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2HistorySchoolFilter").currentText() == "All Schools"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2HistoryClassroomFilter").currentText() == "All Classrooms"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2HistoryCycleFilter").currentText() == "All Statuses"
    history_date_range = page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryDateRangeFilter")
    assert history_date_range.text() == "2026-05-08 - 2026-07-01"
    assert not history_date_range.icon().isNull()
    assert not history_date_range.isEnabled()
    history_search = page.findChild(qt_widgets.QLineEdit, "StaffingV2HistorySearch")
    assert history_search.placeholderText() == "Search assignments..."
    assert history_search.actions()
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryMoreFilters").icon().isNull()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2HistoryTable")
    assert table.rowCount() == 2
    assert table.columnCount() == 10
    assert table.horizontalHeaderItem(9).text() == "Actions"
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert {
        f"A-{open_id:04d}",
        "Quest",
        "Teacher 2",
        "2026-07-01",
        "Open",
        "OPEN POSITION",
        "Warning",
        f"A-{closed_id:04d}",
        "Harmony 1",
        "Teacher 1",
        "2026-05-08",
        "2026-05-20",
        "12",
        "Closed",
        "Emily Carter",
        "Healthy",
    } <= table_text
    for row in range(table.rowCount()):
        view_button = table.cellWidget(row, 9)
        assert isinstance(view_button, qt_widgets.QPushButton)
        assert view_button.text() == "View"
        assert not view_button.icon().isNull()
        assert view_button.isEnabled()
    closed_history_row = next(
        row for row in range(table.rowCount()) if table.item(row, 0).text() == f"A-{closed_id:04d}"
    )
    closed_view = table.cellWidget(closed_history_row, 9)
    assert isinstance(closed_view, qt_widgets.QPushButton)
    closed_view.click()
    app.processEvents()
    closed_detail_panel = page.findChild(qt_widgets.QFrame, "StaffingV2HistoryDetailPanel")
    assert closed_detail_panel.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2AssignmentHistoryDashboard")
    history_detail_scroll = closed_detail_panel.findChild(qt_widgets.QScrollArea, "StaffingV2HistoryDetailPanelScroll")
    assert history_detail_scroll is not None
    assert history_detail_scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    first_history_card = closed_detail_panel.findChild(qt_widgets.QFrame, "StaffingV2HistoryDetailCard")
    assert first_history_card.y() < 160
    assert first_history_card.width() <= history_detail_scroll.viewport().width()
    assert history_detail_scroll.widget().findChild(qt_widgets.QPushButton, "StaffingV2HistoryExportRecord") is None
    history_detail_close = closed_detail_panel.findChild(qt_widgets.QPushButton, "StaffingV2HistoryDetailClose")
    assert history_detail_close is not None
    assert history_detail_close.text() == ""
    assert not history_detail_close.icon().isNull()
    history_detail_footer = closed_detail_panel.findChild(qt_widgets.QWidget, "StaffingV2HistoryDetailPanelFooter")
    assert history_detail_footer is not None
    assert history_detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2HistoryDetailClose") is None
    closed_detail = _widget_text(closed_detail_panel)
    assert f"Assignment ID: A-{closed_id:04d}" in closed_detail
    closed_chip = closed_detail_panel.findChild(qt_widgets.QFrame, "StaffingV2HistoryAssignmentIdChip")
    assert closed_chip is not None
    assert _widget_text(closed_chip) == f"A-{closed_id:04d}"
    for detail_card in closed_detail_panel.findChildren(qt_widgets.QFrame, "StaffingV2HistoryDetailCard"):
        assert detail_card.sizePolicy().verticalPolicy() == qt_widgets.QSizePolicy.Policy.Maximum
    open_employee = closed_detail_panel.findChild(qt_widgets.QPushButton, "StaffingV2HistoryOpenEmployee")
    assert open_employee is not None
    assert open_employee.isEnabled()
    export_record = closed_detail_panel.findChild(qt_widgets.QPushButton, "StaffingV2HistoryExportRecord")
    assert export_record is not None
    assert export_record.isEnabled()
    assert history_detail_footer.findChild(qt_widgets.QPushButton, "StaffingV2HistoryExportRecord") is export_record
    export_record.click()
    app.processEvents()
    export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2HistoryExportRecordDialog")
    assert export_dialog is not None
    export_text = _widget_text(export_dialog)
    assert "Export Record" in export_text
    assert f"A-{closed_id:04d}" in export_text
    assert "Harmony 1" in export_text
    assert "Emily Carter" in export_text
    export_dialog.accept()
    app.processEvents()
    open_employee.click()
    app.processEvents()
    assert page.findChild(qt_widgets.QWidget, "StaffingV2PeopleDashboard") is not None
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2PeopleNavButton").property("staffingV2ActiveNav") is True
    people_detail = page.findChild(qt_widgets.QFrame, "StaffingV2PeopleDetailPanel")
    people_detail_text = _widget_text(people_detail)
    assert "Emily Carter" in people_detail_text
    assert "Teacher" in people_detail_text
    page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryNavButton").click()
    app.processEvents()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2HistoryTable")
    open_history_row = next(
        row for row in range(table.rowCount()) if table.item(row, 0).text() == f"A-{open_id:04d}"
    )
    open_view = table.cellWidget(open_history_row, 9)
    assert isinstance(open_view, qt_widgets.QPushButton)
    open_view.click()
    app.processEvents()
    assert page.findChild(qt_widgets.QLabel, "StaffingV2HistoryResultCount").text() == "Showing 1 to 2 of 2 records"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2HistoryRowsPerPage").currentText() == "10 / page"
    detail = page.findChild(qt_widgets.QFrame, "StaffingV2HistoryDetailPanel")
    detail_text = _widget_text(detail)
    assert f"Assignment ID: A-{open_id:04d}" in detail_text
    assignment_id_chip = detail.findChild(qt_widgets.QFrame, "StaffingV2HistoryAssignmentIdChip")
    assert assignment_id_chip is not None
    assert _widget_text(assignment_id_chip) == f"A-{open_id:04d}"
    assert "Quest" in detail_text
    assert "Lifecycle Events" in detail_text
    assert "Position opened" in detail_text
    lifecycle_rows = page.findChildren(qt_widgets.QFrame, "StaffingV2HistoryLifecycleEventRow")
    assert lifecycle_rows
    assert any(row.property("staffingV2LifecycleEventType") == "opened" for row in lifecycle_rows)
    assert "Validation / Integrity" in detail_text
    assert "Duplicate active cycle" in detail_text
    validation_rows = page.findChildren(qt_widgets.QFrame, "StaffingV2HistoryValidationCheckRow")
    assert len(validation_rows) >= 3
    validation_statuses = {row.property("staffingV2ValidationCheckStatus") for row in validation_rows}
    assert {"pass", "warning"} <= validation_statuses
    view_assignment = page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryViewAssignment")
    assert view_assignment.text() == "View Assignment"
    assert view_assignment.isEnabled()
    assert not view_assignment.icon().isNull()
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryOpenEmployee").icon().isNull()
    assert not page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryExportRecord").icon().isNull()
    view_assignment.click()
    app.processEvents()
    assert page.findChild(qt_widgets.QWidget, "StaffingV2Dashboard") is not None
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2DashboardNavButton").property("staffingV2ActiveNav") is True
    assert page.findChild(qt_widgets.QLabel, "StaffingV2DrawerPositionName").text() == "Teacher 2"
    page.findChild(qt_widgets.QPushButton, "StaffingV2DrawerViewHistory").click()
    app.processEvents()
    history_validation_button = page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryValidationButton")
    history_validation_button.click()
    app.processEvents()
    assert page.findChild(qt_widgets.QWidget, "StaffingV2ValidationDashboard") is not None
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationNavButton").property("staffingV2ActiveNav") is True
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2HistoryNavButton").property("staffingV2ActiveNav") is False
    assert len(store.list_assignment_history()) == len(before_records)
    assert len(store.list_assignments()) == len(before_assignments)
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_validation_dashboard_and_filter_drawer_use_existing_staffing_data(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    db_path = tmp_path / "staffing.sqlite3"
    store = pyside_interview_app.StaffingStore(db_path)
    store.initialize()
    old_open_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Harmony 1",
        position_name="Teacher 1",
        position_type="Teacher",
        status="need_now",
        permit_status="unknown",
    )
    coming_id = store.seed_assignment(
        school="Hawthorne",
        classroom="Tranquility",
        position_name="Teacher 2",
        position_type="Teacher",
        status="coming",
        person_name="James Mitchell",
        permit_status="permit_in_process",
    )
    store.seed_assignment(
        school="Hawthorne",
        classroom="Unity 1",
        position_name="Aide 1",
        position_type="Aide",
        status="filled",
        person_name="Sofia Ramirez",
        permit_status="teacher_permit_approved",
    )
    with store.connect() as conn:
        conn.execute(
            "UPDATE assignments SET current_opened_date = '2026-06-01T00:00:00Z' WHERE id = ?",
            (old_open_id,),
        )
        conn.execute(
            "UPDATE assignments SET start_date = '' WHERE id = ?",
            (coming_id,),
        )
    before_assignments = store.list_assignments()
    before_history = store.list_assignment_history()
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing_seed.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()

    validation_nav = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationNavButton")
    assert validation_nav.isEnabled()
    validation_nav.click()
    app.processEvents()

    assert page.findChild(qt_widgets.QWidget, "StaffingV2ValidationDashboard") is not None
    assert validation_nav.property("staffingV2ActiveNav") is True
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ValidationTitle").text() == "Staffing Validation"
    assert "Review staffing compliance" in page.findChild(qt_widgets.QLabel, "StaffingV2ValidationSubtitle").text()
    validation_export = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationExportButton")
    assert validation_export.text() == "Export Report"
    assert not validation_export.icon().isNull()
    assert validation_export.isEnabled()
    validation_export.click()
    app.processEvents()
    validation_export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ValidationExportDialog")
    assert validation_export_dialog is not None
    validation_export_text = _widget_text(validation_export_dialog)
    assert "Export Validation Report" in validation_export_text
    assert "Total issues 2" in validation_export_text
    assert "Critical 1" in validation_export_text
    assert "Unfilled Need Now position" in validation_export_text
    assert "Coming position missing start date" in validation_export_text
    validation_export_dialog.accept()
    app.processEvents()
    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2ValidationMetricCard"))
    assert "Total Issues 2" in metric_text
    assert "Critical 1" in metric_text
    assert "Warning 1" in metric_text
    assert "Info 0" in metric_text
    assert "Overall Compliance" in metric_text
    validation_tabs = page.findChild(qt_widgets.QFrame, "StaffingV2ValidationTabs")
    assert validation_tabs is not None
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationAllIssuesTab").text() == "All Issues (2)"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationCriticalTab").text() == "Critical (1)"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationWarningsTab").text() == "Warnings (1)"
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationInfoTab").text() == "Info (0)"
    all_issues_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationAllIssuesTab")
    critical_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationCriticalTab")
    warning_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationWarningsTab")
    info_tab = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationInfoTab")
    assert all_issues_tab.property("staffingV2ActiveValidationTab") is True
    assert critical_tab.property("staffingV2ActiveValidationTab") is False
    assert warning_tab.property("staffingV2ActiveValidationTab") is False
    assert info_tab.property("staffingV2ActiveValidationTab") is False
    validation_search = page.findChild(qt_widgets.QLineEdit, "StaffingV2ValidationSearch")
    assert validation_search.placeholderText() == "Search issues..."
    assert validation_search.actions()
    filters = page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationFiltersButton")
    assert filters.text() == "Filters 3"
    assert filters.property("staffingV2FilterActiveCount") == 3
    assert not filters.icon().isNull()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ValidationTable")
    validation_table_width = table.width()
    assert table.rowCount() == 2
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert {
        "Unfilled Need Now position",
        "Coming position missing start date",
        "Harmony 1",
        "Tranquility",
        "Critical",
        "Warning",
    } <= table_text
    for row in range(table.rowCount()):
        view_button = table.cellWidget(row, 6)
        assert isinstance(view_button, qt_widgets.QPushButton)
        assert view_button.text() == "View"
        assert not view_button.icon().isNull()
        assert view_button.isEnabled()
    critical_issue_row = next(
        row for row in range(table.rowCount()) if table.item(row, 0).text() == "Unfilled Need Now position"
    )
    critical_view = table.cellWidget(critical_issue_row, 6)
    assert isinstance(critical_view, qt_widgets.QPushButton)
    critical_view.click()
    app.processEvents()
    drawer = page.findChild(qt_widgets.QFrame, "StaffingV2PositionDrawer")
    assert drawer is not None
    assert drawer.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2Dashboard")
    assert table.width() == validation_table_width
    drawer_scroll = drawer.findChild(qt_widgets.QScrollArea, "StaffingV2PositionDrawerScroll")
    assert drawer_scroll is not None
    assert drawer_scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    assert drawer_scroll.widget().findChild(qt_widgets.QPushButton, "StaffingV2DrawerSaveChanges") is None
    drawer_footer = drawer.findChild(qt_widgets.QWidget, "StaffingV2PositionDrawerFooter")
    assert drawer_footer is not None
    assert drawer_footer.findChild(qt_widgets.QPushButton, "StaffingV2DrawerSaveChanges") is None
    assert drawer.findChild(qt_widgets.QLabel, "StaffingV2DrawerPositionName").text() == "Teacher 1"
    drawer_text = _widget_text(drawer)
    assert "Harmony 1" in drawer_text
    assert "Need Now" in drawer_text
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ValidationResultCount").text() == "Showing 1 to 2 of 2 issues"
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2ValidationRowsPerPage").currentText() == "10 / page"
    critical_tab.click()
    app.processEvents()
    assert table.rowCount() == 1
    assert table.item(0, 3).text() == "Critical"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ValidationResultCount").text() == "Showing 1 to 1 of 1 issues"
    assert all_issues_tab.property("staffingV2ActiveValidationTab") is False
    assert critical_tab.property("staffingV2ActiveValidationTab") is True
    all_issues_tab.click()
    app.processEvents()
    assert table.rowCount() == 2
    assert all_issues_tab.property("staffingV2ActiveValidationTab") is True
    assert critical_tab.property("staffingV2ActiveValidationTab") is False
    right_panel = page.findChild(qt_widgets.QFrame, "StaffingV2ValidationRightPanel")
    assert right_panel.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2ValidationDashboard")
    right_scroll = right_panel.findChild(qt_widgets.QScrollArea, "StaffingV2ValidationRightPanelScroll")
    assert right_scroll is not None
    assert right_scroll.verticalScrollBarPolicy() == qt_core.Qt.ScrollBarPolicy.ScrollBarAlwaysOn
    right_panel_text = _widget_text(right_panel)
    assert "Compliance Summary" in right_panel_text
    assert "Quick Actions" in right_panel_text
    quick_actions = {
        "StaffingV2ValidationRunFullButton": "Run Full Validation",
        "StaffingV2ValidationExportQuickButton": "Export Validation Report",
        "StaffingV2ValidationRulesButton": "View Validation Rules",
    }
    for object_name, text in quick_actions.items():
        action_button = page.findChild(qt_widgets.QPushButton, object_name)
        assert action_button is not None
        assert action_button.text() == text
        assert not action_button.icon().isNull()
        assert action_button.isEnabled()
    page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationExportQuickButton").click()
    app.processEvents()
    quick_export_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ValidationExportDialog")
    assert quick_export_dialog is not None
    assert "Export Validation Report" in _widget_text(quick_export_dialog)
    quick_export_dialog.accept()
    app.processEvents()
    page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationRulesButton").click()
    app.processEvents()
    rules_dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2ValidationRulesDialog")
    assert rules_dialog is not None
    rules_text = _widget_text(rules_dialog)
    assert "Validation Rules" in rules_text
    assert "Coverage" in rules_text
    assert "Permit Status" in rules_text
    assert "Upcoming Start Dates" in rules_text
    rules_dialog.accept()
    app.processEvents()
    with store.connect() as conn:
        conn.execute(
            "UPDATE assignments SET start_date = '2026-07-20' WHERE id = ?",
            (coming_id,),
        )
    page.findChild(qt_widgets.QPushButton, "StaffingV2ValidationRunFullButton").click()
    app.processEvents()
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2ValidationTable")
    refreshed_issues = {
        table.item(row, 0).text()
        for row in range(table.rowCount())
        if table.item(row, 0) is not None
    }
    assert "Coming position missing start date" not in refreshed_issues
    assert refreshed_issues == {"Unfilled Need Now position"}
    assert page.findChild(qt_widgets.QLabel, "StaffingV2ValidationResultCount").text() == "Showing 1 to 1 of 1 issues"
    assert "About Validation" in right_panel_text

    drawer = page.findChild(qt_widgets.QFrame, "StaffingV2FilterDrawer")
    assert drawer is not None
    assert drawer.isHidden()
    filters.click()
    app.processEvents()
    assert not drawer.isHidden()
    assert drawer.parent() is page.findChild(qt_widgets.QWidget, "StaffingV2ValidationDashboard")
    assert table.width() == validation_table_width
    filter_scroll = drawer.findChild(qt_widgets.QScrollArea, "StaffingV2FilterDrawerScroll")
    assert filter_scroll is not None
    assert filter_scroll.widget().findChild(qt_widgets.QPushButton, "StaffingV2FilterApplyButton") is None
    filter_footer = drawer.findChild(qt_widgets.QWidget, "StaffingV2FilterDrawerFooter")
    assert filter_footer is not None
    drawer_text = _widget_text(drawer)
    assert "Filters" in drawer_text
    assert "School" in drawer_text
    assert "Severity" in drawer_text
    assert "Issue Type" in drawer_text
    assert "Detected Date" in drawer_text
    assert page.findChild(qt_widgets.QComboBox, "StaffingV2FilterDetectedDate").currentText() == "Last 30 Days"
    reset_button = page.findChild(qt_widgets.QPushButton, "StaffingV2FilterResetButton")
    assert reset_button.text() == "Reset"
    assert not reset_button.icon().isNull()
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2FilterCloseButton").text() == ""
    apply_button = page.findChild(qt_widgets.QPushButton, "StaffingV2FilterApplyButton")
    assert apply_button.text() == "Apply Filters 3"
    assert apply_button.property("staffingV2FilterActiveCount") == 3
    assert not apply_button.icon().isNull()
    assert filter_footer.findChild(qt_widgets.QPushButton, "StaffingV2FilterApplyButton") is apply_button
    assert page.findChild(qt_widgets.QPushButton, "StaffingV2FilterCancelButton").text() == "Cancel"
    page.findChild(qt_widgets.QPushButton, "StaffingV2FilterCancelButton").click()
    app.processEvents()
    assert drawer.isHidden()
    assert len(store.list_assignments()) == len(before_assignments)
    assert len(store.list_assignment_history()) == len(before_history)
    window.window.close()
    app.processEvents()


def test_pyside_staffing_open_action_refreshes_dashboard(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "dont_need_now"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    table = window.window.findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")
    button = _staffing_button_for_position(table, "Teacher 1")

    assert button.text() == "Open"
    button.click()
    app.processEvents()

    table = window.window.findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")
    labels = [label.text() for label in window.stack.currentWidget().findChildren(qt_widgets.QLabel)]
    row = _staffing_row_for_position(table, "Teacher 1")
    assert table.item(row, 3).text() == "need_now"
    assert table.cellWidget(row, 7).text() == "Mark Coming"
    assert any("Open positions: 1" in text for text in labels)
    window.window.close()
    app.processEvents()

def test_pyside_director_staffing_mode_uses_same_staffing_page_only(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "dont_need_now"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    director_model = pyside_interview_app.build_director_staffing_model(full_model)
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    page = window.stack.widget(0)
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    action_button = table.cellWidget(0, 7)

    assert director_model.navigation == ["Staffing v2"]
    assert isinstance(window.staffing_v2_dashboard, StaffingDashboardV2Page)
    assert window.staffing_v2_dashboard.school_filter == ""
    assert window.window.windowTitle() == "Director Staffing Dashboard"
    staffing_sidebar = page.findChild(qt_widgets.QFrame, "StaffingV2Sidebar")

    assert window.sidebar_panel.isHidden()
    assert staffing_sidebar is not None
    assert "Launch Pad Learning" in _widget_text(staffing_sidebar)
    assert window.sidebar.currentRow() == 0
    assert window.stack.count() == 1
    assert page.objectName() == "PySideStaffingV2Page"
    assert page.findChild(qt_widgets.QLabel, "StaffingV2PageTitle").text() == "Staffing Dashboard"
    assert table is not None
    assert action_button.text() == "Mark Need Now"
    window.window.close()
    app.processEvents()

def test_pyside_director_staffing_mode_filters_to_assigned_school(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Hawthorne Teacher", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Harmony",
                                "positions": [
                                    {"position_name": "Palmdale Teacher", "position_type": "Teacher", "status": "need_now"},
                                    {"position_name": "Palmdale Aide", "position_type": "Aide", "status": "filled", "person": {"name": "Koryn"}},
                                ],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "Palmdale"],
    )

    director_model = pyside_interview_app.build_director_staffing_model(full_model, school="Palmdale")
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    page = window.stack.widget(0)
    selector = page.findChild(qt_widgets.QComboBox, "StaffingV2SchoolFilter")
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2MetricCard"))
    board_text = _widget_text(page)

    assert director_model.director_staffing_school == "Palmdale"
    assert isinstance(window.staffing_v2_dashboard, StaffingDashboardV2Page)
    assert window.staffing_v2_dashboard.school_filter == "Palmdale"
    assert [selector.itemText(index) for index in range(selector.count())] == ["Palmdale"]
    assert selector.currentText() == "Palmdale"
    assert table.rowCount() == 2
    assert "Harmony" in board_text
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert "Palmdale Teacher" in table_text
    assert "OPEN POSITION" in table_text
    assert "Koryn" in table_text
    assert "Tranquility" not in board_text
    assert "Open positions: 1" in metric_text
    window.window.close()
    app.processEvents()

def test_staffing_db_path_for_school_uses_safe_school_specific_names(tmp_path: Path) -> None:
    base_path = tmp_path / "staffing_dashboard.sqlite3"

    assert pyside_interview_app.staffing_db_path_for_school("Hawthorne", base_path=base_path) == (
        tmp_path / "staffing_dashboard_hawthorne.sqlite3"
    )
    assert pyside_interview_app.staffing_db_path_for_school("Long Beach / Bixby", base_path=base_path) == (
        tmp_path / "staffing_dashboard_long_beach_bixby.sqlite3"
    )
    assert pyside_interview_app.staffing_db_path_for_school("", base_path=base_path) == base_path

def test_school_specific_staffing_db_bootstraps_from_existing_base_db(tmp_path: Path) -> None:
    base_path = tmp_path / "staffing_dashboard.sqlite3"
    base_store = pyside_interview_app.StaffingStore(base_path)
    base_store.initialize()
    base_store.seed_assignment(
        school="Palmdale",
        classroom="Harmony",
        position_name="Existing Palmdale Teacher",
        position_type="Teacher",
        status="need_now",
    )
    school_path = pyside_interview_app.staffing_db_path_for_school("Palmdale", base_path=base_path)

    pyside_interview_app._bootstrap_school_staffing_db_from_base("Palmdale", school_path, base_path=base_path)

    school_store = pyside_interview_app.StaffingStore(school_path)
    assignments = school_store.list_assignments()
    assert school_path.exists()
    assert [assignment.position_name for assignment in assignments] == ["Existing Palmdale Teacher"]


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_director_staffing_mode_uses_school_specific_db_when_other_school_locked(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [{"position_name": "Hawthorne Teacher", "position_type": "Teacher"}],
                            }
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Harmony",
                                "positions": [{"position_name": "Palmdale Teacher", "position_type": "Teacher"}],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    base_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", base_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    hawthorne_path = pyside_interview_app.staffing_db_path_for_school("Hawthorne", base_path=base_path)
    hawthorne_path.parent.mkdir(parents=True, exist_ok=True)
    hawthorne_path.with_suffix(hawthorne_path.suffix + ".editing.lock").write_text(
        '{"owner": "hawthorne-director", "created_at": "2099-01-01T00:00:00Z"}',
        encoding="utf-8",
    )
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "Palmdale"],
    )

    director_model = pyside_interview_app.build_director_staffing_model(full_model, school="Palmdale")
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")

    assert window.staffing_store.path == pyside_interview_app.staffing_db_path_for_school("Palmdale", base_path=base_path)
    assert table.rowCount() == 1
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert "Palmdale Teacher" in table_text
    assert "Hawthorne Teacher" not in _widget_text(window.window)
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
def test_interview_finalize_queues_director_referral_without_staffing_db_write(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    staffing_path = tmp_path / "staffing.sqlite3"
    referral_queue_path = tmp_path / "staffing_referrals.pending.jsonl"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_REFERRAL_QUEUE_PATH", referral_queue_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    window.session = SimpleNamespace(
        candidate_name="Queued Candidate",
        school="Palmdale",
        position="Teacher",
        interview_date="2026-07-08",
    )
    staffing_bytes_before = staffing_path.read_bytes()

    window._record_staffing_director_referral_from_finalize_result(
        {
            "history_id": "hist-queued",
            "scoring": {"outcome": "Hire", "interviewer_rating": 8.8},
        }
    )

    assert staffing_path.read_bytes() == staffing_bytes_before
    assert pyside_interview_app.staffing_referral_queue_db_path(queue_path=referral_queue_path).exists()
    queued = pyside_interview_app._pop_staffing_referral_queue_for_school("Palmdale", queue_path=referral_queue_path)
    assert queued[0]["history_id"] == "hist-queued"
    assert queued[0]["school"] == "Palmdale"
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_director_staffing_poll_imports_queued_referral_and_refreshes_gui(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps({"schools": [{"name": "Palmdale", "classrooms": [{"name": "Harmony", "positions": []}]}]}),
        encoding="utf-8",
    )
    base_path = tmp_path / "staffing.sqlite3"
    referral_queue_path = tmp_path / "staffing_referrals.pending.jsonl"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", base_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_REFERRAL_QUEUE_PATH", referral_queue_path)
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Palmdale"],
    )
    director_model = pyside_interview_app.build_director_staffing_model(full_model, school="Palmdale")
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")
    assert table.rowCount() == 0
    pyside_interview_app._append_staffing_referral_queue(
        {
            "history_id": "hist-live",
            "candidate_name": "Live Queue Candidate",
            "school": "Palmdale",
            "position": "Teacher",
            "interviewer_rating": 8.5,
            "interviewer_outcome": "hire",
            "interview_date": "2026-07-09",
            "candidate_email": "",
            "referral_date": "2026-07-09",
        }
    )

    window._poll_staffing_referral_queue()
    app.processEvents()

    assert table.rowCount() == 1
    assert table.item(0, 0).text() == "Live Queue Candidate"
    assert not referral_queue_path.exists()
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_director_staffing_poll_imports_review_score_dismissal_and_removes_pending_referral(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    staffing_path = tmp_path / "staffing.sqlite3"
    referral_queue_path = tmp_path / "staffing_referrals.pending.jsonl"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_REFERRAL_QUEUE_PATH", referral_queue_path)
    school_store = pyside_interview_app.StaffingStore(
        pyside_interview_app.staffing_db_path_for_school("Palmdale", base_path=staffing_path)
    )
    school_store.initialize()
    pyside_interview_app.StaffingService(school_store).upsert_director_candidate_referral(
        history_id="hist-review-remove",
        candidate_name="Remove Candidate",
        school="Palmdale",
        position="Teacher",
        interviewer_rating=7.0,
        interviewer_outcome="borderline",
        interview_date="2026-07-10",
    )
    full_model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Palmdale"],
    )
    director_model = pyside_interview_app.build_director_staffing_model(full_model, school="Palmdale")
    window = pyside_interview_app.PySideInterviewWindow(director_model)
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")
    assert table.rowCount() == 1

    pyside_interview_app._append_staffing_referral_dismissal_queue(
        {
            "history_id": "hist-review-remove",
            "candidate_name": "Remove Candidate",
            "school": "Palmdale",
            "position": "Teacher",
        }
    )
    window._poll_staffing_referral_queue()
    app.processEvents()

    assert table.rowCount() == 0
    assert "hist-review-remove" in school_store.list_dismissed_director_referral_history_ids()
    window.window.close()
    app.processEvents()

def test_pyside_staffing_v2_summary_cards_follow_selected_school(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Hawthorne Teacher", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Harmony",
                                "positions": [
                                    {"position_name": "Palmdale Teacher", "position_type": "Teacher", "status": "need_now"},
                                    {"position_name": "Palmdale Aide", "position_type": "Aide", "status": "need_now"},
                                ],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    selector = page.findChild(qt_widgets.QComboBox, "StaffingV2SchoolFilter")

    selector.setCurrentText("Hawthorne")
    app.processEvents()

    metric_text = " ".join(_widget_text(card) for card in page.findChildren(qt_widgets.QFrame, "StaffingV2MetricCard"))
    assert "Schools: 1" in metric_text
    assert "Open positions: 1" in metric_text
    assert "Open positions: 3" not in metric_text
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_pyside_staffing_v2_director_candidates_follow_admin_school_selector(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony",
                                "positions": [
                                    {
                                        "position_name": "Hawthorne Teacher",
                                        "position_type": "Teacher",
                                        "status": "need_now",
                                    }
                                ],
                            }
                        ],
                    },
                    {
                        "name": "North Long Beach",
                        "classrooms": [
                            {
                                "name": "Unity",
                                "positions": [
                                    {
                                        "position_name": "North Long Beach Teacher",
                                        "position_type": "Teacher",
                                        "status": "need_now",
                                    }
                                ],
                            }
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Quest",
                                "positions": [
                                    {
                                        "position_name": "Palmdale Teacher",
                                        "position_type": "Teacher",
                                        "status": "need_now",
                                    }
                                ],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    staffing_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    store = pyside_interview_app.StaffingStore(staffing_path)
    store.initialize()
    store.import_seed_file(seed_path)
    service = pyside_interview_app.StaffingService(store)
    for school, candidate in [
        ("Hawthorne", "Hawthorne Candidate"),
        ("North Long Beach", "North Long Beach Candidate"),
        ("Palmdale", "Palmdale Candidate"),
    ]:
        service.upsert_director_candidate_referral(
            history_id=f"hist-{school.lower().replace(' ', '-')}",
            candidate_name=candidate,
            school=school,
            position="Teacher",
            interviewer_rating=8.0,
            interviewer_outcome="hire",
            interview_date="2026-07-09",
        )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "North Long Beach", "Palmdale"],
    )
    window = pyside_interview_app.PySideInterviewWindow(model, defer_secondary_pages=True)
    nav_items = [window.sidebar.item(index).text() for index in range(window.sidebar.count())]
    window.sidebar.setCurrentRow(nav_items.index("Staffing v2"))
    app.processEvents()
    page = window.stack.currentWidget()
    selector = page.findChild(qt_widgets.QComboBox, "StaffingV2SchoolFilter")
    table = page.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")

    selector.setCurrentText("Hawthorne")
    app.processEvents()
    assert [table.item(row, 0).text() for row in range(table.rowCount())] == ["Hawthorne Candidate"]
    assert [table.horizontalHeaderItem(column).text() for column in range(table.columnCount())] == [
        "Candidate",
        "Outcome",
        "Score",
        "Interview\nDate",
        "Role",
        "Referral\nDate",
        "Action",
    ]
    assert not table.horizontalHeader().stretchLastSection()
    assert table.horizontalHeader().minimumHeight() >= 54
    assert table.columnWidth(1) == 90
    assert table.columnWidth(2) == 64
    assert table.columnWidth(3) == 100
    assert table.columnWidth(4) == 140
    assert table.columnWidth(6) >= 156
    record_button = table.cellWidget(0, 6)
    assert record_button is not None
    assert record_button.minimumWidth() >= 144
    assert table.columnWidth(6) >= record_button.minimumWidth()

    selector.setCurrentText("North Long Beach")
    app.processEvents()
    assert [table.item(row, 0).text() for row in range(table.rowCount())] == ["North Long Beach Candidate"]

    selector.setCurrentText("Palmdale")
    app.processEvents()
    assert [table.item(row, 0).text() for row in range(table.rowCount())] == ["Palmdale Candidate"]
    window.window.close()
    app.processEvents()

def test_staffing_v2_director_pending_table_uses_compact_readable_columns(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps({"schools": [{"name": "Hawthorne", "classrooms": [{"name": "Harmony", "positions": []}]}]}),
        encoding="utf-8",
    )
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    store.import_seed_file(seed_path)
    service = pyside_interview_app.StaffingService(store)
    service.upsert_director_candidate_referral(
        history_id="hist-compact-columns",
        candidate_name="Adrianna Love",
        school="Hawthorne",
        position="assistant_director_enrollment_specialist",
        interviewer_rating=9.75,
        interviewer_outcome="hire",
        interview_date="2026-07-02",
    )

    page = StaffingDashboardV2Page(
        QtCore=qt_core,
        QtGui=qt_gui,
        QtWidgets=qt_widgets,
        store=store,
        service_factory=lambda: service,
        school_filter="Hawthorne",
        notification_store_path=tmp_path / "notification_rules.sqlite3",
    )
    table = page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")
    status = page.widget.findChild(qt_widgets.QLabel, "StaffingV2DirectorInterviewStatus")
    delete_selected = page.widget.findChild(qt_widgets.QPushButton, "StaffingV2DirectorInterviewDeleteSelected")

    assert [table.item(0, column).text() for column in range(1, 5)] == [
        "Hire",
        "9.75",
        "2026-07-02",
        "Assistant Director",
    ]
    role_badge = table.cellWidget(0, 4)
    assert role_badge is not None
    assert role_badge.objectName() == "StaffingV2RoleBadge"
    assert role_badge.accessibleName() == "Role: Assistant Director"
    assert role_badge.findChild(qt_widgets.QLabel, "StaffingV2RoleBadgeText").text() == "Assistant Director"
    assert [table.columnWidth(column) for column in range(7)] == [210, 90, 64, 100, 140, 105, 160]
    assert sum(table.columnWidth(column) for column in range(table.columnCount())) <= 880
    assert status.text() == "1 pending / 0 completed"
    assert not status.wordWrap()
    assert status.minimumWidth() >= 170
    assert delete_selected.minimumWidth() >= 170
    assert table.item(0, 4).toolTip() == "Assistant Director"
    page.widget.close()
    app.processEvents()


def test_staffing_v2_initial_refresh_automatically_fills_due_coming_position(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Palmdale",
        classroom="Harmony",
        position_name="Teacher 1",
        position_type="Teacher",
    )
    service = pyside_interview_app.StaffingService(store, clock=lambda: "2026-07-15T08:00:00Z")
    service.open_position(assignment_id)
    service.mark_coming(
        assignment_id,
        person_name="Koryn",
        start_date=date.today().isoformat(),
    )

    page = StaffingDashboardV2Page(
        QtCore=qt_core,
        QtGui=qt_gui,
        QtWidgets=qt_widgets,
        store=store,
        service_factory=lambda: service,
        school_filter="Palmdale",
        notification_store_path=tmp_path / "notification_rules.sqlite3",
    )

    assignment = store.get_assignment(assignment_id)
    assert assignment.status == "filled"
    assert assignment.current_filled_date == date.today().isoformat()
    table = page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    assert _staffing_button_for_position(table, "Teacher 1").text() == "Manage Filled"
    page.widget.close()
    app.processEvents()


def test_staffing_v2_date_rollover_automatically_fills_due_coming_position(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    today = date.today()
    tomorrow = today + timedelta(days=1)
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Palmdale",
        classroom="Harmony",
        position_name="Teacher 1",
        position_type="Teacher",
    )
    service = pyside_interview_app.StaffingService(store, clock=lambda: "2026-07-15T08:00:00Z")
    service.open_position(assignment_id)
    service.mark_coming(
        assignment_id,
        person_name="Koryn",
        start_date=tomorrow.isoformat(),
    )
    page = StaffingDashboardV2Page(
        QtCore=qt_core,
        QtGui=qt_gui,
        QtWidgets=qt_widgets,
        store=store,
        service_factory=lambda: service,
        school_filter="Palmdale",
        notification_store_path=tmp_path / "notification_rules.sqlite3",
    )
    assert store.get_assignment(assignment_id).status == "coming"

    class TomorrowDate(date):
        @classmethod
        def today(cls) -> date:
            return tomorrow

    monkeypatch.setattr(sys.modules["staffing_dashboard_v2"], "date", TomorrowDate)
    timer = page.widget.findChild(qt_core.QTimer, "StaffingV2AutomaticFillTimer")
    assert timer is not None
    assert timer.isActive()
    timer.timeout.emit()
    app.processEvents()

    assert store.get_assignment(assignment_id).status == "filled"
    page.widget.close()
    app.processEvents()


def test_staffing_v2_edit_position_reverts_auto_filled_employee_when_delayed(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    original_start = date.today()
    delayed_start = original_start + timedelta(days=3)
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    assignment_id = store.seed_assignment(
        school="Palmdale",
        classroom="Harmony",
        position_name="Teacher 1",
        position_type="Teacher",
    )
    service = pyside_interview_app.StaffingService(store, clock=lambda: "2026-07-15T08:00:00Z")
    service.open_position(assignment_id)
    service.mark_coming(
        assignment_id,
        person_name="Koryn",
        start_date=original_start.isoformat(),
    )
    page = StaffingDashboardV2Page(
        QtCore=qt_core,
        QtGui=qt_gui,
        QtWidgets=qt_widgets,
        store=store,
        service_factory=lambda: service,
        school_filter="Palmdale",
        notification_store_path=tmp_path / "notification_rules.sqlite3",
    )
    assert store.get_assignment(assignment_id).status == "filled"
    table = page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2PositionsTable")
    action_button = _staffing_button_for_position(table, "Teacher 1")
    next(action for action in action_button.menu().actions() if action.text() == "View Details").trigger()
    app.processEvents()
    page.widget.findChild(qt_widgets.QPushButton, "StaffingV2DrawerEditPosition").click()
    app.processEvents()

    dialog = page.widget.findChild(qt_widgets.QDialog, "StaffingV2EditPositionDialog")
    assert dialog is not None
    start_date = dialog.findChild(qt_widgets.QDateEdit, "StaffingV2EditPositionStartDate")
    assert start_date is not None
    assert start_date.isEnabled()
    assert start_date.calendarPopup()
    assert start_date.date() == qt_core.QDate.fromString(original_start.isoformat(), "yyyy-MM-dd")
    start_date.setDate(qt_core.QDate.fromString(delayed_start.isoformat(), "yyyy-MM-dd"))
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2EditPositionSubmit").click()
    app.processEvents()

    assignment = store.get_assignment(assignment_id)
    assert assignment.status == "coming"
    assert assignment.start_date == delayed_start.isoformat()
    assert store.active_history_count(assignment_id) == 1
    page.widget.close()
    app.processEvents()


def test_staffing_v2_page_defers_hidden_subdashboard_work(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps({"schools": [{"name": "Hawthorne", "classrooms": [{"name": "Harmony", "positions": []}]}]}),
        encoding="utf-8",
    )
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    store.import_seed_file(seed_path)
    service = pyside_interview_app.StaffingService(store)

    monkeypatch.setattr(store, "list_people", lambda: pytest.fail("people loaded during initial Staffing v2 refresh"))
    monkeypatch.setattr(store, "list_assignment_history", lambda: pytest.fail("history loaded during initial Staffing v2 refresh"))
    monkeypatch.setattr(store, "list_classrooms", lambda: pytest.fail("classrooms loaded during initial Staffing v2 refresh"))

    page = StaffingDashboardV2Page(
        QtCore=qt_core,
        QtGui=qt_gui,
        QtWidgets=qt_widgets,
        store=store,
        service_factory=lambda: service,
        school_filter="Hawthorne",
        notification_store_path=tmp_path / "notification_rules.sqlite3",
    )

    assert page.widget.findChild(qt_widgets.QWidget, "StaffingV2ClassroomManagementDashboard") is None
    assert page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2PeopleTable") is None
    assert page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2HistoryTable") is None
    assert page.widget.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList") is None
    assert page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2ValidationTable") is None
    page.widget.close()
    app.processEvents()


def test_staffing_v2_lazy_views_build_once_on_navigation(tmp_path: Path) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_gui = pytest.importorskip("PySide6.QtGui")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps({"schools": [{"name": "Hawthorne", "classrooms": [{"name": "Harmony", "positions": []}]}]}),
        encoding="utf-8",
    )
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    store.import_seed_file(seed_path)
    service = pyside_interview_app.StaffingService(store)
    page = StaffingDashboardV2Page(
        QtCore=qt_core,
        QtGui=qt_gui,
        QtWidgets=qt_widgets,
        store=store,
        service_factory=lambda: service,
        school_filter="Hawthorne",
        notification_store_path=tmp_path / "notification_rules.sqlite3",
    )
    built: dict[str, int] = {}

    for name, method_name, button_name in [
        ("classrooms", "_build_classrooms_view", "StaffingV2ClassroomsNavButton"),
        ("people", "_build_people_view", "StaffingV2PeopleNavButton"),
        ("history", "_build_history_view", "StaffingV2HistoryNavButton"),
        ("notifications", "_build_notifications_view", "StaffingV2NotificationsNavButton"),
        ("validation", "_build_validation_view", "StaffingV2ValidationNavButton"),
    ]:
        original = getattr(page, method_name)

        def counted(original=original, name=name):
            built[name] = built.get(name, 0) + 1
            return original()

        setattr(page, method_name, counted)
        button = page.widget.findChild(qt_widgets.QPushButton, button_name)
        button.click()
        button.click()
        app.processEvents()

    assert built == {"classrooms": 1, "people": 1, "history": 1, "notifications": 1, "validation": 1}
    assert page.widget.findChild(qt_widgets.QWidget, "StaffingV2ClassroomManagementDashboard") is not None
    assert page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2PeopleTable") is not None
    assert page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2HistoryTable") is not None
    assert page.widget.findChild(qt_widgets.QListWidget, "StaffingV2NotificationsRuleList") is not None
    assert page.widget.findChild(qt_widgets.QTableWidget, "StaffingV2ValidationTable") is not None
    page.widget.close()
    app.processEvents()

def test_pyside_staffing_action_button_exposes_secondary_actions(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "coming", "start_date": "2026-07-03", "person": {"name": "Jane Doe"}},
                                    {"position_name": "Teacher 2", "position_type": "Teacher", "status": "replace", "person": {"name": "Angie"}},
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    table = window.window.findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")

    filled_menu_labels = [action.text() for action in _staffing_button_for_position(table, "Teacher 1").menu().actions()]
    replace_menu_labels = [action.text() for action in _staffing_button_for_position(table, "Teacher 2").menu().actions()]

    assert "Update Permit" in filled_menu_labels
    assert "Mark Not Needed" in filled_menu_labels
    assert "Clear Replacement" in replace_menu_labels
    assert "Mark Not Needed" in replace_menu_labels
    window.window.close()
    app.processEvents()

def test_pyside_staffing_dashboard_groups_by_school_and_colors_statuses(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {"name": "Hawthorne", "classrooms": [{"name": "Tranquility", "positions": [{"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}]}]},
                    {"name": "Palmdale", "classrooms": [{"name": "Harmony", "positions": [{"position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Angie", "permit_status": "teacher_permit_approved"}}]}]},
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "Palmdale"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    tabs = window.window.findChild(qt_widgets.QTabWidget, "PySideStaffingSchoolTabs")
    first_table = tabs.widget(0).findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")
    second_table = tabs.widget(1).findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")

    assert [tabs.tabText(index) for index in range(tabs.count())] == ["Hawthorne", "Palmdale"]
    assert "Need Now: 1" in tabs.widget(0).findChild(qt_widgets.QLabel, "PySideStaffingSummary").text()
    assert first_table.item(_staffing_row_for_position(first_table, "Teacher 1"), 3).background().color().isValid()
    second_row = _staffing_row_for_position(second_table, "Teacher 1")
    assert second_table.item(second_row, 2).text() == "Angie"
    assert second_table.item(second_row, 2).background().color().isValid()
    window.window.close()
    app.processEvents()

def test_pyside_staffing_dashboard_renders_workbook_layout_tabs_and_actual_names(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "display_order": 1,
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "ratio_group": "3 to 1 (infant units needed)",
                                "licensed_capacity": 12,
                                "slots": [
                                    {"slot_group": "teacher", "position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Angie", "permit_status": "teacher_permit_approved"}},
                                    {"slot_group": "teacher", "position_name": "Teacher 2", "position_type": "Teacher", "status": "need_now", "notes": "?"},
                                ],
                            }
                        ],
                        "support_rows": [
                            {"name": "Infant Floater", "slots": [{"slot_group": "support", "position_name": "Infant Floater", "position_type": "Support", "status": "filled", "person": {"name": "Amy"}, "notes": "Full time"}]}
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "display_order": 2,
                        "classrooms": [
                            {
                                "name": "Harmony",
                                "ratio_group": "4 to 1",
                                "licensed_capacity": 24,
                                "slots": [
                                    {"slot_group": "teacher", "position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Madisan"}},
                                    {"slot_group": "aide", "position_name": "Aide 1", "position_type": "Aide", "status": "coming", "person": {"name": "Koryn"}},
                                ],
                            }
                        ],
                    },
                    {
                        "name": "North Long Beach",
                        "display_order": 3,
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "ratio_group": "3 to 1 (infant units needed)",
                                "licensed_capacity": 16,
                                "slots": [
                                    {"slot_group": "teacher", "position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Naomi*"}},
                                    {"slot_group": "aide", "position_name": "Aide 1", "position_type": "Aide", "status": "filled", "person": {"name": "Ruby"}},
                                ],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    tabs = window.window.findChild(qt_widgets.QTabWidget, "PySideStaffingSchoolTabs")
    selector = window.window.findChild(qt_widgets.QComboBox, "PySideStaffingSchoolSelector")
    hawthorne = tabs.widget(0)
    board = hawthorne.findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")
    labels = [label.text() for label in hawthorne.findChildren(qt_widgets.QLabel)]

    assert [tabs.tabText(index) for index in range(tabs.count())] == ["Hawthorne", "Palmdale", "North Long Beach"]
    assert [selector.itemText(index) for index in range(selector.count())] == ["Hawthorne", "Palmdale", "North Long Beach"]
    selector.setCurrentIndex(2)
    app.processEvents()
    assert tabs.currentIndex() == 2
    assert board is not None
    board_text = {
        board.item(row, column).text()
        for row in range(board.rowCount())
        for column in range(board.columnCount())
        if board.item(row, column) is not None
    }
    assert {"Tranquility", "Angie", "OPEN POSITION", "12", "Infant Floater", "Amy", "Full time"} <= board_text
    assert "need_now" in board_text
    assert any("3 to 1 (infant units needed)" in text for text in labels)
    window.window.close()
    app.processEvents()

def test_pyside_staffing_dashboard_refreshes_partial_seed_and_hides_color_key(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "slots": [
                                    {"slot_group": "teacher", "position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Angie"}},
                                    {"slot_group": "teacher", "position_name": "Teacher 2", "position_type": "Teacher", "status": "need_now"},
                                ],
                            },
                            {
                                "name": "Harmony",
                                "slots": [
                                    {"slot_group": "teacher", "position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Brenda"}}
                                ],
                            },
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Unity",
                                "slots": [
                                    {"slot_group": "teacher", "position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Cora"}}
                                ],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    db_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)

    partial_store = pyside_interview_app.StaffingStore(db_path)
    partial_store.initialize()
    partial_store.seed_assignment(
        school="Hawthorne",
        classroom="Tranquility",
        position_name="Teacher 1",
        position_type="Teacher",
        status="filled",
        person_name="Angie",
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    tabs = window.window.findChild(qt_widgets.QTabWidget, "PySideStaffingSchoolTabs")
    selector = window.window.findChild(qt_widgets.QComboBox, "PySideStaffingSchoolSelector")
    labels = [label.text() for label in window.stack.currentWidget().findChildren(qt_widgets.QLabel)]
    hawthorne_board = tabs.widget(0).findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")

    assert [tabs.tabText(index) for index in range(tabs.count())] == ["Hawthorne", "Palmdale"]
    assert [selector.itemText(index) for index in range(selector.count())] == ["Hawthorne", "Palmdale"]
    board_text = {
        hawthorne_board.item(row, column).text()
        for row in range(hawthorne_board.rowCount())
        for column in range(hawthorne_board.columnCount())
        if hawthorne_board.item(row, column) is not None
    }
    assert {"Tranquility", "Harmony", "Angie", "Brenda", "OPEN POSITION"} <= board_text
    assert "Color Code Key" not in labels
    assert "Need Now - Job Opening" not in labels
    window.window.close()
    app.processEvents()

def test_pyside_staffing_classroom_detail_matches_dashboard_mockup_shell(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "ratio_group": "Preschool",
                                "licensed_capacity": 18,
                                "slots": [
                                    {
                                        "slot_group": "teacher",
                                        "position_name": "Teacher 1",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Imgard M.", "permit_status": "teacher_permit_approved"},
                                        "start_date": "2025-03-03",
                                    },
                                    {
                                        "slot_group": "teacher",
                                        "position_name": "Teacher 2",
                                        "position_type": "Teacher",
                                        "status": "replace",
                                        "person": {"name": "Angie R.", "permit_status": "permit_in_process"},
                                        "start_date": "2025-05-01",
                                    },
                                    {
                                        "slot_group": "teacher",
                                        "position_name": "Teacher 3",
                                        "position_type": "Teacher",
                                        "status": "coming",
                                        "person": {"name": "Denise A.", "permit_status": "teacher_permit_approved"},
                                        "start_date": "2025-05-19",
                                    },
                                    {
                                        "slot_group": "aide",
                                        "position_name": "Aide 1",
                                        "position_type": "Aide",
                                        "status": "need_now",
                                        "current_opened_date": "2025-05-14",
                                    },
                                ],
                            },
                            {
                                "name": "Quest",
                                "ratio_group": "Preschool",
                                "licensed_capacity": 16,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Sarah M."}}
                                ],
                            },
                            {
                                "name": "Unity",
                                "ratio_group": "Preschool",
                                "licensed_capacity": 18,
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "replace", "person": {"name": "Mia"}},
                                ],
                            },
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    page = window.stack.currentWidget()
    school_selector = page.findChild(qt_widgets.QComboBox, "PySideStaffingSchoolSelector")
    classroom_selector = page.findChild(qt_widgets.QComboBox, "PySideStaffingClassroomSelector")
    classroom_list = page.findChild(qt_widgets.QListWidget, "PySideStaffingClassroomList")
    section_splitter = page.findChild(qt_widgets.QSplitter, "PySideStaffingSectionSplitter")
    cards = page.findChildren(qt_widgets.QFrame, "PySideStaffingMetricCard")
    table = page.findChild(qt_widgets.QTableWidget, "PySideStaffingPositionsTable")
    title = page.findChild(qt_widgets.QLabel, "PySideStaffingClassroomTitle")
    priority = page.findChild(qt_widgets.QLabel, "PySideStaffingPriorityBadge")

    assert window.stack.currentWidget().findChild(qt_widgets.QLabel, "Title").text() == "Classroom Detail"
    assert school_selector.currentText() == "Hawthorne"
    assert classroom_selector.currentText() == "Harmony 1"
    assert [classroom_selector.itemText(index) for index in range(classroom_selector.count())] == ["Harmony 1", "Quest", "Unity"]
    assert [classroom_list.item(index).text() for index in range(classroom_list.count())] == [
        "Harmony 1\nNeed: 1 - Replace: 1 - Filled: 2 - Don't Need: 0",
        "Quest\nNeed: 0 - Replace: 0 - Filled: 1 - Don't Need: 0",
        "Unity\nNeed: 0 - Replace: 1 - Filled: 0 - Don't Need: 0",
    ]
    assert section_splitter is not None
    assert section_splitter.count() == 3
    assert not section_splitter.childrenCollapsible()
    assert section_splitter.widget(0).maximumWidth() > 320
    assert section_splitter.widget(2).minimumWidth() == 300
    assert classroom_list.minimumWidth() >= 360
    assert classroom_list.wordWrap()
    assert classroom_list.item(0).background().color().name().upper() == "#FEF08A"
    assert classroom_list.item(1).background().color().name().upper() == "#BBF7D0"
    assert classroom_list.item(2).background().color().name().upper() == "#FF0000"
    assert title.text() == "Harmony 1"
    assert priority.text() == "Need Now"
    assert len(cards) == 6
    assert [
        table.horizontalHeaderItem(index).text()
        for index in range(table.columnCount())
    ] == ["Position", "Person", "Status", "Start Date", "Days Open", "Permit Status", "Action"]
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert {"Teacher 1", "Imgard M.", "Filled", "Teacher Permit Approved", "Aide 1", "OPEN POSITION", "Need Now"} <= table_text
    assert _staffing_button_for_position(table, "Aide 1").text() == "Mark Coming"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_filters_button_focuses_filter_controls(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    filters = window.window.findChild(qt_widgets.QPushButton, "PySideStaffingFiltersButton")
    school_selector = window.window.findChild(qt_widgets.QComboBox, "PySideStaffingSchoolSelector")

    filters.click()
    app.processEvents()

    assert school_selector.currentText() == "Hawthorne"
    assert window.staffing_status_label.text() == "Use the school and classroom controls to filter staffing rows."
    window.window.close()
    app.processEvents()


def test_pyside_staffing_add_position_button_creates_position_and_opens_drawer(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "filled"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(qt_widgets.QInputDialog, "getText", lambda *_args, **_kwargs: ("Float Teacher", True))
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    add_position = window.window.findChild(qt_widgets.QPushButton, "PySideStaffingAddPositionButton")

    add_position.click()
    app.processEvents()

    table = window.window.findChild(qt_widgets.QTableWidget, "PySideStaffingPositionsTable")
    table_text = {
        table.item(row, column).text()
        for row in range(table.rowCount())
        for column in range(table.columnCount())
        if table.item(row, column) is not None
    }
    assert "Float Teacher" in table_text
    assert "Need Now" in table_text
    drawer_text = _widget_text(window.staffing_detail_drawer)
    assert "Position Details" in drawer_text
    assert "Float Teacher" in drawer_text
    assert window.staffing_status_label.text() == "Position added."
    window.window.close()
    app.processEvents()


def test_pyside_staffing_row_click_opens_mockup_detail_drawer(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Quest",
                                "slots": [
                                    {
                                        "position_name": "Teacher 1",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Sarah M.", "permit_status": "teacher_permit_approved"},
                                        "start_date": "2025-01-15",
                                    },
                                    {
                                        "position_name": "Teacher 3",
                                        "position_type": "Teacher",
                                        "status": "need_now",
                                        "current_opened_date": "2025-05-14",
                                    },
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    table = window.staffing_positions_table
    drawer = window.staffing_detail_drawer

    table.cellClicked.emit(_staffing_row_for_position(table, "Teacher 3"), 0)
    app.processEvents()

    drawer_labels = [label.text() for label in drawer.findChildren(qt_widgets.QLabel)]
    drawer_buttons = [button.text() for button in drawer.findChildren(qt_widgets.QPushButton)]
    assert not drawer.isHidden()
    assert "Position Details" in drawer_labels
    assert "OPEN POSITION" in drawer_labels
    assert "Teacher 3" in drawer_labels
    assert "Need Now" in drawer_labels
    assert {"Edit", "Save", "Mark Coming", "Don't Need Now", "Close"} <= set(drawer_buttons)

    position_type = drawer.findChild(qt_widgets.QComboBox, "PySideStaffingDetailPositionType")
    status = drawer.findChild(qt_widgets.QComboBox, "PySideStaffingDetailStatus")
    program = drawer.findChild(qt_widgets.QComboBox, "PySideStaffingDetailProgram")
    start_date = drawer.findChild(qt_widgets.QDateEdit, "PySideStaffingDetailStartDate")
    shift_start = drawer.findChild(qt_widgets.QTimeEdit, "PySideStaffingDetailShiftStart")
    notes = drawer.findChild(qt_widgets.QTextEdit, "PySideStaffingDetailNotes")
    save = drawer.findChild(qt_widgets.QPushButton, "PySideStaffingDetailSave")
    assert position_type.currentText() == "Teacher"
    assert status.currentText() == "Need Now"
    assert start_date.specialValueText() == "-"
    assert shift_start.displayFormat() == "h:mm AP"

    position_type.setCurrentText("Aide")
    program.setCurrentText("Toddler")
    notes.setPlainText("Needs bilingual coverage.")
    save.click()
    app.processEvents()

    updated = window.staffing_store.get_assignment(
        next(row.id for row in window.staffing_store.list_assignments() if row.position_name == "Teacher 3")
    )
    assert updated.position_type == "Aide"
    assert updated.classroom_program == "Toddler"
    assert updated.notes == "Needs bilingual coverage."

    drawer = window.staffing_detail_drawer
    filled_id = next(row.id for row in window.staffing_store.list_assignments() if row.position_name == "Teacher 1")
    window._open_staffing_assignment_details(filled_id)
    app.processEvents()

    drawer_labels = [label.text() for label in drawer.findChildren(qt_widgets.QLabel)]
    drawer_buttons = [button.text() for button in drawer.findChildren(qt_widgets.QPushButton)]
    assert "Person Details" in drawer_labels
    assert "Sarah M." in drawer_labels
    assert "Teacher Permit Approved" in drawer_labels
    assert {"Replace", "Edit", "Save", "Close"} <= set(drawer_buttons)

    person_name = drawer.findChild(qt_widgets.QLineEdit, "PySideStaffingDetailPersonName")
    position_name = drawer.findChild(qt_widgets.QLineEdit, "PySideStaffingDetailPositionName")
    permit = drawer.findChild(qt_widgets.QComboBox, "PySideStaffingDetailPermitStatus")
    filled_notes = drawer.findChild(qt_widgets.QTextEdit, "PySideStaffingDetailNotes")
    filled_save = drawer.findChild(qt_widgets.QPushButton, "PySideStaffingDetailSave")
    person_name.setText("Sara M.")
    position_name.setText("Lead Teacher")
    permit.setCurrentText("permit_in_process")
    filled_notes.setPlainText("Moved to lead slot.")
    filled_save.click()
    app.processEvents()

    filled = window.staffing_store.get_assignment(filled_id)
    assert filled.person_name == "Sara M."
    assert filled.position_name == "Lead Teacher"
    assert filled.permit_status == "permit_in_process"
    assert filled.notes == "Moved to lead slot."
    window.window.close()
    app.processEvents()

def test_pyside_mark_coming_uses_guided_dialog_and_saves_position(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "slots": [
                                    {
                                        "position_name": "Aide 1",
                                        "position_type": "Aide",
                                        "status": "need_now",
                                        "current_opened_date": "2026-07-01",
                                    }
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    assignment_id = next(row.id for row in window.staffing_store.list_assignments() if row.position_name == "Aide 1")
    monkeypatch.setattr(
        window.QtWidgets.QInputDialog,
        "getText",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(AssertionError("legacy prompt")),
    )

    window._mark_staffing_coming(assignment_id)
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "PySideStaffingMarkComingDialog")
    assert dialog is not None
    assert dialog.isVisible()
    assert {
        label.text() for label in dialog.findChildren(qt_widgets.QLabel) if label.text()
    } >= {
        "Mark Coming",
        "Teacher Name *",
        "Start Date *",
        "Role *",
        "Permit Status *",
        "Units",
    }
    dialog.findChild(qt_widgets.QLineEdit, "PySideStaffingComingName").setText("Samantha Lee")
    dialog.findChild(qt_widgets.QLineEdit, "PySideStaffingComingStartDate").setText("2026-08-01")
    dialog.findChild(qt_widgets.QPushButton, "PySideStaffingComingSave").click()
    app.processEvents()

    updated = window.staffing_store.get_assignment(assignment_id)
    assert updated.status == "coming"
    assert updated.person_name == "Samantha Lee"
    assert updated.start_date == "2026-08-01"
    window.window.close()
    app.processEvents()

def test_pyside_replace_employee_uses_guided_dialog_and_saves_replacement(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_test = pytest.importorskip("PySide6.QtTest")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Destiny",
                                "slots": [
                                    {
                                        "position_name": "Teacher 2",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Angie R.", "permit_status": "permit_in_process"},
                                    }
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Palmdale"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    assignment_id = next(row.id for row in window.staffing_store.list_assignments() if row.position_name == "Teacher 2")
    monkeypatch.setattr(
        window.QtWidgets.QInputDialog,
        "getText",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(AssertionError("legacy prompt")),
    )

    window._mark_staffing_replacing(assignment_id)
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "PySideStaffingReplaceDialog")
    assert dialog is not None
    assert dialog.isVisible()
    assert {
        label.text() for label in dialog.findChildren(qt_widgets.QLabel) if label.text()
    } >= {"Replace Employee", "Current Employee", "Notice Given *", "Final Working Day *", "Reason (optional)"}
    notice_field = dialog.findChild(qt_widgets.QDateEdit, "PySideStaffingReplaceNotice")
    final_day_field = dialog.findChild(qt_widgets.QDateEdit, "PySideStaffingReplaceFinalDay")
    reason_field = dialog.findChild(qt_widgets.QComboBox, "PySideStaffingReplaceReason")
    assert notice_field is not None
    assert final_day_field is not None
    assert notice_field.calendarPopup()
    assert final_day_field.calendarPopup()
    assert notice_field.date() == qt_core.QDate.currentDate()
    assert final_day_field.date() == qt_core.QDate.currentDate()
    assert reason_field is not None
    assert [reason_field.itemText(index) for index in range(reason_field.count())] == [
        "Resignation",
        "Termination",
        "Leave of absence",
        "Transfer",
        "Other",
    ]
    save_button = dialog.findChild(qt_widgets.QPushButton, "PySideStaffingReplaceSave")
    qt_test.QTest.mouseClick(notice_field, qt_core.Qt.MouseButton.LeftButton, pos=notice_field.rect().center())
    app.processEvents()
    assert notice_field.calendarWidget().isVisible()
    calendar_parent = notice_field.calendarWidget().parentWidget()
    if calendar_parent is not None:
        calendar_parent.hide()
    assert window.staffing_store.get_assignment(assignment_id).status == "filled"

    notice_field.setDate(qt_core.QDate(2026, 8, 1))
    final_day_field.setDate(qt_core.QDate(2026, 8, 15))
    reason_field.setCurrentText("Transfer")
    save_button.click()
    app.processEvents()

    updated = window.staffing_store.get_assignment(assignment_id)
    assert updated.status == "replace"
    assert updated.person_name == "Angie R."
    assert updated.notice_given == "2026-08-01"
    assert updated.final_working_day == "2026-08-15"
    window.window.close()
    app.processEvents()

def test_pyside_update_permit_uses_guided_dialog_and_saves_people_status(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "slots": [
                                    {
                                        "position_name": "Teacher 3",
                                        "position_type": "Teacher",
                                        "status": "filled",
                                        "person": {"name": "Denise A.", "permit_status": "permit_in_process"},
                                    }
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    assignment_id = next(row.id for row in window.staffing_store.list_assignments() if row.position_name == "Teacher 3")
    monkeypatch.setattr(
        window.QtWidgets.QInputDialog,
        "getItem",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(AssertionError("legacy prompt")),
    )

    window._update_staffing_permit(assignment_id)
    app.processEvents()

    dialog = window.window.findChild(qt_widgets.QDialog, "PySideStaffingPermitDialog")
    assert dialog is not None
    assert dialog.isVisible()
    assert {
        label.text() for label in dialog.findChildren(qt_widgets.QLabel) if label.text()
    } >= {"Update Permit Status", "Permit Status", "Units", "Effective Date"}
    permit_combo = dialog.findChild(qt_widgets.QComboBox, "PySideStaffingPermitStatus")
    permit_combo.setCurrentText("teacher_permit_approved")
    dialog.findChild(qt_widgets.QPushButton, "PySideStaffingPermitSave").click()
    app.processEvents()

    updated = window.staffing_store.get_assignment(assignment_id)
    assert updated.permit_status == "teacher_permit_approved"
    window.window.close()
    app.processEvents()

@pytest.mark.pyside_gui
@pytest.mark.visual_inspection
def test_pyside_staffing_dashboard_real_seed_visual_scenario(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    visual_test_databases: VisualTestDatabaseRegistry,
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    configure_visual_test_app(app)
    repo_root = Path(__file__).resolve().parents[1]
    seed_path = repo_root / "config" / "staffing_seed.json"
    seed_data = json.loads(seed_path.read_text(encoding="utf-8"))
    expected_schools = [school["name"] for school in seed_data["schools"]]
    expected_assignment_count = sum(
        len(classroom.get("slots", classroom.get("positions", [])))
        for school in seed_data["schools"]
        for classroom in school.get("classrooms", [])
    ) + sum(
        len(row.get("slots", row.get("positions", [])))
        for school in seed_data["schools"]
        for row in school.get("support_rows", [])
    )
    db_path = visual_test_databases.database("staffing_dashboard.sqlite3")
    visual_test_databases.expect_seeded(db_path, table="assignments")
    partial_store = pyside_interview_app.StaffingStore(db_path)
    partial_store.initialize()
    partial_store.seed_assignment(
        school="Hawthorne",
        classroom="Tranquility",
        position_name="Teacher 1",
        position_type="Teacher",
        status="filled",
        person_name="Angie",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", db_path)
    monkeypatch.chdir(repo_root / "src")
    history_path = visual_test_databases.database("staffing_interview_history.sqlite3")
    visual_test_databases.expect_seeded(history_path, table="interview_history")
    InterviewHistoryStore(history_path).append(
        {
            "history_id": "staffing-visual-fixture",
            "candidate_name": "Ghjypq Avery",
            "school": "Hawthorne",
            "position": "Teacher",
            "interview_date": "2026-07-15",
            "outcome": "Hire",
            "score": "100.0%",
        }
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=history_path,
        school_options=["Hawthorne"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    window.window.resize(1600, 1100)
    window.window.show()
    app.processEvents()
    tabs = window.window.findChild(qt_widgets.QTabWidget, "PySideStaffingSchoolTabs")
    selector = window.window.findChild(qt_widgets.QComboBox, "PySideStaffingSchoolSelector")
    labels = [label.text() for label in window.stack.currentWidget().findChildren(qt_widgets.QLabel)]
    rendered = window.stack.currentWidget().grab()
    screenshot_path = tmp_path / "staffing_dashboard_visual.png"
    assert rendered.save(str(screenshot_path))

    assert [tabs.tabText(index) for index in range(tabs.count())] == expected_schools
    assert [selector.itemText(index) for index in range(selector.count())] == expected_schools
    assert len(partial_store.list_assignments()) == expected_assignment_count
    assert "Color Code Key" not in labels
    assert rendered.width() >= 600
    assert rendered.height() >= 400
    image = rendered.toImage()
    sample_points = [
        image.pixelColor(x, y).name()
        for x in range(0, image.width(), max(1, image.width() // 8))
        for y in range(0, image.height(), max(1, image.height() // 8))
    ]
    assert len(set(sample_points)) > 3
    window.window.close()
    app.processEvents()

def test_pyside_staffing_action_surfaces_exact_service_error(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", tmp_path / "missing.json")
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")

    window._run_staffing_action(lambda service: service.open_position(999), "unused")

    assert "Assignment not found." in window.staffing_status_label.text()
    window.window.close()
    app.processEvents()

def test_pyside_staffing_actions_use_notification_service(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "dont_need_now"}
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    notifications = []

    class FakeNotifications:
        def emit_event(self, event_type, payload, idempotency_key):
            notifications.append((event_type, payload, idempotency_key))
            return []

    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    monkeypatch.setattr(
        pyside_interview_app,
        "notification_service_from_email_account_settings",
        lambda **_kwargs: FakeNotifications(),
    )
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    table = window.window.findChild(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")

    _staffing_button_for_position(table, "Teacher 1").click()
    app.processEvents()

    assert notifications
    assert notifications[0][0] == "staffing.assignment.need_now"
    window.window.close()
    app.processEvents()

def test_pyside_staffing_uses_single_draggable_colored_workbook_table(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    qt_core = pytest.importorskip("PySide6.QtCore")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {"name": "Hawthorne", "classrooms": [{"name": "Tranquility", "positions": [{"position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Angie"}}]}]},
                    {"name": "Palmdale", "classrooms": [{"name": "Harmony", "positions": [{"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}]}]},
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne", "Palmdale"],
    )

    window = _pyside_window_on_page(model, "Staffing")
    tabs = window.window.findChild(qt_widgets.QTabWidget, "PySideStaffingSchoolTabs")
    workbook_tables = window.window.findChildren(qt_widgets.QTableWidget, "PySideStaffingWorkbookBoard")
    flat_tables = window.window.findChildren(qt_widgets.QTableWidget, "PySideStaffingAssignments")

    assert len(workbook_tables) == 2
    assert flat_tables == []
    assert all(table.dragEnabled() for table in workbook_tables)
    assert all(table.acceptDrops() for table in workbook_tables)
    assert workbook_tables[0].dragDropMode() == qt_widgets.QAbstractItemView.DragDropMode.DragDrop
    assert [
        workbook_tables[0].horizontalHeaderItem(index).text()
        for index in range(workbook_tables[0].columnCount())
    ] == ["Ratio", "Classroom", "Person", "Status", "Capacity", "Permit Status", "Details", "Action"]
    assert isinstance(workbook_tables[0].cellWidget(_staffing_row_for_position(workbook_tables[0], "Teacher 1"), 5), qt_widgets.QComboBox)
    assert tabs.tabBar().tabTextColor(0) != tabs.tabBar().tabTextColor(1)
    assert tabs.tabBar().tabTextColor(0).isValid()
    assert workbook_tables[0].item(_staffing_row_for_position(workbook_tables[0], "Teacher 1"), 4).flags() & qt_core.Qt.ItemFlag.ItemIsDragEnabled
    window.window.close()
    app.processEvents()

def test_pyside_staffing_confirm_move_updates_source_and_target(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Tranquility",
                                "positions": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "filled", "person": {"name": "Angie"}},
                                    {"position_name": "Teacher 2", "position_type": "Teacher", "status": "need_now"},
                                ],
                            }
                        ],
                    }
                ]
            }
        ),
        encoding="utf-8",
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    model = build_interview_redesign_model(
        rubric_path=_write_test_rubric(tmp_path),
        overrides_path=_write_test_overrides(tmp_path),
        history_path=tmp_path / "interview_history.sqlite3",
        school_options=["Hawthorne"],
    )
    window = _pyside_window_on_page(model, "Staffing")
    yes = window.QtWidgets.QMessageBox.StandardButton.Yes
    monkeypatch.setattr(window.QtWidgets.QMessageBox, "question", lambda *_args, **_kwargs: yes)
    assignments = window.staffing_store.list_assignments()
    source_id = next(row.id for row in assignments if row.position_name == "Teacher 1")
    target_id = next(row.id for row in assignments if row.position_name == "Teacher 2")

    assert window._confirm_staffing_move(source_id, target_id) is True

    source = window.staffing_store.get_assignment(source_id)
    target = window.staffing_store.get_assignment(target_id)
    assert source.status == "need_now"
    assert source.person_name == ""
    assert target.status == "filled"
    assert target.person_name == "Angie"
    window.window.close()
    app.processEvents()

def test_staffing_v2_director_interviews_sync_pending_history_and_record_completion(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    seed_path = tmp_path / "staffing_seed.json"
    seed_path.write_text(
        json.dumps(
            {
                "schools": [
                    {
                        "name": "Hawthorne",
                        "classrooms": [
                            {
                                "name": "Harmony 1",
                                "slots": [
                                    {"position_name": "Teacher 2", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    },
                    {
                        "name": "Palmdale",
                        "classrooms": [
                            {
                                "name": "Destiny",
                                "slots": [
                                    {"position_name": "Teacher 1", "position_type": "Teacher", "status": "need_now"}
                                ],
                            }
                        ],
                    },
                ]
            }
        ),
        encoding="utf-8",
    )
    history_path = tmp_path / "interview_history.sqlite3"
    history_store = InterviewHistoryStore(history_path)
    history_store.append(
        {
            "history_id": "hist-old-hawthorne",
            "candidate_name": "Past Candidate",
            "candidate_email": "past@example.org",
            "school": "Hawthorne",
            "position": "Teacher",
            "interview_date": "2026-07-05",
            "outcome": "Hire",
            "score": "8.5",
        }
    )
    history_store.append(
        {
            "history_id": "hist-palmdale",
            "candidate_name": "Riley Park",
            "school": "Palmdale",
            "position": "Teacher",
            "interview_date": "2026-07-05",
            "outcome": "Borderline",
            "score": "7.5",
        }
    )
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", tmp_path / "staffing.sqlite3")
    monkeypatch.setattr(pyside_interview_app, "STAFFING_SEED_PATH", seed_path)
    store = pyside_interview_app.StaffingStore(tmp_path / "staffing.sqlite3")
    store.initialize()
    pyside_interview_app.StaffingService(store).upsert_director_candidate_referral(
        history_id="new-hawthorne",
        candidate_name="Jordan Lee",
        school="Hawthorne",
        position="Teacher",
        interviewer_rating=8.5,
        interviewer_outcome="hire",
        interview_date="2026-07-07",
    )
    model = pyside_interview_app.build_director_staffing_model(
        build_interview_redesign_model(
            rubric_path=_write_test_rubric(tmp_path),
            overrides_path=_write_test_overrides(tmp_path),
            history_path=history_path,
            school_options=["Hawthorne", "Palmdale"],
        ),
        school="Hawthorne",
    )

    window = pyside_interview_app.PySideInterviewWindow(model)
    panel = window.window.findChild(qt_widgets.QFrame, "StaffingV2DirectorInterviewPanel")
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")

    assert panel is not None
    assert table.rowCount() == 1
    assert table.item(0, 0).text() == "Jordan Lee"
    rendered_candidates = {
        table.item(row, 0).text()
        for row in range(table.rowCount())
        if table.item(row, 0) is not None
    }
    assert "Past Candidate" not in rendered_candidates
    assert "Riley Park" not in rendered_candidates

    table.cellWidget(0, table.columnCount() - 1).click()
    app.processEvents()
    dialog = window.window.findChild(qt_widgets.QDialog, "StaffingV2DirectorInterviewDialog")
    assert dialog is not None
    dialog.findChild(qt_widgets.QLineEdit, "StaffingV2DirectorInterviewDirectorName").setText("Avery Director")
    dialog.findChild(qt_widgets.QDoubleSpinBox, "StaffingV2DirectorInterviewRating").setValue(9.0)
    decision = dialog.findChild(qt_widgets.QComboBox, "StaffingV2DirectorInterviewDecision")
    shift_start = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2DirectorInterviewShiftStartText")
    shift_end = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2DirectorInterviewShiftEndText")
    classroom = dialog.findChild(qt_widgets.QComboBox, "StaffingV2DirectorInterviewClassroom")
    candidate_email = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2DirectorInterviewCandidateEmail")
    candidate_phone = dialog.findChild(qt_widgets.QLineEdit, "StaffingV2DirectorInterviewCandidatePhone")
    assert candidate_email is not None
    assert candidate_phone is not None
    decision.setCurrentText("No-Hire")
    app.processEvents()
    assert not shift_start.isVisible()
    assert not shift_end.isVisible()
    assert not classroom.isVisible()
    assert not candidate_email.isVisible()
    assert not candidate_phone.isVisible()
    decision.setCurrentText("Hire")
    app.processEvents()
    assert shift_start.isVisible()
    assert shift_end.isVisible()
    assert classroom.isVisible()
    assert candidate_email.isVisible()
    assert candidate_phone.isVisible()
    shift_start.setText("8:00 AM")
    shift_end.setText("5:00 PM")
    classroom.setCurrentText("Harmony 1")
    candidate_email.setText("jordan@example.org")
    dialog.findChild(qt_widgets.QTextEdit, "StaffingV2DirectorInterviewNotes").setPlainText("Strong classroom presence.")
    dialog.findChild(qt_widgets.QPushButton, "StaffingV2DirectorInterviewSave").click()
    app.processEvents()

    assert table.rowCount() == 0
    history_table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewHistoryTable")
    assert history_table.rowCount() == 1
    assert history_table.item(0, 0).text() == ""
    assert history_table.cellWidget(0, 0).text() == "Jordan Lee"
    assignment = next(row for row in window.staffing_store.list_assignments() if row.school == "Hawthorne")
    assert assignment.status == "need_now"
    assert assignment.person_name == ""
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_staffing_v2_director_interviews_backfill_passed_history_rows(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_core = pytest.importorskip("PySide6.QtCore")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    history_store = InterviewHistoryStore(history_path)
    history_store.append(
        {
            "history_id": "hist-hire",
            "candidate_name": "Hire Candidate",
            "candidate_email": "hire@example.org",
            "school": "Palmdale",
            "position": "Teacher",
            "interview_date": "2026-07-08",
            "outcome": "Hire",
            "score": "88%",
        }
    )
    history_store.append(
        {
            "history_id": "hist-borderline",
            "candidate_name": "Borderline Candidate",
            "school": "Palmdale",
            "position": "Assistant Teacher",
            "interview_date": "2026-07-08",
            "outcome": "Borderline",
            "score": "72%",
        }
    )
    history_store.append(
        {
            "history_id": "hist-no-hire",
            "candidate_name": "No Hire Candidate",
            "school": "Palmdale",
            "position": "Teacher",
            "interview_date": "2026-07-08",
            "outcome": "No Hire",
            "score": "42%",
        }
    )
    history_store.append(
        {
            "history_id": "hist-other-school",
            "candidate_name": "Other School Candidate",
            "school": "Hawthorne",
            "position": "Teacher",
            "interview_date": "2026-07-08",
            "outcome": "Hire",
            "score": "90%",
        }
    )
    staffing_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    store = pyside_interview_app.StaffingStore(
        pyside_interview_app.staffing_db_path_for_school("Palmdale", base_path=staffing_path)
    )
    store.initialize()
    model = pyside_interview_app.build_director_staffing_model(
        build_interview_redesign_model(
            rubric_path=_write_test_rubric(tmp_path),
            overrides_path=_write_test_overrides(tmp_path),
            history_path=history_path,
            school_options=["Palmdale", "Hawthorne"],
        ),
        school="Palmdale",
    )

    window = pyside_interview_app.PySideInterviewWindow(model)
    window._sync_staffing_v2_director_referrals_after_first_paint()
    app.processEvents()
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")
    delete_selected = window.window.findChild(qt_widgets.QPushButton, "StaffingV2DirectorInterviewDeleteSelected")

    rendered_candidates = [
        table.item(row, 0).text()
        for row in range(table.rowCount())
        if table.item(row, 0) is not None
    ]
    assert rendered_candidates == ["Borderline Candidate", "Hire Candidate"]
    monkeypatch.setattr(
        qt_widgets.QMessageBox,
        "question",
        lambda *_args, **_kwargs: qt_widgets.QMessageBox.StandardButton.Yes,
    )
    selection_model = table.selectionModel()
    for row in range(table.rowCount()):
        selection_model.select(
            table.model().index(row, 0),
            qt_core.QItemSelectionModel.SelectionFlag.Select | qt_core.QItemSelectionModel.SelectionFlag.Rows,
        )
    delete_selected.click()
    app.processEvents()

    assert table.rowCount() == 0
    dismissed_history_ids = store.list_dismissed_director_referral_history_ids()
    assert {"hist-borderline", "hist-hire"} <= dismissed_history_ids
    assert "hist-no-hire" not in dismissed_history_ids
    assert "hist-other-school" not in dismissed_history_ids
    assert len(InterviewHistoryStore(history_path).load()) == 4
    window.window.close()
    app.processEvents()


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_director_staffing_launch_queues_history_backfill_when_edit_lock_exists(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    history_store = InterviewHistoryStore(history_path)
    history_store.append(
        {
            "history_id": "hist-locked",
            "candidate_name": "Locked Candidate",
            "school": "Palmdale",
            "position": "Teacher",
            "interview_date": "2026-07-08",
            "outcome": "Hire",
            "score": "88%",
        }
    )
    staffing_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    school_staffing_path = pyside_interview_app.staffing_db_path_for_school("Palmdale", base_path=staffing_path)
    store = pyside_interview_app.StaffingStore(school_staffing_path)
    store.initialize()
    lock_path = school_staffing_path.with_suffix(school_staffing_path.suffix + ".editing.lock")
    lock_path.write_text(
        json.dumps(
            {
                "owner": "other-director",
                "created_at": "2099-01-01T00:00:00Z",
                "database": str(school_staffing_path),
            }
        ),
        encoding="utf-8",
    )
    model = pyside_interview_app.build_director_staffing_model(
        build_interview_redesign_model(
            rubric_path=_write_test_rubric(tmp_path),
            overrides_path=_write_test_overrides(tmp_path),
            history_path=history_path,
            school_options=["Palmdale"],
        ),
        school="Palmdale",
    )

    window = pyside_interview_app.PySideInterviewWindow(model)
    window._sync_staffing_v2_director_referrals_after_first_paint()
    app.processEvents()
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")

    assert table is not None
    assert table.rowCount() == 0
    window.window.close()
    app.processEvents()
    assert store.pending_operations_path.exists()
    lock_path.unlink()
    assert pyside_interview_app.StaffingService(store).flush_pending_operations() == 1
    pending = store.list_director_candidate_referrals(school="Palmdale")
    assert [candidate.candidate_name for candidate in pending] == ["Locked Candidate"]


@pytest.mark.pyside_gui
@pytest.mark.slow_pyside
def test_staffing_v2_director_interviews_delete_checked_pending_rows(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")
    qt_widgets = pytest.importorskip("PySide6.QtWidgets")
    app = qt_widgets.QApplication.instance() or qt_widgets.QApplication([])
    history_path = tmp_path / "interview_history.sqlite3"
    staffing_path = tmp_path / "staffing.sqlite3"
    monkeypatch.setattr(pyside_interview_app, "STAFFING_DB_PATH", staffing_path)
    store = pyside_interview_app.StaffingStore(
        pyside_interview_app.staffing_db_path_for_school("Palmdale", base_path=staffing_path)
    )
    store.initialize()
    service = pyside_interview_app.StaffingService(store)
    for history_id, candidate_name in [
        ("hist-delete-1", "Adrianna Love"),
        ("hist-delete-2", "Michelle Marti"),
        ("hist-keep", "Catherine Gray"),
    ]:
        service.upsert_director_candidate_referral(
            history_id=history_id,
            candidate_name=candidate_name,
            school="Palmdale",
            position="Teacher",
            interviewer_rating=8.0,
            interviewer_outcome="hire",
            interview_date="2026-07-08",
        )
    model = pyside_interview_app.build_director_staffing_model(
        build_interview_redesign_model(
            rubric_path=_write_test_rubric(tmp_path),
            overrides_path=_write_test_overrides(tmp_path),
            history_path=history_path,
            school_options=["Palmdale"],
        ),
        school="Palmdale",
    )

    window = pyside_interview_app.PySideInterviewWindow(model)
    table = window.window.findChild(qt_widgets.QTableWidget, "StaffingV2DirectorInterviewPendingTable")
    delete_selected = window.window.findChild(qt_widgets.QPushButton, "StaffingV2DirectorInterviewDeleteSelected")

    rendered_rows = {
        table.item(row, 0).text(): row
        for row in range(table.rowCount())
        if table.item(row, 0) is not None
    }
    assert table.selectionMode() == qt_widgets.QAbstractItemView.SelectionMode.MultiSelection
    assert set(rendered_rows) == {"Adrianna Love", "Michelle Marti", "Catherine Gray"}
    adrianna_checkbox = table.cellWidget(rendered_rows["Adrianna Love"], 0)
    michelle_checkbox = table.cellWidget(rendered_rows["Michelle Marti"], 0)
    assert adrianna_checkbox is not None
    assert michelle_checkbox is not None
    assert adrianna_checkbox.text() == "Adrianna Love"
    assert michelle_checkbox.text() == "Michelle Marti"
    adrianna_checkbox.setChecked(True)
    michelle_checkbox.setChecked(True)
    monkeypatch.setattr(
        qt_widgets.QMessageBox,
        "question",
        lambda *_args, **_kwargs: qt_widgets.QMessageBox.StandardButton.Yes,
    )
    delete_selected.click()
    app.processEvents()

    assert table.rowCount() == 1
    assert table.item(0, 0).text() == "Catherine Gray"
    assert [candidate.candidate_name for candidate in service.list_pending_director_interviews(school="Palmdale")] == [
        "Catherine Gray"
    ]
    window.window.close()
    app.processEvents()


def _staffing_row_for_position(table, position_name: str) -> int:
    for row in range(table.rowCount()):
        for column in range(table.columnCount()):
            item = table.item(row, column)
            if item is not None and item.toolTip() == position_name:
                return row
    raise AssertionError(f"Missing staffing row: {position_name}")


def _staffing_button_for_position(table, position_name: str):
    row = _staffing_row_for_position(table, position_name)
    button = table.cellWidget(row, table.columnCount() - 1)
    assert button is not None
    return button


def _write_test_overrides(tmp_path: Path) -> Path:
    overrides_path = tmp_path / "question_overrides.json"
    overrides_path.write_text(
        json.dumps(
            {
                "track_trait_order": {},
                "trait_question_overrides": {},
                "custom_questions": {
                    "preschool": [{"id": "Why-LPL", "text": "Why Launch Pad Learning?", "order": 1}]
                },
                "track_question_flow": {
                    "preschool": [
                        {"type": "custom", "id": "Why-LPL"},
                        {"type": "trait", "id": "trait_1"},
                    ]
                },
            }
        ),
        encoding="utf-8",
    )
    return overrides_path


def _write_workflow_test_overrides(tmp_path: Path) -> Path:
    overrides_path = tmp_path / "question_overrides.json"
    overrides_path.write_text(
        json.dumps(
            {
                "track_trait_order": {},
                "trait_question_overrides": {},
                "custom_questions": {
                    "preschool": [
                        {"id": "Why-ECE", "text": "Why early childhood education?", "order": 1},
                        {"id": "Why-LPL", "text": "Why Launch Pad Learning?", "order": 2},
                        {"id": "FT-or-PT", "text": "Full-time or part-time?", "order": 3},
                        {"id": "Not-Avail", "text": "Any hours unavailable?", "order": 4},
                        {"id": "Pay", "text": "What pay are you looking for?", "order": 5},
                        {"id": "Start", "text": "When could you start?", "order": 6},
                    ]
                },
                "track_question_flow": {
                    "preschool": [
                        {"type": "custom", "id": "Why-ECE"},
                        {"type": "custom", "id": "Why-LPL"},
                        {"type": "trait", "id": "trait_1"},
                        {"type": "custom", "id": "FT-or-PT"},
                        {"type": "custom", "id": "Not-Avail"},
                        {"type": "custom", "id": "Pay"},
                        {"type": "custom", "id": "Start"},
                    ]
                },
            }
        ),
        encoding="utf-8",
    )
    return overrides_path
