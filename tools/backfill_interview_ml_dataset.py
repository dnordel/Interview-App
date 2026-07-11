from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))

from data_store import InterviewHistoryStore, InterviewMLDatasetStore, ml_dataset_path_for_history_path
from platform_services import INTERVIEW_HISTORY_PATH, USER_ARTIFACTS_DIR


def _load_json(path: Path) -> dict[str, Any]:
    try:
        payload = json.loads(Path(path).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    return payload if isinstance(payload, dict) else {}


def _candidate_from_history(row: dict[str, Any]) -> dict[str, Any]:
    qualification = row.get("qualification", {}) if isinstance(row.get("qualification"), dict) else {}
    return {
        "name": str(row.get("candidate_name") or row.get("candidate") or row.get("name") or "").strip(),
        "email": str(row.get("candidate_email") or row.get("email") or "").strip(),
        "phone": str(row.get("candidate_phone") or row.get("phone") or "").strip(),
        "school": str(row.get("school") or "").strip(),
        "track": str(row.get("track") or row.get("position") or "").strip(),
        "position": str(row.get("position") or row.get("track") or "").strip(),
        "interview_date": str(row.get("interview_date") or row.get("date") or "").strip(),
        "qualification": qualification,
    }


def _candidate_docx_paths(row: dict[str, Any], notes_dirs: list[Path]) -> list[Path]:
    paths: list[Path] = []
    for key in ("interview_notes_path", "saved_report_path", "notes_path", "report_path"):
        raw = str(row.get(key, "") or "").strip()
        if raw:
            paths.append(Path(raw))
    candidate = _normalize_match(str(row.get("candidate_name") or row.get("candidate") or ""))
    interview_date = str(row.get("interview_date") or row.get("date") or "").strip()
    for notes_dir in notes_dirs:
        if not notes_dir.is_dir():
            continue
        for path in notes_dir.glob("*.docx"):
            name = _normalize_match(path.name)
            if candidate and candidate not in name:
                continue
            if interview_date and interview_date not in path.name:
                continue
            paths.append(path)
    output: list[Path] = []
    seen: set[str] = set()
    for path in paths:
        key = str(path)
        if key in seen:
            continue
        seen.add(key)
        output.append(path)
    return output


def _normalize_match(value: str) -> str:
    return re.sub(r"\s+", " ", str(value or "").strip()).casefold()


def _docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""
    try:
        doc = Document(path)
    except Exception:
        return ""
    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks).strip()


def recover_flow_transcript_from_docx(row: dict[str, Any], notes_dirs: list[Path]) -> tuple[list[dict[str, Any]], Path | None]:
    for path in _candidate_docx_paths(row, notes_dirs):
        if not path.exists() or path.suffix.lower() != ".docx":
            continue
        text = _docx_text(path)
        if not text:
            continue
        sections = _split_docx_transcript_sections(text)
        if sections:
            return sections, path
        return [
            {
                "flow_index": 1,
                "type": "custom",
                "id": "docx_transcript_recovered",
                "question": "Recovered transcript text from interview notes DOCX",
                "candidate_transcript": text,
                "transcript_recovered_from_docx": True,
            }
        ], path
    return [], None


def _split_docx_transcript_sections(text: str) -> list[dict[str, Any]]:
    pattern = re.compile(r"(?im)^(?:Q(?:uestion)?\s*(?P<num>\d+)[\.:)]?\s*)(?P<title>.*)$")
    matches = list(pattern.finditer(text))
    if not matches:
        return []
    rows: list[dict[str, Any]] = []
    for index, match in enumerate(matches):
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        body = text[start:end].strip(" \n:-")
        if not body:
            continue
        flow_index = int(match.group("num"))
        rows.append(
            {
                "flow_index": flow_index,
                "type": "custom",
                "id": f"docx_q{flow_index}",
                "question": match.group("title").strip() or f"Question {flow_index}",
                "candidate_transcript": body,
                "transcript_recovered_from_docx": True,
            }
        )
    return rows


def _job_path_for_row(row: dict[str, Any], history_path: Path, base_dir: Path) -> Path | None:
    history_id = str(row.get("history_id") or "").strip()
    candidates: list[Path] = []
    stored = str(row.get("deepseek_job_path", "") or "").strip()
    if stored:
        candidates.append(Path(stored))
    if history_id:
        name = f"deepseek-finalize-{history_id}.json"
        candidates.extend(
            [
                history_path.parent / "deepseek_jobs" / name,
                history_path.parent / "interviews" / "deepseek_jobs" / name,
                base_dir / "deepseek_jobs" / name,
            ]
        )
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def backfill_ml_dataset(
    *,
    history_path: Path,
    ml_path: Path | None = None,
    base_dir: Path | None = None,
    notes_dirs: list[Path] | None = None,
    export_dir: Path | None = None,
) -> dict[str, Any]:
    history_path = Path(history_path)
    base_dir = Path(base_dir) if base_dir is not None else USER_ARTIFACTS_DIR / "interviews"
    notes_dirs = list(notes_dirs or [])
    store = InterviewMLDatasetStore(Path(ml_path) if ml_path is not None else ml_dataset_path_for_history_path(history_path))
    history = InterviewHistoryStore(history_path)
    rows = history.load()
    imported = 0
    recovered_docx = 0
    trace_count = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        job_path = _job_path_for_row(row, history_path, base_dir)
        job = _load_json(job_path) if job_path is not None else {}
        payload = job.get("payload", {}) if isinstance(job.get("payload"), dict) else {}
        scoring = job.get("scoring", {}) if isinstance(job.get("scoring"), dict) else {}
        if not payload:
            payload = {"candidate": _candidate_from_history(row), "flow_transcript": []}
        payload.setdefault("candidate", _candidate_from_history(row))
        if not any(str(item.get("candidate_transcript") or "").strip() for item in payload.get("flow_transcript", []) or [] if isinstance(item, dict)):
            recovered, source_docx = recover_flow_transcript_from_docx(row, notes_dirs)
            if recovered:
                payload["flow_transcript"] = recovered
                payload["transcript_recovered_from_docx_path"] = str(source_docx or "")
                recovered_docx += 1
        store.upsert_interview(row, payload, scoring, source_job_path=job_path or "")
        traces = [item for item in job.get("deepseek_trace_events", []) or [] if isinstance(item, dict)]
        if traces:
            trace_count += store.record_deepseek_traces(str(row.get("history_id") or ""), traces, source_path=job_path or "")
        imported += 1
    exports = store.export_dataset(export_dir or history_path.parent)
    return {
        "history_rows": len(rows),
        "imported": imported,
        "docx_transcripts_recovered": recovered_docx,
        "deepseek_traces": trace_count,
        "ml_dataset_path": str(store.db_path),
        "pending_ai_analysis_path": str(exports["pending_ai_analysis"]),
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Backfill interview ML dataset without running DeepSeek.")
    parser.add_argument("--history-path", type=Path, default=INTERVIEW_HISTORY_PATH)
    parser.add_argument("--ml-path", type=Path, default=None)
    parser.add_argument("--base-dir", type=Path, default=USER_ARTIFACTS_DIR / "interviews")
    parser.add_argument("--notes-dir", type=Path, action="append", default=[])
    parser.add_argument("--export-dir", type=Path, default=None)
    args = parser.parse_args(argv)
    result = backfill_ml_dataset(
        history_path=args.history_path,
        ml_path=args.ml_path,
        base_dir=args.base_dir,
        notes_dirs=args.notes_dir,
        export_dir=args.export_dir,
    )
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
