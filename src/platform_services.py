from __future__ import annotations

import argparse
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
import faulthandler
from importlib import import_module
import json
import logging
import os
import platform
from pathlib import Path
import re
from logging.handlers import RotatingFileHandler
import runpy
import sys
import tempfile
from types import ModuleType
from types import FrameType
from types import TracebackType
import threading
import time
import traceback
from statistics import median
from typing import Any, Collection

_APP_LOG_PATH: Path | None = None
_INITIALIZED = False

LOG_FORMAT = "%(asctime)s %(levelname)s %(name)s %(message)s"
_TRACE_EVENTS = {"call", "exception", "return"}

APP_TITLE = "Structured Preschool Interview Tool"
APP_DIR = Path(__file__).resolve().parent
REPO_ROOT = APP_DIR.parent
CONFIG_DIR = REPO_ROOT / "config"
USER_ARTIFACTS_DIR = REPO_ROOT / "user_artifacts"

DEFAULT_RUBRIC_PATH = CONFIG_DIR / "rubric.json"
DEFAULT_SIGNALS_PATH = CONFIG_DIR / "disqualifier_signals.json"

# Stores GUI edits (trait order, trait question overrides, custom questions, and mixed flow)
QUESTIONS_OVERRIDE_PATH = CONFIG_DIR / "question_overrides.json"
INTERVIEW_HISTORY_DB_PATH = USER_ARTIFACTS_DIR / "interview_history.sqlite3"
INTERVIEW_HISTORY_LEGACY_JSON_PATH = USER_ARTIFACTS_DIR / "interview_history.json"
INTERVIEW_HISTORY_PATH = INTERVIEW_HISTORY_DB_PATH
SCHOOL_OFFER_SETTINGS_PATH = USER_ARTIFACTS_DIR / "school_offer_settings.json"
SCHOOL_EMAIL_TEMPLATE_SETTINGS_PATH = USER_ARTIFACTS_DIR / "school_email_template_settings.json"
INTERVIEW_APP_SETTINGS_PATH = USER_ARTIFACTS_DIR / "interview_app_settings.json"

DEFAULT_BASE_DIR = USER_ARTIFACTS_DIR / "interviews"

DEFAULT_FONT_SIZE = 10
MIN_FONT_SIZE = 8
MAX_FONT_SIZE = 18

INTRO_BODY_FONT_SPEC = ("Segoe UI", 10, "normal")
INTRO_HEADING_FONT_SPEC = ("Segoe UI", 11, "bold")

DEFAULT_SCHOOL_OPTIONS = ["Hawthorne", "Palmdale", "North Long Beach"]

INTRO_SCRIPT_TEMPLATE = (
    "Let me go ahead and share a little bit of info about our company, program and benefits.\n\n"
    "Company Statement\n"
    "At Launch Pad Learning, we empower every child, starting as young as six weeks, to explore, play, and grow "
    "through a creative curriculum that builds confidence, curiosity, and early social and academic skills in a safe, "
    "inclusive community.\n\n"
    "Program Structure\n"
    "Our preschool runs year-round on a state-subsidized model, so we don’t have summer or winter breaks, and\n"
    "{hours_line}\n\n"
    "We have a really tight-knit team teachers and support staff, and we also try to keep our classes small, so three infants per "
    "teacher, four toddlers per teacher, and eight preschool-age children per teacher.\n\n"
    "Benefits\n"
    "On the benefits side, we offer a 401(k) with employer matching contributions, medical, dental, and vision insurance, "
    "provide life & AD&D insurance, and give you ten paid holidays plus one week of PTO."
    "We also have an on-site chef who whips up homemade meals not just for the kids, but for you too, "
    "so meals become another chance to sit down together and build community.\n\n"
    "Let me go ahead and start the recording, and you should receive a notification."
)

SCHOOL_HOURS_LINE: dict[str, str] = {
    "Hawthorne": "• Hawthorne is open weekdays from 6 AM to 8 PM, closed on the weekends, and we are caring for about 100 children.",
    "North Long Beach": "• North Long Beach is open weekdays from 6 AM to 6 PM, closed on the weekends, and we are caring for about 100 children.",
    "Palmdale": "• Palmdale is open weekdays from 5:30 AM to 7 PM, closed on the weekends, and we are licensed for about 140.",
}


def compose_intro_script(school: str) -> str:
    school_clean = (school or "").strip()
    hours_line = SCHOOL_HOURS_LINE.get(school_clean)
    if not hours_line:
        hours_line = "• (Select a school above to show hours here)"
    return INTRO_SCRIPT_TEMPLATE.format(hours_line=hours_line)


SCHOOL_INFO: dict[str, str] = {
    "Hawthorne": "Open weekdays: 6:00 AM to 8:00 PM",
    "North Long Beach": "Open weekdays: 6:00 AM to 6:00 PM",
    "Palmdale": "Open weekdays: 5:30 AM to 7:00 PM",
}

NO_EXAMPLE_PHRASES = (
    "that's never happened to me",
    "thats never happened to me",
    "i can't think of an example",
    "i cant think of an example",
    "i've never experienced that",
    "ive never experienced that",
    "that's never come up for me",
    "thats never come up for me",
    "that doesn't really happen",
    "that doesnt really happen",
    "i've never really gotten feedback like that",
    "ive never really gotten feedback like that",
    "i've never really had challenging behavior",
    "ive never really had challenging behavior",
    "i've never had a conflict like that",
    "ive never had a conflict like that",
)

NEVER_HAPPENED_GLOBAL_SCRIPT = (
    "That makes sense. Let's zoom out a bit. Even if it wasn't a big incident, "
    "can you tell me about the closest situation you can remember?"
)

NEVER_HAPPENED_BY_TRAIT: dict[str, dict[str, Any]] = {
    "trait_1": {
        "title": "Empathy and Respect for Children",
        "followups": [
            "Can you tell me about a time a child was upset, even in a small way?",
            "How do you usually know when a child is having an emotional need?",
            "What signs do you look for when a child is struggling emotionally?",
        ],
        "scoring": [
            "Score 5-4 if they can generalize empathy, attunement, and child-centered thinking.",
            "Score 3 if response is surface-level but appropriate.",
            "Score 2 if they minimize emotional experiences.",
            "Score 1 if they deny emotional needs or frame children as manipulative.",
        ],
        "concerns": [
            "Denial that children have meaningful emotional distress.",
            "Statements implying emotions are insignificant or exaggerated.",
        ],
    },
    "trait_2": {
        "title": "Emotional Regulation Under Stress",
        "followups": [
            "What does stress look like for you at work, even on a mild day?",
            "How do you usually notice when you're starting to feel dysregulated?",
            "What do you do to stay calm during busy or chaotic moments?",
        ],
        "scoring": [
            "Score 5-4 if they show self-awareness and proactive regulation.",
            "Score 3 if they describe basic coping without reflection.",
            "Score 2 if they minimize stress or normalize frustration.",
            "Score 1 if they deny stress entirely or imply loss of control is inevitable.",
        ],
        "concerns": [
            "Claiming they never feel stress combined with rigidity or emotional flatness.",
            "Statements suggesting emotional regulation is unnecessary.",
        ],
    },
    "trait_3": {
        "title": "Respect for Children's Rights and Safety",
        "followups": [
            "How do you think about children's boundaries in everyday routines?",
            "What does respectful handling mean to you?",
            "What helps you stay safety-focused even during routine tasks?",
        ],
        "scoring": [
            "Score 5-4 if they articulate safety and dignity proactively.",
            "Score 3 if they reference rules without deeper understanding.",
            "Score 2 if safety is framed as an inconvenience.",
            "Score 1 if they minimize safety or boundaries.",
        ],
        "concerns": [
            "Dismissing the importance of boundaries or safety rules.",
        ],
    },
    "trait_4": {
        "title": "Coachability and Openness to Feedback",
        "followups": [
            "What kind of feedback do you usually receive?",
            "How do you prefer feedback to be given?",
            "How do you know when you need to adjust something in your work?",
        ],
        "scoring": [
            "Score 5-4 if they show openness and self-correction.",
            "Score 3 if they accept feedback passively.",
            "Score 2 if they subtly resist or deflect.",
            "Score 1 if they position themselves as beyond feedback.",
        ],
        "concerns": [
            "Statements implying supervision is unnecessary or unwelcome.",
        ],
    },
    "trait_5": {
        "title": "Reliability and Accountability",
        "followups": [
            "When things do go wrong, how do you usually respond?",
            "What do you do if you make a small mistake?",
            "How do you hold yourself accountable day to day?",
        ],
        "scoring": [
            "Score 5-4 if they describe ownership and corrective action.",
            "Score 3 if accountability is vague.",
            "Score 2 if responsibility is externalized.",
            "Score 1 if accountability is rejected.",
        ],
        "concerns": [],
    },
    "trait_6": {
        "title": "Team Orientation and Collaboration",
        "followups": [
            "How do you usually communicate with coworkers?",
            "What do you do if you and another adult see things differently?",
            "How do you contribute to a team environment?",
        ],
        "scoring": [
            "Score 5-4 if collaboration is proactive and respectful.",
            "Score 3 if teamwork is neutral.",
            "Score 2 if independence is emphasized over collaboration.",
            "Score 1 if teamwork is dismissed.",
        ],
        "concerns": [],
    },
    "trait_7": {
        "title": "Curiosity and Willingness to Learn",
        "followups": [
            "What do you feel most confident about in your approach?",
            "What do you still want to get better at?",
            "How do you usually learn new things at work?",
        ],
        "scoring": [
            "Score 5-4 if curiosity and reflection emerge.",
            "Score 3 if learning is passive.",
            "Score 2 if growth is minimized.",
            "Score 1 if beliefs are rigid.",
        ],
        "concerns": [],
    },
    "trait_8": {
        "title": "Gentleness and Physical Awareness (Infant/Toddler)",
        "followups": [
            "What helps you stay gentle during repetitive or tiring routines?",
            "How do you respond when a baby resists a routine?",
            "What cues tell you to slow down?",
        ],
        "scoring": [
            "Lack of reflection here is high risk.",
            "If they cannot articulate gentleness, score conservatively.",
            "Score 1 language remains automatic no hire.",
        ],
        "concerns": [],
    },
    "trait_9": {
        "title": "Patience with Nonverbal Communication (Infant/Toddler)",
        "followups": [
            "What signals do you usually look for?",
            "What do you do if your first guess is wrong?",
            "How do you stay patient during repeated crying?",
        ],
        "scoring": [],
        "concerns": [],
    },
    "trait_10": {
        "title": "Positive Behavior Guidance (Preschool)",
        "followups": [
            "How do you define challenging behavior?",
            "What do you do when a child does not follow expectations?",
            "How do you teach behavior skills?",
        ],
        "scoring": [],
        "concerns": [
            "Denial of challenging behavior in preschool settings can signal lack of insight or overly punitive environments.",
        ],
    },
    "trait_11": {
        "title": "Structure and Flexibility (Preschool)",
        "followups": [
            "How do you usually respond when children move at a different pace?",
            "What do you do if an activity is not working?",
            "How comfortable are you changing plans mid-day?",
        ],
        "scoring": [],
        "concerns": [],
    },
}


def text_suggests_no_example(text: str) -> bool:
    """Heuristic check: does the response resemble a 'no example' statement."""
    t = (text or "").strip().lower()
    if not t:
        return False
    return any(p in t for p in NO_EXAMPLE_PHRASES)


def sanitize_filename(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name or "Unknown"


def is_valid_date_yyyy_mm_dd(value: str) -> bool:
    try:
        datetime.strptime(value, "%Y-%m-%d")
        return True
    except ValueError:
        return False


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")

MAX_CONFIG_BYTES = 2_000_000

CONFIG_ASSET_REGISTRY: dict[str, dict[str, Any]] = {
    "rubric.json": {
        "owner": "interview_runtime_service",
        "consumers": ["data_store", "pyside_interview_app"],
        "schema": "object(metadata, scoring, tracks, traits[], absolute_disqualifiers[])",
    },
    "disqualifier_signals.json": {
        "owner": "interview_runtime_service",
        "consumers": ["data_store", "pyside_interview_app"],
        "schema": "object(questions[])",
    },
    "question_overrides.json": {
        "owner": "interview_runtime_service",
        "consumers": ["data_store", "pyside_interview_app", "question_settings_service"],
        "schema": "object(track_trait_order, trait_question_overrides, custom_questions, track_question_flow)",
    },
    "interview_output.schema.json": {
        "owner": "interview_runtime_service",
        "consumers": ["scoring_reporting", "interview_runtime"],
        "schema": "json-schema draft 2020-12",
    },
    "cues.json": {
        "owner": "interview_runtime_service",
        "consumers": [],
        "schema": "object(version, scoring_scale, behavior_flags, final_outcomes, cases[])",
    },
    "sample_draft.json": {
        "owner": "interview_runtime_service",
        "consumers": [],
        "schema": "object(candidate, current_index, trait_inputs)",
    },
}

_ENV_DEBUG_VALUES = {"1", "true", "yes", "on", "debug"}
_DEFAULT_LEVEL = logging.INFO
_DEBUG_LEVEL = logging.DEBUG
_MAX_LOG_BYTES = 5 * 1024 * 1024
_BACKUP_COUNT = 5

_SENSITIVE_KEYS = {
    "candidate_name",
    "candidate_first_name",
    "candidate_last_name",
    "first_name",
    "last_name",
    "full_name",
    "email",
    "phone",
    "address",
}

_STANDARD_RECORD_KEYS = {
    "name",
    "msg",
    "args",
    "levelname",
    "levelno",
    "pathname",
    "filename",
    "module",
    "exc_info",
    "exc_text",
    "stack_info",
    "lineno",
    "funcName",
    "created",
    "msecs",
    "relativeCreated",
    "thread",
    "threadName",
    "processName",
    "process",
    "message",
}

_EMAIL_PATTERN = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
_PHONE_PATTERN = re.compile(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)")

SAFE_EXTENSIONS = {".wav", ".mp3", ".json", ".jsonl", ".npy", ".npz", ".pt", ".bin", ".tmp", ".chunk"}
SAFE_NAME_TOKENS = ("candidate_", "_transcript", "embedding", "chunk", "recording")

EVENT_TASK_CREATED = "task_created"
EVENT_TASK_COMPLETED = "task_completed"
EVENT_TASK_OVERDUE = "task_overdue"
EVENT_REMINDER_SENT = "reminder_sent"
EVENT_INTERVIEW_FINALIZED = "interview_finalized"

UX_EVENT_KINDS = ("view", "click", "validation_error", "completion")
UX_REQUIRED_FIELDS_BY_KIND = {"view": ("target",), "click": ("target",), "validation_error": ("error_type",), "completion": ("outcome",)}
_BLOCKED_UX_FIELDS = {
    "candidate_name",
    "employee_name",
    "notes",
    "free_text",
    "email",
    "phone",
    "address",
    "resume_path",
}
_EMAIL_LIKE_VALUE_MARKERS = ("@",)

ONBOARDING_EVENT_REGISTRY: dict[str, dict[str, type[Any]]] = {
    "ux.onboarding.add_employee_form.view": {
        "entry_point": str,
        "time_from_screen_open_ms": int,
    },
    "ux.onboarding.add_employee_form.validation_error": {
        "error_type": str,
        "required_fields_missing_count": int,
    },
    "ux.onboarding.urgent_filter.click": {
        "time_to_filter_ms": int,
        "result_count": int,
    },
    "ux.onboarding.reminder_mode.click": {
        "mode": str,
        "time_to_mode_select_ms": int,
        "changed_from_default": bool,
    },
    "ux.onboarding.reminder_run.completion": {
        "mode": str,
        "recipient_count": int,
        "skipped_count": int,
        "warning_count": int,
        "sent_count": int,
        "failed_count": int,
        "blocked_count": int,
    },
    "ux.onboarding.sender_email.validation_error": {
        "error_reason": str,
        "attempt_count": int,
    },
    "ux.onboarding.sender_email.completion": {
        "attempts_before_success": int,
        "domain_type": str,
    },
}

LEGACY_TO_CANONICAL_UX_EVENTS = {
    "onboarding_add_employee_opened": "ux.onboarding.add_employee_form.view",
    "onboarding_employee_save_error": "ux.onboarding.add_employee_form.validation_error",
    "onboarding_urgent_filter_applied": "ux.onboarding.urgent_filter.click",
    "onboarding_reminder_mode_selected": "ux.onboarding.reminder_mode.click",
    "onboarding_reminder_dry_run_completed": "ux.onboarding.reminder_run.completion",
    "onboarding_reminder_live_run_completed": "ux.onboarding.reminder_run.completion",
    "onboarding_sender_email_validation_error": "ux.onboarding.sender_email.validation_error",
    "onboarding_sender_email_updated": "ux.onboarding.sender_email.completion",
}

SCOPE_EVENT_MONTH = "event_month"
SCOPE_CREATED_MONTH = "created_month"
SUMMARY_SCOPES = (SCOPE_EVENT_MONTH, SCOPE_CREATED_MONTH)


class RedactionFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        redacted_extras: dict[str, Any] = {}
        for key, value in record.__dict__.items():
            if key in _STANDARD_RECORD_KEYS:
                continue
            redacted_extras[key] = _redact_field(key, value)

        for key, value in redacted_extras.items():
            setattr(record, key, value)

        if isinstance(record.msg, str):
            record.msg = _redact_text(record.msg)

        if isinstance(record.args, tuple):
            record.args = tuple(_redact_value(value) for value in record.args)
        elif isinstance(record.args, dict):
            record.args = {key: _redact_value(value) for key, value in record.args.items()}

        return True


class JsonLogFormatter(logging.Formatter):
    def format(self, record: logging.LogRecord) -> str:
        payload: dict[str, Any] = {
            "timestamp": self.formatTime(record, self.datefmt),
            "level": record.levelname,
            "logger": record.name,
            "thread": record.threadName,
            "message": record.getMessage(),
        }

        extras = self._collect_extras(record)
        if extras:
            payload["context"] = extras

        if record.exc_info:
            payload["exception"] = self.formatException(record.exc_info)
        if record.stack_info:
            payload["stack"] = self.formatStack(record.stack_info)

        return json.dumps(payload, ensure_ascii=False)

    def _collect_extras(self, record: logging.LogRecord) -> dict[str, Any]:
        extras: dict[str, Any] = {}
        for key, value in record.__dict__.items():
            if key in _STANDARD_RECORD_KEYS:
                continue
            extras[key] = _serialize(value)
        return extras


class ConfigValidationError(ValueError):
    """Validation failed for untrusted config payloads."""


@dataclass(slots=True)
class MonthlyMetricsSummary:
    month_key: str
    scope: str
    on_time_completion_pct: float
    overdue_count: int
    previous_overdue_count: int
    overdue_change_pct: float | None
    median_days_by_task_type: dict[str, float]


class UxMetricsLogger:
    def __init__(self, root_dir: str | Path, filename: str = "ux_metrics.jsonl") -> None:
        self.root_dir = Path(root_dir)
        self.root_dir.mkdir(parents=True, exist_ok=True)
        self.path = self.root_dir / filename
        self._overdue_event_keys = self._load_overdue_event_keys()

    def log_event(self, event_type: str, **fields: Any) -> None:
        canonical_type = LEGACY_TO_CANONICAL_UX_EVENTS.get(event_type, event_type)
        if not canonical_type:
            return
        safe_fields = dict(fields)
        if canonical_type.startswith("ux."):
            safe_fields = _sanitize_ux_payload(safe_fields)
        if canonical_type in ONBOARDING_EVENT_REGISTRY:
            safe_fields = _coerce_canonical_onboarding_fields(canonical_type, safe_fields)
        payload = {
            "event_type": canonical_type,
            "timestamp": _utc_now_iso(),
            **safe_fields,
        }
        self.path.parent.mkdir(parents=True, exist_ok=True)
        with self.path.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps(payload, ensure_ascii=False) + "\n")

    def log_ux_view(self, *, app: str, surface: str, target: str = "", **fields: Any) -> None:
        self._log_ux_event("view", app=app, surface=surface, target=target, **fields)

    def log_ux_click(self, *, app: str, surface: str, target: str, **fields: Any) -> None:
        self._log_ux_event("click", app=app, surface=surface, target=target, **fields)

    def log_ux_validation_error(
        self,
        *,
        app: str,
        surface: str,
        error_type: str,
        field_name: str = "",
        required_fields_present: bool = True,
        **fields: Any,
    ) -> None:
        self._log_ux_event(
            "validation_error",
            app=app,
            surface=surface,
            error_type=error_type,
            field_name=field_name,
            required_fields_present=required_fields_present,
            **fields,
        )

    def log_ux_completion(self, *, app: str, surface: str, outcome: str, **fields: Any) -> None:
        self._log_ux_event("completion", app=app, surface=surface, outcome=outcome, **fields)

    def log_onboarding_canonical_event(self, event_type: str, **fields: Any) -> None:
        canonical_type = LEGACY_TO_CANONICAL_UX_EVENTS.get(event_type, event_type)
        if canonical_type not in ONBOARDING_EVENT_REGISTRY:
            return
        payload = _coerce_canonical_onboarding_fields(canonical_type, _sanitize_ux_payload(fields))
        self.log_event(canonical_type, **payload)

    def _log_ux_event(self, kind: str, **fields: Any) -> None:
        if kind not in UX_EVENT_KINDS:
            return
        app = _sanitize_event_token(fields.pop("app", ""))
        surface = _sanitize_event_token(fields.pop("surface", ""))
        if not app or not surface:
            return
        payload = _sanitize_ux_payload(fields)
        required = UX_REQUIRED_FIELDS_BY_KIND.get(kind, ())
        for field in required:
            value = str(payload.get(field, "")).strip()
            if not value:
                return
        event_type = f"ux.{app}.{surface}.{kind}"
        self.log_event(event_type, **payload)

    def log_overdue_once(self, *, task_id: str, due_date: str | None, **fields: Any) -> bool:
        event_key = f"{task_id}|{due_date or ''}|{EVENT_TASK_OVERDUE}"
        if not task_id or event_key in self._overdue_event_keys:
            return False
        self._overdue_event_keys.add(event_key)
        self.log_event(EVENT_TASK_OVERDUE, task_id=task_id, due_date=due_date, event_key=event_key, **fields)
        return True

    def read_events(self) -> list[dict[str, Any]]:
        if not self.path.exists():
            return []

        events: list[dict[str, Any]] = []
        with self.path.open("r", encoding="utf-8") as handle:
            for line in handle:
                line_clean = line.strip()
                if not line_clean:
                    continue
                try:
                    payload = json.loads(line_clean)
                except json.JSONDecodeError:
                    continue
                if isinstance(payload, dict):
                    events.append(payload)
        return events

    def _load_overdue_event_keys(self) -> set[str]:
        keys: set[str] = set()
        for event in self.read_events():
            if str(event.get("event_type") or "") != EVENT_TASK_OVERDUE:
                continue
            key = str(event.get("event_key") or "").strip()
            if key:
                keys.add(key)
        return keys


def initialize_app_logging(*, app_root: Path | None = None) -> Path:
    global _APP_LOG_PATH, _INITIALIZED
    if _INITIALIZED and _APP_LOG_PATH is not None:
        return _APP_LOG_PATH

    root = app_root or Path(__file__).resolve().parent.parent
    log_dir = root / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)

    log_path = log_dir / "interview-app.log"
    handler = RotatingFileHandler(
        log_path,
        maxBytes=_MAX_LOG_BYTES,
        backupCount=_BACKUP_COUNT,
        encoding="utf-8",
    )
    handler.setFormatter(JsonLogFormatter())
    handler.addFilter(RedactionFilter())

    root_logger = logging.getLogger()
    root_logger.setLevel(_resolve_log_level())
    _attach_handler_once(root_logger, handler)

    _APP_LOG_PATH = log_path
    _INITIALIZED = True
    install_uncaught_exception_hooks()
    logging.getLogger(__name__).info(
        "app_logging_initialized",
        extra={"log_path": str(log_path), "log_level": logging.getLevelName(root_logger.level)},
    )
    return log_path


def get_configured_log_path() -> Path | None:
    return _APP_LOG_PATH


def write_crash_report(
    *,
    source: str,
    exc_type: type[BaseException],
    exc_value: BaseException,
    exc_traceback: TracebackType | None,
    app_root: Path | None = None,
) -> Path | None:
    root = app_root or Path(__file__).resolve().parent.parent
    crash_dir = root / "logs" / "crash-reports"
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    file_path = crash_dir / f"crash-{stamp}.log"
    payload = _build_crash_payload(source, exc_type, exc_value, exc_traceback)
    try:
        crash_dir.mkdir(parents=True, exist_ok=True)
        file_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return file_path
    except Exception:
        logging.getLogger(__name__).exception("crash_report_write_failed")
        return None


def _build_crash_payload(
    source: str,
    exc_type: type[BaseException],
    exc_value: BaseException,
    exc_traceback: TracebackType | None,
) -> dict[str, Any]:
    origin = _traceback_origin(exc_traceback)
    return {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "source": source,
        "error_type": exc_type.__name__,
        "error_message": str(exc_value),
        "origin": origin,
        "python": sys.version,
        "platform": platform.platform(),
        "cwd": str(Path.cwd()),
        "traceback": "".join(traceback.format_exception(exc_type, exc_value, exc_traceback)).strip(),
    }


def _traceback_origin(exc_traceback: TracebackType | None) -> dict[str, Any]:
    if exc_traceback is None:
        return {"function": "<unknown>", "line": None, "file": "<unknown>"}
    last_tb = exc_traceback
    while last_tb.tb_next is not None:
        last_tb = last_tb.tb_next
    frame = last_tb.tb_frame
    code = frame.f_code
    return {
        "function": code.co_name,
        "line": int(last_tb.tb_lineno),
        "file": str(code.co_filename),
    }


def install_uncaught_exception_hooks() -> None:
    def _sys_hook(exc_type: type[BaseException], exc_value: BaseException, exc_traceback: TracebackType | None) -> None:
        logging.getLogger(__name__).error(
            "uncaught_main_exception",
            exc_info=(exc_type, exc_value, exc_traceback),
        )

    def _thread_hook(args: threading.ExceptHookArgs) -> None:
        logging.getLogger(__name__).error(
            "uncaught_thread_exception",
            extra={"thread_name": args.thread.name if args.thread else "unknown"},
            exc_info=(args.exc_type, args.exc_value, args.exc_traceback),
        )

    sys.excepthook = _sys_hook
    threading.excepthook = _thread_hook


def _attach_handler_once(root_logger: logging.Logger, handler: logging.Handler) -> None:
    target_path = getattr(handler, "baseFilename", None)
    for existing in root_logger.handlers:
        existing_path = getattr(existing, "baseFilename", None)
        if target_path is not None and existing_path == target_path:
            return
    root_logger.addHandler(handler)


def _resolve_log_level() -> int:
    value = os.getenv("INTERVIEW_APP_LOG_LEVEL", "").strip()
    if value:
        parsed_level = logging.getLevelName(value.upper())
        if isinstance(parsed_level, int):
            return parsed_level
    debug_enabled = os.getenv("INTERVIEW_APP_DEBUG", "").strip().lower() in _ENV_DEBUG_VALUES
    if debug_enabled:
        return _DEBUG_LEVEL
    return _DEFAULT_LEVEL


def _redact_field(key: str, value: Any) -> Any:
    if key.lower() in _SENSITIVE_KEYS:
        return "[REDACTED]"
    return _redact_value(value)


def _redact_value(value: Any) -> Any:
    if isinstance(value, str):
        return _redact_text(value)
    if isinstance(value, dict):
        return {str(k): _redact_value(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_redact_value(item) for item in value]
    if isinstance(value, tuple):
        return tuple(_redact_value(item) for item in value)
    return value


def _redact_text(text: str) -> str:
    masked = _EMAIL_PATTERN.sub("[REDACTED_EMAIL]", text)
    return _PHONE_PATTERN.sub("[REDACTED_PHONE]", masked)


def _serialize(value: Any) -> Any:
    if isinstance(value, (str, int, float, bool)) or value is None:
        return value
    if isinstance(value, dict):
        return {str(k): _serialize(v) for k, v in value.items()}
    if isinstance(value, (list, tuple)):
        return [_serialize(item) for item in value]
    return str(value)


def _is_within(base_dir: Path, path: Path) -> bool:
    try:
        path.resolve().relative_to(base_dir.resolve())
    except ValueError:
        return False
    return True


def _safe_to_delete(base_dir: Path, path: Path) -> bool:
    if not path.exists() or not path.is_file():
        return False
    if not _is_within(base_dir, path):
        return False
    if path.suffix.lower() in SAFE_EXTENSIONS:
        return True
    name = path.name.lower()
    return any(token in name for token in SAFE_NAME_TOKENS)


def extract_artifact_paths(flow_recordings: dict[int, dict[str, Any]] | dict[str, Any]) -> list[Path]:
    out: list[Path] = []
    for entry in (flow_recordings or {}).values():
        item = dict(entry or {})
        attempts = item.get("attempts")
        if not isinstance(attempts, list):
            attempts = [item]
        for attempt in attempts:
            rec = dict(attempt or {})
            for key in ("mic_wav", "sys_wav", "transcript_txt", "transcript_jsonl"):
                raw = str(rec.get(key) or "").strip()
                if raw:
                    out.append(Path(raw).expanduser())
            output_dir = str(rec.get("output_dir") or "").strip()
            base_name = str(rec.get("base_name") or "").strip()
            if not output_dir or not base_name:
                continue
            root = Path(output_dir).expanduser()
            for ext in (".mp3", ".json", ".npy", ".npz", ".pt", ".bin", ".tmp", ".chunk"):
                out.append(root / f"{base_name}{ext}")
    return out


def delete_recording_artifacts(base_dir: Path, flow_recordings: dict[int, dict[str, Any]] | dict[str, Any]) -> list[Path]:
    deleted: list[Path] = []
    for path in extract_artifact_paths(flow_recordings):
        if not _safe_to_delete(base_dir, path):
            continue
        try:
            path.unlink(missing_ok=True)
        except OSError:
            continue
        deleted.append(path)
    return deleted


def cleanup_stale_artifacts(base_dir: Path) -> list[Path]:
    if not base_dir.exists():
        return []
    deleted: list[Path] = []
    for path in base_dir.glob("*"):
        if not path.is_file():
            continue
        lower_name = path.name.lower()
        stale_name = lower_name.startswith("candidate_") or any(token in lower_name for token in SAFE_NAME_TOKENS)
        if not stale_name:
            continue
        if not _safe_to_delete(base_dir, path):
            continue
        try:
            path.unlink(missing_ok=True)
        except OSError:
            continue
        deleted.append(path)
    return deleted


def build_monthly_summary(
    *,
    month: date,
    scope: str,
    employees: list[Any],
    events: list[dict[str, Any]],
    grace_hours: int = 24,
) -> MonthlyMetricsSummary:
    selected_month = date(month.year, month.month, 1)
    previous_month = _prev_month(selected_month)
    records = _build_task_records(employees)

    selected = _records_for_scope(records, selected_month, scope)
    on_time_pct = _on_time_completion_pct(selected, grace_hours=grace_hours)
    median_days = _median_days_by_type(selected)

    overdue_count = _overdue_count(selected_month, events, records)
    previous_overdue = _overdue_count(previous_month, events, records)
    change_pct = _pct_change(overdue_count, previous_overdue)

    return MonthlyMetricsSummary(
        month_key=selected_month.strftime("%Y-%m"),
        scope=scope,
        on_time_completion_pct=on_time_pct,
        overdue_count=overdue_count,
        previous_overdue_count=previous_overdue,
        overdue_change_pct=change_pct,
        median_days_by_task_type=median_days,
    )


def _build_task_records(employees: list[Any]) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    for employee in employees:
        for task in getattr(employee, "tasks", []):
            created = _parse_datetime(getattr(task, "created_at", None))
            completed = _parse_datetime(getattr(task, "completed_at", None))
            due = _parse_datetime(getattr(task, "due_date", None))
            if created is None:
                created = due or completed
            records.append(
                {
                    "task_id": getattr(task, "id", ""),
                    "task_type": str(getattr(task, "template_id", "") or getattr(task, "title", "unknown")),
                    "created_at": created,
                    "completed_at": completed,
                    "due_at": due,
                    "completed": bool(getattr(task, "completed", False)),
                }
            )
    return records


def _records_for_scope(records: list[dict[str, Any]], month: date, scope: str) -> list[dict[str, Any]]:
    if scope == SCOPE_CREATED_MONTH:
        return [item for item in records if _in_month(item.get("created_at"), month)]
    return [item for item in records if _in_month(item.get("completed_at"), month)]


def _on_time_completion_pct(records: list[dict[str, Any]], grace_hours: int) -> float:
    completed_records = [item for item in records if item.get("completed_at")]
    if not completed_records:
        return 0.0

    grace_delta = timedelta(hours=grace_hours)
    on_time_count = 0
    for item in completed_records:
        due_at = item.get("due_at")
        completed_at = item.get("completed_at")
        if due_at is None or completed_at is None:
            continue
        if completed_at <= due_at + grace_delta:
            on_time_count += 1

    return round((on_time_count / len(completed_records)) * 100.0, 1)


def _median_days_by_type(records: list[dict[str, Any]]) -> dict[str, float]:
    by_type: dict[str, list[float]] = {}
    for item in records:
        created = item.get("created_at")
        completed = item.get("completed_at")
        task_type = str(item.get("task_type") or "unknown")
        if created is None or completed is None:
            continue
        day_span = (completed - created).total_seconds() / 86400.0
        by_type.setdefault(task_type, []).append(day_span)

    medians: dict[str, float] = {}
    for task_type, values in by_type.items():
        if not values:
            continue
        medians[task_type] = round(float(median(values)), 1)
    return medians


def _overdue_count(month: date, events: list[dict[str, Any]], records: list[dict[str, Any]]) -> int:
    logged_count = sum(1 for item in events if _event_matches_month(item, EVENT_TASK_OVERDUE, month))
    if logged_count:
        return logged_count

    fallback = 0
    for item in records:
        due_at = item.get("due_at")
        completed_at = item.get("completed_at")
        if due_at is None:
            continue
        if not _in_month(due_at, month):
            continue
        if completed_at is None or completed_at > due_at:
            fallback += 1
    return fallback


def _event_matches_month(event: dict[str, Any], event_type: str, month: date) -> bool:
    if str(event.get("event_type") or "") != event_type:
        return False
    ts = _parse_datetime(event.get("timestamp"))
    return _in_month(ts, month)


def _pct_change(current: int, previous: int) -> float | None:
    if previous == 0:
        if current == 0:
            return 0.0
        return None
    return round(((current - previous) / previous) * 100.0, 1)


def _parse_datetime(value: Any) -> datetime | None:
    if not isinstance(value, str):
        return None
    raw = value.strip()
    if not raw:
        return None
    if len(raw) == 10:
        raw = f"{raw}T00:00:00"
    if raw.endswith("Z"):
        raw = raw[:-1] + "+00:00"
    try:
        parsed = datetime.fromisoformat(raw)
    except ValueError:
        return None
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=timezone.utc)
    return parsed.astimezone(timezone.utc)


def _in_month(value: datetime | None, month: date) -> bool:
    if value is None:
        return False
    return value.year == month.year and value.month == month.month


def _prev_month(value: date) -> date:
    if value.month == 1:
        return date(value.year - 1, 12, 1)
    return date(value.year, value.month - 1, 1)


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _collect_fieldnames(events: list[dict[str, Any]]) -> list[str]:
    ordered = ["timestamp", "event_type", "task_id", "task_type", "employee_id", "employee_name", "due_date"]
    dynamic = set()
    for item in events:
        dynamic.update(item.keys())
    for field in ordered:
        dynamic.discard(field)
    return ordered + sorted(dynamic)


def _sanitize_event_token(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    cleaned = value.strip().lower().replace(" ", "_")
    return "".join(char for char in cleaned if char.isalnum() or char in {"_", "-"})


def _sanitize_ux_payload(fields: dict[str, Any]) -> dict[str, Any]:
    sanitized: dict[str, Any] = {}
    for key, value in fields.items():
        if key in _BLOCKED_UX_FIELDS or key.endswith("_notes"):
            continue
        if key.endswith("_name"):
            continue
        normalized_key = _sanitize_event_token(key)
        if not normalized_key:
            continue
        safe_value = _sanitize_ux_value(value)
        if safe_value is None and value is not None:
            continue
        if safe_value is None:
            sanitized[normalized_key] = None
            continue
        sanitized[normalized_key] = safe_value
    return sanitized


def _sanitize_ux_value(value: Any) -> str | int | float | bool | None:
    if isinstance(value, bool):
        return value
    if isinstance(value, int):
        return value
    if isinstance(value, float):
        return value
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return ""
    if any(marker in text for marker in _EMAIL_LIKE_VALUE_MARKERS):
        return None
    if len(text) > 120:
        return text[:120]
    return text


def _coerce_canonical_onboarding_fields(event_type: str, fields: dict[str, Any]) -> dict[str, Any]:
    schema = ONBOARDING_EVENT_REGISTRY.get(event_type, {})
    payload: dict[str, Any] = {}
    for key, expected_type in schema.items():
        value = fields.get(key)
        coerced = _coerce_value(value, expected_type)
        if coerced is None:
            continue
        payload[key] = coerced
    return payload


def _coerce_value(value: Any, expected_type: type[Any]) -> Any:
    if value is None:
        return None
    if expected_type is bool:
        if isinstance(value, bool):
            return value
        text = str(value).strip().lower()
        if text in {"true", "1", "yes"}:
            return True
        if text in {"false", "0", "no"}:
            return False
        return None
    if expected_type is int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return None
    if expected_type is str:
        text = str(value).strip().lower()
        return _sanitize_event_token(text)
    return value


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Runtime wrapper for GUI entrypoints")
    parser.add_argument("--target", required=True, help="Path to the target Python script to execute")
    parser.add_argument("--app-root", default="", help="Optional repository/app root for logs")
    parser.add_argument("--debug", action="store_true", help="Enable deep call tracing")
    args, target_args = parser.parse_known_args(argv)
    args.target_args = target_args
    return args


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    target_path = Path(args.target).expanduser().resolve()
    if not target_path.exists():
        raise FileNotFoundError(f"Target script not found: {target_path}")

    app_root = _resolve_app_root(target_path, args.app_root)
    log_paths = _configure_runtime_logging(app_root)
    _install_global_exception_hooks(app_root, log_paths["runtime_log"])
    _enable_faulthandler(log_paths["fault_log"])

    debug_enabled = bool(args.debug or _env_debug_enabled())
    if debug_enabled:
        _install_trace_logging(app_root, log_paths["trace_log"])

    logger = logging.getLogger("runtime_wrapper")
    logger.info("runtime_wrapper_start", extra={"target": str(target_path), "debug": debug_enabled})

    previous_argv = sys.argv[:]
    sys.argv = [str(target_path), *list(getattr(args, "target_args", []))]
    try:
        runpy.run_path(str(target_path), run_name="__main__")
    except SystemExit as exc:
        code = exc.code if isinstance(exc.code, int) else 0 if exc.code is None else 1
        if code == 0:
            logger.info("runtime_wrapper_exit")
            return 0
        report_path = write_wrapper_crash_report(
            app_root=app_root,
            source="runtime_wrapper_main",
            exc_type=type(exc),
            exc_value=exc,
            exc_traceback=exc.__traceback__,
        )
        logger.error("runtime_wrapper_fatal_exception", exc_info=(type(exc), exc, exc.__traceback__))
        logger.error("runtime_wrapper_crash_report", extra={"path": str(report_path) if report_path else ""})
        raise
    except BaseException as exc:  # noqa: BLE001
        report_path = write_wrapper_crash_report(
            app_root=app_root,
            source="runtime_wrapper_main",
            exc_type=type(exc),
            exc_value=exc,
            exc_traceback=exc.__traceback__,
        )
        logger.error("runtime_wrapper_fatal_exception", exc_info=(type(exc), exc, exc.__traceback__))
        logger.error("runtime_wrapper_crash_report", extra={"path": str(report_path) if report_path else ""})
        raise
    finally:
        sys.argv = previous_argv

    logger.info("runtime_wrapper_exit")
    return 0


def _resolve_app_root(target_path: Path, arg_root: str) -> Path:
    if arg_root.strip():
        return Path(arg_root).expanduser().resolve()
    candidate = target_path.parent.parent
    if (candidate / "src").exists():
        return candidate
    return target_path.parent


def _configure_runtime_logging(app_root: Path) -> dict[str, Path]:
    log_dir = app_root / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)

    runtime_log = log_dir / "runtime_wrapper.log"
    fault_log = log_dir / "runtime_faults.log"
    trace_log = log_dir / "runtime_trace.log"

    logging.basicConfig(
        level=logging.DEBUG if _env_debug_enabled() else logging.INFO,
        format=LOG_FORMAT,
        handlers=[logging.FileHandler(runtime_log, encoding="utf-8")],
        force=True,
    )
    return {"runtime_log": runtime_log, "fault_log": fault_log, "trace_log": trace_log}


def _enable_faulthandler(fault_log_path: Path) -> None:
    stream = fault_log_path.open("a", encoding="utf-8")
    faulthandler.enable(file=stream, all_threads=True)


def _install_global_exception_hooks(app_root: Path, runtime_log_path: Path) -> None:
    logger = logging.getLogger("runtime_wrapper")

    def _sys_hook(exc_type: type[BaseException], exc_value: BaseException, exc_traceback: TracebackType | None) -> None:
        report = write_wrapper_crash_report(
            app_root=app_root,
            source="sys_excepthook",
            exc_type=exc_type,
            exc_value=exc_value,
            exc_traceback=exc_traceback,
        )
        logger.error("uncaught_main_exception", exc_info=(exc_type, exc_value, exc_traceback))
        logger.error("crash_report_written", extra={"path": str(report) if report else "", "runtime_log": str(runtime_log_path)})

    def _thread_hook(args: threading.ExceptHookArgs) -> None:
        report = write_wrapper_crash_report(
            app_root=app_root,
            source="thread_excepthook",
            exc_type=args.exc_type,
            exc_value=args.exc_value,
            exc_traceback=args.exc_traceback,
        )
        logger.error(
            "uncaught_thread_exception",
            extra={"thread_name": args.thread.name if args.thread else "unknown", "report": str(report) if report else ""},
            exc_info=(args.exc_type, args.exc_value, args.exc_traceback),
        )

    sys.excepthook = _sys_hook
    threading.excepthook = _thread_hook


def _install_trace_logging(app_root: Path, trace_path: Path) -> None:
    trace_handle = trace_path.open("a", encoding="utf-8")
    trace_handle.write(f"\n=== TRACE START {datetime.now().isoformat(timespec='seconds')} ===\n")

    def _trace(frame: FrameType, event: str, arg: Any) -> Any:
        if event not in _TRACE_EVENTS:
            return _trace
        code = frame.f_code
        file_path = Path(code.co_filename)
        if not _is_under_root(file_path, app_root):
            return _trace
        if event == "return" and code.co_name == "<module>":
            return _trace
        line = frame.f_lineno
        trace_handle.write(f"{time.time():.3f} {event:<9} {file_path}:{line} {code.co_name}\n")
        return _trace

    sys.settrace(_trace)
    threading.settrace(_trace)


def _is_under_root(file_path: Path, app_root: Path) -> bool:
    try:
        file_path.resolve().relative_to(app_root.resolve())
    except Exception:
        return False
    return True


def _env_debug_enabled() -> bool:
    return str(os.getenv("INTERVIEW_APP_DEBUG", "")).strip().lower() in {"1", "true", "yes", "on", "debug"}


def write_wrapper_crash_report(
    *,
    app_root: Path,
    source: str,
    exc_type: type[BaseException],
    exc_value: BaseException,
    exc_traceback: TracebackType | None,
) -> Path | None:
    crash_dir = app_root / "logs" / "crash-reports"
    crash_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    report_path = crash_dir / f"wrapper-crash-{stamp}.json"
    payload = {
        "timestamp": datetime.now().isoformat(timespec="seconds"),
        "source": source,
        "error_type": exc_type.__name__,
        "error_message": str(exc_value),
        "origin": traceback_origin(exc_traceback),
        "python": sys.version,
        "cwd": str(Path.cwd()),
        "traceback": "".join(traceback.format_exception(exc_type, exc_value, exc_traceback)).strip(),
    }
    try:
        report_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except Exception:
        logging.getLogger("runtime_wrapper").exception("wrapper_crash_report_write_failed")
        return None
    return report_path


def traceback_origin(exc_traceback: TracebackType | None) -> dict[str, Any]:
    if exc_traceback is None:
        return {"function": "<unknown>", "line": None, "file": "<unknown>"}
    current = exc_traceback
    while current.tb_next is not None:
        current = current.tb_next
    code = current.tb_frame.f_code
    return {"function": code.co_name, "line": int(current.tb_lineno), "file": str(code.co_filename)}


def inventory_config_assets(config_dir: Path) -> list[dict[str, Any]]:
    assets: list[dict[str, Any]] = []
    for name, metadata in CONFIG_ASSET_REGISTRY.items():
        path = Path(config_dir) / name
        assets.append(
            {
                "asset": name,
                "path": str(path),
                "exists": path.exists(),
                "owner": metadata["owner"],
                "consumers": list(metadata["consumers"]),
                "schema": metadata["schema"],
            }
        )
    return assets


def load_json_dict(
    path: Path,
    *,
    required: bool,
    context: str,
    default: dict[str, Any] | None = None,
    max_bytes: int = MAX_CONFIG_BYTES,
) -> dict[str, Any]:
    source_path = Path(path)
    if not source_path.exists():
        if required:
            raise FileNotFoundError(f"{context} file not found")
        return deepcopy(default or {})

    size = source_path.stat().st_size
    if size > max_bytes:
        raise ConfigValidationError(f"{context} exceeds the safe size limit")

    try:
        with source_path.open("r", encoding="utf-8") as handle:
            payload = json.load(handle)
    except (json.JSONDecodeError, OSError) as exc:
        raise ConfigValidationError(f"{context} is not valid JSON") from exc

    if not isinstance(payload, dict):
        raise ConfigValidationError(f"{context} must be a JSON object")
    return payload


def validate_rubric_config(payload: dict[str, Any]) -> None:
    _require_keys(
        payload,
        required=["metadata", "scoring", "tracks", "traits", "absolute_disqualifiers"],
        context="rubric.json",
    )
    _expect_type(payload["metadata"], dict, "metadata")
    _expect_type(payload["scoring"], dict, "scoring")
    _expect_type(payload["tracks"], dict, "tracks")
    _expect_type(payload["absolute_disqualifiers"], list, "absolute_disqualifiers")
    if "gateway_requirements" in payload["scoring"]:
        _expect_str_list(payload["scoring"]["gateway_requirements"], "scoring.gateway_requirements")
    for track_key, track in payload["tracks"].items():
        _expect_type(track, dict, f"tracks.{track_key}")
        if "gateway_requirements" in track:
            _expect_str_list(track["gateway_requirements"], f"tracks.{track_key}.gateway_requirements")

    traits = payload["traits"]
    _expect_type(traits, list, "traits")
    if not traits:
        raise ConfigValidationError("rubric.json field 'traits' must not be empty")

    for index, trait in enumerate(traits):
        _expect_type(trait, dict, f"traits[{index}]")
        _require_keys(
            trait,
            required=[
                "id",
                "name",
                "priority",
                "weight",
                "primary_question",
                "descriptors",
                "sample_answers",
                "applicable_tracks",
            ],
            context=f"traits[{index}]",
        )
        _expect_non_empty_str(trait["id"], f"traits[{index}].id")
        _expect_non_empty_str(trait["name"], f"traits[{index}].name")
        _expect_non_empty_str(trait["primary_question"], f"traits[{index}].primary_question")
        _expect_type(trait["priority"], str, f"traits[{index}].priority")
        _expect_type(trait["descriptors"], dict, f"traits[{index}].descriptors")
        _expect_type(trait["sample_answers"], dict, f"traits[{index}].sample_answers")
        _expect_str_list(trait["applicable_tracks"], f"traits[{index}].applicable_tracks")
        if "follow_up_probes" in trait:
            _expect_str_list(trait["follow_up_probes"], f"traits[{index}].follow_up_probes")

        weight = trait["weight"]
        if not isinstance(weight, (int, float)):
            raise ConfigValidationError(f"traits[{index}].weight must be numeric")
        if weight <= 0 or weight > 10:
            raise ConfigValidationError(f"traits[{index}].weight must be between 0 and 10")


def validate_disqualifier_config(payload: dict[str, Any]) -> None:
    questions = payload.get("questions", [])
    _expect_type(questions, list, "questions")
    for index, item in enumerate(questions):
        _expect_type(item, dict, f"questions[{index}]")
        trait_id = str(item.get("trait_id", "")).strip()
        if not trait_id:
            raise ConfigValidationError(f"questions[{index}].trait_id is required")


def normalize_question_overrides_config(payload: dict[str, Any]) -> dict[str, Any]:
    top_level = {
        "track_trait_order": _normalize_track_trait_order(payload.get("track_trait_order", {})),
        "trait_question_overrides": _normalize_trait_overrides(payload.get("trait_question_overrides", {})),
        "custom_questions": _normalize_custom_questions(payload.get("custom_questions", {})),
        "track_question_flow": _normalize_question_flow(payload.get("track_question_flow", {})),
    }
    return top_level


def _normalize_track_trait_order(value: Any) -> dict[str, list[str]]:
    _expect_type(value, dict, "track_trait_order")
    normalized: dict[str, list[str]] = {}
    for track, trait_ids in value.items():
        track_key = str(track).strip()
        if not track_key:
            continue
        _expect_type(trait_ids, list, f"track_trait_order.{track_key}")
        normalized[track_key] = [str(item).strip() for item in trait_ids if str(item).strip()]
    return normalized


def _normalize_trait_overrides(value: Any) -> dict[str, str]:
    _expect_type(value, dict, "trait_question_overrides")
    normalized: dict[str, str] = {}
    for trait_id, text in value.items():
        clean_trait = str(trait_id).strip()
        clean_text = str(text).strip()
        if clean_trait and clean_text:
            normalized[clean_trait] = clean_text
    return normalized


def _normalize_custom_questions(value: Any) -> dict[str, list[dict[str, Any]]]:
    _expect_type(value, dict, "custom_questions")
    normalized: dict[str, list[dict[str, Any]]] = {}
    for track, items in value.items():
        track_key = str(track).strip()
        if not track_key:
            continue
        _expect_type(items, list, f"custom_questions.{track_key}")
        normalized_items: list[dict[str, Any]] = []
        for idx, item in enumerate(items):
            _expect_type(item, dict, f"custom_questions.{track_key}[{idx}]")
            item_id = str(item.get("id", "")).strip()
            text = str(item.get("text", "")).strip()
            order = item.get("order", idx)
            if not item_id or not text:
                continue
            if not isinstance(order, int) or order < 0:
                raise ConfigValidationError(f"custom_questions.{track_key}[{idx}].order must be >= 0")
            normalized_items.append({"id": item_id, "text": text, "order": order})
        normalized[track_key] = normalized_items
    return normalized


def _normalize_question_flow(value: Any) -> dict[str, list[dict[str, str]]]:
    _expect_type(value, dict, "track_question_flow")
    normalized: dict[str, list[dict[str, str]]] = {}
    for track, items in value.items():
        track_key = str(track).strip()
        if not track_key:
            continue
        _expect_type(items, list, f"track_question_flow.{track_key}")
        cleaned_items: list[dict[str, str]] = []
        for idx, item in enumerate(items):
            _expect_type(item, dict, f"track_question_flow.{track_key}[{idx}]")
            kind = str(item.get("type", "")).strip().lower()
            item_id = str(item.get("id", "")).strip()
            if kind not in {"trait", "custom"}:
                continue
            if not item_id:
                continue
            cleaned_items.append({"type": kind, "id": item_id})
        normalized[track_key] = cleaned_items
    return normalized


def _expect_str_list(value: Any, field: str) -> None:
    _expect_type(value, list, field)
    for item in value:
        if not isinstance(item, str):
            raise ConfigValidationError(f"{field} must contain only strings")


def _expect_non_empty_str(value: Any, field: str) -> None:
    if not isinstance(value, str) or not value.strip():
        raise ConfigValidationError(f"{field} must be a non-empty string")


def _expect_type(value: Any, expected_type: type, field: str) -> None:
    if not isinstance(value, expected_type):
        raise ConfigValidationError(f"{field} must be of type {expected_type.__name__}")


def _require_keys(payload: dict[str, Any], *, required: list[str], context: str) -> None:
    for key in required:
        if key not in payload:
            raise ConfigValidationError(f"{context} missing required key: {key}")


def atomic_write_json(path: Path, payload: Any, *, indent: int = 2, ensure_ascii: bool = False) -> None:
    """Write JSON atomically using temp file + flush/fsync + replace."""
    target_path = Path(path)
    target_path.parent.mkdir(parents=True, exist_ok=True)

    fd, temp_path_str = tempfile.mkstemp(
        prefix=f".{target_path.name}.",
        suffix=".tmp",
        dir=target_path.parent,
    )
    temp_path = Path(temp_path_str)
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as file_obj:
            json.dump(payload, file_obj, indent=indent, ensure_ascii=ensure_ascii)
            file_obj.flush()
            os.fsync(file_obj.fileno())
        replace_delay_seconds = 0.05
        replaced = False
        for replace_attempt in range(4):
            try:
                temp_path.replace(target_path)
                replaced = True
                break
            except PermissionError as exc:
                if target_path.exists():
                    try:
                        target_path.chmod(0o666)
                    except OSError:
                        pass
                if replace_attempt == 3:
                    raise
                time.sleep(replace_delay_seconds)
                replace_delay_seconds *= 2
        if not replaced:
            raise PermissionError(f"Could not replace {target_path}")
    except Exception:
        if temp_path.exists():
            temp_path.unlink()
        raise


def safe_read_json(path: Path, default: Any, expected_type: type[Any] | None = None) -> Any:
    """Read JSON with fallback default if missing/invalid/wrong type."""
    target_path = Path(path)
    if not target_path.exists():
        return default
    try:
        with target_path.open("r", encoding="utf-8") as file_obj:
            payload = json.load(file_obj)
    except (json.JSONDecodeError, OSError):
        return default

    if expected_type is None:
        return payload
    if isinstance(payload, expected_type):
        return payload
    return default


def validate_existing_file_path(
    path_text: str,
    *,
    allowed_suffixes: Collection[str] | None = None,
) -> tuple[str, str]:
    text = str(path_text or "").strip()
    if not text:
        return "", "empty"

    candidate = Path(text).expanduser()
    if not candidate.exists():
        return "", "missing"
    if not candidate.is_file():
        return "", "not_file"

    normalized_suffixes = _normalized_suffixes(allowed_suffixes)
    if normalized_suffixes and candidate.suffix.lower() not in normalized_suffixes:
        return "", "unsupported_extension"
    return str(candidate), ""


def _normalized_suffixes(allowed_suffixes: Collection[str] | None) -> set[str]:
    if not allowed_suffixes:
        return set()
    return {str(suffix).strip().lower() for suffix in allowed_suffixes if str(suffix).strip()}


try:  # pragma: no cover - exercised indirectly in environments with python-docx
    from docx import Document as Document  # type: ignore[attr-defined]

    BACKEND = "python-docx"
except Exception:  # pragma: no cover - fallback used in constrained environments
    BACKEND = "fallback"

    @dataclass
    class _Paragraph:
        text: str = ""

    class _Cell:
        def __init__(self, text: str = "") -> None:
            self.paragraphs: list[_Paragraph] = [_Paragraph(text)]

        @property
        def text(self) -> str:
            return "\n".join(p.text for p in self.paragraphs)

        @text.setter
        def text(self, value: str) -> None:
            self.paragraphs = [_Paragraph(value)]

    class _Row:
        def __init__(self, cols: int) -> None:
            self.cells: list[_Cell] = [_Cell() for _ in range(cols)]

    class _Table:
        def __init__(self, rows: int, cols: int) -> None:
            self._cols = cols
            self.rows: list[_Row] = [_Row(cols) for _ in range(rows)]

        def add_row(self) -> _Row:
            row = _Row(self._cols)
            self.rows.append(row)
            return row

    class Document:  # noqa: D401 - mirrors python-docx API
        """Small subset of python-docx's ``Document`` interface."""

        def __init__(self, path: str | Path | None = None) -> None:
            self.paragraphs: list[_Paragraph] = []
            self.tables: list[_Table] = []
            if path is not None:
                self._load(path)

        def add_heading(self, text: str, level: int = 1) -> _Paragraph:
            prefix = "#" * max(1, int(level))
            return self.add_paragraph(f"{prefix} {text}")

        def add_paragraph(self, text: str = "") -> _Paragraph:
            paragraph = _Paragraph(text)
            self.paragraphs.append(paragraph)
            return paragraph

        def add_table(self, rows: int, cols: int) -> _Table:
            table = _Table(rows, cols)
            self.tables.append(table)
            return table

        def save(self, path: str | Path) -> None:
            payload = {
                "paragraphs": [p.text for p in self.paragraphs],
                "tables": [[[cell.text for cell in row.cells] for row in table.rows] for table in self.tables],
            }
            Path(path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

        def _load(self, path: str | Path) -> None:
            target = Path(path)
            if not target.exists():
                return
            try:
                data: dict[str, Any] = json.loads(target.read_text(encoding="utf-8"))
            except Exception:
                self.paragraphs = []
                self.tables = []
                return
            self.paragraphs = [_Paragraph(str(text)) for text in data.get("paragraphs", [])]
            self.tables = []
            for table_rows in data.get("tables", []):
                cols = len(table_rows[0]) if table_rows else 0
                table = _Table(0, cols)
                table.rows = []
                for row_data in table_rows:
                    row = _Row(cols)
                    for idx, value in enumerate(row_data):
                        if idx < len(row.cells):
                            row.cells[idx].text = str(value)
                    table.rows.append(row)
                self.tables.append(table)

_COMPAT_MODULES: tuple[str, ...] = (
    "app_logging",
    "data_store",
    "runtime_wrapper",
)

_WRAPPER_POLICY = (
    "Legacy platform modules are compatibility wrappers during flattening. "
    "New production imports should prefer platform_services."
)


def available_modules() -> tuple[str, ...]:
    return _COMPAT_MODULES


def module_ownership() -> dict[str, str]:
    return {module_name: "platform_services" for module_name in _COMPAT_MODULES}


def wrapper_policy() -> str:
    return _WRAPPER_POLICY


def load_compat_module(module_name: str) -> ModuleType:
    if module_name not in _COMPAT_MODULES:
        raise AttributeError(f"{module_name!r} is not part of platform_services")
    return import_module(module_name)


def public_symbols(module_name: str | None = None) -> tuple[str, ...]:
    module_names = (module_name,) if module_name is not None else _COMPAT_MODULES
    symbols: set[str] = set()
    for compat_name in module_names:
        module = load_compat_module(compat_name)
        symbols.update(name for name in dir(module) if not name.startswith("_"))
    return tuple(sorted(symbols))


def resolve_compat_symbol(symbol_name: str) -> Any:
    for module_name in _COMPAT_MODULES:
        module = import_module(module_name)
        if hasattr(module, symbol_name):
            return getattr(module, symbol_name)
    raise AttributeError(f"platform_services has no attribute {symbol_name!r}")


def __getattr__(name: str) -> Any:
    if name.startswith("__"):
        raise AttributeError(f"platform_services has no attribute {name!r}")
    if name in _COMPAT_MODULES:
        return import_module(name)
    return resolve_compat_symbol(name)


def __dir__() -> list[str]:
    return sorted(set(globals()) | set(_COMPAT_MODULES) | set(public_symbols()))
