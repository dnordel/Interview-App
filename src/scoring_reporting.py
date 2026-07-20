from __future__ import annotations

from dataclasses import dataclass, replace
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal, ROUND_HALF_UP
from importlib import import_module
import importlib.util
import json
import os
from pathlib import Path
import re
import subprocess
import sys
from types import ModuleType
from typing import Any, Final, Iterable, Optional
from urllib import error, request
from urllib.parse import quote, urlparse

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import yaml

from platform_services import Document, sanitize_filename


TITLE_OPTIONS: tuple[str, str] = ("Mr.", "Ms.")
DEFAULT_CANDIDATE_TITLE = "Ms."

CANONICAL_DEGREE_TYPES: tuple[str, ...] = (
    "AA",
    "AS",
    "BA",
    "BS",
    "MA",
    "MS",
    "MBA",
    "PhD",
    "EdD",
)

_EMAIL_PATTERN = re.compile(r"^[A-Za-z0-9.!#$%&'*+/=?^_`{|}~-]+@[A-Za-z0-9-]+(?:\.[A-Za-z0-9-]+)+$")
_PUBLIC_EMAIL_DOMAINS = {
    "gmail.com",
    "yahoo.com",
    "outlook.com",
    "hotmail.com",
    "icloud.com",
}

REQUIRED_PACKET_DOCS: Final[tuple[tuple[str, str], ...]] = (
    ("resume_path", "Resume"),
    ("interview_notes_document_path", "Interview notes document"),
)

OPTIONAL_PACKET_DOCS: Final[tuple[tuple[str, str], ...]] = (
    ("interview_notes_path", "Interview notes (legacy)"),
    ("transcript_path", "Transcript"),
)

ALL_PACKET_DOCS: Final[tuple[tuple[str, str], ...]] = REQUIRED_PACKET_DOCS + OPTIONAL_PACKET_DOCS

ALLOWED_DOC_EXTENSIONS: Final[set[str]] = {
    ".doc",
    ".docx",
    ".pdf",
    ".txt",
    ".rtf",
}

PLACEHOLDER_PATTERN = re.compile(r"\{([A-Za-z_][A-Za-z0-9_]*)\}|\[([A-Za-z_][A-Za-z0-9_]*)\]")
UNKNOWN_PLACEHOLDER_PRESERVE = "preserve"
UNKNOWN_PLACEHOLDER_EMPTY = "empty"
UNKNOWN_PLACEHOLDER_ERROR = "error"
UNKNOWN_PLACEHOLDER_POLICIES = {
    UNKNOWN_PLACEHOLDER_PRESERVE,
    UNKNOWN_PLACEHOLDER_EMPTY,
    UNKNOWN_PLACEHOLDER_ERROR,
}


@dataclass(frozen=True)
class PlaceholderMeta:
    key: str
    label: str
    description: str

    @property
    def token(self) -> str:
        return f"[{self.key}]"

    @property
    def alternate_token(self) -> str:
        return f"{{{self.key}}}"


PLACEHOLDER_METADATA: dict[str, PlaceholderMeta] = {
    "candidate_name": PlaceholderMeta("candidate_name", "Candidate name", "Interview candidate full name."),
    "first_name": PlaceholderMeta("first_name", "First name", "Candidate or employee first name."),
    "last_name": PlaceholderMeta("last_name", "Last name", "Candidate or employee last name."),
    "school": PlaceholderMeta("school", "School", "School name tied to the candidate or reminder run."),
    "track": PlaceholderMeta("track", "Track", "Interview track selected for the candidate."),
    "interview_date": PlaceholderMeta("interview_date", "Interview date", "Interview date value for the candidate."),
    "offer_path": PlaceholderMeta("offer_path", "Offer file path", "Generated offer file path, when available."),
    "employee_summary": PlaceholderMeta("employee_summary", "Employees summary", "Distinct employee names in the reminder batch."),
    "task_summary": PlaceholderMeta("task_summary", "Task summary", "Rendered task line items for reminder/escalation emails."),
    "due_date_summary": PlaceholderMeta("due_date_summary", "Due date summary", "Comma-separated due date list for reminder/escalation emails."),
    "count": PlaceholderMeta("count", "Count", "Number of reminder items in the current batch."),
}

PLACEHOLDERS_BY_CONTEXT: dict[str, tuple[str, ...]] = {
    "interview": ("candidate_name", "first_name", "last_name", "school", "track", "interview_date"),
    "director": ("candidate_name", "first_name", "last_name", "school", "track", "interview_date"),
    "offer": ("candidate_name", "first_name", "last_name", "school", "track", "interview_date", "offer_path"),
    "welcome": ("candidate_name", "first_name", "last_name", "school", "track", "interview_date", "offer_path"),
    "onboarding_reminder": ("employee_summary", "first_name", "last_name", "task_summary", "due_date_summary", "school", "count"),
    "escalation": ("employee_summary", "first_name", "last_name", "task_summary", "due_date_summary", "school", "count"),
}


def normalize_candidate_title(value: Any) -> str:
    text = str(value or "").strip()
    if text in TITLE_OPTIONS:
        return text
    return DEFAULT_CANDIDATE_TITLE


@dataclass
class CandidateQualification:
    has_degree: bool | None = None
    degree_type: str = ""
    degree_in_ece: bool = False
    ece_units_completed: int | None = None
    infant_toddler_class_completed: bool = False
    total_units_completed: int | None = None
    years_experience: int | None = None

    def to_dict(self) -> dict[str, Any]:
        return {
            "has_degree": self.has_degree,
            "degree_type": self.degree_type,
            "degree_in_ece": self.degree_in_ece,
            "ece_units_completed": self.ece_units_completed,
            "infant_toddler_class_completed": self.infant_toddler_class_completed,
            "total_units_completed": self.total_units_completed,
            "years_experience": self.years_experience,
        }

    @classmethod
    def from_dict(cls, payload: dict[str, Any]) -> "CandidateQualification":
        has_degree_raw = payload.get("has_degree", None)
        has_degree = has_degree_raw if isinstance(has_degree_raw, bool) else None
        degree_type = normalize_degree_type(payload.get("degree_type", ""))
        degree_in_ece = bool(payload.get("degree_in_ece", False))
        ece_units_completed = coerce_non_negative_int(payload.get("ece_units_completed"))
        infant_toddler_class_completed = bool(payload.get("infant_toddler_class_completed", False))
        total_units_completed = coerce_non_negative_int(payload.get("total_units_completed"))
        years_experience = coerce_non_negative_int(payload.get("years_experience"))
        return cls(
            has_degree=has_degree,
            degree_type=degree_type,
            degree_in_ece=degree_in_ece,
            ece_units_completed=ece_units_completed,
            infant_toddler_class_completed=infant_toddler_class_completed,
            total_units_completed=total_units_completed,
            years_experience=years_experience,
        )


def normalize_degree_type(value: Any) -> str:
    text = str(value or "").strip().upper()
    if text in CANONICAL_DEGREE_TYPES:
        return text
    return ""


def coerce_non_negative_int(value: Any) -> int | None:
    if isinstance(value, int) and value >= 0:
        return value
    if isinstance(value, str) and value.isdigit():
        return int(value)
    return None


def validate_candidate_qualification(
    has_degree_raw: str,
    degree_type_raw: str,
    degree_in_ece: bool,
    ece_units_raw: str,
    total_units_raw: str,
    infant_toddler_class_completed: bool,
    years_experience_raw: str,
) -> tuple[bool, str, CandidateQualification]:
    has_degree = parse_yes_no(has_degree_raw)
    if has_degree is None:
        return False, "Please confirm whether the candidate has a degree.", CandidateQualification()

    ece_units_completed = coerce_non_negative_int(ece_units_raw)
    if not degree_in_ece and ece_units_completed is None:
        return False, "ECE units completed is required and must be a non-negative whole number unless the degree is in ECE.", CandidateQualification()

    degree_type = ""
    total_units_completed: int | None = None

    if has_degree:
        degree_type = normalize_degree_type(degree_type_raw)
        if not degree_type:
            allowed = ", ".join(CANONICAL_DEGREE_TYPES)
            return False, f"Degree type is required and must be one of: {allowed}.", CandidateQualification()
    else:
        total_units_completed = coerce_non_negative_int(total_units_raw)
        if total_units_completed is None:
            return False, "Total units completed is required when no degree is reported.", CandidateQualification()

    years_experience = coerce_non_negative_int(years_experience_raw)
    if years_experience is None:
        return False, "Years of experience is required and must be a non-negative whole number.", CandidateQualification()

    qualification = CandidateQualification(
        has_degree=has_degree,
        degree_type=degree_type,
        degree_in_ece=degree_in_ece,
        ece_units_completed=ece_units_completed,
        infant_toddler_class_completed=infant_toddler_class_completed,
        total_units_completed=total_units_completed,
        years_experience=years_experience,
    )
    return True, "", qualification


def parse_yes_no(value: str) -> bool | None:
    text = str(value or "").strip().lower()
    if text == "yes":
        return True
    if text == "no":
        return False
    return None


def sanitize_email_subject(subject: str) -> str:
    text = str(subject or "")
    return text.replace("\r", " ").replace("\n", " ").strip()


def is_valid_email_address(value: str) -> bool:
    return bool(_EMAIL_PATTERN.match(str(value or "").strip()))


def sender_email_error_reason(value: str) -> str | None:
    email = str(value or "").strip()
    if not email:
        return "missing"
    if not is_valid_email_address(email):
        return "invalid_format"
    return None


def sender_email_domain_type(value: str) -> str:
    email = str(value or "").strip().lower()
    if "@" not in email:
        return "unknown"
    domain = email.split("@", 1)[1]
    if domain in _PUBLIC_EMAIL_DOMAINS:
        return "public"
    if domain.endswith(".edu"):
        return "education"
    return "organization"


def normalize_referral_packet(packet: dict[str, str] | None) -> dict[str, str]:
    source = packet or {}
    normalized = {key: "" for key, _ in ALL_PACKET_DOCS}
    for key, _ in ALL_PACKET_DOCS:
        normalized[key] = str(source.get(key, "") or "").strip()

    canonical_notes = normalized["interview_notes_document_path"]
    if not canonical_notes:
        canonical_notes = normalized["interview_notes_path"] or normalized["transcript_path"]
        normalized["interview_notes_document_path"] = canonical_notes
    if canonical_notes and not normalized["interview_notes_path"]:
        normalized["interview_notes_path"] = canonical_notes

    return normalized


def missing_required_docs(packet: dict[str, str] | None) -> list[str]:
    normalized = normalize_referral_packet(packet)
    missing: list[str] = []
    for key, label in REQUIRED_PACKET_DOCS:
        if normalized.get(key, ""):
            continue
        missing.append(label)
    return missing


def validate_referral_packet(packet: dict[str, str] | None) -> tuple[bool, list[str]]:
    missing = missing_required_docs(packet)
    return (len(missing) == 0), missing


def is_supported_document_path(path_text: str) -> bool:
    suffix = Path(path_text).suffix.lower()
    return suffix in ALLOWED_DOC_EXTENSIONS


def placeholder_meta_for_context(context: str) -> list[PlaceholderMeta]:
    keys = PLACEHOLDERS_BY_CONTEXT.get(context, ())
    return [PLACEHOLDER_METADATA[key] for key in keys if key in PLACEHOLDER_METADATA]


def placeholder_tokens_for_context(context: str) -> list[str]:
    return [meta.token for meta in placeholder_meta_for_context(context)]


def placeholder_picker_options(contexts: Iterable[str]) -> list[str]:
    options: list[str] = []
    seen: set[str] = set()
    for context in contexts:
        for meta in placeholder_meta_for_context(context):
            if meta.key in seen:
                continue
            options.append(f"{meta.token} / {meta.alternate_token} — {meta.label}: {meta.description}")
            seen.add(meta.key)
    return options


def token_from_picker_label(label: str) -> str:
    prefix = label.split(" ", 1)[0].strip()
    token = prefix.split("/", 1)[0].strip()
    return token if token_to_key(token) else ""


def token_to_key(token: str) -> str:
    text = str(token or "").strip()
    if len(text) < 3:
        return ""
    open_char, close_char = text[0], text[-1]
    if (open_char, close_char) not in {("[", "]"), ("{", "}")}:
        return ""
    key = text[1:-1].strip()
    return key if re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", key) else ""


def _match_key(match: re.Match[str]) -> str:
    return match.group(1) or match.group(2) or ""


def extract_placeholders(text: str) -> set[str]:
    return {_match_key(match) for match in PLACEHOLDER_PATTERN.finditer(str(text or "")) if _match_key(match)}


def missing_placeholder_keys(text: str, values: dict[str, object], context: str) -> list[str]:
    allowlist = set(PLACEHOLDERS_BY_CONTEXT.get(context, ()))
    needed = sorted(extract_placeholders(text) & allowlist)
    return [key for key in needed if not str(values.get(key, "")).strip()]


def find_unknown_placeholders(text: str, context: str) -> set[str]:
    allowlist = set(PLACEHOLDERS_BY_CONTEXT.get(context, ()))
    return {name for name in extract_placeholders(text) if name not in allowlist}


def validate_template_map(templates: dict[str, str], context_by_key: dict[str, str]) -> dict[str, set[str]]:
    unknown_by_key: dict[str, set[str]] = {}
    for key, value in templates.items():
        context = context_by_key.get(key, "")
        unknown = find_unknown_placeholders(value, context)
        if unknown:
            unknown_by_key[key] = unknown
    return unknown_by_key


def render_template(
    template: str,
    values: dict[str, object],
    *,
    context: str | None = None,
    unknown_policy: str = UNKNOWN_PLACEHOLDER_PRESERVE,
) -> str:
    text = str(template or "")
    if not text:
        return ""
    policy = unknown_policy if unknown_policy in UNKNOWN_PLACEHOLDER_POLICIES else UNKNOWN_PLACEHOLDER_PRESERVE
    unknown = find_unknown_placeholders(text, context or "") if context else set()
    if unknown and policy == UNKNOWN_PLACEHOLDER_ERROR:
        names = ", ".join(sorted(unknown))
        raise ValueError(f"Unknown placeholders for context '{context}': {names}")

    allowed = set(PLACEHOLDERS_BY_CONTEXT.get(context or "", ()))

    def _replace(match: re.Match[str]) -> str:
        key = _match_key(match)
        if context and key not in allowed:
            if policy == UNKNOWN_PLACEHOLDER_EMPTY:
                return ""
            return match.group(0)
        if key not in values:
            return match.group(0)
        return str(values.get(key, ""))

    return PLACEHOLDER_PATTERN.sub(_replace, text)


def insert_token_into_widget(widget: Any, token: str) -> bool:
    insert = getattr(widget, "insert", None)
    if not callable(insert):
        return False
    index = getattr(widget, "index", None)
    cursor = "insert"
    if callable(index):
        try:
            cursor = index("insert")
        except Exception:
            cursor = "insert"
    insert(cursor, token)
    focus_set = getattr(widget, "focus_set", None)
    if callable(focus_set):
        focus_set()
    return True


def insert_token_into_focused_widget(root: Any, token: str, allowed_widgets: Iterable[Any]) -> bool:
    focus_get = getattr(root, "focus_get", None)
    if not callable(focus_get):
        return False
    focused = focus_get()
    if focused in set(allowed_widgets):
        return insert_token_into_widget(focused, token)
    return False


class ReportingValidationError(ValueError):
    """Raised when a report cannot be scored or exported due to invalid draft data."""


class ScoringEngine:
    """
    Computes:
    - weighted totals
    - percent of max
    - critical trait override flags
    - absolute disqualifier lock
    - final outcome: Hire / Borderline / No Hire
    """

    @staticmethod
    def _coerce_raw_score(value: Any) -> tuple[Optional[int], int]:
        if isinstance(value, int) and value in {1, 2, 3, 4, 5}:
            return value, value
        if isinstance(value, str) and value.isdigit():
            v = int(value)
            if v in {1, 2, 3, 4, 5}:
                return v, v
        return None, 0

    @staticmethod
    def _get_track_config(rubric: dict[str, Any], track_key: Any) -> dict[str, Any]:
        tracks = rubric.get("tracks", {}) or {}
        resolved_track_key = ScoringEngine._resolve_track_key_for_scoring(rubric, track_key)
        track_cfg = tracks[resolved_track_key]
        return track_cfg

    @staticmethod
    def _resolve_track_key_for_scoring(rubric: dict[str, Any], track_key: Any) -> str:
        tracks = rubric.get("tracks", {}) or {}
        if isinstance(track_key, str) and track_key in tracks:
            return track_key
        if tracks:
            return next(iter(tracks))
        raise ReportingValidationError(
            "Invalid track key in draft. This track is missing from the current rubric."
        )

    @staticmethod
    def _calculate_percent(weighted_total: int, denominator: int) -> tuple[Optional[Decimal], float]:
        if denominator <= 0:
            return None, 0.0

        pct = (Decimal(weighted_total) * Decimal("100")) / Decimal(denominator)
        pct_rounded = pct.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
        return pct, float(pct_rounded)

    @staticmethod
    def _is_critical_priority(priority: Any) -> bool:
        return isinstance(priority, str) and priority.strip().lower() == "critical"

    @staticmethod
    def evaluate(rubric: dict[str, Any], track_key: Any, trait_results: dict[str, dict[str, Any]]) -> dict[str, Any]:
        resolved_track_key = ScoringEngine._resolve_track_key_for_scoring(rubric, track_key)
        track_cfg = ScoringEngine._get_track_config(rubric, resolved_track_key)
        traits = [
            t for t in rubric["traits"]
            if "all" in t["applicable_tracks"] or resolved_track_key in t["applicable_tracks"]
        ]

        trait_q_overrides = rubric.get("trait_question_overrides", {}) or {}

        rows: list[dict[str, Any]] = []
        weighted_total = 0
        weighted_max_possible_included_traits = 0
        skipped_traits_count = 0
        scored_traits_count = 0

        critical_eq_1 = False
        critical_lt_3 = False
        disqualifier_present = False

        for trait in traits:
            tid = trait["id"]
            state = trait_results.get(tid, {}) or {}

            dq = bool(state.get("absolute_disqualifier", False))
            disqualifier_present = disqualifier_present or dq

            raw_display, raw_for_math = ScoringEngine._coerce_raw_score(state.get("raw_score", None))
            skipped = bool(state.get("skipped", False)) or raw_display is None

            weight = int(trait["weight"])
            weighted = 0
            if skipped:
                skipped_traits_count += 1
            else:
                scored_traits_count += 1
                weighted_max_possible_included_traits += 5 * weight
                weighted = raw_for_math * weight
                weighted_total += weighted

            is_critical = ScoringEngine._is_critical_priority(trait.get("priority"))
            has_scored_value = raw_display is not None
            if is_critical and not skipped and has_scored_value:
                if raw_for_math == 1:
                    critical_eq_1 = True
                if raw_for_math < 3:
                    critical_lt_3 = True

            pq = trait_q_overrides.get(tid) or trait["primary_question"]
            selected_signal_ids = list(state.get("selected_signal_ids", []) or [])

            rows.append(
                {
                    "trait_id": tid,
                    "trait_name": trait["name"],
                    "priority": trait["priority"],
                    "weight": weight,
                    "skipped": skipped,
                    "raw_score": raw_display,
                    "raw_score_math": raw_for_math,
                    "weighted_score": weighted,
                    "question_notes": state.get("question_notes", ""),
                    "trait_notes": state.get("trait_notes", ""),
                    "verbatim_notes": state.get("verbatim_notes", ""),
                    "no_example_after_followups": bool(state.get("no_example_after_followups", False)),
                    "absolute_disqualifier": dq,
                    "primary_question": pq,
                    "system_checkbox_score": weighted,
                    "selected_signal_ids": selected_signal_ids,
                }
            )

        configured_max_weighted = int(track_cfg["max_weighted_total"])
        effective_max_weighted = weighted_max_possible_included_traits or configured_max_weighted

        logic_denominator = effective_max_weighted
        pct, percent_of_max = ScoringEngine._calculate_percent(weighted_total, logic_denominator)
        logic_pct, _logic_percent_of_max = ScoringEngine._calculate_percent(weighted_total, logic_denominator)

        percent_label = f"{percent_of_max}%"
        if weighted_max_possible_included_traits == 0:
            percent_label = "N/A (all questions skipped)"

        locked_rule: Optional[str] = None
        pct_for_logic = float(logic_pct) if logic_pct is not None else 0.0

        missing_required_scores = any(
            not row.get("skipped", False) and row.get("raw_score") is None
            for row in rows
        )
        if missing_required_scores:
            outcome = "Incomplete"
            locked_rule = "One or more applicable traits are missing final raw scores"
        elif disqualifier_present:
            outcome = "No Hire"
            locked_rule = "Any Absolute Disqualifier observed => Immediate NO HIRE"
        elif critical_eq_1:
            outcome = "No Hire"
            locked_rule = "Any Critical trait raw score = 1 => Immediate NO HIRE"
        elif critical_lt_3:
            outcome = "No Hire"
            locked_rule = "Any Critical trait raw score < 3 => Cannot assign HIRE"
        elif pct_for_logic >= 80:
            outcome = "Hire"
        elif pct_for_logic >= 65:
            outcome = "Borderline"
        else:
            outcome = "No Hire"

        return {
            "rows": rows,
            "weighted_total": weighted_total,
            "configured_max_weighted_total": configured_max_weighted,
            "max_weighted_total": effective_max_weighted,
            "max_weighted_total_included_traits": weighted_max_possible_included_traits,
            "percent_of_max": percent_of_max,
            "percent_of_max_label": percent_label,
            "skipped_traits_count": skipped_traits_count,
            "scored_traits_count": scored_traits_count,
            "critical_eq_1": critical_eq_1,
            "critical_lt_3": critical_lt_3,
            "disqualifier_present": disqualifier_present,
            "locked_rule": locked_rule,
            "outcome": outcome,
        }


class DraftManager:
    """Saves and loads interview drafts as JSON under <base>/drafts."""

    def __init__(self, base_dir: Path):
        self.base_dir = Path(base_dir)
        self.drafts_dir = self.base_dir / "drafts"
        self.final_dir = self.base_dir / "final"
        self.drafts_dir.mkdir(parents=True, exist_ok=True)
        self.final_dir.mkdir(parents=True, exist_ok=True)

    def save_draft(self, payload: dict[str, Any]) -> Path:
        candidate = payload.get("candidate", {}).get("name", "Unknown")
        safe = sanitize_filename(candidate or "Unknown")
        stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        path = self.drafts_dir / f"draft-{stamp}-{safe}.json"
        with path.open("w", encoding="utf-8") as f:
            json.dump(payload, f, indent=2, ensure_ascii=False)
        return path

    def load_draft(self, path: Path) -> dict[str, Any]:
        with Path(path).open("r", encoding="utf-8") as f:
            return json.load(f)


class DocxExporter:
    """Exports a finalized interview report to a single .docx file (one per candidate)."""

    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _require_candidate(payload: dict[str, Any]) -> dict[str, Any]:
        candidate = payload.get("candidate")
        if not isinstance(candidate, dict):
            raise ReportingValidationError("Draft is missing candidate details; unable to export report.")
        return candidate

    @staticmethod
    def _require_candidate_field(candidate: dict[str, Any], field_name: str) -> str:
        value = candidate.get(field_name)
        if isinstance(value, str) and value.strip():
            return value.strip()
        raise ReportingValidationError(f"Draft is missing required candidate field: '{field_name}'.")

    @staticmethod
    def _extract_full_candidate_transcript(payload: dict[str, Any]) -> str:
        transcript_segments: list[str] = []
        flow_transcript = payload.get("flow_transcript", []) or []
        for item in flow_transcript:
            if not isinstance(item, dict):
                continue
            if str(item.get("type") or "").strip().lower() == "intro":
                continue
            tx = str(item.get("candidate_transcript") or "").strip()
            if tx:
                transcript_segments.append(tx)

        if transcript_segments:
            return "\n\n".join(transcript_segments).strip()

        audio_recording = payload.get("audio_recording", {}) or {}
        if isinstance(audio_recording, list):
            for item in audio_recording:
                if not isinstance(item, dict):
                    continue
                tx = str(item.get("candidate_transcript") or "").strip()
                if tx:
                    transcript_segments.append(tx)
        elif isinstance(audio_recording, dict):
            tx = str(audio_recording.get("candidate_transcript") or "").strip()
            if tx:
                transcript_segments.append(tx)

        return "\n\n".join(transcript_segments).strip()

    def export_basic_interview_notes(self, rubric: dict[str, Any], payload: dict[str, Any], scoring: dict[str, Any]) -> Path:
        candidate = self._require_candidate(payload)
        cname = self._require_candidate_field(candidate, "name")
        interview_date = self._require_candidate_field(candidate, "interview_date")
        track_key = self._require_candidate_field(candidate, "track")
        school = str(candidate.get("school", "") or "").strip()
        qualification = candidate.get("qualification", {}) or {}
        track_cfg = ScoringEngine._get_track_config(rubric, track_key)
        track_label = str(track_cfg.get("label") or track_key)

        body_font = "Aptos"
        navy = "1F4E79"
        green = "385723"
        label_fill = "EEF2F7"
        score_header_fill = "9DC3E6"
        answer_header_fill = "A9D18E"
        box_fill = "F8FAFC"
        border_color = "D9E2F3"
        dark_text = "1F2937"
        content_width_inches = 7.2

        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.6)
            section.right_margin = Inches(0.65)
            section.bottom_margin = Inches(0.55)
            section.left_margin = Inches(0.65)
            section.header_distance = Inches(0.25)
            section.footer_distance = Inches(0.25)

        for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
            style = doc.styles[style_name]
            style.font.name = body_font
            style.font.size = Pt(12)
            style.font.color.rgb = RGBColor.from_string(dark_text)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.12
            if style_name == "Title":
                style.font.size = Pt(18)
                style.font.bold = True
                style.font.color.rgb = RGBColor.from_string(navy)
                style.paragraph_format.space_after = Pt(6)
            elif style_name == "Subtitle":
                style.font.size = Pt(12)
                style.font.color.rgb = RGBColor.from_string("6B7280")
                style.paragraph_format.space_after = Pt(4)
            elif style_name in {"Heading 1", "Heading 2", "Heading 3"}:
                style.font.bold = True
                style.font.color.rgb = RGBColor.from_string(navy if style_name != "Heading 3" else green)
                style.paragraph_format.space_before = Pt(10)
                style.paragraph_format.space_after = Pt(4)

        def _replace_child(parent: Any, tag: str, child: Any) -> None:
            for existing in list(parent):
                if existing.tag == tag:
                    parent.remove(existing)
            parent.append(child)

        def set_cell_shading(cell: Any, fill: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            _replace_child(tc_pr, qn("w:shd"), shd)

        def set_cell_borders(cell: Any, color: str = border_color) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:val"), "single")
                node.set(qn("w:sz"), "4")
                node.set(qn("w:space"), "0")
                node.set(qn("w:color"), color)
                borders.append(node)
            _replace_child(tc_pr, qn("w:tcBorders"), borders)

        def set_cell_margins(cell: Any, *, top: int = 80, bottom: int = 80, start: int = 90, end: int = 90) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for margin_name, margin_value in {
                "top": top,
                "bottom": bottom,
                "start": start,
                "end": end,
            }.items():
                node = tc_mar.find(qn(f"w:{margin_name}"))
                if node is None:
                    node = OxmlElement(f"w:{margin_name}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(margin_value))
                node.set(qn("w:type"), "dxa")

        def set_cell_width(cell: Any, width_dxa: int) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")

        def set_table_geometry(table: Any, widths: list[float], *, indent_dxa: int = 0) -> None:
            width_dxa = [int(width * 1440) for width in widths]
            table.autofit = False
            tbl_pr = table._tbl.tblPr
            tbl_w = OxmlElement("w:tblW")
            tbl_w.set(qn("w:w"), str(sum(width_dxa)))
            tbl_w.set(qn("w:type"), "dxa")
            _replace_child(tbl_pr, qn("w:tblW"), tbl_w)
            tbl_ind = OxmlElement("w:tblInd")
            tbl_ind.set(qn("w:w"), str(indent_dxa))
            tbl_ind.set(qn("w:type"), "dxa")
            _replace_child(tbl_pr, qn("w:tblInd"), tbl_ind)
            tbl_grid = OxmlElement("w:tblGrid")
            for width in width_dxa:
                grid_col = OxmlElement("w:gridCol")
                grid_col.set(qn("w:w"), str(width))
                tbl_grid.append(grid_col)
            for existing in list(table._tbl):
                if existing.tag == qn("w:tblGrid"):
                    table._tbl.remove(existing)
            table._tbl.insert(1, tbl_grid)
            for row in table.rows:
                for index, cell in enumerate(row.cells):
                    set_cell_width(cell, width_dxa[min(index, len(width_dxa) - 1)])

        def normalize_paragraph(paragraph: Any, *, size: float = 12, color: str = dark_text, bold: bool = False) -> None:
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            for run in paragraph.runs:
                run.font.name = body_font
                run.font.size = Pt(size)
                run.font.color.rgb = RGBColor.from_string(color)
                run.bold = bold

        def set_cell_text(
            cell: Any,
            text: Any,
            *,
            bold: bool = False,
            fill: str | None = None,
            color: str = dark_text,
            align: Any | None = None,
            size: float = 12,
        ) -> None:
            cell.text = str(text)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell)
            if fill:
                set_cell_shading(cell, fill)
            for paragraph in cell.paragraphs:
                if align is not None:
                    paragraph.alignment = align
                normalize_paragraph(paragraph, size=size, color=color, bold=bold)

        def add_note(text: str) -> None:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.12
            run = paragraph.add_run(text)
            run.italic = True
            run.font.name = body_font
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor.from_string("6B7280")

        def add_heading(text: str) -> None:
            paragraph = doc.add_heading(text, level=1)
            normalize_paragraph(paragraph, size=14, color=navy, bold=True)

        def add_key_value_table(rows: list[tuple[str, Any]]) -> Any:
            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            set_table_geometry(table, [2.25, content_width_inches - 2.25])
            for label, value in rows:
                cells = table.add_row().cells
                set_cell_text(cells[0], label, bold=True, fill=label_fill, size=12)
                set_cell_text(cells[1], value if str(value).strip() else "Not provided", size=12)
            return table

        def add_spacer(points: float = 4) -> None:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(points)

        def format_bool(value: Any) -> str:
            if value is True:
                return "Yes"
            if value is False:
                return "No"
            return "Not provided"

        def display(value: Any) -> str:
            text = str(value).strip() if value is not None else ""
            return text or "Not provided"

        def custom_answer_by_id() -> dict[str, str]:
            answers: dict[str, str] = {}
            for item in payload.get("custom_answers", []) or []:
                if not isinstance(item, dict):
                    continue
                item_id = str(item.get("id") or "").strip()
                answer = str(item.get("answer") or "").strip()
                if item_id and answer:
                    answers[item_id] = answer
            return answers

        custom_answers = custom_answer_by_id()

        def answer_for_flow_item(item: dict[str, Any], *, seen_scored_question: bool) -> str:
            item_type = str(item.get("type") or "").strip().lower()
            item_id = str(item.get("id") or "").strip()
            use_manual_notes = seen_scored_question and item_type in {"custom", "qualification"}
            if use_manual_notes:
                return str(item.get("evaluator_notes") or custom_answers.get(item_id) or "").strip()
            return str(item.get("candidate_transcript") or item.get("answer") or item.get("evaluator_notes") or custom_answers.get(item_id) or "").strip()

        def first_answer_matching(*needles: str) -> str:
            seen_scored_question = False
            for item in payload.get("flow_transcript", []) or []:
                if not isinstance(item, dict):
                    continue
                if str(item.get("type") or "").strip().lower() == "intro":
                    continue
                prompt = str(item.get("question") or item.get("prompt") or item.get("title") or "").lower()
                answer = answer_for_flow_item(item, seen_scored_question=seen_scored_question)
                if any(needle in prompt for needle in needles) and answer:
                    return answer
                if str(item.get("type") or "").strip().lower() == "trait":
                    seen_scored_question = True
            return ""

        def row_for_trait(item: dict[str, Any]) -> dict[str, Any]:
            trait_id = str(item.get("id") or item.get("trait_id") or "").strip()
            for row in scoring.get("rows", []) or []:
                if not isinstance(row, dict):
                    continue
                if trait_id and trait_id in {str(row.get("trait_id") or ""), str(row.get("id") or "")}:
                    return row
            return {}

        def raw_score_text(row: dict[str, Any], item: dict[str, Any]) -> str:
            if row.get("skipped", False):
                return "Skipped"
            raw_score = row.get("raw_score", item.get("raw_score"))
            return "N/A" if raw_score is None else str(raw_score)

        def first_present_bool(*values: Any) -> Any:
            for value in values:
                if value is True or value is False:
                    return value
            return None

        header = doc.sections[0].header.paragraphs[0]
        header.text = "Launch Pad Learning | Structured Interview Notes Template"
        normalize_paragraph(header, size=8, color="6B7280", bold=True)

        footer = doc.sections[0].footer.paragraphs[0]
        footer.text = "Basic Interview Notes"
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        normalize_paragraph(footer, size=8, color="6B7280")

        title = doc.add_paragraph(style="Title")
        title.add_run("Structured Behavioral Interview Notes")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        normalize_paragraph(title, size=18, color=navy, bold=True)
        subtitle = doc.add_paragraph("Master Template for Python-Generated Candidate Reports", style="Subtitle")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        normalize_paragraph(subtitle, size=12, color="6B7280")

        add_heading("1. Candidate Snapshot")
        add_key_value_table(
            [
                ("Candidate Name", cname),
                ("Interview Date", interview_date),
                ("School / Location", school or "Unknown school"),
                ("Position / Track", track_label),
                ("Final Outcome", scoring.get("outcome", "Incomplete")),
            ]
        )
        add_spacer()

        add_heading("2. Candidate Education and Experience Summary")
        add_key_value_table(
            [
                ("Has Degree", format_bool(qualification.get("has_degree")) if isinstance(qualification, dict) else "Not provided"),
                ("Degree Type", display(qualification.get("degree_type")) if isinstance(qualification, dict) else "Not provided"),
                ("Degree in ECE", format_bool(qualification.get("degree_in_ece")) if isinstance(qualification, dict) else "Not provided"),
                ("ECE Units Completed", display(qualification.get("ece_units_completed")) if isinstance(qualification, dict) else "Not provided"),
                (
                    "Infant / Toddler Class Completed",
                    format_bool(qualification.get("infant_toddler_class_completed")) if isinstance(qualification, dict) else "Not provided",
                ),
                (
                    "Total Units Completed if No Degree",
                    display(qualification.get("total_units_completed")) if isinstance(qualification, dict) else "Not provided",
                ),
                ("Years of Experience", display(qualification.get("years_experience")) if isinstance(qualification, dict) else "Not provided"),
                ("Availability Summary", first_answer_matching("availability", "available")),
                ("Pay Expectation", first_answer_matching("pay", "compensation", "salary")),
                ("Earliest Start Date", first_answer_matching("start")),
            ]
        )
        add_spacer()

        add_heading("3. Score Summary")
        add_note(
            'Skipped scored questions display as "Skipped" or "N/A" and are excluded from the max score when scoring data marks them skipped.'
        )
        rating_rows = [row for row in scoring.get("rows", []) or [] if isinstance(row, dict) and not row.get("skipped", False)]
        skipped_rows = [row for row in scoring.get("rows", []) or [] if isinstance(row, dict) and row.get("skipped", False)]
        if rating_rows:
            score_table = doc.add_table(rows=1, cols=5)
            score_table.style = "Table Grid"
            set_table_geometry(score_table, [2.1, 1.0, 0.75, 1.45, 1.9], indent_dxa=360)
            score_headers = ["Trait", "Priority", "Weight", "Raw Score", "Weighted Score"]
            for cell, header_text in zip(score_table.rows[0].cells, score_headers):
                set_cell_text(cell, header_text, bold=True, fill=score_header_fill, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
            for row in rating_rows:
                cells = score_table.add_row().cells
                set_cell_text(cells[0], row.get("trait_name") or row.get("trait_id") or "")
                set_cell_text(cells[1], row.get("priority") or "")
                set_cell_text(cells[2], row.get("weight") or "")
                raw_score = row.get("raw_score")
                set_cell_text(cells[3], "N/A" if raw_score is None else raw_score, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_text(cells[4], row.get("weighted_score", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            total_cells = score_table.add_row().cells
            set_cell_text(total_cells[0], "Weighted Total", bold=True)
            set_cell_text(total_cells[1], "")
            set_cell_text(total_cells[2], "")
            set_cell_text(total_cells[3], "")
            set_cell_text(total_cells[4], f"{scoring.get('weighted_total', 0)} / {scoring.get('max_weighted_total', 0)}", bold=True)
            skipped_cells = score_table.add_row().cells
            set_cell_text(skipped_cells[0], "Skipped Scored Questions", bold=True)
            set_cell_text(skipped_cells[1], "")
            set_cell_text(skipped_cells[2], "")
            set_cell_text(skipped_cells[3], ", ".join(str(row.get("trait_name") or row.get("trait_id") or "Question") for row in skipped_rows) or "None")
            set_cell_text(skipped_cells[4], "")
        else:
            add_key_value_table(
                [
                    ("Scored Ratings", "No scored trait ratings were recorded."),
                    ("Skipped Scored Questions", str(len(skipped_rows)) if skipped_rows else "None"),
                ]
            )
        add_spacer()

        add_heading("4. Candidate Answers")
        add_note(
            "This section preserves the sequence of the interview. It includes scored and non-scored questions with full-width answer blocks."
        )
        flow_transcript = [
            item
            for item in payload.get("flow_transcript", []) or []
            if isinstance(item, dict) and str(item.get("type") or "").strip().lower() != "intro"
        ]
        if flow_transcript:
            seen_scored_question = False
            for index, item in enumerate(flow_transcript, start=1):
                question = str(item.get("question") or item.get("prompt") or item.get("title") or "Question").strip()
                item_type = str(item.get("type") or "").strip().lower()
                row = row_for_trait(item)
                answer = answer_for_flow_item(item, seen_scored_question=seen_scored_question)
                if item_type == "trait":
                    question_table = doc.add_table(rows=2, cols=4)
                    question_table.style = "Table Grid"
                    set_table_geometry(question_table, [1.25, 2.45, 1.15, 2.35])
                    for cell, header_text in zip(question_table.rows[0].cells, ["Question", "Type / Trait", "Raw Score", "Flags"]):
                        set_cell_text(cell, header_text, bold=True, fill=answer_header_fill, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
                    flag_text = (
                        "No example after follow-ups: "
                        f"{format_bool(first_present_bool(item.get('no_example_after_followups'), row.get('no_example_after_followups')))}\n"
                        "Absolute disqualifier checked: "
                        f"{format_bool(first_present_bool(item.get('absolute_disqualifier_checked'), row.get('absolute_disqualifier')))}"
                    )
                    body_cells = question_table.rows[1].cells
                    set_cell_text(body_cells[0], item.get("flow_index") or index, bold=True)
                    set_cell_text(body_cells[1], row.get("trait_name") or item_type.title() or "Question")
                    set_cell_text(body_cells[2], raw_score_text(row, item), align=WD_ALIGN_PARAGRAPH.CENTER)
                    set_cell_text(body_cells[3], flag_text, size=12)

                q_label = doc.add_paragraph()
                q_label.paragraph_format.space_before = Pt(2)
                q_label.paragraph_format.space_after = Pt(2)
                q_run = q_label.add_run("Question Text")
                q_run.bold = True
                q_run.font.name = body_font
                q_run.font.size = Pt(12)
                q_run.font.color.rgb = RGBColor.from_string(navy)
                question_box = doc.add_table(rows=1, cols=1)
                question_box.style = "Table Grid"
                set_table_geometry(question_box, [content_width_inches])
                set_cell_text(question_box.rows[0].cells[0], question, fill=box_fill, size=12)

                a_label = doc.add_paragraph()
                a_label.paragraph_format.space_before = Pt(2)
                a_label.paragraph_format.space_after = Pt(2)
                a_run = a_label.add_run("Candidate Answer")
                a_run.bold = True
                a_run.font.name = body_font
                a_run.font.size = Pt(12)
                a_run.font.color.rgb = RGBColor.from_string(navy)
                answer_box = doc.add_table(rows=1, cols=1)
                answer_box.style = "Table Grid"
                set_table_geometry(answer_box, [content_width_inches])
                set_cell_text(answer_box.rows[0].cells[0], answer or "No transcript captured.", fill=box_fill, size=12)
                add_spacer(6)
                if item_type == "trait":
                    seen_scored_question = True
        else:
            transcript = self._extract_full_candidate_transcript(payload)
            answer_box = doc.add_table(rows=1, cols=1)
            answer_box.style = "Table Grid"
            set_table_geometry(answer_box, [content_width_inches])
            set_cell_text(answer_box.rows[0].cells[0], transcript or "No transcript captured.", fill=box_fill, size=12)

        school_part = sanitize_filename(school) if school else "UnknownSchool"
        filename = f"{interview_date} - {school_part} - {sanitize_filename(cname)} - Basic Interview Notes.docx"
        out_path = self.output_dir / filename
        doc.save(out_path)
        return out_path

    def export(self, rubric: dict[str, Any], payload: dict[str, Any], scoring: dict[str, Any]) -> Path:
        return self.export_basic_interview_notes(rubric, payload, scoring)
class DirectorEmailDraftError(RuntimeError):
    pass


def open_outlook_draft(*, subject: str, body: str, attachments: list[str], to_recipients: str = "") -> None:
    if not sys.platform.startswith("win"):
        raise DirectorEmailDraftError("Outlook draft is only supported on Windows.")

    existing_files = [str(Path(path).expanduser()) for path in attachments if Path(path).expanduser().exists()]
    escaped_subject = _ps_quote(sanitize_email_subject(subject))
    escaped_body = _ps_quote(body)
    escaped_to = _ps_quote(to_recipients)
    attachment_script = "\n".join(
        [f"$mail.Attachments.Add('{_ps_quote(path)}') | Out-Null" for path in existing_files]
    )

    script = (
        "$ErrorActionPreference = 'Stop'\n"
        "try {\n"
        "  $outlook = New-Object -ComObject Outlook.Application\n"
        "  if ($null -eq $outlook) { throw 'Outlook COM automation is unavailable.' }\n"
        "  $mail = $outlook.CreateItem(0)\n"
        "  if ($null -eq $mail) { throw 'Unable to create an Outlook draft item.' }\n"
        f"  $mail.Subject = '{escaped_subject}'\n"
        f"  $mail.Body = '{escaped_body}'\n"
        f"  $mail.To = '{escaped_to}'\n"
        f"  {attachment_script}\n"
        "  $mail.Display()\n"
        "} catch {\n"
        "  Write-Error $_.Exception.Message\n"
        "  exit 1\n"
        "}\n"
    )

    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
            check=True,
            capture_output=True,
            text=True,
        )
    except subprocess.CalledProcessError as exc:
        stderr = (exc.stderr or "").strip()
        raise DirectorEmailDraftError(f"Could not open Outlook draft. {stderr}".strip()) from exc


def build_mailto_url(*, subject: str, body: str, to_recipients: str = "") -> str:
    recipient_value = _normalize_mailto_recipients(to_recipients)
    query_parts: list[str] = []
    subject_value = sanitize_email_subject(subject)
    body_value = str(body or "").strip()
    if subject_value:
        query_parts.append(f"subject={quote(subject_value)}")
    if body_value:
        query_parts.append(f"body={quote(body_value)}")
    query = "&".join(query_parts)
    if not query:
        return f"mailto:{recipient_value}"
    return f"mailto:{recipient_value}?{query}"


def _normalize_mailto_recipients(to_recipients: str) -> str:
    raw_recipients = str(to_recipients or "").strip()
    if not raw_recipients:
        return ""

    uses_semicolon_separator = ";" in raw_recipients
    separator = ";" if uses_semicolon_separator else ","
    recipients = [token.strip() for token in raw_recipients.replace(";", ",").split(",")]
    encoded_recipients = [quote(recipient, safe="@") for recipient in recipients if recipient]
    return separator.join(encoded_recipients)


def _ps_quote(value: str) -> str:
    return str(value or "").replace("'", "''")


class DirectorReferralError(RuntimeError):
    pass


def _allowed_referral_hosts_from_env() -> set[str]:
    raw_hosts = str(os.environ.get("DIRECTOR_REFERRAL_ALLOWED_HOSTS", "")).strip()
    if not raw_hosts:
        return set()
    hosts = [part.strip().lower() for part in raw_hosts.split(",")]
    return {host for host in hosts if host}


def _validate_referral_endpoint(endpoint: str, allowed_hosts: set[str] | None = None) -> str:
    parsed = urlparse(endpoint)
    host = (parsed.hostname or "").strip().lower()
    if parsed.scheme.lower() != "https":
        raise DirectorReferralError("Director referral endpoint must use HTTPS.")
    if not host:
        raise DirectorReferralError("Director referral endpoint host is missing.")

    enforced_hosts = allowed_hosts if allowed_hosts is not None else _allowed_referral_hosts_from_env()
    if enforced_hosts and host not in {h.lower() for h in enforced_hosts if h}:
        raise DirectorReferralError("Director referral endpoint host is not in the allowlist.")

    return endpoint


def build_director_packet(
    *,
    payload: dict[str, Any],
    scoring: dict[str, Any],
    report_path: Path,
    integration_path: Path,
    referral_packet: dict[str, str],
    generated_transcript_path: Path | None = None,
) -> dict[str, Any]:
    candidate = payload.get("candidate", {}) or {}
    documents = {
        "resume_path": str(referral_packet.get("resume_path", "")).strip(),
        "interview_notes_path": str(referral_packet.get("interview_notes_path", "")).strip(),
        "transcript_path": str(generated_transcript_path or referral_packet.get("transcript_path", "")).strip(),
        "final_report_path": str(report_path),
        "integration_export_path": str(integration_path),
    }
    return {
        "event": "director_referral_packet",
        "sent_at": datetime.now(timezone.utc).isoformat(),
        "candidate": {
            "name": str(candidate.get("name", "")).strip(),
            "interview_date": str(candidate.get("interview_date", "")).strip(),
            "school": str(candidate.get("school", "")).strip(),
            "track": str(candidate.get("track", "")).strip(),
        },
        "scoring": {
            "outcome": scoring.get("outcome"),
            "percent_of_max": scoring.get("percent_of_max"),
            "weighted_total": scoring.get("weighted_total"),
            "max_weighted_total": scoring.get("max_weighted_total"),
        },
        "documents": documents,
    }


def send_director_packet(
    packet: dict[str, Any],
    endpoint: str,
    *,
    timeout_seconds: int = 12,
    allowed_hosts: set[str] | None = None,
) -> dict[str, Any]:
    endpoint_clean = str(endpoint or "").strip()
    if not endpoint_clean:
        raise DirectorReferralError("Director referral endpoint is not configured.")
    _validate_referral_endpoint(endpoint_clean, allowed_hosts)

    req = request.Request(
        endpoint_clean,
        data=json.dumps(packet).encode("utf-8"),
        method="POST",
        headers={"Content-Type": "application/json"},
    )
    try:
        with request.urlopen(req, timeout=timeout_seconds) as resp:
            response_body = resp.read().decode("utf-8", errors="replace")
            return {
                "status_code": int(getattr(resp, "status", 200) or 200),
                "response": response_body,
            }
    except error.HTTPError as exc:
        raise DirectorReferralError(f"Referral endpoint rejected packet ({exc.code}): {exc.reason}") from exc
    except error.URLError as exc:
        raise DirectorReferralError(f"Failed to reach referral endpoint: {exc.reason}") from exc


def default_referral_endpoint() -> str:
    return str(os.environ.get("DIRECTOR_REFERRAL_ENDPOINT", "")).strip()


def append_communication_log(base_dir: Path, event: dict[str, Any], *, candidate_name: str) -> Path:
    comm_dir = Path(base_dir).expanduser() / "communications"
    comm_dir.mkdir(parents=True, exist_ok=True)
    safe_candidate = sanitize_filename(candidate_name or "Unknown")
    out_path = comm_dir / f"director-referral-{safe_candidate}.jsonl"
    line = json.dumps(event, ensure_ascii=False)
    with out_path.open("a", encoding="utf-8") as f:
        f.write(line + "\n")
    return out_path


def normalize_outcome_label(outcome: Any) -> str:
    raw = str(outcome or "").strip().lower()
    mapping = {
        "hire": "hire",
        "borderline": "borderline",
        "no hire": "no_hire",
        "no_hire": "no_hire",
        "nohire": "no_hire",
    }
    return mapping.get(raw, "borderline")


def build_integration_payload(
    payload: dict[str, Any],
    scoring: dict[str, Any],
    *,
    include_flow_slices: bool = True,
) -> dict[str, Any]:
    candidate = payload.get("candidate", {}) or {}
    rows = scoring.get("rows", []) or []
    custom_answers = [
        item
        for item in payload.get("custom_answers", []) or []
        if isinstance(item, dict)
        and str(item.get("type") or "").strip().lower() != "intro"
        and str(item.get("id") or "").strip().lower() != "intro_script"
    ]
    flow_transcript = payload.get("flow_transcript", []) or []

    trait_notes = [_trait_note(row) for row in rows]
    referral_packet = payload.get("referral_packet", {}) or {}
    communication_log = payload.get("communication_log", []) or []

    export_payload: dict[str, Any] = {
        "candidate": {
            "name": str(candidate.get("name", "")).strip(),
            "interview_date": str(candidate.get("interview_date", "")).strip(),
            "school": str(candidate.get("school", "")).strip(),
            "track": str(candidate.get("track", "")).strip(),
            "qualification": {
                "has_degree": candidate.get("qualification", {}).get("has_degree", None),
                "degree_type": str(candidate.get("qualification", {}).get("degree_type", "")).strip(),
                "degree_in_ece": bool(candidate.get("qualification", {}).get("degree_in_ece", False)),
                "ece_units_completed": candidate.get("qualification", {}).get("ece_units_completed", None),
                "infant_toddler_class_completed": bool(
                    candidate.get("qualification", {}).get("infant_toddler_class_completed", False)
                ),
                "total_units_completed": candidate.get("qualification", {}).get("total_units_completed", None),
                "years_experience": candidate.get("qualification", {}).get("years_experience", None),
            },
        },
        "percent_of_max": float(scoring.get("percent_of_max", 0.0) or 0.0),
        "decision": normalize_outcome_label(scoring.get("outcome")),
        "interview_notes": {
            "traits": trait_notes,
            "custom_answers": custom_answers,
        },
        "referral_packet": {
            "resume_path": str(referral_packet.get("resume_path", "")).strip(),
            "interview_notes_path": str(referral_packet.get("interview_notes_path", "")).strip(),
            "transcript_path": str(referral_packet.get("transcript_path", "")).strip(),
        },
        "communication_log": list(communication_log),
    }

    slices = _flow_slices(flow_transcript)
    if include_flow_slices and slices:
        export_payload["flow_transcript_slices"] = slices

    summary_fields = {
        "executive_summary": str(payload.get("executive_summary") or "").strip(),
        "interview_highlights": [
            str(item or "").strip()
            for item in payload.get("interview_highlights", []) or []
            if str(item or "").strip()
        ],
        "answer_summaries": [
            item
            for item in payload.get("answer_summaries", []) or []
            if isinstance(item, dict) and str(item.get("summary") or "").strip()
        ],
        "summary_status": str(payload.get("summary_status") or "").strip(),
        "summary_warnings": [
            str(item or "").strip()
            for item in payload.get("summary_warnings", []) or []
            if str(item or "").strip()
        ],
    }
    for key, value in summary_fields.items():
        if value:
            export_payload[key] = value

    return export_payload


def serialize_integration_payload(
    base_output_dir: Path,
    export_payload: dict[str, Any],
    *,
    candidate_name: str,
) -> Path:
    base_dir = Path(base_output_dir).expanduser().resolve()
    export_dir = (base_dir / "integration_exports").resolve()
    if base_dir not in {export_dir, *export_dir.parents}:
        raise ValueError("Refusing to write integration export outside base output directory")

    export_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_candidate = sanitize_filename(candidate_name or "Unknown")
    out_path = export_dir / f"integration-{stamp}-{safe_candidate}.json"

    with out_path.open("w", encoding="utf-8") as f:
        json.dump(export_payload, f, indent=2, ensure_ascii=False)

    return out_path


def _trait_note(row: dict[str, Any]) -> dict[str, Any]:
    return {
        "trait_id": row.get("trait_id", ""),
        "trait_name": row.get("trait_name", ""),
        "raw_score": row.get("raw_score"),
        "question_notes": row.get("question_notes", ""),
        "trait_notes": row.get("trait_notes", ""),
        "verbatim_notes": row.get("verbatim_notes", ""),
    }


def _flow_slices(flow_transcript: list[dict[str, Any]]) -> list[dict[str, Any]]:
    slices: list[dict[str, Any]] = []
    for item in flow_transcript:
        if str(item.get("type") or "").strip().lower() == "intro":
            continue
        candidate_tx = str(item.get("candidate_transcript", "")).strip()
        if not candidate_tx:
            continue

        slices.append(
            {
                "type": item.get("type", ""),
                "id": item.get("id", ""),
                "question": item.get("question", ""),
                "candidate_transcript": candidate_tx,
            }
        )

    return slices


POSITION_OPTIONS = [
    "Director",
    "Assistant Director",
    "Lead Teacher",
    "Teacher",
    "Teacher/Floater",
    "Cook",
]


@dataclass(frozen=True)
class OfferSchedule:
    gross_daily_hours: Decimal
    net_daily_hours: Decimal
    weekly_hours: Decimal
    employment_type: str


def derive_offer_schedule(start_time_12h: str, end_time_12h: str) -> OfferSchedule:
    start = parse_clock_12h(start_time_12h)
    end = parse_clock_12h(end_time_12h)
    elapsed_minutes = Decimal(int((end - start).total_seconds() // 60))
    if elapsed_minutes <= 0:
        raise ValueError("Shift end must be later than shift start on the same day.")
    gross = elapsed_minutes / Decimal(60)
    net = gross if gross <= Decimal(6) else gross - Decimal(1)
    return OfferSchedule(
        gross_daily_hours=gross,
        net_daily_hours=net,
        weekly_hours=net * Decimal(5),
        employment_type="part_time" if gross < Decimal(6) else "full_time",
    )


def parse_requested_hourly_pay(value: str) -> Decimal | None:
    text = str(value or "").strip()
    if not text:
        return None
    dollar_values = re.findall(r"\$\s*(\d{1,3}(?:\.\d{1,2})?)", text)
    candidates = dollar_values
    if not candidates:
        candidates = re.findall(
            r"(?<!\d)(\d{1,3}(?:\.\d{1,2})?)\s*(?:/\s*(?:hour|hr)|per\s+(?:hour|hr)|hourly)\b",
            text,
            flags=re.IGNORECASE,
        )
    if not candidates and re.fullmatch(r"\s*\d{1,3}(?:\.\d{1,2})?\s*", text):
        candidates = [text.strip()]
    values = {Decimal(item).quantize(Decimal("0.01")) for item in candidates}
    if len(values) != 1:
        return None
    result = values.pop()
    return result if Decimal("0") < result <= Decimal("500") else None


@dataclass(frozen=True)
class OfferInput:
    first_name: str
    last_name: str
    city: str
    position: str
    start_date: date
    start_time_12h: str
    end_time_12h: str
    hourly_pay: float
    hours: Decimal | float | int
    created_on: date
    title: str = ""

    @property
    def pto(self) -> Decimal:
        return Decimal(str(self.hours)) * Decimal(2)

    @property
    def pto2(self) -> Decimal:
        return Decimal(str(self.hours)) * Decimal(4)

    @property
    def offer_deadline(self) -> date:
        return self.created_on + timedelta(days=3)


def build_approval_offer_input(
    *,
    first_name: str,
    last_name: str,
    city: str,
    position: str,
    approval_date: date,
    terms: dict[str, Any],
) -> OfferInput:
    """Convert immutable director offer terms without losing display precision."""
    hours = Decimal(str(terms.get("weekly_hours") or terms.get("hours_week") or 0))
    return OfferInput(
        first_name=first_name,
        last_name=last_name,
        city=city,
        position=position,
        start_date=approval_date,
        start_time_12h=str(terms.get("start_time") or "08:00 AM"),
        end_time_12h=str(terms.get("end_time") or "05:00 PM"),
        hourly_pay=float(terms.get("hourly_pay") or 0),
        hours=hours,
        created_on=approval_date,
        title=str(terms.get("honorific") or terms.get("title") or DEFAULT_CANDIDATE_TITLE),
    )


class OfferTemplateError(ValueError):
    pass


class OfferLetterService:
    ALLOWED_TEMPLATE_SUFFIXES = {".docx", ".docm"}
    PLACEHOLDER_ORDER = [
        "[OfferDate]",
        "[Title]",
        "[First Name]",
        "[Last Name]",
        "[City]",
        "[Position]",
        "[StartDate]",
        "[StartTime]",
        "[EndTime]",
        "[HourlyPay]",
        "[Hours]",
        "[PTO]",
        "[PTO2]",
        "[OfferDeadline]",
    ]

    @staticmethod
    def classify_employment_type(hours: int) -> str:
        return "full_time" if hours >= 30 else "part_time"

    @classmethod
    def validate_template_path(cls, path: Path) -> None:
        if path.suffix.lower() not in cls.ALLOWED_TEMPLATE_SUFFIXES:
            raise OfferTemplateError("Offer template must be a .docx or .docm file.")
        if not path.exists() or not path.is_file():
            raise OfferTemplateError(f"Template not found: {path}")

    @classmethod
    def build_replacements(cls, data: OfferInput) -> dict[str, str]:
        def display_number(value: Decimal | float | int) -> str:
            return format(Decimal(str(value)).normalize(), "f")

        return {
            "[OfferDate]": f"{data.created_on.strftime('%B')} {data.created_on.day}, {data.created_on.year}",
            "[Title]": data.title.strip(),
            "[First Name]": data.first_name.strip(),
            "[Last Name]": data.last_name.strip(),
            "[City]": data.city.strip(),
            "[Position]": data.position.strip(),
            "[StartDate]": data.start_date.strftime("%m/%d/%Y"),
            "[StartTime]": data.start_time_12h.strip(),
            "[EndTime]": data.end_time_12h.strip(),
            "[HourlyPay]": f"{data.hourly_pay:.2f}",
            "[Hours]": display_number(data.hours),
            "[PTO]": display_number(data.pto),
            "[PTO2]": display_number(data.pto2),
            "[OfferDeadline]": data.offer_deadline.strftime("%m/%d/%Y"),
        }

    @classmethod
    def render_offer(cls, template_path: Path, output_path: Path, data: OfferInput) -> Path:
        cls.validate_template_path(template_path)
        replacements = cls.build_replacements(data)

        doc = Document(str(template_path))
        cls._replace_document_text(doc, replacements)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    @classmethod
    def render_approved_offer(
        cls,
        template_path: Path,
        output_path: Path,
        data: OfferInput,
        *,
        approval_date: date,
    ) -> Path:
        from hiring_pipeline import calculate_offer_approval_dates

        cls.validate_template_path(template_path)
        dates = calculate_offer_approval_dates(approval_date)
        approved_data = replace(data, created_on=dates.offer_date, start_date=dates.start_date)
        replacements = cls.build_replacements(approved_data)
        doc = Document(str(template_path))
        has_placeholder = any("[OfferDate]" in paragraph.text for paragraph in doc.paragraphs)
        legacy_date_paragraph = None
        if not has_placeholder:
            date_pattern = re.compile(
                r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}$"
            )
            candidates = [
                paragraph
                for paragraph in doc.paragraphs[:12]
                if date_pattern.fullmatch(paragraph.text.strip())
            ]
            if len(candidates) != 1:
                raise OfferTemplateError("Offer template must contain [OfferDate] or one unambiguous top offer date.")
            legacy_date_paragraph = candidates[0]
        cls._replace_document_text(doc, replacements)
        if legacy_date_paragraph is not None:
            legacy_date_paragraph.text = replacements["[OfferDate]"]
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    @classmethod
    def _replace_document_text(cls, doc: Document, replacements: dict[str, str]) -> None:
        for paragraph in doc.paragraphs:
            cls._replace_in_paragraph(paragraph, replacements)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        cls._replace_in_paragraph(paragraph, replacements)

    @staticmethod
    def _replace_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
        if not paragraph.runs:
            return
        text = "".join(run.text for run in paragraph.runs)
        for token, value in replacements.items():
            text = text.replace(token, value)
        if not paragraph.runs:
            return
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""


def build_offer_filename(first_name: str, last_name: str, created_on: date) -> str:
    date_part = created_on.strftime("%Y-%m-%d")
    name_part = sanitize_filename(f"{first_name.strip()}_{last_name.strip()}")
    return f"{date_part} - Offer - {name_part}.docx"


def build_school_offer_filename(school: str, candidate_full_name: str) -> str:
    normalized_school = str(school or "").strip().casefold()
    prefix = {
        "palmdale": "Launch Pad Learning PMD Offer of Employment to",
        "north long beach": "Launch Pad Learning NLB Offer of Employment to",
        "long beach": "Launch Pad Learning NLB Offer of Employment to",
        "hawthorne": "Preschool Partners, LLC Offer of Employment to",
    }.get(normalized_school)
    if prefix is None:
        raise ValueError(f"Offer filename is not configured for school: {school}")
    safe_name = re.sub(r"[\x00-\x1f]", "", sanitize_filename(candidate_full_name)).rstrip(" .")
    if not safe_name:
        raise ValueError("Candidate name is required for offer filename.")
    return f"{prefix} {safe_name}.docx"


def next_available_offer_path(output_dir: Path, filename: str) -> Path:
    safe_filename = Path(str(filename or "")).name
    if not safe_filename or safe_filename != str(filename) or Path(safe_filename).suffix.casefold() != ".docx":
        raise ValueError("Offer filename must be a plain .docx filename.")
    output_dir = Path(output_dir)
    candidate = output_dir / safe_filename
    if not candidate.exists():
        return candidate
    stem = candidate.stem
    revision = 2
    while True:
        revised = output_dir / f"{stem} ({revision}){candidate.suffix}"
        if not revised.exists():
            return revised
        revision += 1


def parse_clock_12h(value: str) -> datetime:
    return datetime.strptime(value.strip(), "%I:%M %p")


DEFAULT_EXTENDED_GROUP_LABEL = "Extended Signals"
DEFAULT_ENGINE_MODULE_CONTRACT = Path("contracts/trait_based_scoring_engine.contract.yaml")
DEFAULT_ENGINE_RUNTIME_CONTRACT = Path("Trait-Based Scoring/trait_based_scoring_contract.yaml")
RUBRIC_TRAIT_ID_PATTERN = re.compile(r"trait_(\d+)", re.IGNORECASE)
RUNTIME_TRAIT_ID_PATTERN = re.compile(r"T(\d+)(?:_[A-Za-z0-9_]+)?")
SignalUIDefinition = dict[str, Any]
DEFAULT_CORE_DISPLAY_STYLE = "checkbox"
DEFAULT_EXTENDED_DISPLAY_STYLE = "checkbox"
DEFAULT_CORE_SECTION_LABEL = "Core Signals (Most Important)"
DEFAULT_EXTENDED_SECTION_LABEL = "Additional Observations"
DEFAULT_EXTENDED_COLLAPSIBLE = True
DEFAULT_EXTENDED_DEFAULT_COLLAPSED = True
TRAIT_SELECTION_FIELDS = ("selected_signal_ids", "selected_signals", "signal_selections")
SELECTION_COLLECTION_TYPES = (list, tuple, set)

CanonicalSignal = dict[str, Any]
CanonicalSignalGroup = dict[str, Any]

CANONICAL_SIGNAL_COMPARISON_KEYS = (
    "signal_id",
    "label",
    "group_label",
    "weight",
    "is_critical",
)


class CanonicalSignalRecord(dict[str, Any]):
    def __eq__(self, other: object) -> bool:
        if not isinstance(other, dict):
            return super().__eq__(other)
        return _comparison_view(self) == _comparison_view(other)


class CanonicalSignalGroupRecord(dict[str, Any]):
    def __eq__(self, other: object) -> bool:
        if not isinstance(other, dict):
            return super().__eq__(other)
        return {
            "group_id": self.get("group_id"),
            "group_label": self.get("group_label"),
            "signals": self.get("signals", []),
        } == {
            "group_id": other.get("group_id"),
            "group_label": other.get("group_label"),
            "signals": other.get("signals", []),
        }


def _comparison_view(payload: dict[str, Any]) -> dict[str, Any]:
    return {key: payload.get(key) for key in CANONICAL_SIGNAL_COMPARISON_KEYS}


def build_signal_dictionary_index(signal_dictionary: dict[str, Any]) -> dict[str, dict[str, Any]]:
    signals = signal_dictionary.get("signals", []) if isinstance(signal_dictionary, dict) else []
    index: dict[str, dict[str, Any]] = {}
    for signal in signals or []:
        if not isinstance(signal, dict):
            continue
        signal_id = str(signal.get("id", "") or "").strip()
        if signal_id:
            index[signal_id] = signal
    return index


def normalize_trait_signal(signal: dict[str, Any], *, default_group_label: str) -> CanonicalSignal:
    signal_id = resolve_trait_signal_selection_id(signal)
    return CanonicalSignalRecord(
        {
            "signal_id": signal_id,
            "runtime_signal_id": resolve_trait_signal_runtime_id(signal),
            "selection_aliases": signal_selection_aliases(signal),
            "label": resolve_trait_signal_label(signal, fallback=signal_id),
            "group_label": str(signal.get("group", "") or default_group_label),
            "weight": resolve_trait_signal_weight(signal),
            "is_critical": bool(signal.get("is_critical", False)),
        }
    )


def resolve_trait_signal_selection_id(signal: dict[str, Any]) -> str:
    if str(signal.get("id", "") or "").strip() and signal.get("maps_to"):
        return str(signal.get("id") or "").strip()
    if str(signal.get("ref", "") or "").strip():
        return str(signal.get("ref") or "").strip()
    return resolve_trait_signal_runtime_id(signal)


def resolve_trait_signal_runtime_id(signal: dict[str, Any]) -> str:
    return str(signal.get("ref") or signal.get("id") or "").strip()


def signal_selection_aliases(signal: dict[str, Any]) -> list[str]:
    aliases: list[str] = []
    question_signal_id = str(signal.get("id", "") or "").strip()
    if question_signal_id:
        aliases.append(question_signal_id)
    runtime_signal_id = resolve_trait_signal_runtime_id(signal)
    if runtime_signal_id:
        aliases.append(runtime_signal_id)
    for mapped_signal_id in signal.get("maps_to", []) or []:
        alias = str(mapped_signal_id or "").strip()
        if alias:
            aliases.append(alias)
    return list(dict.fromkeys(aliases))


def resolve_trait_signal_label(signal: dict[str, Any], *, fallback: str = "") -> str:
    return str(signal.get("label", "") or fallback).strip()


def resolve_trait_signal_weight(signal: dict[str, Any]) -> float:
    raw_weight = signal.get("weight", signal.get("base_weight", signal.get("default_weight", 0)))
    if isinstance(raw_weight, str) and raw_weight.strip().upper() == "AUTO_NO_HIRE":
        return 0.0
    return float(raw_weight or 0)


def normalize_core_signals(core_signals: list[dict[str, Any]]) -> list[CanonicalSignal]:
    return _normalize_signal_collection(core_signals, default_group_label="Core")


def normalize_extended_signal_groups(
    trait_definition: dict[str, Any],
    *,
    signal_dictionary_index: dict[str, dict[str, Any]] | None = None,
) -> list[CanonicalSignalGroup]:
    explicit_groups = trait_definition.get("extended_signal_groups", []) or []
    if explicit_groups:
        return _normalize_explicit_extended_groups(explicit_groups)
    return _normalize_runtime_extended_signals(
        trait_definition.get("extended_signals", []) or [],
        signal_dictionary_index=signal_dictionary_index or {},
    )


def iter_trait_schema_signals(trait_definition: dict[str, Any]) -> list[dict[str, Any]]:
    signals = list(trait_definition.get("core_signals", []) or [])
    for group in trait_definition.get("extended_signal_groups", []) or []:
        if isinstance(group, dict):
            signals.extend(group.get("signals", []) or [])
    signals.extend(trait_definition.get("extended_signals", []) or [])
    return [signal for signal in signals if isinstance(signal, dict)]


def _normalize_signal_collection(signals: list[dict[str, Any]], *, default_group_label: str) -> list[CanonicalSignal]:
    normalized_signals: list[CanonicalSignal] = []
    for signal in signals or []:
        if not isinstance(signal, dict):
            continue
        normalized = normalize_trait_signal(signal, default_group_label=default_group_label)
        if normalized["signal_id"]:
            normalized_signals.append(normalized)
    return normalized_signals


def _normalize_explicit_extended_groups(groups: list[dict[str, Any]]) -> list[CanonicalSignalGroup]:
    normalized_groups: list[CanonicalSignalGroup] = []
    for index, group in enumerate(groups or [], start=1):
        if not isinstance(group, dict):
            continue
        group_label = str(group.get("group_label", "") or f"Group {index}").strip()
        normalized_groups.append(
            CanonicalSignalGroupRecord(
                {
                    "group_id": str(group.get("group_id", "") or f"group_{index}").strip(),
                    "group_label": group_label,
                    "signals": _normalize_signal_collection(
                        group.get("signals", []) or [],
                        default_group_label=group_label,
                    ),
                }
            )
        )
    return normalized_groups


def _normalize_runtime_extended_signals(
    signals: list[dict[str, Any]],
    *,
    signal_dictionary_index: dict[str, dict[str, Any]],
) -> list[CanonicalSignalGroup]:
    grouped: dict[str, list[CanonicalSignal]] = {}
    ordered_labels: list[str] = []
    for signal in signals or []:
        if not isinstance(signal, dict):
            continue
        group_label = _runtime_extended_group_label(signal, signal_dictionary_index)
        if group_label not in grouped:
            grouped[group_label] = []
            ordered_labels.append(group_label)
        normalized = normalize_trait_signal(signal, default_group_label=group_label)
        if normalized["signal_id"]:
            grouped[group_label].append(normalized)
    return [
        CanonicalSignalGroupRecord(
            {
                "group_id": _group_id_from_label(group_label),
                "group_label": group_label,
                "signals": grouped[group_label],
            }
        )
        for group_label in ordered_labels
        if grouped[group_label]
    ]


def _runtime_extended_group_label(signal: dict[str, Any], signal_dictionary_index: dict[str, dict[str, Any]]) -> str:
    explicit_group = str(signal.get("group", "") or "").strip()
    if explicit_group:
        return explicit_group
    signal_category = str(signal.get("signal_category", "") or signal.get("category", "") or "").strip()
    if signal_category:
        return signal_category
    for mapped_signal_id in signal.get("maps_to", []) or []:
        dictionary_signal = signal_dictionary_index.get(str(mapped_signal_id).strip(), {})
        category = str(dictionary_signal.get("category", "") or "").strip()
        if category:
            return category
    return DEFAULT_EXTENDED_GROUP_LABEL


def _group_id_from_label(group_label: str) -> str:
    words = [part.lower() for part in str(group_label).split() if part.strip()]
    if not words:
        return "extended_signals"
    return "_".join(words)


def load_trait_definitions_from_runtime_bundle(runtime_bundle: dict[str, Any]) -> list[dict[str, Any]]:
    bundled_traits = runtime_bundle.get("traits")
    if isinstance(bundled_traits, list):
        bundled_trait_definitions = [_normalize_trait_definition(item) for item in bundled_traits if isinstance(item, dict)]
        if bundled_trait_definitions:
            return bundled_trait_definitions

    resolved_trait_dir = _resolve_traits_dir_from_bundle(runtime_bundle)
    if resolved_trait_dir is not None:
        return load_trait_definitions_from_dir(resolved_trait_dir)

    runtime_contract_path = runtime_bundle.get("runtime_contract_path")
    if runtime_contract_path:
        return load_trait_definitions_from_contract(runtime_contract_path)
    return []


def load_trait_definitions_from_contract(runtime_contract_path: str | Path) -> list[dict[str, Any]]:
    contract_payload = _load_yaml(Path(runtime_contract_path))
    trait_dir = _resolve_traits_dir_from_contract_payload(contract_payload, runtime_contract_path)
    return load_trait_definitions_from_dir(trait_dir)


def load_trait_definitions_from_dir(traits_dir: str | Path) -> list[dict[str, Any]]:
    resolved_dir = Path(traits_dir).expanduser().resolve()
    if not resolved_dir.exists() or not resolved_dir.is_dir():
        return []

    trait_definitions: list[dict[str, Any]] = []
    for trait_path in sorted(resolved_dir.glob("T*.json")):
        trait_payload = _load_json(trait_path)
        if isinstance(trait_payload, dict):
            trait_definitions.append(_normalize_trait_definition(trait_payload))
    return trait_definitions


def canonical_trait_id(trait_id: Any) -> str:
    candidate = str(trait_id or "").strip()
    if not candidate:
        return ""
    rubric_match = RUBRIC_TRAIT_ID_PATTERN.fullmatch(candidate)
    if rubric_match:
        return f"trait_{int(rubric_match.group(1))}"
    runtime_match = RUNTIME_TRAIT_ID_PATTERN.fullmatch(candidate)
    if runtime_match:
        return f"trait_{int(runtime_match.group(1))}"
    return candidate


def trait_id_aliases(trait_id: Any) -> list[str]:
    candidate = str(trait_id or "").strip()
    if not candidate:
        return []
    canonical_id = canonical_trait_id(candidate)
    aliases = [canonical_id]
    if candidate != canonical_id:
        aliases.append(candidate)
    return list(dict.fromkeys(alias for alias in aliases if alias))


def _normalize_trait_definition(trait_definition: dict[str, Any]) -> dict[str, Any]:
    canonical_id = canonical_trait_id(trait_definition.get("trait_id"))
    normalized = dict(trait_definition)
    normalized["trait_id"] = canonical_id or str(trait_definition.get("trait_id", "") or "").strip()
    normalized["trait_aliases"] = trait_id_aliases(trait_definition.get("trait_id"))
    return normalized


def _resolve_traits_dir_from_bundle(runtime_bundle: dict[str, Any]) -> Path | None:
    resolved_paths = runtime_bundle.get("resolved_paths")
    if not isinstance(resolved_paths, dict):
        return None

    traits_dir = resolved_paths.get("traits_dir")
    if not traits_dir:
        return None
    return Path(traits_dir)


def _resolve_traits_dir_from_contract_payload(
    contract_payload: dict[str, Any],
    runtime_contract_path: str | Path,
) -> Path:
    paths = contract_payload.get("paths")
    if not isinstance(paths, dict):
        return Path(runtime_contract_path).expanduser().resolve().parent

    traits_dir = paths.get("traits_dir")
    if not isinstance(traits_dir, str) or not traits_dir.strip():
        return Path(runtime_contract_path).expanduser().resolve().parent

    base_dir = Path(runtime_contract_path).expanduser().resolve().parent
    return (base_dir / traits_dir).resolve()


def _load_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if isinstance(payload, dict):
        return payload
    return {}


def _load_yaml(path: Path) -> dict[str, Any]:
    resolved_path = path.expanduser().resolve()
    if not resolved_path.exists():
        return {}

    payload = yaml.safe_load(resolved_path.read_text(encoding="utf-8"))
    if isinstance(payload, dict):
        return payload
    return {}


def default_signal_ui_definition(trait_id: str) -> SignalUIDefinition:
    return {
        "trait_id": trait_id,
        "core_display_style": DEFAULT_CORE_DISPLAY_STYLE,
        "extended_display_style": DEFAULT_EXTENDED_DISPLAY_STYLE,
        "core_section_label": DEFAULT_CORE_SECTION_LABEL,
        "extended_section_label": DEFAULT_EXTENDED_SECTION_LABEL,
        "extended_collapsible": DEFAULT_EXTENDED_COLLAPSIBLE,
        "extended_default_collapsed": DEFAULT_EXTENDED_DEFAULT_COLLAPSED,
        "core_signals": [],
        "extended_groups": [],
        "valid_signal_ids": [],
    }


def load_trait_signal_ui_definition(
    trait_id: str,
    *,
    engine_module_contract_path: str | Path = DEFAULT_ENGINE_MODULE_CONTRACT,
    engine_runtime_contract_path: str | Path = DEFAULT_ENGINE_RUNTIME_CONTRACT,
) -> SignalUIDefinition:
    runtime_bundle = load_module_contract_runtime_bundle(
        engine_module_contract_path=engine_module_contract_path,
        engine_runtime_contract_path=engine_runtime_contract_path,
    )
    validate_runtime_bundle_metadata(runtime_bundle)
    trait_definition = _find_trait_definition(runtime_bundle.get("trait_definitions", []), trait_id)
    if not trait_definition:
        return _empty_signal_ui_definition(trait_id)
    return _build_signal_ui_definition(runtime_bundle, trait_definition)


def ensure_trait_signal_ui_definition(
    trait_id: str,
    *,
    engine_module_contract_path: str | Path = DEFAULT_ENGINE_MODULE_CONTRACT,
    engine_runtime_contract_path: str | Path = DEFAULT_ENGINE_RUNTIME_CONTRACT,
) -> SignalUIDefinition:
    definition = load_trait_signal_ui_definition(
        trait_id,
        engine_module_contract_path=engine_module_contract_path,
        engine_runtime_contract_path=engine_runtime_contract_path,
    )
    if definition.get("valid_signal_ids"):
        return definition
    raise ReportingValidationError(
        f"Trait scoring configuration mismatch: rubric trait '{trait_id}' is missing a runtime signal definition."
    )


def normalize_trait_signal_selection_state(trait_state: dict[str, Any] | None, valid_signal_ids: list[str]) -> list[str]:
    normalized_state = normalize_trait_state_item(trait_state)
    allowed = set(valid_signal_ids)
    selected_signal_ids = normalized_state.get("selected_signal_ids", []) or []
    if not allowed:
        return []
    return [signal_id for signal_id in selected_signal_ids if signal_id in allowed]


def write_canonical_selected_signal_ids(trait_state: dict[str, Any], selected_signal_ids: list[str]) -> None:
    canonical_ids = list(dict.fromkeys(signal_id for signal_id in selected_signal_ids if str(signal_id).strip()))
    trait_state["selected_signal_ids"] = canonical_ids
    trait_state.pop("selected_signals", None)
    trait_state.pop("signal_selections", None)


def count_selected_trait_checkbox_entries(state: dict[str, Any], trait_id: str) -> int:
    selection_value = resolve_trait_selection_value(state)
    if selection_value is None:
        return 0
    return _count_selected_entries(selection_value, trait_id)


def resolve_trait_selection_value(state: dict[str, Any]) -> Any:
    for field_name in TRAIT_SELECTION_FIELDS:
        if field_name in state:
            return state[field_name]
    return None


def trait_requires_signal_selection(raw_state: dict[str, Any], normalized_trait_state: dict[str, Any], trait_id: str) -> bool:
    if normalized_trait_state["skipped"]:
        return False
    if normalized_trait_state["absolute_disqualifier"]:
        return False
    return count_selected_trait_checkbox_entries(raw_state, trait_id) == 0


def _count_selected_entries(selection_value: Any, trait_id: str) -> int:
    if isinstance(selection_value, dict):
        return _count_selected_mapping_entries(selection_value, trait_id)
    if isinstance(selection_value, SELECTION_COLLECTION_TYPES):
        return _count_selected_sequence_entries(selection_value, trait_id)
    raise ValueError(
        f"Trait '{trait_id}' has malformed trait checkbox selections: expected mapping or list-like value."
    )


def _count_selected_mapping_entries(selection_value: dict[str, Any], trait_id: str) -> int:
    if _is_boolean_mapping(selection_value):
        return sum(1 for is_selected in selection_value.values() if is_selected)
    if _is_grouped_selection_mapping(selection_value):
        return sum(_count_selected_sequence_entries(group_value, trait_id) for group_value in selection_value.values())
    raise ValueError(
        f"Trait '{trait_id}' has malformed trait checkbox selections: mapping entries must be booleans or list-like groups."
    )


def _count_selected_sequence_entries(selection_value: Any, trait_id: str) -> int:
    count = 0
    for item in selection_value:
        count += _count_selected_sequence_item(item, trait_id)
    return count


def _count_selected_sequence_item(item: Any, trait_id: str) -> int:
    if isinstance(item, bool):
        return int(item)
    if isinstance(item, str):
        if item.strip():
            return 1
        raise ValueError(f"Trait '{trait_id}' has malformed trait checkbox selections: blank signal reference.")
    if isinstance(item, dict):
        return _count_selected_item_mapping(item, trait_id)
    raise ValueError(
        f"Trait '{trait_id}' has malformed trait checkbox selections: list items must be strings, booleans, or mappings."
    )


def _count_selected_item_mapping(item: dict[str, Any], trait_id: str) -> int:
    if "selected" not in item:
        raise ValueError(
            f"Trait '{trait_id}' has malformed trait checkbox selections: mapping items must include 'selected'."
        )
    if not isinstance(item.get("selected"), bool):
        raise ValueError(
            f"Trait '{trait_id}' has malformed trait checkbox selections: 'selected' must be a boolean."
        )
    return int(item["selected"])


def _is_boolean_mapping(selection_value: dict[str, Any]) -> bool:
    return bool(selection_value) and all(isinstance(value, bool) for value in selection_value.values())


def _is_grouped_selection_mapping(selection_value: dict[str, Any]) -> bool:
    return bool(selection_value) and all(
        isinstance(value, dict) or isinstance(value, SELECTION_COLLECTION_TYPES)
        for value in selection_value.values()
    )


def _find_trait_definition(trait_definitions: list[dict[str, Any]], trait_id: str) -> dict[str, Any]:
    candidate_ids = set(trait_id_aliases(trait_id))
    canonical_id = canonical_trait_id(trait_id)
    if canonical_id:
        candidate_ids.add(canonical_id)
    if not candidate_ids:
        return {}
    for trait_definition in trait_definitions:
        definition_ids = set(trait_id_aliases(trait_definition.get("trait_id")))
        definition_ids.update(str(alias).strip() for alias in trait_definition.get("trait_aliases", []) or [])
        if candidate_ids.intersection(definition_ids):
            return trait_definition
    runtime_prefix = _runtime_trait_id_alias(str(trait_id or ""))
    if not runtime_prefix:
        return {}
    for trait_definition in trait_definitions:
        candidate = str(trait_definition.get("trait_id", "") or "").strip()
        if candidate.startswith(runtime_prefix):
            return trait_definition
    return {}


def _runtime_trait_id_alias(trait_id: str) -> str:
    match = re.fullmatch(r"trait_(\d+)", trait_id.strip().lower())
    if not match:
        return ""
    numeric_prefix = f"T{int(match.group(1))}_"
    return numeric_prefix


def _empty_signal_ui_definition(trait_id: str) -> SignalUIDefinition:
    return default_signal_ui_definition(trait_id)


def _build_signal_ui_definition(runtime_bundle: dict[str, Any], trait_definition: dict[str, Any]) -> SignalUIDefinition:
    ui_config = runtime_bundle.get("config", {}).get("ui", {})
    core_config = ui_config.get("core_signals", {})
    extended_config = ui_config.get("extended_signals", {})
    signal_dictionary_index = build_signal_dictionary_index(runtime_bundle.get("signal_dictionary", {}))
    core_signals = normalize_core_signals(trait_definition.get("core_signals", []))
    extended_groups = normalize_extended_signal_groups(
        trait_definition,
        signal_dictionary_index=signal_dictionary_index,
    )
    core_signals = [_normalize_ui_signal(signal) for signal in core_signals]
    extended_groups = [_normalize_ui_group(group) for group in extended_groups]
    valid_signal_ids = _collect_valid_signal_ids(core_signals)
    for group in extended_groups:
        valid_signal_ids.extend(_collect_valid_signal_ids(group.get("signals", [])))
    return {
        "trait_id": str(trait_definition.get("trait_id", "") or ""),
        "core_display_style": str(core_config.get("display_style", DEFAULT_CORE_DISPLAY_STYLE) or DEFAULT_CORE_DISPLAY_STYLE),
        "extended_display_style": str(
            extended_config.get("display_style", DEFAULT_EXTENDED_DISPLAY_STYLE) or DEFAULT_EXTENDED_DISPLAY_STYLE
        ),
        "core_section_label": str(core_config.get("section_label", DEFAULT_CORE_SECTION_LABEL) or DEFAULT_CORE_SECTION_LABEL),
        "extended_section_label": str(
            extended_config.get("section_label", DEFAULT_EXTENDED_SECTION_LABEL) or DEFAULT_EXTENDED_SECTION_LABEL
        ),
        "extended_collapsible": bool(extended_config.get("collapsible", DEFAULT_EXTENDED_COLLAPSIBLE)),
        "extended_default_collapsed": bool(
            extended_config.get("default_collapsed", DEFAULT_EXTENDED_DEFAULT_COLLAPSED)
        ),
        "core_signals": core_signals,
        "extended_groups": extended_groups,
        "valid_signal_ids": list(dict.fromkeys(valid_signal_ids)),
    }


def _normalize_ui_group(group: dict[str, Any]) -> dict[str, Any]:
    normalized_group = dict(group)
    normalized_group["signals"] = [_normalize_ui_signal(signal) for signal in group.get("signals", [])]
    return normalized_group


def _normalize_ui_signal(signal: dict[str, Any]) -> dict[str, Any]:
    normalized_signal = dict(signal)
    aliases = [str(alias).strip() for alias in signal.get("selection_aliases", []) if str(alias).strip()]
    if aliases:
        normalized_signal["signal_id"] = _preferred_ui_signal_id(aliases)
        normalized_signal["selection_aliases"] = aliases
        return normalized_signal
    signal_id = str(signal.get("signal_id", "") or "").strip()
    normalized_signal["selection_aliases"] = [signal_id] if signal_id else []
    return normalized_signal


def _preferred_ui_signal_id(aliases: list[str]) -> str:
    for alias in aliases:
        if alias.startswith("Q"):
            return alias
    if len(aliases) == 1:
        return aliases[0]
    return aliases[1]


def _collect_valid_signal_ids(signals: list[dict[str, Any]]) -> list[str]:
    valid_signal_ids: list[str] = []
    for signal in signals:
        aliases = signal.get("selection_aliases", [])
        if aliases:
            valid_signal_ids.extend(str(alias).strip() for alias in aliases if str(alias).strip())
            continue
        signal_id = str(signal.get("signal_id", "") or "").strip()
        if signal_id:
            valid_signal_ids.append(signal_id)
    return valid_signal_ids

VALID_RAW_SCORES = {1, 2, 3, 4, 5}
CANONICAL_TRAIT_STATE_SCHEMA_VERSION = "1.2.0"
OUTPUT_SCHEMA_VERSION = "1.0.0"
DECISION_LABELS = {
    "strong_hire": "Hire",
    "hire": "Hire",
    "borderline": "Borderline",
    "no_hire": "No Hire",
}


def build_trait_scoring_payload(
    rubric: dict[str, Any],
    track_key: Any,
    trait_state: dict[str, dict[str, Any]] | None,
    *,
    engine_module_contract_path: str | Path = DEFAULT_ENGINE_MODULE_CONTRACT,
    engine_runtime_contract_path: str | Path = DEFAULT_ENGINE_RUNTIME_CONTRACT,
) -> dict[str, Any]:
    runtime_bundle = load_module_contract_runtime_bundle(
        engine_module_contract_path=engine_module_contract_path,
        engine_runtime_contract_path=engine_runtime_contract_path,
    )
    normalized_state = normalize_app_trait_state(trait_state)
    validate_normalized_state(normalized_state)
    engine_output = invoke_scoring_engine(
        rubric,
        track_key,
        normalized_state,
        runtime_bundle=runtime_bundle,
        engine_runtime_contract_path=runtime_bundle["runtime_contract_path"],
    )
    return map_engine_output_to_normalized_shape(
        rubric=rubric,
        track_key=track_key,
        normalized_state=normalized_state,
        engine_output=engine_output,
        runtime_bundle=runtime_bundle,
    )


def load_module_contract_runtime_bundle(
    *,
    engine_module_contract_path: str | Path,
    engine_runtime_contract_path: str | Path,
) -> dict[str, Any]:
    module_contract_path = _resolve_contract_path(engine_module_contract_path)
    module_contract = _load_yaml(module_contract_path)
    runtime_contract_path = _resolve_contract_path(engine_runtime_contract_path)
    runtime_bundle = _load_runtime_bundle(runtime_contract_path)
    runtime_bundle_with_path = {**runtime_bundle, "runtime_contract_path": str(runtime_contract_path)}
    trait_definitions = load_trait_definitions(runtime_bundle_with_path)
    runtime_error = runtime_bundle.get("runtime_error")
    metadata = {
        "module_contract": module_contract,
        "runtime_contract_path": str(runtime_contract_path),
        "runtime_bundle_loaded": bool(runtime_bundle) and not runtime_error,
        "runtime_error": runtime_error,
        "trait_definitions": trait_definitions,
    }
    return {**runtime_bundle, **metadata}


def load_trait_definitions(runtime_bundle: dict[str, Any]) -> list[dict[str, Any]]:
    return load_trait_definitions_from_runtime_bundle(runtime_bundle)


def validate_runtime_bundle_metadata(runtime_bundle: dict[str, Any]) -> None:
    runtime_error = str(runtime_bundle.get("runtime_error", "") or "").strip()
    if runtime_error:
        raise ReportingValidationError(f"Unable to load scoring runtime bundle: {runtime_error}")
    if runtime_bundle.get("runtime_bundle_loaded"):
        return
    runtime_contract_path = runtime_bundle.get("runtime_contract_path", "runtime bundle")
    raise ReportingValidationError(f"Unable to load scoring runtime bundle: {runtime_contract_path}")


def normalize_skipped(value: Any) -> bool:
    return _normalize_bool(value)


def normalize_absolute_disqualifier(value: Any) -> bool:
    return _normalize_bool(value)


def coerce_raw_score(value: Any) -> int | None:
    if value in {None, ""}:
        return None
    if isinstance(value, bool):
        return None
    if isinstance(value, int) and value in VALID_RAW_SCORES:
        return value
    if isinstance(value, str):
        stripped = value.strip()
        if stripped.isdigit():
            parsed = int(stripped)
            if parsed in VALID_RAW_SCORES:
                return parsed
        return None
    return None


def normalize_verbatim_notes(value: Any) -> str:
    return str(value or "").strip()


def normalize_trait_state_item(state: dict[str, Any] | None) -> dict[str, Any]:
    source = state if isinstance(state, dict) else {}
    skipped = normalize_skipped(source.get("skipped", False))
    raw_score = coerce_raw_score(source.get("raw_score"))
    selected_signal_ids = _normalize_selected_signal_ids(source)
    normalized = {
        "schema_version": CANONICAL_TRAIT_STATE_SCHEMA_VERSION,
        "raw_score": raw_score,
        "raw_score_invalid": _is_invalid_raw_score_input(source.get("raw_score")),
        "selected_signal_ids": selected_signal_ids,
        "skipped": skipped,
        "absolute_disqualifier": normalize_absolute_disqualifier(source.get("absolute_disqualifier", False)),
        "no_example_after_followups": _normalize_bool(source.get("no_example_after_followups", False)),
        "verbatim_notes": normalize_verbatim_notes(source.get("verbatim_notes")),
    }
    if skipped:
        normalized["selected_signal_ids"] = []
        normalized["raw_score"] = None
    return normalized


def normalize_app_trait_state(trait_state: dict[str, dict[str, Any]] | None) -> dict[str, dict[str, Any]]:
    if not isinstance(trait_state, dict):
        return {}

    normalized: dict[str, dict[str, Any]] = {}
    for trait_id, state in trait_state.items():
        if not isinstance(state, dict):
            continue
        normalized[canonical_trait_id(trait_id)] = normalize_trait_state_item(state)
    return normalized


def validate_normalized_state(normalized_state: dict[str, dict[str, Any]]) -> None:
    for trait_id, state in normalized_state.items():
        if state.get("raw_score_invalid"):
            raise ReportingValidationError(f"Trait '{trait_id}' has invalid raw_score '{state.get('raw_score')}'.")
        raw_score = state.get("raw_score")
        if raw_score is not None and raw_score not in VALID_RAW_SCORES:
            raise ReportingValidationError(f"Trait '{trait_id}' has invalid raw_score '{raw_score}'.")
        if state.get("skipped"):
            continue
        suggested_raw_score = state.get("suggested_raw_score")
        if (
            suggested_raw_score is not None
            and raw_score is not None
            and suggested_raw_score != raw_score
            and not state.get("adjustment_reason")
        ):
            raise ReportingValidationError(
                f"Trait '{trait_id}' final raw score differs from suggested raw score but adjustment_reason is missing."
            )
        if state.get("absolute_disqualifier") and not state.get("verbatim_notes"):
            raise ReportingValidationError(
                f"Trait '{trait_id}' has disqualifier checked but no verbatim notes."
            )


def invoke_scoring_engine(
    rubric: dict[str, Any],
    track_key: Any,
    normalized_state: dict[str, dict[str, Any]],
    *,
    runtime_bundle: dict[str, Any],
    engine_runtime_contract_path: str | Path,
) -> dict[str, Any]:
    runtime_contract_path = _resolve_contract_path(engine_runtime_contract_path)
    trait_definitions = load_trait_definitions(runtime_bundle)
    validate_runtime_bundle_metadata(
        {
            "runtime_bundle_loaded": bool(runtime_bundle) and not runtime_bundle.get("runtime_error"),
            "runtime_error": runtime_bundle.get("runtime_error"),
            "runtime_contract_path": str(runtime_contract_path),
        }
    )
    _validate_trait_scoring_configuration(rubric, track_key, trait_definitions, normalized_state)

    engine = _build_trait_engine(runtime_bundle, runtime_contract_path)
    selections = _build_trait_selections(trait_definitions, normalized_state)
    resolved_track_key = ScoringEngine._resolve_track_key_for_scoring(rubric, track_key)
    session_result = engine.score_session(trait_definitions, selections, track=resolved_track_key)
    return _build_compatibility_engine_output(
        rubric=rubric,
        track_key=track_key,
        normalized_state=normalized_state,
        trait_definitions=trait_definitions,
        runtime_bundle=runtime_bundle,
        session_result=session_result,
    )


def _validate_trait_scoring_configuration(
    rubric: dict[str, Any],
    track_key: Any,
    trait_definitions: list[dict[str, Any]],
    normalized_state: dict[str, dict[str, Any]],
) -> None:
    runtime_trait_ids = _trait_ids_from_runtime_definitions(trait_definitions)
    input_trait_ids = _trait_ids_from_normalized_state(normalized_state)
    resolved_track_key = ScoringEngine._resolve_track_key_for_scoring(rubric, track_key)
    rubric_trait_ids = _trait_ids_from_rubric(rubric, resolved_track_key)

    _raise_for_missing_trait_overlap(input_trait_ids, runtime_trait_ids)
    _raise_for_rubric_runtime_mismatch(resolved_track_key, rubric_trait_ids, runtime_trait_ids)


def _trait_ids_from_runtime_definitions(trait_definitions: list[dict[str, Any]]) -> set[str]:
    return {
        canonical_trait_id(item.get("trait_id"))
        for item in trait_definitions
        if canonical_trait_id(item.get("trait_id"))
    }


def _trait_ids_from_normalized_state(normalized_state: dict[str, dict[str, Any]]) -> set[str]:
    return {canonical_trait_id(trait_id) for trait_id in normalized_state if canonical_trait_id(trait_id)}


def _trait_ids_from_rubric(rubric: dict[str, Any], resolved_track_key: str) -> set[str]:
    rubric_traits = _rubric_trait_map(rubric, resolved_track_key)
    return {trait_id for trait_id in rubric_traits if trait_id}


def _raise_for_missing_trait_overlap(input_trait_ids: set[str], runtime_trait_ids: set[str]) -> None:
    overlap = sorted(input_trait_ids.intersection(runtime_trait_ids))
    if overlap:
        return
    input_list = ", ".join(sorted(input_trait_ids)) or "<none>"
    runtime_list = ", ".join(sorted(runtime_trait_ids)) or "<none>"
    raise ReportingValidationError(
        "Trait scoring configuration mismatch: finalized trait inputs do not overlap "
        f"the trait runtime bundle. Input traits: {input_list}. Runtime traits: {runtime_list}."
    )


def _raise_for_rubric_runtime_mismatch(
    resolved_track_key: str,
    rubric_trait_ids: set[str],
    runtime_trait_ids: set[str],
) -> None:
    if not rubric_trait_ids:
        raise ReportingValidationError(
            "Trait scoring configuration mismatch: rubric track "
            f"'{resolved_track_key}' does not define any trait-scoring entries."
        )
    missing_runtime_traits = sorted(rubric_trait_ids.difference(runtime_trait_ids))
    if not missing_runtime_traits:
        return
    runtime_list = ", ".join(sorted(runtime_trait_ids)) or "<none>"
    missing_list = ", ".join(missing_runtime_traits)
    raise ReportingValidationError(
        "Trait scoring configuration mismatch: rubric track "
        f"'{resolved_track_key}' includes traits missing from the runtime bundle: {missing_list}. "
        f"Runtime traits: {runtime_list}."
    )


def map_engine_output_to_normalized_shape(
    *,
    rubric: dict[str, Any],
    track_key: Any,
    normalized_state: dict[str, dict[str, Any]],
    engine_output: dict[str, Any],
    runtime_bundle: dict[str, Any],
) -> dict[str, Any]:
    rows = list(engine_output.get("rows", []) or [])
    traits = [_map_trait_row(row) for row in rows]
    resolved_track_key = ScoringEngine._resolve_track_key_for_scoring(rubric, track_key)
    summary = _build_summary(engine_output)
    return {
        "schema_version": OUTPUT_SCHEMA_VERSION,
        "track_key": resolved_track_key,
        "summary": summary,
        "traits": traits,
        "rows": rows,
        "normalized_state": normalized_state,
        "engine_metadata": runtime_bundle,
        **summary,
    }


def _build_summary(engine_output: dict[str, Any]) -> dict[str, Any]:
    configured_max = int(engine_output.get("configured_max_weighted_total", 0) or 0)
    included_max = int(engine_output.get("max_weighted_total_included_traits", 0) or 0)
    denominator = _resolve_scoring_denominator(engine_output, included_max, configured_max)
    percent_label = _resolve_percent_label(engine_output, denominator)
    return {
        "weighted_total": int(engine_output.get("weighted_total", 0) or 0),
        "configured_max_weighted_total": configured_max,
        "max_weighted_total": int(engine_output.get("max_weighted_total", denominator) or denominator),
        "max_weighted_total_included_traits": included_max,
        "percent_denominator": denominator,
        "percent_of_max": float(engine_output.get("percent_of_max", 0.0) or 0.0),
        "percent_of_max_label": percent_label,
        "percent_label": percent_label,
        "outcome": str(engine_output.get("outcome", "") or ""),
        "critical_eq_1": bool(engine_output.get("critical_eq_1", False)),
        "critical_lt_3": bool(engine_output.get("critical_lt_3", False)),
        "any_critical_selected": bool(engine_output.get("any_critical_selected", False)),
        "disqualifier_present": bool(engine_output.get("disqualifier_present", False)),
        "auto_no_hire_present": bool(engine_output.get("auto_no_hire_present", False)),
        "triggered_critical": bool(engine_output.get("triggered_critical", False)),
        "locked_rule": engine_output.get("locked_rule"),
        "override_rationale": engine_output.get("override_rationale"),
        "skipped_traits_count": int(engine_output.get("skipped_traits_count", 0) or 0),
        "scored_traits_count": int(engine_output.get("scored_traits_count", 0) or 0),
    }


def _resolve_scoring_denominator(engine_output: dict[str, Any], included_max: int, configured_max: int) -> int:
    explicit_denominator = engine_output.get("percent_denominator")
    if explicit_denominator is not None:
        return int(explicit_denominator or 0)
    if included_max > 0:
        return included_max
    return configured_max


def _resolve_percent_label(engine_output: dict[str, Any], denominator: int) -> str:
    label = str(engine_output.get("percent_of_max_label", "") or "").strip()
    if label:
        return label
    if denominator <= 0:
        return "N/A (all questions skipped)"
    percent_value = float(engine_output.get("percent_of_max", 0.0) or 0.0)
    return f"{percent_value}%"


def _map_trait_row(row: dict[str, Any]) -> dict[str, Any]:
    signal_counts = row.get("signal_counts", {}) or {}
    return {
        "trait_id": str(row.get("trait_id", "") or ""),
        "trait_name": str(row.get("trait_name", "") or ""),
        "priority": row.get("priority"),
        "weight": int(row.get("weight", 0) or 0),
        "primary_question": str(row.get("primary_question", "") or ""),
        "score": {
            "raw": row.get("raw_score"),
            "raw_for_math": int(row.get("raw_score_math", 0) or 0),
            "weighted": int(row.get("weighted_score", 0) or 0),
            "skipped": bool(row.get("skipped", False)),
        },
        "flags": {
            "absolute_disqualifier": bool(row.get("absolute_disqualifier", False)),
            "no_example_after_followups": bool(row.get("no_example_after_followups", False)),
        },
        "notes": {
            "verbatim": str(row.get("verbatim_notes", "") or ""),
            "question_notes": str(row.get("question_notes", "") or ""),
            "trait_notes": str(row.get("trait_notes", "") or ""),
        },
        "signal_counts": {
            "core": int(signal_counts.get("core", 0) or 0),
            "extended": int(signal_counts.get("extended", 0) or 0),
        },
        "session_trait_outcome": str(row.get("session_trait_outcome", "") or ""),
    }


def _is_invalid_raw_score_input(value: Any) -> bool:
    if value in {None, ""}:
        return False
    if isinstance(value, str) and not value.strip():
        return False
    return coerce_raw_score(value) is None


def _normalize_selected_signal_ids(source: dict[str, Any]) -> list[str]:
    for field_name in ("selected_signal_ids", "selected_signals", "signal_selections"):
        if field_name not in source:
            continue
        normalized = _normalize_selection_variant(source[field_name])
        return normalized
    return []


def _normalize_selection_variant(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, dict):
        return _normalize_selection_mapping(value)
    if isinstance(value, (list, tuple, set)):
        return _normalize_selection_sequence(value)
    return []


def _normalize_selection_mapping(value: dict[str, Any]) -> list[str]:
    if not value:
        return []
    if all(isinstance(item, bool) for item in value.values()):
        return [key.strip() for key, is_selected in value.items() if is_selected and str(key).strip()]

    selected_signal_ids: list[str] = []
    for group_value in value.values():
        selected_signal_ids.extend(_normalize_selection_variant(group_value))
    return _dedupe_signal_ids(selected_signal_ids)


def _normalize_selection_sequence(value: list[Any] | tuple[Any, ...] | set[Any]) -> list[str]:
    selected_signal_ids: list[str] = []
    for item in value:
        signal_id = _normalize_selection_item(item)
        if signal_id:
            selected_signal_ids.append(signal_id)
    return _dedupe_signal_ids(selected_signal_ids)


def _normalize_selection_item(item: Any) -> str:
    if isinstance(item, str):
        return item.strip()
    if not isinstance(item, dict):
        return ""
    if not item.get("selected"):
        return ""
    for field_name in ("signal_id", "ref", "id", "value"):
        candidate = str(item.get(field_name, "") or "").strip()
        if candidate:
            return candidate
    return ""


def _dedupe_signal_ids(signal_ids: list[str]) -> list[str]:
    return list(dict.fromkeys(signal_id for signal_id in signal_ids if signal_id))


def _normalize_bool(value: Any) -> bool:
    if isinstance(value, str):
        stripped = value.strip().lower()
        if stripped in {"true", "1", "yes", "on"}:
            return True
        if stripped in {"false", "0", "no", "off", ""}:
            return False
    return bool(value)


def _contract_resolution_base_dir() -> Path:
    return Path(__file__).resolve().parent.parent


def _resolve_contract_path_from_base(path: str | Path, *, base_dir: Path) -> Path:
    candidate = Path(path).expanduser()
    if candidate.is_absolute():
        return candidate.resolve()
    return (base_dir / candidate).resolve()


def _resolve_contract_path(path: str | Path) -> Path:
    return _resolve_contract_path_from_base(path, base_dir=_contract_resolution_base_dir())


def _load_yaml(path: Path) -> dict[str, Any]:
    resolved = path.expanduser().resolve()
    if not resolved.exists():
        return {}
    payload = yaml.safe_load(resolved.read_text(encoding="utf-8"))
    if isinstance(payload, dict):
        return payload
    return {}


def _load_runtime_bundle(runtime_contract_path: Path) -> dict[str, Any]:
    resolved_path = runtime_contract_path.expanduser().resolve()
    if not resolved_path.exists():
        return {}

    try:
        engine_class = _load_trait_engine_class(resolved_path)
        config, signal_dictionary, traits, resolved_paths = engine_class.load_runtime_bundle(resolved_path)
    except (FileNotFoundError, ImportError, KeyError, PermissionError, TypeError, ValueError) as exc:
        return {"runtime_error": str(exc), "traits": []}

    return {
        "config": config,
        "signal_dictionary": signal_dictionary,
        "traits": traits,
        "resolved_paths": {key: str(value) for key, value in resolved_paths.items()},
    }


def _load_trait_engine_class(runtime_contract_path: Path) -> type[Any]:
    engine_module_path = runtime_contract_path.resolve().parent / "trait_based_scoring_engine.py"
    spec = importlib.util.spec_from_file_location("trait_based_scoring_engine_adapter", engine_module_path)
    if spec is None or spec.loader is None:
        raise ImportError(f"Unable to load scoring engine module from {engine_module_path}.")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module.ScoringEngine


def _has_trait_definition_overlap(
    trait_definitions: list[dict[str, Any]],
    normalized_state: dict[str, dict[str, Any]],
) -> bool:
    trait_ids = {canonical_trait_id(item.get("trait_id")) for item in trait_definitions if canonical_trait_id(item.get("trait_id"))}
    normalized_ids = {canonical_trait_id(trait_id) for trait_id in normalized_state if canonical_trait_id(trait_id)}
    return bool(trait_ids.intersection(normalized_ids))


def _build_trait_engine(runtime_bundle: dict[str, Any], runtime_contract_path: Path) -> Any:
    engine_class = _load_trait_engine_class(Path(runtime_contract_path))
    return engine_class(runtime_bundle["config"], runtime_bundle["signal_dictionary"])


def _build_trait_selections(
    trait_definitions: list[dict[str, Any]],
    normalized_state: dict[str, dict[str, Any]],
) -> dict[str, list[str]]:
    selections: dict[str, list[str]] = {}
    for trait_definition in trait_definitions:
        trait_id = canonical_trait_id(trait_definition.get("trait_id"))
        if not trait_id:
            continue
        state = normalized_state.get(trait_id) or {}
        selections[trait_id] = _select_signal_refs_for_state(trait_definition, state)
    return selections


def _select_signal_refs_for_state(trait_definition: dict[str, Any], state: dict[str, Any]) -> list[str]:
    if state.get("skipped"):
        return []
    selected_signal_ids = {
        str(signal_id).strip()
        for signal_id in (state.get("selected_signal_ids", []) or [])
        if str(signal_id).strip()
    }
    resolved_refs: list[str] = []
    for signal in _iter_trait_signals(trait_definition):
        runtime_signal_id = resolve_trait_signal_runtime_id(signal)
        if not runtime_signal_id:
            continue
        is_auto_no_hire = (
            bool(signal.get("is_auto_no_hire", False))
            or bool(signal.get("auto_no_hire", False))
            or str(signal.get("signal_category", "") or "").strip().lower() == "automatic_no_hire"
            or str(signal.get("base_weight", "") or "").strip().upper() == "AUTO_NO_HIRE"
        )
        if resolve_trait_signal_weight(signal) == 0 and not is_auto_no_hire:
            continue
        if not selected_signal_ids.intersection(signal_selection_aliases(signal)):
            continue
        resolved_refs.append(runtime_signal_id)
    return resolved_refs


def _positive_signal_refs(trait_definition: dict[str, Any]) -> list[str]:
    return _signal_refs_by_weight(trait_definition, positive=True)


def _negative_signal_refs(trait_definition: dict[str, Any]) -> list[str]:
    return _signal_refs_by_weight(trait_definition, positive=False)


def _signal_refs_by_weight(trait_definition: dict[str, Any], *, positive: bool) -> list[str]:
    refs: list[str] = []
    for signal in _iter_trait_signals(trait_definition):
        weight = resolve_trait_signal_weight(signal)
        signal_id = resolve_trait_signal_selection_id(signal)
        if positive and weight > 0 and signal_id:
            refs.append(signal_id)
        if (not positive) and weight < 0 and signal_id:
            refs.append(signal_id)
    return refs


def _iter_trait_signals(trait_definition: dict[str, Any]) -> list[dict[str, Any]]:
    return iter_trait_schema_signals(trait_definition)


def _build_compatibility_engine_output(
    *,
    rubric: dict[str, Any],
    track_key: Any,
    normalized_state: dict[str, dict[str, Any]],
    trait_definitions: list[dict[str, Any]],
    runtime_bundle: dict[str, Any],
    session_result: dict[str, Any],
) -> dict[str, Any]:
    session_traits = {
        canonical_trait_id(item.get("trait_id")): item
        for item in session_result.get("traits", []) or []
        if canonical_trait_id(item.get("trait_id"))
    }
    resolved_track_key = ScoringEngine._resolve_track_key_for_scoring(rubric, track_key)
    rubric_trait_map = _rubric_trait_map(rubric, resolved_track_key)
    rows = _build_rows(trait_definitions, rubric_trait_map, normalized_state, session_traits, session_result, runtime_bundle)
    weighted_total = sum(int(row.get("weighted_score", 0) or 0) for row in rows)
    configured_max_weighted_total = _configured_max_weighted_total(rubric, resolved_track_key, trait_definitions)
    included_max_weighted_total = _max_weighted_total(trait_definitions, normalized_state)
    percent_denominator = _resolve_percent_denominator(included_max_weighted_total, configured_max_weighted_total)
    percent_of_max = _percent_of_max(weighted_total, percent_denominator)
    percent_label = _percent_label(percent_of_max, included_max_weighted_total)
    missing_required_scores = any(
        not row.get("skipped", False) and row.get("raw_score") is None
        for row in rows
    )
    disqualifier_present = any(bool(state.get("absolute_disqualifier")) for state in normalized_state.values())
    any_critical_selected = bool(session_result.get("any_critical_selected", False))
    triggered_critical = bool(session_result.get("triggered_critical", False))
    critical_eq_1 = any(
        str(row.get("priority") or "").strip().lower() == "critical"
        and row.get("raw_score") == 1
        and not row.get("skipped", False)
        for row in rows
    )
    critical_lt_3 = any(
        str(row.get("priority") or "").strip().lower() == "critical"
        and row.get("raw_score") is not None
        and int(row.get("raw_score") or 0) < 3
        and not row.get("skipped", False)
        for row in rows
    )
    locked_rule = session_result.get("locked_rule")
    override_rationale = session_result.get("override_rationale")
    if missing_required_scores:
        outcome = "Incomplete"
        locked_rule = "One or more applicable traits are missing final raw scores"
    elif disqualifier_present:
        outcome = "No Hire"
        locked_rule = "Any Absolute Disqualifier observed => Immediate NO HIRE"
    elif critical_eq_1:
        outcome = "No Hire"
        locked_rule = "Any Critical trait raw score = 1 => Immediate NO HIRE"
    elif critical_lt_3:
        outcome = "No Hire"
        locked_rule = "Any Critical trait raw score < 3 => Cannot assign HIRE"
    elif percent_of_max >= 80:
        outcome = "Hire"
    elif percent_of_max >= 65:
        outcome = "Borderline"
    else:
        outcome = "No Hire"
    return {
        "rows": rows,
        "weighted_total": weighted_total,
        "configured_max_weighted_total": configured_max_weighted_total,
        "max_weighted_total": percent_denominator,
        "max_weighted_total_included_traits": included_max_weighted_total,
        "percent_denominator": percent_denominator,
        "percent_of_max": percent_of_max,
        "percent_of_max_label": percent_label,
        "percent_label": percent_label,
        "skipped_traits_count": sum(1 for state in normalized_state.values() if state.get("skipped")),
        "scored_traits_count": sum(1 for state in normalized_state.values() if not state.get("skipped")),
        "critical_eq_1": critical_eq_1,
        "critical_lt_3": critical_lt_3,
        "any_critical_selected": any_critical_selected,
        "disqualifier_present": disqualifier_present,
        "triggered_critical": triggered_critical,
        "locked_rule": locked_rule,
        "override_rationale": override_rationale,
        "outcome": outcome,
        "session_result": session_result,
    }


def _rubric_trait_map(rubric: dict[str, Any], resolved_track_key: str) -> dict[str, dict[str, Any]]:
    rubric_traits = rubric.get("traits", []) or []
    mapping: dict[str, dict[str, Any]] = {}
    for trait in rubric_traits:
        tracks = trait.get("applicable_tracks", []) or []
        if tracks and "all" not in tracks and resolved_track_key not in tracks:
            continue
        trait_id = canonical_trait_id(trait.get("id"))
        if trait_id:
            mapping[trait_id] = trait
    return mapping


def _build_rows(
    trait_definitions: list[dict[str, Any]],
    rubric_trait_map: dict[str, dict[str, Any]],
    normalized_state: dict[str, dict[str, Any]],
    session_traits: dict[str, dict[str, Any]],
    session_result: dict[str, Any],
    runtime_bundle: dict[str, Any],
) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for trait_definition in trait_definitions:
        trait_id = canonical_trait_id(trait_definition.get("trait_id"))
        if not trait_id:
            continue
        state = normalized_state.get(trait_id) or {}
        session_trait = session_traits.get(trait_id, {})
        rubric_trait = rubric_trait_map.get(trait_id, {})
        rows.append(_build_trait_row(trait_definition, rubric_trait, state, session_trait, session_result, runtime_bundle))
    return rows


def _build_trait_row(
    trait_definition: dict[str, Any],
    rubric_trait: dict[str, Any],
    state: dict[str, Any],
    session_trait: dict[str, Any],
    session_result: dict[str, Any],
    runtime_bundle: dict[str, Any],
) -> dict[str, Any]:
    raw_score = state.get("raw_score")
    skipped = bool(state.get("skipped", False)) or raw_score is None
    canonical_id = canonical_trait_id(trait_definition.get("trait_id"))
    trait_label = str(rubric_trait.get("id") or canonical_id or "")
    weight = int(rubric_trait.get("weight", 0) or 0)
    system_checkbox_score = 0 if skipped else int(raw_score or 0) * weight
    return {
        "trait_id": canonical_id,
        "trait_name": str(rubric_trait.get("name") or trait_label),
        "priority": rubric_trait.get("priority"),
        "weight": weight,
        "skipped": skipped,
        "raw_score": raw_score,
        "raw_score_math": int(raw_score or 0),
        "weighted_score": system_checkbox_score,
        "system_checkbox_score": system_checkbox_score,
        "verbatim_notes": str(state.get("verbatim_notes", "") or ""),
        "no_example_after_followups": bool(state.get("no_example_after_followups", False)),
        "absolute_disqualifier": bool(state.get("absolute_disqualifier", False)),
        "primary_question": str(rubric_trait.get("primary_question") or trait_definition.get("question", "") or ""),
        "question_notes": "",
        "trait_notes": "",
        "signal_counts": {
            "core": len(session_trait.get("selected_core", []) or []),
            "extended": len(session_trait.get("selected_extended", []) or []),
        },
        "selected_signal_ids": list(state.get("selected_signal_ids", []) or []),
        "session_trait_outcome": str(session_result.get("decision", "") or ""),
        "trait_aliases": trait_id_aliases(trait_definition.get("trait_id")),
    }


def _max_weighted_total(
    trait_definitions: list[dict[str, Any]],
    normalized_state: dict[str, dict[str, Any]],
) -> int:
    total = 0.0
    for trait_definition in trait_definitions:
        trait_id = canonical_trait_id(trait_definition.get("trait_id"))
        state = normalized_state.get(trait_id) or {}
        if state.get("skipped") or state.get("raw_score") is None:
            continue
        total += _max_trait_final_score(trait_definition)
    return int(round(total))


def _max_trait_final_score(trait_definition: dict[str, Any]) -> float:
    core_total = sum(
        max(resolve_trait_signal_weight(signal), 0.0)
        for signal in trait_definition.get("core_signals", []) or []
        if isinstance(signal, dict)
    )
    extended_total = sum(max(resolve_trait_signal_weight(signal), 0.0) for signal in _iter_extended_trait_signals(trait_definition))
    return core_total + extended_total


def _iter_extended_trait_signals(trait_definition: dict[str, Any]) -> list[dict[str, Any]]:
    return iter_trait_schema_signals(
        {
            "extended_signal_groups": trait_definition.get("extended_signal_groups", []),
            "extended_signals": trait_definition.get("extended_signals", []),
        }
    )


def _configured_max_weighted_total(
    rubric: dict[str, Any],
    resolved_track_key: str,
    trait_definitions: list[dict[str, Any]],
) -> int:
    track_cfg = ((rubric.get("tracks") or {}).get(resolved_track_key) or {})
    configured_value = track_cfg.get("max_weighted_total")
    if configured_value is not None:
        return int(configured_value or 0)
    return int(round(sum(_max_trait_final_score(trait_definition) for trait_definition in trait_definitions)))


def _resolve_percent_denominator(included_max_weighted_total: int, configured_max_weighted_total: int) -> int:
    if included_max_weighted_total > 0:
        return included_max_weighted_total
    return configured_max_weighted_total


def _percent_of_max(weighted_total: int, max_weighted_total: int) -> float:
    if max_weighted_total <= 0:
        return 0.0
    return round((float(weighted_total) * 100.0) / float(max_weighted_total), 2)


def _percent_label(percent_of_max: float, included_max_weighted_total: int) -> str:
    if included_max_weighted_total <= 0:
        return "N/A (all questions skipped)"
    return f"{percent_of_max}%"
_COMPAT_MODULES: tuple[str, ...] = (
    "trait_scoring_adapter",
)

_WRAPPER_POLICY = (
    "Legacy scoring/reporting modules are compatibility wrappers during flattening. "
    "New production imports should prefer scoring_reporting."
)


def available_modules() -> tuple[str, ...]:
    return _COMPAT_MODULES


def module_ownership() -> dict[str, str]:
    return {module_name: "scoring_reporting" for module_name in _COMPAT_MODULES}


def wrapper_policy() -> str:
    return _WRAPPER_POLICY


def load_compat_module(module_name: str) -> ModuleType:
    if module_name not in _COMPAT_MODULES:
        raise AttributeError(f"{module_name!r} is not part of scoring_reporting")
    return import_module(module_name)


def public_symbols(module_name: str | None = None) -> tuple[str, ...]:
    module_names = (module_name,) if module_name is not None else _COMPAT_MODULES
    symbols: set[str] = set()
    for compat_name in module_names:
        module = load_compat_module(compat_name)
        symbols.update(name for name in dir(module) if not name.startswith("_"))
    return tuple(sorted(symbols))


def resolve_compat_symbol(symbol_name: str) -> Any:
    for module_name in _COMPAT_MODULES:
        module = import_module(module_name)
        if hasattr(module, symbol_name):
            return getattr(module, symbol_name)
    raise AttributeError(f"scoring_reporting has no attribute {symbol_name!r}")


def __getattr__(name: str) -> Any:
    if name.startswith("__"):
        raise AttributeError(f"scoring_reporting has no attribute {name!r}")
    if name in _COMPAT_MODULES:
        return import_module(name)
    return resolve_compat_symbol(name)


def __dir__() -> list[str]:
    return sorted(set(globals()) | set(_COMPAT_MODULES) | set(public_symbols()))

