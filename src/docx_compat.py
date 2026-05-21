"""Compatibility layer for Word document operations.

Prefers ``python-docx`` when available. Falls back to a lightweight
in-repo implementation for test and constrained environments where a
conflicting legacy ``docx`` package is installed.
"""

from __future__ import annotations

from dataclasses import dataclass
import json
from pathlib import Path
from typing import Any


try:  # pragma: no cover - exercised indirectly in environments with python-docx
    from docx import Document as Document  # type: ignore[attr-defined]

    BACKEND = "python-docx"
except Exception:  # pragma: no cover - fallback used in constrained environments
    BACKEND = "fallback"

    @dataclass
    class _Paragraph:
        text: str = ""

    class _Cell:
        def __init__(self, text: str = "") -> None:
            self.paragraphs: list[_Paragraph] = [_Paragraph(text)]

        @property
        def text(self) -> str:
            return "\n".join(p.text for p in self.paragraphs)

        @text.setter
        def text(self, value: str) -> None:
            self.paragraphs = [_Paragraph(value)]

    class _Row:
        def __init__(self, cols: int) -> None:
            self.cells: list[_Cell] = [_Cell() for _ in range(cols)]

    class _Table:
        def __init__(self, rows: int, cols: int) -> None:
            self._cols = cols
            self.rows: list[_Row] = [_Row(cols) for _ in range(rows)]

        def add_row(self) -> _Row:
            row = _Row(self._cols)
            self.rows.append(row)
            return row

    class Document:  # noqa: D401 - mirrors python-docx API
        """Small subset of python-docx's ``Document`` interface."""

        def __init__(self, path: str | Path | None = None) -> None:
            self.paragraphs: list[_Paragraph] = []
            self.tables: list[_Table] = []
            if path is not None:
                self._load(path)

        def add_heading(self, text: str, level: int = 1) -> _Paragraph:
            prefix = "#" * max(1, int(level))
            return self.add_paragraph(f"{prefix} {text}")

        def add_paragraph(self, text: str = "") -> _Paragraph:
            paragraph = _Paragraph(text)
            self.paragraphs.append(paragraph)
            return paragraph

        def add_table(self, rows: int, cols: int) -> _Table:
            table = _Table(rows, cols)
            self.tables.append(table)
            return table

        def save(self, path: str | Path) -> None:
            payload = {
                "paragraphs": [p.text for p in self.paragraphs],
                "tables": [[[cell.text for cell in row.cells] for row in table.rows] for table in self.tables],
            }
            Path(path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

        def _load(self, path: str | Path) -> None:
            target = Path(path)
            if not target.exists():
                return
            try:
                data: dict[str, Any] = json.loads(target.read_text(encoding="utf-8"))
            except Exception:
                self.paragraphs = []
                self.tables = []
                return
            self.paragraphs = [_Paragraph(str(text)) for text in data.get("paragraphs", [])]
            self.tables = []
            for table_rows in data.get("tables", []):
                cols = len(table_rows[0]) if table_rows else 0
                table = _Table(0, cols)
                table.rows = []
                for row_data in table_rows:
                    row = _Row(cols)
                    for idx, value in enumerate(row_data):
                        if idx < len(row.cells):
                            row.cells[idx].text = str(value)
                    table.rows.append(row)
                self.tables.append(table)
